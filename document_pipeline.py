from __future__ import annotations

from copy import deepcopy
from datetime import datetime, timezone
from html import escape
import json
from pathlib import Path
import re
from typing import Any, Dict, Iterable, List, Optional
from urllib.parse import urlencode

from complaint_phases import ComplaintPhase
from complaint_phases.intake_claim_registry import (
    normalize_claim_type,
    registry_element_for_claim_type,
)
from claim_support_review import summarize_claim_reasoning_review
from document_optimization import (
    AgenticDocumentOptimizer,
    _build_claim_reasoning_theorem_export_metadata,
)
from intake_status import (
    build_intake_case_review_summary,
    build_intake_status_summary,
    summarize_temporal_issue_registry,
)
from workflow_phase_guidance import (
    build_drafting_document_generation_phase_guidance,
    build_graph_analysis_phase_guidance,
    build_workflow_phase_plan,
    build_workflow_phase_warning_entries,
)


DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parent / "tmp" / "generated_documents"
DEFAULT_RELIEF = [
    "Compensatory damages in an amount to be proven at trial.",
    "Pre- and post-judgment interest as allowed by law.",
    "Reasonable attorney's fees and costs where authorized.",
    "Injunctive and declaratory relief sufficient to stop the unlawful conduct.",
    "Such other and further relief as the Court deems just and proper.",
]

STATE_DEFAULT_RELIEF = [
    "General and special damages according to proof.",
    "Costs of suit incurred herein.",
    "Such other and further relief as the Court deems just and proper.",
]
ACTOR_CRITIC_PHASE_FOCUS_ORDER = ("graph_analysis", "document_generation", "intake_questioning")
DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS = {
    "empathy": 0.22,
    "question_quality": 0.58,
    "information_extraction": 0.40,
    "coverage": 0.40,
    "efficiency": 0.62,
}
_CONFIRMATION_PLACEHOLDER_PATTERN = re.compile(
    r"\b(?:needs?\s+confirmation|to\s+be\s+confirmed|confirm(?:ed|ation)?\s+pending|tbd|unknown|not\s+sure|unclear|pending)\b",
    flags=re.IGNORECASE,
)


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _slugify(value: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "-", str(value or "document").strip().lower())
    return text.strip("-") or "document"


def _unique_preserving_order(values: Iterable[str]) -> List[str]:
    seen = set()
    unique: List[str] = []
    for value in values:
        text = str(value or "").strip()
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(text)
    return unique


def _coerce_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if value is None:
        return []
    return [value]


def _dedupe_text_values(values: Iterable[Any]) -> List[str]:
    seen = set()
    normalized_values: List[str] = []
    for value in values:
        text = str(value or "").strip()
        if not text:
            continue
        if text in seen:
            continue
        seen.add(text)
        normalized_values.append(text)
    return normalized_values


def _merge_chip_labels(existing: Any, additions: Any) -> List[str]:
    merged: List[str] = []
    seen = set()
    for source in (existing, additions):
        if not isinstance(source, list):
            continue
        for item in source:
            label = str(item or "").strip()
            if not label or label in seen:
                continue
            seen.add(label)
            merged.append(label)
    return merged


def _humanize_checklist_label(value: Any) -> str:
    text = str(value or "").strip().replace("_", " ")
    return text.title() if text else ""


def _build_claim_checklist_chip_labels(claim: Dict[str, Any]) -> List[str]:
    if not isinstance(claim, dict):
        return []

    chip_labels: List[str] = []
    claim_status = _humanize_checklist_label(claim.get("status"))
    if claim_status:
        chip_labels.append(f"claim status: {claim_status}")

    temporal_gap_hint_count = int(claim.get("temporal_gap_hint_count") or 0)
    if temporal_gap_hint_count > 0:
        chip_labels.append(f"chronology gaps: {temporal_gap_hint_count}")

    proof_gap_count = int(claim.get("proof_gap_count") or 0)
    if proof_gap_count > 0:
        chip_labels.append(f"proof gaps: {proof_gap_count}")

    unresolved_element_count = int(claim.get("unresolved_element_count") or 0)
    if unresolved_element_count > 0:
        chip_labels.append(f"unresolved elements: {unresolved_element_count}")

    contradiction_candidate_count = int(claim.get("contradiction_candidate_count") or 0)
    if contradiction_candidate_count > 0:
        chip_labels.append(f"contradiction candidates: {contradiction_candidate_count}")

    claim_unresolved_temporal_issue_count = int(claim.get("claim_unresolved_temporal_issue_count") or 0)
    if claim_unresolved_temporal_issue_count > 0:
        chip_labels.append(f"unresolved temporal issues: {claim_unresolved_temporal_issue_count}")

    claim_missing_temporal_predicates = _extract_text_candidates(claim.get("claim_missing_temporal_predicates"))
    if claim_missing_temporal_predicates:
        chip_labels.append(f"missing temporal predicates: {len(claim_missing_temporal_predicates)}")

    claim_required_provenance_kinds = _extract_text_candidates(claim.get("claim_required_provenance_kinds"))
    if claim_required_provenance_kinds:
        chip_labels.append(f"required provenance kinds: {len(claim_required_provenance_kinds)}")

    return chip_labels


def _collect_temporal_registry_identifiers(records: Any, *keys: str) -> List[str]:
    identifiers: List[str] = []
    for record in records if isinstance(records, list) else []:
        if not isinstance(record, dict):
            continue
        for key in keys:
            value = str(record.get(key) or "").strip()
            if value:
                identifiers.append(value)
                break
    return _dedupe_text_values(identifiers)


def _collect_unresolved_temporal_issue_identifiers(records: Any) -> List[str]:
    unresolved_issue_ids: List[str] = []
    for record in records if isinstance(records, list) else []:
        if not isinstance(record, dict):
            continue
        status = str(record.get("current_resolution_status") or record.get("status") or "open").strip().lower()
        if status in {"resolved", "closed", "complete", "completed"}:
            continue
        for key in ("temporal_issue_id", "issue_id", "timeline_issue_id"):
            value = str(record.get(key) or "").strip()
            if value:
                unresolved_issue_ids.append(value)
                break
    return _dedupe_text_values(unresolved_issue_ids)


def _normalize_identifier_list(values: Any) -> List[str]:
    identifiers: List[str] = []
    seen = set()
    for value in _coerce_list(values):
        text = str(value or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        identifiers.append(text)
    return identifiers


def _format_timeline_date(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        return datetime.strptime(text, "%Y-%m-%d").strftime("%B %-d, %Y")
    except ValueError:
        return text


def _chronology_fact_label(fact: Dict[str, Any]) -> str:
    event_label = str((fact if isinstance(fact, dict) else {}).get("event_label") or "").strip()
    if event_label:
        return event_label
    predicate_family = str((fact if isinstance(fact, dict) else {}).get("predicate_family") or "").strip().replace("_", " ")
    if predicate_family:
        return predicate_family.title()
    return "Event"


def _join_chronology_segments(segments: List[str]) -> str:
    if not segments:
        return ""
    if len(segments) == 1:
        return segments[0]
    if len(segments) == 2:
        return f"{segments[0]} and {segments[1]}"
    return f"{', '.join(segments[:-1])}, and {segments[-1]}"


def _build_anchored_chronology_summary_from_case_file(intake_case_file: Dict[str, Any], *, limit: int = 3) -> List[str]:
    case_file = intake_case_file if isinstance(intake_case_file, dict) else {}
    facts = [dict(item) for item in list(case_file.get("canonical_facts") or []) if isinstance(item, dict)]
    relations = [dict(item) for item in list(case_file.get("timeline_relations") or []) if isinstance(item, dict)]
    if not facts or not relations:
        return []

    fact_by_id = {
        str(fact.get("fact_id") or "").strip(): fact
        for fact in facts
        if str(fact.get("fact_id") or "").strip()
    }
    relation_records = []
    for relation in relations:
        if str(relation.get("relation_type") or "").strip().lower() != "before":
            continue
        source_id = str(relation.get("source_fact_id") or "").strip()
        target_id = str(relation.get("target_fact_id") or "").strip()
        source_fact = fact_by_id.get(source_id)
        target_fact = fact_by_id.get(target_id)
        if not source_fact or not target_fact:
            continue
        source_date = _format_timeline_date((source_fact.get("temporal_context") or {}).get("start_date") or relation.get("source_start_date"))
        target_date = _format_timeline_date((target_fact.get("temporal_context") or {}).get("start_date") or relation.get("target_start_date"))
        if not source_date or not target_date:
            continue
        relation_records.append(
            {
                "key": (source_id, target_id),
                "source_id": source_id,
                "target_id": target_id,
                "source_fact": source_fact,
                "target_fact": target_fact,
                "source_date": source_date,
                "target_date": target_date,
            }
        )
    if not relation_records:
        return []

    outgoing: Dict[str, List[Dict[str, Any]]] = {}
    incoming_count: Dict[str, int] = {}
    for record in relation_records:
        outgoing.setdefault(record["source_id"], []).append(record)
        incoming_count[record["target_id"]] = incoming_count.get(record["target_id"], 0) + 1
        incoming_count.setdefault(record["source_id"], incoming_count.get(record["source_id"], 0))

    lines: List[str] = []
    seen = set()
    used_keys = set()
    for record in relation_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        if incoming_count.get(record["source_id"], 0) != 0 or len(outgoing.get(record["source_id"], [])) != 1:
            continue
        chain = [record]
        next_id = record["target_id"]
        temp_used = {record["key"]}
        while len(outgoing.get(next_id, [])) == 1 and incoming_count.get(next_id, 0) == 1:
            next_record = outgoing[next_id][0]
            if next_record["key"] in temp_used:
                break
            chain.append(next_record)
            temp_used.add(next_record["key"])
            next_id = next_record["target_id"]
        if len(chain) < 2:
            continue
        segments = [
            f"{_chronology_fact_label(chain[0]['source_fact'])} on {chain[0]['source_date']}"
        ]
        segments.extend(
            f"{_chronology_fact_label(item['target_fact'])} on {item['target_date']}"
            for item in chain
        )
        line = f"{_join_chronology_segments(segments)} occurred in sequence."
        last_target = chain[-1]["target_fact"]
        target_context = last_target.get("temporal_context") if isinstance(last_target.get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f". The later date is currently derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        used_keys.update(temp_used)
        lines.append(line)

    for record in relation_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        source_label = _chronology_fact_label(record["source_fact"])
        target_label = _chronology_fact_label(record["target_fact"])
        line = f"{source_label} on {record['source_date']} preceded {target_label.lower()} on {record['target_date']}."
        target_context = record["target_fact"].get("temporal_context") if isinstance(record["target_fact"].get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f". The later date is currently derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        lines.append(line)
    return lines


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return float(default)


def _coerce_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    text = str(value or "").strip().lower()
    if not text:
        return default
    if text in {"1", "true", "yes", "y", "on", "available", "enabled"}:
        return True
    if text in {"0", "false", "no", "n", "off", "unavailable", "disabled"}:
        return False
    return default


def _extract_text_candidates(value: Any) -> List[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        results: List[str] = []
        for item in value:
            results.extend(_extract_text_candidates(item))
        return results
    if isinstance(value, dict):
        keys = (
            "fact",
            "fact_text",
            "text",
            "summary",
            "description",
            "name",
            "parsed_text_preview",
            "claim_element",
            "claim_element_text",
            "answer",
            "question",
            "title",
            "relevance",
        )
        results = []
        for key in keys:
            if key in value and value.get(key):
                results.extend(_extract_text_candidates(value.get(key)))
        return results
    return []


def _contains_date_anchor(value: Any) -> bool:
    text = str(value or "")
    return bool(
        re.search(
            r"\b(?:\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{1,2}(?:,\s+\d{2,4})?)\b",
            text,
            flags=re.IGNORECASE,
        )
    )


def _contains_actor_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "who at hacc",
            "caseworker",
            "housing specialist",
            "program manager",
            "hearing officer",
            "staff",
            "supervisor",
            "director",
            "coordinator",
            "name",
            "title",
        )
    )


def _contains_causation_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    if not lowered:
        return False
    return (
        any(marker in lowered for marker in ("because", "as a result", "after", "following", "in retaliation", "retaliat", "days after", "weeks after", "shortly after"))
        and any(marker in lowered for marker in ("complained", "reported", "grievance", "appeal", "protected activity", "requested accommodation"))
        and any(marker in lowered for marker in ("adverse action", "termination", "denial", "loss of assistance", "retaliat"))
    )


def _contains_hearing_timing_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    if not lowered:
        return False
    timing_markers = (
        "hearing request",
        "requested a hearing",
        "requested review",
        "review request",
        "informal hearing request",
        "informal review request",
        "requested an informal review",
        "requested a grievance hearing",
        "requested an appeal",
        "grievance request",
        "grievance hearing",
        "submitted a grievance",
        "submitted a request for review",
        "review was scheduled",
        "hearing was scheduled",
        "informal review was scheduled",
        "hearing on ",
        "review on ",
    )
    return any(marker in lowered for marker in timing_markers) and (
        _contains_date_anchor(value)
        or any(marker in lowered for marker in ("after", "before", "within", "days", "weeks"))
    )


def _contains_response_date_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    if not lowered or not _contains_date_anchor(value):
        return False
    return any(
        marker in lowered
        for marker in (
            "response date",
            "responded on",
            "response was",
            "review decision",
            "hearing outcome",
            "notice date",
            "decision date",
            "denial notice",
            "written notice",
            "sent plaintiff",
            "issued",
            "mailed",
            "dated",
            "decision denying assistance",
            "final decision",
        )
    )


def _contains_staff_identity_marker(value: Any) -> bool:
    text = str(value or "").strip()
    lowered = text.lower()
    if not lowered:
        return False
    references_actor = "hacc" in lowered or "housing authority" in lowered
    references_role = any(
        marker in lowered
        for marker in (
            "staff",
            "caseworker",
            "manager",
            "officer",
            "specialist",
            "director",
            "coordinator",
            "supervisor",
            "hearing officer",
        )
    )
    has_named_person = bool(
        re.search(r"\b[A-Z][a-z]+ [A-Z][a-z]+\b", text)
    )
    return references_actor and (
        (references_role and has_named_person)
        or ("signed by" in lowered and has_named_person)
        or ("issued by" in lowered and has_named_person)
    )


def _contains_sequence_timing_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "days after",
            "weeks after",
            "before",
            "after",
            "thereafter",
            "shortly after",
            "in close sequence",
        )
    )


def _extract_date_anchor_text(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    match = re.search(
        r"\b(\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{1,2}(?:,\s+\d{2,4})?)\b",
        text,
        flags=re.IGNORECASE,
    )
    return match.group(1).strip() if match else ""


def _extract_dated_event_clause(value: Any) -> tuple[str, str]:
    text = re.sub(r"\s+", " ", str(value or "")).strip().rstrip(".")
    if not text:
        return "", ""
    match = re.match(
        r"^(?:On\s+)?(\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2}(?:,\s+\d{2,4})?),?\s+(.*)$",
        text,
        flags=re.IGNORECASE,
    )
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return _extract_date_anchor_text(text), text


def _contains_confirmation_placeholder(value: Any) -> bool:
    return bool(_CONFIRMATION_PLACEHOLDER_PATTERN.search(str(value or "")))


def _extract_latest_adversarial_priority_findings(value: Any) -> List[str]:
    findings: List[str] = []
    if isinstance(value, str):
        text = value.strip()
        if text:
            findings.append(text)
        return findings
    if isinstance(value, list):
        for item in value:
            findings.extend(_extract_latest_adversarial_priority_findings(item))
        return findings
    if isinstance(value, dict):
        candidate_keys = (
            "summary",
            "finding",
            "message",
            "reason",
            "title",
            "description",
            "priority",
            "text",
        )
        for key in candidate_keys:
            if key in value and value.get(key):
                findings.extend(_extract_latest_adversarial_priority_findings(value.get(key)))
        for nested_key in (
            "priorities",
            "priority_findings",
            "findings",
            "issues",
            "gaps",
            "latest_priorities",
            "latest_adversarial_priorities",
            "critical_findings",
            "latest_batch_findings",
            "latest_batch_priorities",
            "latest_adversarial_batch",
            "latest_adversarial_batch_summary",
        ):
            if nested_key in value and value.get(nested_key):
                findings.extend(_extract_latest_adversarial_priority_findings(value.get(nested_key)))
    return findings


def _has_chronology_gap_priority(findings: Iterable[str]) -> bool:
    lowered = " ".join(str(item or "").strip().lower() for item in findings if str(item or "").strip())
    if not lowered:
        return False
    chronology_markers = (
        "chronology",
        "exact date",
        "date gap",
        "response timing",
        "response date",
        "sequence",
        "timeline",
        "follow up",
        "follow-up",
        "critical chronology",
    )
    return any(marker in lowered for marker in chronology_markers)


def _has_decision_or_document_precision_priority(findings: Iterable[str]) -> bool:
    lowered = " ".join(str(item or "").strip().lower() for item in findings if str(item or "").strip())
    if not lowered:
        return False
    precision_markers = (
        "decision-maker",
        "decision maker",
        "adverse action",
        "documentary artifact",
        "document artifact",
        "documentary evidence",
        "specific decision",
        "who made",
        "who decided",
        "named actor",
        "artifact precision",
    )
    return any(marker in lowered for marker in precision_markers)


def _roman(index: int) -> str:
    numerals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    value = max(1, int(index))
    result = []
    for number, symbol in numerals:
        while value >= number:
            result.append(symbol)
            value -= number
    return "".join(result)


def _safe_call(target: Any, method_name: str, *args: Any, **kwargs: Any) -> Any:
    method = getattr(target, method_name, None)
    if not callable(method):
        return None
    try:
        return method(*args, **kwargs)
    except Exception:
        return None


def _merge_status(current: str, candidate: str) -> str:
    order = {
        "ready": 0,
        "warning": 1,
        "blocked": 2,
        "critical": 3,
    }
    current_status = str(current or "ready")
    candidate_status = str(candidate or "ready")
    return candidate_status if order.get(candidate_status, 0) > order.get(current_status, 0) else current_status


def _build_runtime_workflow_optimization_guidance(
    *,
    mediator: Any,
    drafting_readiness: Dict[str, Any],
    workflow_phase_plan: Dict[str, Any],
    document_optimization: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
    existing_guidance = (
        optimization_report.get("workflow_optimization_guidance")
        if isinstance(optimization_report.get("workflow_optimization_guidance"), dict)
        else {}
    )
    if existing_guidance:
        guidance = dict(existing_guidance)
        if workflow_phase_plan and "workflow_phase_plan" not in guidance:
            guidance["workflow_phase_plan"] = dict(workflow_phase_plan)
        intake_case_summary = build_intake_case_review_summary(mediator)
        document_drafting_next_action = (
            intake_case_summary.get("document_drafting_next_action")
            if isinstance(intake_case_summary.get("document_drafting_next_action"), dict)
            else {}
        )
        workflow_targeting_summary = (
            optimization_report.get("workflow_targeting_summary")
            if isinstance(optimization_report.get("workflow_targeting_summary"), dict)
            else intake_case_summary.get("workflow_targeting_summary")
        )
        document_workflow_execution_summary = (
            optimization_report.get("document_workflow_execution_summary")
            if isinstance(optimization_report.get("document_workflow_execution_summary"), dict)
            else intake_case_summary.get("document_workflow_execution_summary")
        )
        document_execution_drift_summary = (
            optimization_report.get("document_execution_drift_summary")
            if isinstance(optimization_report.get("document_execution_drift_summary"), dict)
            else intake_case_summary.get("document_execution_drift_summary")
        )
        document_grounding_improvement_summary = (
            optimization_report.get("document_grounding_improvement_summary")
            if isinstance(optimization_report.get("document_grounding_improvement_summary"), dict)
            else intake_case_summary.get("document_grounding_improvement_summary")
        )
        document_grounding_lane_outcome_summary = (
            optimization_report.get("document_grounding_lane_outcome_summary")
            if isinstance(optimization_report.get("document_grounding_lane_outcome_summary"), dict)
            else intake_case_summary.get("document_grounding_lane_outcome_summary")
        )
        evidence_workflow_action_queue = (
            intake_case_summary.get("evidence_workflow_action_queue")
            if isinstance(intake_case_summary.get("evidence_workflow_action_queue"), list)
            else []
        )
        evidence_workflow_action_summary = (
            intake_case_summary.get("evidence_workflow_action_summary")
            if isinstance(intake_case_summary.get("evidence_workflow_action_summary"), dict)
            else {}
        )
        if "evidence_workflow_action_queue" not in guidance:
            guidance["evidence_workflow_action_queue"] = list(evidence_workflow_action_queue)
        if "evidence_workflow_action_summary" not in guidance:
            guidance["evidence_workflow_action_summary"] = dict(evidence_workflow_action_summary)
        if (
            "workflow_targeting_summary" not in guidance
            and isinstance(workflow_targeting_summary, dict)
        ):
            guidance["workflow_targeting_summary"] = dict(workflow_targeting_summary)
        if (
            "document_workflow_execution_summary" not in guidance
            and isinstance(document_workflow_execution_summary, dict)
        ):
            guidance["document_workflow_execution_summary"] = dict(document_workflow_execution_summary)
        if (
            "document_execution_drift_summary" not in guidance
            and isinstance(document_execution_drift_summary, dict)
        ):
            guidance["document_execution_drift_summary"] = dict(document_execution_drift_summary)
        if (
            "document_grounding_improvement_summary" not in guidance
            and isinstance(document_grounding_improvement_summary, dict)
        ):
            guidance["document_grounding_improvement_summary"] = dict(document_grounding_improvement_summary)
        if (
            "document_grounding_lane_outcome_summary" not in guidance
            and isinstance(document_grounding_lane_outcome_summary, dict)
        ):
            guidance["document_grounding_lane_outcome_summary"] = dict(document_grounding_lane_outcome_summary)
        if (
            "document_drafting_next_action" not in guidance
            and isinstance(document_drafting_next_action, dict)
            and document_drafting_next_action
        ):
            guidance["document_drafting_next_action"] = dict(document_drafting_next_action)
        existing_queue = (
            guidance.get("workflow_action_queue")
            if isinstance(guidance.get("workflow_action_queue"), list)
            else []
        )
        drift_flag = bool((document_execution_drift_summary or {}).get("drift_flag"))
        if drift_flag:
            top_targeted_claim_element = str(
                (document_execution_drift_summary or {}).get("top_targeted_claim_element") or ""
            ).strip()
            first_executed_claim_element = str(
                (document_execution_drift_summary or {}).get("first_executed_claim_element") or ""
            ).strip()
            drift_action = {
                "rank": 1,
                "phase_name": "document_generation",
                "status": "warning",
                "action": "Realign drafting to the top targeted claim element before further revisions.",
                "focus_areas": [
                    item
                    for item in [top_targeted_claim_element, first_executed_claim_element]
                    if item
                ][:3],
                "top_targeted_claim_element": top_targeted_claim_element,
                "first_executed_claim_element": first_executed_claim_element,
            }
            if not any(
                isinstance(item, dict)
                and str(item.get("action") or "").strip().lower()
                == drift_action["action"].strip().lower()
                for item in existing_queue
            ):
                guidance["workflow_action_queue"] = [drift_action, *existing_queue]
                existing_queue = guidance["workflow_action_queue"]
        if evidence_workflow_action_queue and not existing_queue:
            guidance["workflow_action_queue"] = [
                {
                    "rank": int(action.get("rank") or index + 1),
                    "phase_name": str(action.get("phase_name") or "graph_analysis").strip() or "graph_analysis",
                    "status": str(action.get("status") or "warning").strip().lower() or "warning",
                    "action": str(action.get("action") or "").strip(),
                    "focus_areas": list(action.get("focus_areas") or [])[:3],
                }
                for index, action in enumerate(evidence_workflow_action_queue)
                if isinstance(action, dict) and str(action.get("action") or "").strip()
            ]
        return guidance

    intake_case_summary = build_intake_case_review_summary(mediator)
    optimization_workflow_targeting_summary = (
        optimization_report.get("workflow_targeting_summary")
        if isinstance(optimization_report.get("workflow_targeting_summary"), dict)
        else {}
    )
    intake_sections = (
        intake_case_summary.get("intake_sections")
        if isinstance(intake_case_summary.get("intake_sections"), dict)
        else {}
    )
    question_summary = (
        intake_case_summary.get("question_candidate_summary")
        if isinstance(intake_case_summary.get("question_candidate_summary"), dict)
        else {}
    )
    claim_packet_summary = (
        intake_case_summary.get("claim_support_packet_summary")
        if isinstance(intake_case_summary.get("claim_support_packet_summary"), dict)
        else {}
    )
    evidence_workflow_action_queue = (
        intake_case_summary.get("evidence_workflow_action_queue")
        if isinstance(intake_case_summary.get("evidence_workflow_action_queue"), list)
        else []
    )
    evidence_workflow_action_summary = (
        intake_case_summary.get("evidence_workflow_action_summary")
        if isinstance(intake_case_summary.get("evidence_workflow_action_summary"), dict)
        else {}
    )
    workflow_targeting_summary = (
        optimization_workflow_targeting_summary
        if optimization_workflow_targeting_summary
        else (
            intake_case_summary.get("workflow_targeting_summary")
            if isinstance(intake_case_summary.get("workflow_targeting_summary"), dict)
            else {}
        )
    )
    document_workflow_execution_summary = (
        optimization_report.get("document_workflow_execution_summary")
        if isinstance(optimization_report.get("document_workflow_execution_summary"), dict)
        else (
            intake_case_summary.get("document_workflow_execution_summary")
            if isinstance(intake_case_summary.get("document_workflow_execution_summary"), dict)
            else {}
        )
    )
    document_execution_drift_summary = (
        optimization_report.get("document_execution_drift_summary")
        if isinstance(optimization_report.get("document_execution_drift_summary"), dict)
        else (
            intake_case_summary.get("document_execution_drift_summary")
            if isinstance(intake_case_summary.get("document_execution_drift_summary"), dict)
            else {}
        )
    )
    document_grounding_improvement_summary = (
        optimization_report.get("document_grounding_improvement_summary")
        if isinstance(optimization_report.get("document_grounding_improvement_summary"), dict)
        else (
            intake_case_summary.get("document_grounding_improvement_summary")
            if isinstance(intake_case_summary.get("document_grounding_improvement_summary"), dict)
            else {}
        )
    )
    document_grounding_lane_outcome_summary = (
        optimization_report.get("document_grounding_lane_outcome_summary")
        if isinstance(optimization_report.get("document_grounding_lane_outcome_summary"), dict)
        else (
            intake_case_summary.get("document_grounding_lane_outcome_summary")
            if isinstance(intake_case_summary.get("document_grounding_lane_outcome_summary"), dict)
            else {}
        )
    )
    document_drafting_next_action = (
        intake_case_summary.get("document_drafting_next_action")
        if isinstance(intake_case_summary.get("document_drafting_next_action"), dict)
        else {}
    )
    candidate_claims = _coerce_list(intake_case_summary.get("candidate_claims"))
    claim_types = _unique_preserving_order(
        str((claim or {}).get("claim_type") or "").strip()
        for claim in candidate_claims
        if isinstance(claim, dict)
    )
    intake_focus_areas = _unique_preserving_order(
        str(section_name)
        for section_name, payload in intake_sections.items()
        if isinstance(payload, dict) and str(payload.get("status") or "").strip().lower() != "complete"
    )
    graph_focus_areas = _unique_preserving_order(
        [
            *claim_types,
            *[
                "claim_support_packets"
                if int(claim_packet_summary.get("unsupported_element_count") or 0) > 0
                else ""
            ],
            *[
                str((action or {}).get("claim_element_label") or (action or {}).get("claim_element_id") or "").strip()
                for action in evidence_workflow_action_queue
                if isinstance(action, dict)
            ],
        ]
    )
    document_focus_areas = _unique_preserving_order(
        [
            str(section_name)
            for section_name, payload in dict(drafting_readiness.get("sections") or {}).items()
            if isinstance(payload, dict) and str(payload.get("status") or "").strip().lower() != "ready"
        ]
    )
    cross_phase_findings = []
    if intake_focus_areas and graph_focus_areas:
        cross_phase_findings.append(
            "Intake follow-up gaps remain linked to graph support gaps, so unresolved intake sections should be closed before final drafting."
        )
    if graph_focus_areas and document_focus_areas:
        cross_phase_findings.append(
            "Graph support issues are still affecting drafting readiness, especially in claims-for-relief and chronology-dependent allegations."
        )
    workflow_action_queue: List[Dict[str, Any]] = []
    if bool(document_execution_drift_summary.get("drift_flag")):
        top_targeted_claim_element = str(
            document_execution_drift_summary.get("top_targeted_claim_element") or ""
        ).strip()
        first_executed_claim_element = str(
            document_execution_drift_summary.get("first_executed_claim_element") or ""
        ).strip()
        workflow_action_queue.append(
            {
                "rank": 1,
                "phase_name": "document_generation",
                "status": "warning",
                "action": "Realign drafting to the top targeted claim element before further revisions.",
                "focus_areas": [
                    item
                    for item in [top_targeted_claim_element, first_executed_claim_element]
                    if item
                ][:3],
                "top_targeted_claim_element": top_targeted_claim_element,
                "first_executed_claim_element": first_executed_claim_element,
            }
        )
    workflow_action_queue.extend(
        [
            *[
                {
                    "rank": int(action.get("rank") or index + 1),
                    "phase_name": str(action.get("phase_name") or "graph_analysis").strip() or "graph_analysis",
                    "status": str(action.get("status") or "warning").strip().lower() or "warning",
                    "action": str(action.get("action") or "").strip(),
                    "focus_areas": list(action.get("focus_areas") or [])[:3],
                    "claim_type": str(action.get("claim_type") or "").strip(),
                    "claim_element_id": str(action.get("claim_element_id") or "").strip(),
                }
                for index, action in enumerate(evidence_workflow_action_queue[:2])
                if isinstance(action, dict) and str(action.get("action") or "").strip()
            ],
            *[
                item
                for item in [
                    {
                        "rank": 100,
                        "phase_name": "graph_analysis",
                        "status": "warning" if graph_focus_areas else "ready",
                        "action": "Close graph and claim-support gaps before final drafting.",
                        "focus_areas": graph_focus_areas[:3],
                    },
                    {
                        "rank": 101,
                        "phase_name": "intake_questioning",
                        "status": "warning" if intake_focus_areas else "ready",
                        "action": "Target remaining intake gaps that block graph-supported drafting.",
                        "focus_areas": intake_focus_areas[:3],
                    },
                    {
                        "rank": 102,
                        "phase_name": "document_generation",
                        "status": str(drafting_readiness.get("status") or "ready").strip().lower() or "ready",
                        "action": "Revise complaint sections still flagged by drafting readiness.",
                        "focus_areas": document_focus_areas[:3],
                    },
                ]
                if not (
                    item["phase_name"] == "graph_analysis"
                    and evidence_workflow_action_queue
                )
            ],
        ]
    )
    return {
        "workflow_phase_plan": dict(workflow_phase_plan or {}),
        "phase_scorecards": {
            "intake_questioning": {
                "status": "warning" if intake_focus_areas else "ready",
                "focus_areas": intake_focus_areas,
                "question_candidate_count": int(question_summary.get("count") or 0),
            },
            "graph_analysis": {
                "status": "warning" if graph_focus_areas else "ready",
                "focus_areas": graph_focus_areas,
                "unsupported_element_count": int(claim_packet_summary.get("unsupported_element_count") or 0),
                "evidence_workflow_action_count": int(evidence_workflow_action_summary.get("count") or 0),
            },
            "document_generation": {
                "status": str(drafting_readiness.get("status") or "ready").strip().lower() or "ready",
                "focus_areas": document_focus_areas,
                "warning_count": int(drafting_readiness.get("warning_count") or 0),
            },
        },
        "cross_phase_findings": cross_phase_findings,
        "workflow_action_queue": workflow_action_queue,
        "evidence_workflow_action_queue": list(evidence_workflow_action_queue),
        "evidence_workflow_action_summary": dict(evidence_workflow_action_summary),
        "workflow_targeting_summary": dict(workflow_targeting_summary),
        "document_workflow_execution_summary": dict(document_workflow_execution_summary),
        "document_execution_drift_summary": dict(document_execution_drift_summary),
        "document_grounding_improvement_summary": dict(document_grounding_improvement_summary),
        "document_grounding_lane_outcome_summary": dict(document_grounding_lane_outcome_summary),
        "document_drafting_next_action": dict(document_drafting_next_action),
        "complaint_type_generalization_summary": {
            "complaint_types": claim_types,
            "complaint_type_count": len(claim_types),
        },
        "document_handoff_summary": {
            "ready_for_document_optimization": str(drafting_readiness.get("status") or "").strip().lower() == "ready",
            "drafting_status": str(drafting_readiness.get("status") or "ready").strip().lower() or "ready",
            "blocking_warning_count": int(drafting_readiness.get("warning_count") or 0),
        },
    }


def _build_workflow_optimization_warning_entries(
    workflow_optimization_guidance: Dict[str, Any],
) -> List[Dict[str, Any]]:
    guidance = workflow_optimization_guidance if isinstance(workflow_optimization_guidance, dict) else {}
    phase_scorecards = guidance.get("phase_scorecards") if isinstance(guidance.get("phase_scorecards"), dict) else {}
    document_execution_drift_summary = (
        guidance.get("document_execution_drift_summary")
        if isinstance(guidance.get("document_execution_drift_summary"), dict)
        else {}
    )
    warnings: List[Dict[str, Any]] = []
    phase_labels = {
        "intake_questioning": "intake questioning",
        "graph_analysis": "graph analysis",
        "document_generation": "document generation",
    }
    for phase_name, label in phase_labels.items():
        scorecard = phase_scorecards.get(phase_name) if isinstance(phase_scorecards.get(phase_name), dict) else {}
        status = str(scorecard.get("status") or "ready").strip().lower()
        if status == "ready":
            continue
        focus_areas = _dedupe_text_values(scorecard.get("focus_areas") or [])
        focus_suffix = f" Focus areas: {', '.join(focus_areas[:3])}." if focus_areas else ""
        warnings.append(
            {
                "code": f"workflow_{phase_name}_optimization_warning",
                "severity": "warning" if status == "warning" else status,
                "message": f"Workflow optimization still flags {label} as {status}.{focus_suffix}",
                "phase": phase_name,
                "focus_areas": focus_areas,
            }
        )

    for finding in list(guidance.get("cross_phase_findings") or [])[:2]:
        text = str(finding or "").strip()
        if not text:
            continue
        warnings.append(
            {
                "code": "workflow_cross_phase_optimization_warning",
                "severity": "warning",
                "message": text,
            }
        )
    if bool(document_execution_drift_summary.get("drift_flag")):
        top_targeted_claim_element = str(
            document_execution_drift_summary.get("top_targeted_claim_element") or ""
        ).strip()
        first_executed_claim_element = str(
            document_execution_drift_summary.get("first_executed_claim_element") or ""
        ).strip()
        focus_areas = _dedupe_text_values(
            [top_targeted_claim_element, first_executed_claim_element]
        )
        warnings.append(
            {
                "code": "workflow_document_execution_drift_warning",
                "severity": "warning",
                "message": (
                    "Document optimization is revising the wrong claim element first; "
                    "realign drafting to the top targeted legal element before further revisions."
                ),
                "phase": "document_generation",
                "focus_areas": focus_areas,
                "top_targeted_claim_element": top_targeted_claim_element,
                "first_executed_claim_element": first_executed_claim_element,
            }
        )
    return warnings


class FormalComplaintDocumentBuilder:
    def __init__(self, mediator: Any):
        self.mediator = mediator

    def _load_email_timeline_handoff(self, value: Any) -> Dict[str, Any]:
        if isinstance(value, dict):
            return deepcopy(value)
        path_text = str(value or "").strip()
        if not path_text:
            return {}
        payload = json.loads(Path(path_text).read_text(encoding="utf-8"))
        return payload if isinstance(payload, dict) else {}

    def _load_email_authority_enrichment(self, value: Any) -> Dict[str, Any]:
        if isinstance(value, dict):
            return deepcopy(value)
        path_text = str(value or "").strip()
        if not path_text:
            return {}
        payload = json.loads(Path(path_text).read_text(encoding="utf-8"))
        return payload if isinstance(payload, dict) else {}

    def build_package(
        self,
        *,
        user_id: Optional[str] = None,
        court_name: str = "United States District Court",
        district: str = "",
        county: Optional[str] = None,
        division: Optional[str] = None,
        court_header_override: Optional[str] = None,
        case_number: Optional[str] = None,
        lead_case_number: Optional[str] = None,
        related_case_number: Optional[str] = None,
        assigned_judge: Optional[str] = None,
        courtroom: Optional[str] = None,
        title_override: Optional[str] = None,
        plaintiff_names: Optional[List[str]] = None,
        defendant_names: Optional[List[str]] = None,
        requested_relief: Optional[List[str]] = None,
        jury_demand: Optional[bool] = None,
        jury_demand_text: Optional[str] = None,
        signer_name: Optional[str] = None,
        signer_title: Optional[str] = None,
        signer_firm: Optional[str] = None,
        signer_bar_number: Optional[str] = None,
        signer_contact: Optional[str] = None,
        additional_signers: Optional[List[Dict[str, str]]] = None,
        declarant_name: Optional[str] = None,
        service_method: Optional[str] = None,
        service_recipients: Optional[List[str]] = None,
        service_recipient_details: Optional[List[Dict[str, str]]] = None,
        signature_date: Optional[str] = None,
        verification_date: Optional[str] = None,
        service_date: Optional[str] = None,
        affidavit_title: Optional[str] = None,
        affidavit_intro: Optional[str] = None,
        affidavit_facts: Optional[List[str]] = None,
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]] = None,
        affidavit_include_complaint_exhibits: Optional[bool] = None,
        affidavit_venue_lines: Optional[List[str]] = None,
        affidavit_jurat: Optional[str] = None,
        affidavit_notary_block: Optional[List[str]] = None,
        enable_agentic_optimization: bool = False,
        optimization_max_iterations: int = 2,
        optimization_target_score: float = 0.9,
        optimization_provider: Optional[str] = None,
        optimization_model_name: Optional[str] = None,
        optimization_llm_config: Optional[Dict[str, Any]] = None,
        optimization_persist_artifacts: bool = False,
        email_timeline_handoff_path: Optional[str] = None,
        email_authority_enrichment_path: Optional[str] = None,
        output_dir: Optional[str] = None,
        output_formats: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        resolved_user_id = self._resolve_user_id(user_id)
        formats = self._normalize_formats(output_formats)
        intake_status = build_intake_status_summary(self.mediator)
        raw_status = {}
        get_three_phase_status = getattr(self.mediator, "get_three_phase_status", None)
        if callable(get_three_phase_status):
            candidate_raw_status = get_three_phase_status()
            if isinstance(candidate_raw_status, dict):
                raw_status = candidate_raw_status
        document_drafting_next_action = (
            intake_status.get("document_drafting_next_action")
            if isinstance(intake_status.get("document_drafting_next_action"), dict)
            else {}
        )
        document_grounding_improvement_next_action = (
            intake_status.get("document_grounding_improvement_next_action")
            if isinstance(intake_status.get("document_grounding_improvement_next_action"), dict)
            else {}
        )
        if not document_drafting_next_action and isinstance(raw_status.get("document_drafting_next_action"), dict):
            document_drafting_next_action = dict(raw_status.get("document_drafting_next_action") or {})
        if not document_grounding_improvement_next_action and isinstance(raw_status.get("document_grounding_improvement_next_action"), dict):
            document_grounding_improvement_next_action = dict(raw_status.get("document_grounding_improvement_next_action") or {})
        draft = self.build_draft(
            user_id=resolved_user_id,
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            court_header_override=court_header_override,
            case_number=case_number,
            lead_case_number=lead_case_number,
            related_case_number=related_case_number,
            assigned_judge=assigned_judge,
            courtroom=courtroom,
            title_override=title_override,
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
            requested_relief=requested_relief,
            jury_demand=jury_demand,
            jury_demand_text=jury_demand_text,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            declarant_name=declarant_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            signature_date=signature_date,
            verification_date=verification_date,
            service_date=service_date,
            affidavit_title=affidavit_title,
            affidavit_intro=affidavit_intro,
            affidavit_facts=affidavit_facts,
            affidavit_supporting_exhibits=affidavit_supporting_exhibits,
            affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
            affidavit_venue_lines=affidavit_venue_lines,
            affidavit_jurat=affidavit_jurat,
            affidavit_notary_block=affidavit_notary_block,
        )
        loaded_email_timeline_handoff = self._load_email_timeline_handoff(email_timeline_handoff_path)
        if loaded_email_timeline_handoff:
            source_context = dict(draft.get("source_context") or {})
            source_context["email_timeline_handoff"] = loaded_email_timeline_handoff
            draft["source_context"] = source_context
        loaded_email_authority_enrichment = self._load_email_authority_enrichment(email_authority_enrichment_path)
        if loaded_email_authority_enrichment:
            source_context = dict(draft.get("source_context") or {})
            source_context["email_authority_enrichment"] = loaded_email_authority_enrichment
            draft["source_context"] = source_context
        draft = self._apply_document_drafting_focus(
            draft,
            document_drafting_next_action=document_drafting_next_action,
            document_grounding_improvement_next_action=document_grounding_improvement_next_action,
        )
        self._attach_allegation_references(draft)
        document_optimization = None
        if enable_agentic_optimization:
            draft, document_optimization = self._optimize_draft(
                draft,
                user_id=resolved_user_id,
                max_iterations=optimization_max_iterations,
                target_score=optimization_target_score,
                provider=optimization_provider,
                model_name=optimization_model_name,
                llm_config=optimization_llm_config,
                persist_artifacts=optimization_persist_artifacts,
            )
        drafting_readiness = self._build_drafting_readiness(
            user_id=resolved_user_id,
            draft=draft,
        )
        intake_summary_handoff = self._build_intake_summary_handoff(document_optimization)
        if intake_summary_handoff:
            drafting_readiness["intake_summary_handoff"] = dict(intake_summary_handoff)
            handoff_payload = (
                dict(drafting_readiness.get("drafting_handoff") or {})
                if isinstance(drafting_readiness.get("drafting_handoff"), dict)
                else {}
            )
            handoff_payload["intake_summary_handoff_available"] = True
            handoff_payload["intake_summary_handoff_keys"] = sorted(
                str(key)
                for key in intake_summary_handoff.keys()
                if str(key).strip()
            )
            confirmation = (
                intake_summary_handoff.get("complainant_summary_confirmation")
                if isinstance(intake_summary_handoff.get("complainant_summary_confirmation"), dict)
                else {}
            )
            confirmation_snapshot = (
                confirmation.get("confirmed_summary_snapshot")
                if isinstance(confirmation.get("confirmed_summary_snapshot"), dict)
                else {}
            )
            priority_summary = (
                confirmation_snapshot.get("adversarial_intake_priority_summary")
                if isinstance(confirmation_snapshot.get("adversarial_intake_priority_summary"), dict)
                else {}
            )
            handoff_payload["uncovered_intake_objectives"] = _dedupe_text_values(
                priority_summary.get("uncovered_objectives") or []
            )
            drafting_readiness["drafting_handoff"] = handoff_payload
        workflow_phase_plan = self._build_runtime_workflow_phase_plan(
            drafting_readiness=drafting_readiness,
            document_optimization=document_optimization,
        )
        if not workflow_phase_plan:
            workflow_phase_plan = (
                dict(drafting_readiness.get("workflow_phase_plan") or {})
                if isinstance(drafting_readiness.get("workflow_phase_plan"), dict)
                else {}
            )
        if workflow_phase_plan:
            drafting_readiness["workflow_phase_plan"] = workflow_phase_plan
            self._refresh_drafting_readiness_workflow_warnings(
                drafting_readiness=drafting_readiness,
                workflow_phase_plan=workflow_phase_plan,
            )
            phase_entries = (
                dict(workflow_phase_plan.get("phases") or {})
                if isinstance(workflow_phase_plan.get("phases"), dict)
                else {}
            )
            document_phase = (
                dict(phase_entries.get("document_generation") or {})
                if isinstance(phase_entries.get("document_generation"), dict)
                else {}
            )
            document_phase_signals = (
                dict(document_phase.get("signals") or {})
                if isinstance(document_phase.get("signals"), dict)
                else {}
            )
            if document_phase_signals:
                derived_blockers: List[str] = []
                if _coerce_bool(document_phase_signals.get("gate_on_graph_completeness"), default=False):
                    derived_blockers.append("graph_analysis_not_ready")
                if int(document_phase_signals.get("unresolved_factual_gap_count", 0) or 0) > 0:
                    derived_blockers.append("unresolved_factual_gaps_not_closed")
                if int(document_phase_signals.get("unresolved_legal_gap_count", 0) or 0) > 0:
                    derived_blockers.append("unresolved_legal_gaps_not_closed")
                if int(document_phase_signals.get("uncovered_intake_objective_count", 0) or 0) > 0 or int(
                    document_phase_signals.get("missing_required_intake_objective_count", 0) or 0
                ) > 0:
                    derived_blockers.append("uncovered_intake_objectives")
                if int(document_phase_signals.get("structured_intake_handoff_gap_count", 0) or 0) > 0:
                    derived_blockers.append("structured_intake_handoff_incomplete")
                if list(document_phase_signals.get("targeted_weak_complaint_types") or []):
                    derived_blockers.append("weak_complaint_type_generalization_needed")
                if list(document_phase_signals.get("targeted_weak_evidence_modalities") or []):
                    derived_blockers.append("weak_evidence_modality_support_needed")
                if not _coerce_bool(document_phase_signals.get("ready_for_formalization"), default=False):
                    derived_blockers.append("document_generation_not_ready")

                existing_blockers = [
                    str(item).strip()
                    for item in list(drafting_readiness.get("blockers") or [])
                    if str(item).strip()
                ]
                drafting_readiness["blockers"] = _dedupe_text_values(existing_blockers + derived_blockers)
                if _coerce_bool(document_phase_signals.get("gate_on_graph_completeness"), default=False):
                    drafting_readiness["phase_status"] = _merge_status(
                        str(drafting_readiness.get("phase_status") or "ready"),
                        "blocked",
                    )
                elif not _coerce_bool(document_phase_signals.get("ready_for_formalization"), default=False):
                    drafting_readiness["phase_status"] = _merge_status(
                        str(drafting_readiness.get("phase_status") or "ready"),
                        "warning",
                    )
                if list(document_phase_signals.get("unresolved_factual_gaps") or []) and not list(
                    drafting_readiness.get("unresolved_factual_gaps") or []
                ):
                    drafting_readiness["unresolved_factual_gaps"] = list(document_phase_signals.get("unresolved_factual_gaps") or [])
                if list(document_phase_signals.get("unresolved_legal_gaps") or []) and not list(
                    drafting_readiness.get("unresolved_legal_gaps") or []
                ):
                    drafting_readiness["unresolved_legal_gaps"] = list(document_phase_signals.get("unresolved_legal_gaps") or [])
                if list(document_phase_signals.get("uncovered_intake_objectives") or []) and not list(
                    drafting_readiness.get("uncovered_intake_objectives") or []
                ):
                    drafting_readiness["uncovered_intake_objectives"] = list(document_phase_signals.get("uncovered_intake_objectives") or [])
                if list(document_phase_signals.get("weak_complaint_types") or []) and not list(
                    drafting_readiness.get("weak_complaint_types") or []
                ):
                    drafting_readiness["weak_complaint_types"] = list(document_phase_signals.get("weak_complaint_types") or [])
                if list(document_phase_signals.get("weak_evidence_modalities") or []) and not list(
                    drafting_readiness.get("weak_evidence_modalities") or []
                ):
                    drafting_readiness["weak_evidence_modalities"] = list(document_phase_signals.get("weak_evidence_modalities") or [])
                drafting_handoff_payload = (
                    dict(drafting_readiness.get("drafting_handoff") or {})
                    if isinstance(drafting_readiness.get("drafting_handoff"), dict)
                    else {}
                )
                drafting_handoff_payload["ready_for_formalization"] = _coerce_bool(
                    document_phase_signals.get("ready_for_formalization"),
                    default=False,
                )
                drafting_handoff_payload["gate_on_graph_completeness"] = _coerce_bool(
                    document_phase_signals.get("gate_on_graph_completeness"),
                    default=_coerce_bool(drafting_handoff_payload.get("gate_on_graph_completeness"), default=False),
                )
                drafting_handoff_payload["graph_phase_status"] = str(
                    document_phase_signals.get("graph_phase_status")
                    or drafting_handoff_payload.get("graph_phase_status")
                    or "ready"
                ).strip().lower() or "ready"
                drafting_handoff_payload["graph_remaining_gap_count"] = int(
                    document_phase_signals.get(
                        "graph_remaining_gap_count",
                        drafting_handoff_payload.get("graph_remaining_gap_count", 0),
                    )
                    or 0
                )
                drafting_handoff_payload["unresolved_factual_gaps"] = list(
                    document_phase_signals.get("unresolved_factual_gaps")
                    or drafting_handoff_payload.get("unresolved_factual_gaps")
                    or []
                )[:6]
                drafting_handoff_payload["unresolved_legal_gaps"] = list(
                    document_phase_signals.get("unresolved_legal_gaps")
                    or drafting_handoff_payload.get("unresolved_legal_gaps")
                    or []
                )[:6]
                drafting_handoff_payload["uncovered_intake_objectives"] = list(
                    document_phase_signals.get("uncovered_intake_objectives")
                    or drafting_handoff_payload.get("uncovered_intake_objectives")
                    or []
                )[:8]
                drafting_handoff_payload["targeted_weak_complaint_types"] = list(
                    document_phase_signals.get("targeted_weak_complaint_types")
                    or drafting_handoff_payload.get("targeted_weak_complaint_types")
                    or []
                )[:4]
                drafting_handoff_payload["targeted_weak_evidence_modalities"] = list(
                    document_phase_signals.get("targeted_weak_evidence_modalities")
                    or drafting_handoff_payload.get("targeted_weak_evidence_modalities")
                    or []
                )[:4]
                drafting_handoff_payload["blockers"] = list(drafting_readiness.get("blockers") or [])
                drafting_readiness["drafting_handoff"] = drafting_handoff_payload
        workflow_optimization_guidance = _build_runtime_workflow_optimization_guidance(
            mediator=self.mediator,
            drafting_readiness=drafting_readiness,
            workflow_phase_plan=workflow_phase_plan,
            document_optimization=document_optimization,
        )
        workflow_targeting_summary = (
            workflow_optimization_guidance.get("workflow_targeting_summary")
            if isinstance(workflow_optimization_guidance.get("workflow_targeting_summary"), dict)
            else (
                document_optimization.get("workflow_targeting_summary")
                if isinstance(document_optimization, dict)
                and isinstance(document_optimization.get("workflow_targeting_summary"), dict)
                else {}
            )
        )
        document_workflow_execution_summary = (
            workflow_optimization_guidance.get("document_workflow_execution_summary")
            if isinstance(workflow_optimization_guidance.get("document_workflow_execution_summary"), dict)
            else (
                document_optimization.get("document_workflow_execution_summary")
                if isinstance(document_optimization, dict)
                and isinstance(document_optimization.get("document_workflow_execution_summary"), dict)
                else {}
            )
        )
        document_execution_drift_summary = (
            workflow_optimization_guidance.get("document_execution_drift_summary")
            if isinstance(workflow_optimization_guidance.get("document_execution_drift_summary"), dict)
            else (
                document_optimization.get("document_execution_drift_summary")
                if isinstance(document_optimization, dict)
                and isinstance(document_optimization.get("document_execution_drift_summary"), dict)
                else {}
            )
        )
        document_grounding_improvement_summary = (
            workflow_optimization_guidance.get("document_grounding_improvement_summary")
            if isinstance(workflow_optimization_guidance.get("document_grounding_improvement_summary"), dict)
            else (
                document_optimization.get("document_grounding_improvement_summary")
                if isinstance(document_optimization, dict)
                and isinstance(document_optimization.get("document_grounding_improvement_summary"), dict)
                else {}
            )
        )
        document_grounding_lane_outcome_summary = (
            workflow_optimization_guidance.get("document_grounding_lane_outcome_summary")
            if isinstance(workflow_optimization_guidance.get("document_grounding_lane_outcome_summary"), dict)
            else (
                document_optimization.get("document_grounding_lane_outcome_summary")
                if isinstance(document_optimization, dict)
                and isinstance(document_optimization.get("document_grounding_lane_outcome_summary"), dict)
                else {}
            )
        )
        filing_checklist = self._build_filing_checklist(drafting_readiness)
        self._annotate_filing_checklist_review_links(
            filing_checklist=filing_checklist,
            drafting_readiness=drafting_readiness,
            user_id=resolved_user_id,
        )
        draft["drafting_readiness"] = drafting_readiness
        if workflow_phase_plan:
            draft["workflow_phase_plan"] = workflow_phase_plan
        if workflow_optimization_guidance:
            draft["workflow_optimization_guidance"] = workflow_optimization_guidance
        if workflow_targeting_summary:
            draft["workflow_targeting_summary"] = dict(workflow_targeting_summary)
        if document_workflow_execution_summary:
            draft["document_workflow_execution_summary"] = dict(document_workflow_execution_summary)
        if document_execution_drift_summary:
            draft["document_execution_drift_summary"] = dict(document_execution_drift_summary)
        if document_grounding_improvement_summary:
            draft["document_grounding_improvement_summary"] = dict(document_grounding_improvement_summary)
        if document_grounding_lane_outcome_summary:
            draft["document_grounding_lane_outcome_summary"] = dict(document_grounding_lane_outcome_summary)
        if document_drafting_next_action:
            draft["document_drafting_next_action"] = dict(document_drafting_next_action)
        if document_grounding_improvement_next_action:
            draft["document_grounding_improvement_next_action"] = dict(document_grounding_improvement_next_action)
        if isinstance(draft.get("document_provenance_summary"), dict) and draft.get("document_provenance_summary"):
            draft["document_provenance_summary"] = dict(draft.get("document_provenance_summary") or {})
        draft["filing_checklist"] = filing_checklist
        draft["affidavit"] = self._build_affidavit(draft)
        intake_case_summary = (
            document_optimization.get("intake_case_summary")
            if isinstance(document_optimization, dict)
            and isinstance(document_optimization.get("intake_case_summary"), dict)
            else build_intake_case_review_summary(self.mediator)
        )
        claim_support_temporal_handoff = self._build_claim_support_temporal_handoff(document_optimization)
        email_timeline_handoff = self._resolve_email_timeline_handoff(
            draft=draft,
            document_optimization=document_optimization,
        )
        email_authority_enrichment = self._resolve_email_authority_enrichment(
            draft=draft,
            document_optimization=document_optimization,
        )
        if email_timeline_handoff:
            claim_support_temporal_handoff = self._merge_claim_support_temporal_handoff(
                claim_support_temporal_handoff,
                email_timeline_handoff,
            )
            draft = self._enrich_draft_with_email_timeline_handoff(
                draft,
                email_timeline_handoff=email_timeline_handoff,
            )
        if email_authority_enrichment:
            draft = self._enrich_draft_with_email_authority_enrichment(
                draft,
                email_authority_enrichment=email_authority_enrichment,
            )
        if email_timeline_handoff or email_authority_enrichment:
            draft = self._specialize_generic_claims_from_email_support(
                draft,
                email_timeline_handoff=email_timeline_handoff or {},
                email_authority_enrichment=email_authority_enrichment or {},
            )
            claims_for_relief = (
                list(draft.get("claims_for_relief") or [])
                if isinstance(draft.get("claims_for_relief"), list)
                else []
            )
            if claims_for_relief:
                summary_fact_entries = [
                    dict(entry)
                    for entry in _coerce_list(draft.get("summary_of_fact_entries"))
                    if isinstance(entry, dict)
                ]
                non_generic_entries = [
                    entry
                    for entry in summary_fact_entries
                    if not self._is_generic_claim_support_text(entry.get("text"))
                ]
                if non_generic_entries:
                    summary_fact_entries = non_generic_entries
                    draft["summary_of_fact_entries"] = summary_fact_entries
                    draft["summary_of_facts"] = [
                        str(entry.get("text") or "").strip()
                        for entry in summary_fact_entries
                        if str(entry.get("text") or "").strip()
                    ]
                summary_of_facts = [
                    line
                    for line in self._normalize_text_lines(draft.get("summary_of_facts", []))
                    if not self._is_generic_claim_support_text(line)
                ]
                if summary_of_facts:
                    draft["summary_of_facts"] = summary_of_facts
                draft["factual_allegation_entries"] = self._build_factual_allegation_entries(
                    summary_fact_entries=summary_fact_entries,
                    claims_for_relief=claims_for_relief,
                )
                draft["factual_allegations"] = self._build_factual_allegations(
                    summary_of_facts=summary_of_facts,
                    claims_for_relief=claims_for_relief,
                )
                source_context = draft.get("source_context") if isinstance(draft.get("source_context"), dict) else {}
                draft["source_context"] = {
                    **source_context,
                    "claim_types": [
                        str(claim.get("claim_type") or "").strip()
                        for claim in claims_for_relief
                        if isinstance(claim, dict) and str(claim.get("claim_type") or "").strip()
                    ],
                }
                self._attach_allegation_references(draft)
                inferred_jurisdiction = str(source_context.get("jurisdiction") or "").strip().lower()
                if not inferred_jurisdiction:
                    inferred_jurisdiction = self._infer_forum_type(
                        classification={},
                        court_name=str(draft.get("court_header") or ""),
                    )
                draft["nature_of_action"] = self._build_nature_of_action(
                    claim_types=[
                        str(claim.get("count_title") or claim.get("claim_type") or "").strip()
                        for claim in claims_for_relief
                        if isinstance(claim, dict) and str(claim.get("count_title") or claim.get("claim_type") or "").strip()
                    ],
                    classification={
                        "jurisdiction": inferred_jurisdiction,
                    },
                    statutes=[],
                    court_name=str(draft.get("court_header") or ""),
                )
        claim_reasoning_review = self._build_claim_reasoning_review(document_optimization)
        chronology_blocker_summary = self._build_chronology_blocker_summary(
            intake_case_summary=intake_case_summary,
            claim_support_temporal_handoff=claim_support_temporal_handoff,
        )
        formalization_gate = self._build_formalization_gate_payload(drafting_readiness)
        source_context = draft.get("source_context") if isinstance(draft.get("source_context"), dict) else {}
        enriched_source_context = dict(source_context)
        if claim_support_temporal_handoff:
            enriched_source_context["claim_support_temporal_handoff"] = claim_support_temporal_handoff
        if email_timeline_handoff:
            enriched_source_context["email_timeline_handoff"] = email_timeline_handoff
        if email_authority_enrichment:
            enriched_source_context["email_authority_enrichment"] = email_authority_enrichment
        if claim_reasoning_review:
            enriched_source_context["claim_reasoning_review"] = claim_reasoning_review
        if chronology_blocker_summary:
            enriched_source_context["chronology_blocker_summary"] = chronology_blocker_summary
        if formalization_gate:
            enriched_source_context["formalization_gate"] = formalization_gate
        if enriched_source_context:
            draft["source_context"] = enriched_source_context
        if formalization_gate:
            draft["formalization_gate"] = formalization_gate
        drafting_handoff = (
            dict(drafting_readiness.get("drafting_handoff") or {})
            if isinstance(drafting_readiness.get("drafting_handoff"), dict)
            else {}
        )
        if drafting_handoff:
            draft["drafting_handoff"] = drafting_handoff
        draft["affidavit"] = self._build_affidavit(draft)
        draft["draft_text"] = self._render_draft_text(draft)
        artifacts = self.render_artifacts(
            draft,
            output_dir=output_dir,
            output_formats=formats,
        )
        package_payload = {
            "draft": draft,
            "drafting_readiness": drafting_readiness,
            "filing_checklist": filing_checklist,
            "artifacts": artifacts,
            "document_optimization": document_optimization,
            "workflow_optimization_guidance": workflow_optimization_guidance,
            "workflow_targeting_summary": dict(workflow_targeting_summary),
            "document_workflow_execution_summary": dict(document_workflow_execution_summary),
            "document_execution_drift_summary": dict(document_execution_drift_summary),
            "document_grounding_improvement_summary": dict(document_grounding_improvement_summary),
            "document_grounding_lane_outcome_summary": dict(document_grounding_lane_outcome_summary),
            "document_drafting_next_action": dict(document_drafting_next_action),
            "document_provenance_summary": (
                dict(draft.get("document_provenance_summary") or {})
                if isinstance(draft.get("document_provenance_summary"), dict)
                else {}
            ),
            "intake_summary_handoff": intake_summary_handoff,
            "output_formats": formats,
            "generated_at": _utcnow().isoformat(),
        }
        if formalization_gate:
            package_payload["formalization_gate"] = formalization_gate
        if workflow_phase_plan:
            package_payload["workflow_phase_plan"] = workflow_phase_plan
        if drafting_handoff:
            package_payload["drafting_handoff"] = drafting_handoff
        if claim_support_temporal_handoff:
            package_payload["claim_support_temporal_handoff"] = claim_support_temporal_handoff
        if email_timeline_handoff:
            package_payload["email_timeline_handoff"] = email_timeline_handoff
        if email_authority_enrichment:
            package_payload["email_authority_enrichment"] = email_authority_enrichment
        if claim_reasoning_review:
            package_payload["claim_reasoning_review"] = claim_reasoning_review
        if chronology_blocker_summary:
            package_payload["chronology_blocker_summary"] = chronology_blocker_summary
        return package_payload

    def build_draft(
        self,
        *,
        user_id: str,
        court_name: str,
        district: str,
        county: Optional[str],
        division: Optional[str],
        court_header_override: Optional[str],
        case_number: Optional[str],
        lead_case_number: Optional[str],
        related_case_number: Optional[str],
        assigned_judge: Optional[str],
        courtroom: Optional[str],
        title_override: Optional[str],
        plaintiff_names: Optional[List[str]],
        defendant_names: Optional[List[str]],
        requested_relief: Optional[List[str]],
        jury_demand: Optional[bool],
        jury_demand_text: Optional[str],
        signer_name: Optional[str],
        signer_title: Optional[str],
        signer_firm: Optional[str],
        signer_bar_number: Optional[str],
        signer_contact: Optional[str],
        additional_signers: Optional[List[Dict[str, str]]],
        declarant_name: Optional[str],
        service_method: Optional[str],
        service_recipients: Optional[List[str]],
        service_recipient_details: Optional[List[Dict[str, str]]],
        signature_date: Optional[str],
        verification_date: Optional[str],
        service_date: Optional[str],
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        affidavit_overrides = self._build_affidavit_overrides(
            affidavit_title=affidavit_title,
            affidavit_intro=affidavit_intro,
            affidavit_facts=affidavit_facts,
            affidavit_supporting_exhibits=affidavit_supporting_exhibits,
            affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
            affidavit_venue_lines=affidavit_venue_lines,
            affidavit_jurat=affidavit_jurat,
            affidavit_notary_block=affidavit_notary_block,
        )
        canonical_generate = getattr(self.mediator, "generate_formal_complaint", None)
        if callable(canonical_generate):
            try:
                result = canonical_generate(
                    user_id=user_id,
                    court_name=court_name,
                    district=district,
                    county=county,
                    division=division,
                    court_header_override=court_header_override,
                    case_number=case_number,
                    lead_case_number=lead_case_number,
                    related_case_number=related_case_number,
                    assigned_judge=assigned_judge,
                    courtroom=courtroom,
                    title_override=title_override,
                    plaintiff_names=plaintiff_names,
                    defendant_names=defendant_names,
                    requested_relief=requested_relief,
                    jury_demand=jury_demand,
                    jury_demand_text=jury_demand_text,
                    signer_name=signer_name,
                    signer_title=signer_title,
                    signer_firm=signer_firm,
                    signer_bar_number=signer_bar_number,
                    signer_contact=signer_contact,
                    additional_signers=additional_signers,
                    declarant_name=declarant_name,
                    service_method=service_method,
                    service_recipients=service_recipients,
                    service_recipient_details=service_recipient_details,
                    signature_date=signature_date,
                    verification_date=verification_date,
                    service_date=service_date,
                    affidavit_title=affidavit_title,
                    affidavit_intro=affidavit_intro,
                    affidavit_facts=affidavit_facts,
                    affidavit_supporting_exhibits=affidavit_supporting_exhibits,
                    affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
                    affidavit_venue_lines=affidavit_venue_lines,
                    affidavit_jurat=affidavit_jurat,
                    affidavit_notary_block=affidavit_notary_block,
                )
            except TypeError:
                result = None
            except Exception:
                result = None
            if isinstance(result, dict) and isinstance(result.get("formal_complaint"), dict):
                draft = self._adapt_formal_complaint_to_package_draft(result["formal_complaint"])
                draft["affidavit_overrides"] = affidavit_overrides
                draft["affidavit"] = self._build_affidavit(draft)
                draft["draft_text"] = self._render_draft_text(draft)
                return draft

        return self._build_legacy_draft(
            user_id=user_id,
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            court_header_override=court_header_override,
            case_number=case_number,
            lead_case_number=lead_case_number,
            related_case_number=related_case_number,
            assigned_judge=assigned_judge,
            courtroom=courtroom,
            title_override=title_override,
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
            requested_relief=requested_relief,
            jury_demand=jury_demand,
            jury_demand_text=jury_demand_text,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            declarant_name=declarant_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            signature_date=signature_date,
            verification_date=verification_date,
            service_date=service_date,
            affidavit_title=affidavit_title,
            affidavit_intro=affidavit_intro,
            affidavit_facts=affidavit_facts,
            affidavit_supporting_exhibits=affidavit_supporting_exhibits,
            affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
            affidavit_venue_lines=affidavit_venue_lines,
            affidavit_jurat=affidavit_jurat,
            affidavit_notary_block=affidavit_notary_block,
        )

    def _build_legacy_draft(
        self,
        *,
        user_id: str,
        court_name: str,
        district: str,
        county: Optional[str],
        division: Optional[str],
        court_header_override: Optional[str],
        case_number: Optional[str],
        lead_case_number: Optional[str],
        related_case_number: Optional[str],
        assigned_judge: Optional[str],
        courtroom: Optional[str],
        title_override: Optional[str],
        plaintiff_names: Optional[List[str]],
        defendant_names: Optional[List[str]],
        requested_relief: Optional[List[str]],
        jury_demand: Optional[bool],
        jury_demand_text: Optional[str],
        signer_name: Optional[str],
        signer_title: Optional[str],
        signer_firm: Optional[str],
        signer_bar_number: Optional[str],
        signer_contact: Optional[str],
        additional_signers: Optional[List[Dict[str, str]]],
        declarant_name: Optional[str],
        service_method: Optional[str],
        service_recipients: Optional[List[str]],
        service_recipient_details: Optional[List[Dict[str, str]]],
        signature_date: Optional[str],
        verification_date: Optional[str],
        service_date: Optional[str],
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        state = getattr(self.mediator, "state", None)
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        generated_complaint = self._get_existing_formal_complaint()
        classification = getattr(state, "legal_classification", {}) or {}
        statutes = _coerce_list(getattr(state, "applicable_statutes", []) or [])
        requirements = getattr(state, "summary_judgment_requirements", {}) or {}
        support_summary = _safe_call(self.mediator, "summarize_claim_support", user_id=user_id) or {}
        support_claims = support_summary.get("claims", {}) if isinstance(support_summary, dict) else {}
        claim_types = self._derive_claim_types(
            generated_complaint,
            classification,
            support_claims,
            requirements,
            user_id=user_id,
        )
        plaintiffs, defendants = self._derive_parties(
            generated_complaint,
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
        )
        title = title_override or generated_complaint.get("title") or self._derive_title(plaintiffs, defendants)
        exhibits = self._collect_exhibits(user_id=user_id, claim_types=claim_types, support_claims=support_claims)
        fact_entries = self._build_summary_fact_entries(
            user_id=user_id,
            generated_complaint=generated_complaint,
            classification=classification,
            state=state,
        )
        fact_entries = self._annotate_entries_with_exhibits(fact_entries, exhibits)
        facts = [str(entry.get("text") or "").strip() for entry in fact_entries if str(entry.get("text") or "").strip()]
        claims_for_relief = self._build_claims_for_relief(
            user_id=user_id,
            claim_types=claim_types,
            requirements=requirements,
            statutes=statutes,
            support_claims=support_claims,
            exhibits=exhibits,
        )
        factual_allegation_entries = self._build_factual_allegation_entries(
            summary_fact_entries=fact_entries,
            claims_for_relief=claims_for_relief,
        )
        factual_allegations = self._build_factual_allegations(
            summary_of_facts=facts,
            claims_for_relief=claims_for_relief,
        )
        relief_items = _unique_preserving_order(
            list(requested_relief or [])
            + list(generated_complaint.get("prayer_for_relief", []) or [])
            + self._extract_requested_relief_from_facts(facts)
            + self._build_claim_specific_relief(claim_types=claim_types, facts=facts)
            + (STATE_DEFAULT_RELIEF if str(classification.get("jurisdiction") or "").strip().lower() == "state" else DEFAULT_RELIEF)
        )
        jury_demand_block = self._build_jury_demand(jury_demand=jury_demand, jury_demand_text=jury_demand_text)
        court_header = self._build_court_header(
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            override=court_header_override,
        )
        jurisdiction_statement = self._build_jurisdiction_statement(
            classification=classification,
            statutes=statutes,
            court_name=court_name,
        )
        venue_statement = self._build_venue_statement(
            district=district,
            county=county,
            division=division,
            classification=classification,
            court_name=court_name,
        )
        nature_of_action = self._build_nature_of_action(
            claim_types=claim_types,
            classification=classification,
            statutes=statutes,
            court_name=court_name,
        )
        legal_standards = self._build_legal_standards_summary(statutes=statutes, requirements=requirements)
        signature_block = self._build_signature_block(
            plaintiffs,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            signature_date=signature_date,
        )
        verification = self._build_verification(
            plaintiffs,
            declarant_name=declarant_name,
            signer_name=signer_name,
            verification_date=verification_date,
            jurisdiction=classification.get("jurisdiction"),
        )
        certificate_of_service = self._build_certificate_of_service(
            plaintiffs,
            defendants,
            signer_name=signer_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            service_date=service_date,
            jurisdiction=classification.get("jurisdiction"),
        )

        draft = {
            "court_header": court_header,
            "case_caption": {
                "plaintiffs": plaintiffs,
                "defendants": defendants,
                "case_number": case_number or "________________",
                "county": county.strip().upper() if isinstance(county, str) and county.strip() else None,
                "lead_case_number": lead_case_number.strip() if isinstance(lead_case_number, str) and lead_case_number.strip() else None,
                "related_case_number": related_case_number.strip() if isinstance(related_case_number, str) and related_case_number.strip() else None,
                "assigned_judge": assigned_judge.strip() if isinstance(assigned_judge, str) and assigned_judge.strip() else None,
                "courtroom": courtroom.strip() if isinstance(courtroom, str) and courtroom.strip() else None,
                "jury_demand_notice": "JURY TRIAL DEMANDED" if jury_demand_block else None,
                "document_title": "COMPLAINT",
            },
            "title": title,
            "nature_of_action": nature_of_action,
            "parties": {
                "plaintiffs": plaintiffs,
                "defendants": defendants,
            },
            "jurisdiction_statement": jurisdiction_statement,
            "venue_statement": venue_statement,
            "factual_allegations": factual_allegations,
            "factual_allegation_entries": factual_allegation_entries,
            "summary_of_facts": facts,
            "summary_of_fact_entries": fact_entries,
            "anchored_chronology_summary": _build_anchored_chronology_summary_from_case_file(
                intake_case_file if isinstance(intake_case_file, dict) else {}
            ),
            "claims_for_relief": claims_for_relief,
            "legal_standards": legal_standards,
            "requested_relief": relief_items,
            "jury_demand": jury_demand_block,
            "exhibits": exhibits,
            "signature_block": signature_block,
            "verification": verification,
            "certificate_of_service": certificate_of_service,
            "source_context": {
                "user_id": user_id,
                "claim_types": claim_types,
                "district": district,
                "jurisdiction": classification.get("jurisdiction", "unknown"),
                "generated_at": _utcnow().isoformat(),
            },
            "affidavit_overrides": self._build_affidavit_overrides(
                affidavit_title=affidavit_title,
                affidavit_intro=affidavit_intro,
                affidavit_facts=affidavit_facts,
                affidavit_supporting_exhibits=affidavit_supporting_exhibits,
                affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
                affidavit_venue_lines=affidavit_venue_lines,
                affidavit_jurat=affidavit_jurat,
                affidavit_notary_block=affidavit_notary_block,
            ),
        }
        self._annotate_claim_temporal_gap_hints(draft)
        self._attach_allegation_references(draft)
        self._annotate_case_caption_display(draft)
        draft["affidavit"] = self._build_affidavit(draft)
        draft["draft_text"] = self._render_draft_text(draft)
        return draft

    def _build_anchored_chronology_summary(self) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        return _build_anchored_chronology_summary_from_case_file(
            intake_case_file if isinstance(intake_case_file, dict) else {}
        )

    def _optimize_draft(
        self,
        draft: Dict[str, Any],
        *,
        user_id: Optional[str],
        max_iterations: int,
        target_score: float,
        provider: Optional[str],
        model_name: Optional[str],
        llm_config: Optional[Dict[str, Any]],
        persist_artifacts: bool,
    ) -> tuple[Dict[str, Any], Dict[str, Any]]:
        optimizer = AgenticDocumentOptimizer(
            self.mediator,
            builder=self,
            provider=provider,
            model_name=model_name,
            max_iterations=max_iterations,
            target_score=target_score,
            persist_artifacts=persist_artifacts,
        )
        report = optimizer.optimize_draft(
            draft=draft,
            user_id=user_id,
            drafting_readiness={},
            config={
                "provider": provider,
                "model_name": model_name,
                "max_iterations": max_iterations,
                "target_score": target_score,
                "persist_artifacts": persist_artifacts,
                "llm_config": dict(llm_config or {}),
            },
        )
        optimized_draft = report.get("draft") or dict(draft)
        original_claim_index = {
            str(claim.get("claim_type") or "").strip().lower(): claim
            for claim in _coerce_list(draft.get("claims_for_relief"))
            if isinstance(claim, dict) and str(claim.get("claim_type") or "").strip()
        }
        optimized_draft["summary_of_facts"] = self._normalize_text_lines(optimized_draft.get("summary_of_facts", []))
        optimized_draft["summary_of_fact_entries"] = self._align_entries_to_lines(
            optimized_draft.get("summary_of_fact_entries") or draft.get("summary_of_fact_entries"),
            optimized_draft.get("summary_of_facts", []),
        )
        optimized_draft["factual_allegations"] = self._expand_allegation_sources(
            optimized_draft.get("factual_allegations", []),
            limit=24,
        ) or self._expand_allegation_sources(draft.get("factual_allegations", []), limit=24)
        optimized_draft["factual_allegation_entries"] = self._align_entries_to_lines(
            optimized_draft.get("factual_allegation_entries") or draft.get("factual_allegation_entries"),
            optimized_draft.get("factual_allegations", []),
        )
        for claim in _coerce_list(optimized_draft.get("claims_for_relief")):
            if not isinstance(claim, dict):
                continue
            original_claim = original_claim_index.get(str(claim.get("claim_type") or "").strip().lower(), {})
            if not claim.get("supporting_fact_entries") and isinstance(original_claim, dict):
                claim["supporting_fact_entries"] = deepcopy(original_claim.get("supporting_fact_entries") or [])
            claim["supporting_facts"] = self._expand_allegation_sources(
                claim.get("supporting_facts", []),
                limit=10,
            ) or self._normalize_text_lines(claim.get("supporting_facts", []))
        self._annotate_claim_temporal_gap_hints(optimized_draft)
        self._attach_allegation_references(optimized_draft)
        self._annotate_case_caption_display(optimized_draft)
        optimized_draft["affidavit"] = self._build_affidavit(optimized_draft)
        optimized_draft["draft_text"] = self._render_draft_text(optimized_draft)
        return optimized_draft, report

    def _build_intake_summary_handoff(self, document_optimization: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        optimization_handoff = optimization_report.get("intake_summary_handoff")
        if isinstance(optimization_handoff, dict) and optimization_handoff:
            return dict(optimization_handoff)

        intake_status = build_intake_status_summary(self.mediator)
        status_handoff = intake_status.get("intake_summary_handoff") if isinstance(intake_status, dict) else None
        if isinstance(status_handoff, dict) and status_handoff:
            return dict(status_handoff)

        intake_case_summary = build_intake_case_review_summary(self.mediator)
        case_handoff = (
            intake_case_summary.get("intake_summary_handoff")
            if isinstance(intake_case_summary, dict)
            else None
        )
        if isinstance(case_handoff, dict) and case_handoff:
            return dict(case_handoff)

        return {}

    def _build_claim_support_temporal_handoff(self, document_optimization: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        optimization_handoff = optimization_report.get("claim_support_temporal_handoff")
        if isinstance(optimization_handoff, dict) and optimization_handoff:
            return dict(optimization_handoff)

        intake_case_summary = optimization_report.get("intake_case_summary")
        if not isinstance(intake_case_summary, dict) or not intake_case_summary:
            intake_case_summary = build_intake_case_review_summary(self.mediator)
        if not isinstance(intake_case_summary, dict) or not intake_case_summary:
            return {}

        packet_summary = intake_case_summary.get("claim_support_packet_summary")
        packet_summary = packet_summary if isinstance(packet_summary, dict) else {}
        alignment_tasks = intake_case_summary.get("alignment_evidence_tasks")
        alignment_tasks = alignment_tasks if isinstance(alignment_tasks, list) else []
        temporal_issue_registry_summary = summarize_temporal_issue_registry(
            intake_case_summary.get("temporal_issue_registry_summary")
        )

        unresolved_temporal_issue_ids = _dedupe_text_values(
            packet_summary.get("claim_support_unresolved_temporal_issue_ids") or []
        )
        event_ids: List[str] = []
        temporal_fact_ids: List[str] = []
        temporal_relation_ids: List[str] = []
        timeline_anchor_ids: List[str] = []
        timeline_issue_ids: List[str] = []
        temporal_issue_ids: List[str] = []
        missing_temporal_predicates: List[str] = []
        required_provenance_kinds: List[str] = []
        temporal_proof_bundle_ids: List[str] = []
        temporal_proof_objectives: List[str] = []

        for task in alignment_tasks:
            if not isinstance(task, dict):
                continue
            event_ids.extend(_dedupe_text_values(task.get("event_ids") or []))
            temporal_fact_ids.extend(_dedupe_text_values(task.get("temporal_fact_ids") or []))
            temporal_relation_ids.extend(_dedupe_text_values(task.get("temporal_relation_ids") or []))
            timeline_anchor_ids.extend(_dedupe_text_values(task.get("anchor_ids") or task.get("timeline_anchor_ids") or []))
            timeline_issue_ids.extend(_dedupe_text_values(task.get("timeline_issue_ids") or []))
            temporal_issue_ids.extend(_dedupe_text_values(task.get("temporal_issue_ids") or []))
            missing_temporal_predicates.extend(_dedupe_text_values(task.get("missing_temporal_predicates") or []))
            required_provenance_kinds.extend(_dedupe_text_values(task.get("required_provenance_kinds") or []))
            proof_bundle_id = str(task.get("temporal_proof_bundle_id") or "").strip()
            if proof_bundle_id:
                temporal_proof_bundle_ids.append(proof_bundle_id)
            proof_objective = str(task.get("temporal_proof_objective") or "").strip()
            if proof_objective:
                temporal_proof_objectives.append(proof_objective)

        raw_event_ids = _collect_temporal_registry_identifiers(
            intake_case_summary.get("event_ledger"),
            "event_id",
        )
        raw_temporal_fact_ids = _collect_temporal_registry_identifiers(
            intake_case_summary.get("temporal_fact_registry"),
            "temporal_fact_id",
            "fact_id",
            "event_id",
        )
        if not raw_temporal_fact_ids:
            raw_temporal_fact_ids = _collect_temporal_registry_identifiers(
                intake_case_summary.get("event_ledger"),
                "temporal_fact_id",
                "event_id",
                "fact_id",
            )
        raw_temporal_relation_ids = _collect_temporal_registry_identifiers(
            intake_case_summary.get("temporal_relation_registry"),
            "temporal_relation_id",
            "relation_id",
        )
        if not raw_temporal_relation_ids:
            raw_temporal_relation_ids = _collect_temporal_registry_identifiers(
                intake_case_summary.get("timeline_relations"),
                "temporal_relation_id",
                "relation_id",
            )
        raw_temporal_issue_ids = _collect_temporal_registry_identifiers(
            intake_case_summary.get("temporal_issue_registry"),
            "temporal_issue_id",
            "issue_id",
            "timeline_issue_id",
        )
        raw_timeline_anchor_ids = _collect_temporal_registry_identifiers(
            intake_case_summary.get("timeline_anchors"),
            "anchor_id",
            "timeline_anchor_id",
        )
        raw_unresolved_temporal_issue_ids = _collect_unresolved_temporal_issue_identifiers(
            intake_case_summary.get("temporal_issue_registry")
        )
        summary_issue_ids = _dedupe_text_values(temporal_issue_registry_summary.get("issue_ids") or [])
        summary_missing_temporal_predicates = _dedupe_text_values(
            temporal_issue_registry_summary.get("missing_temporal_predicates") or []
        )
        summary_required_provenance_kinds = _dedupe_text_values(
            temporal_issue_registry_summary.get("required_provenance_kinds") or []
        )

        unresolved_temporal_issue_count = int(
            packet_summary.get("claim_support_unresolved_temporal_issue_count", 0) or 0
        )
        if not unresolved_temporal_issue_count and raw_unresolved_temporal_issue_ids:
            unresolved_temporal_issue_count = len(raw_unresolved_temporal_issue_ids)
        if not unresolved_temporal_issue_count:
            unresolved_temporal_issue_count = int(temporal_issue_registry_summary.get("unresolved_count") or 0)
        if not unresolved_temporal_issue_ids:
            unresolved_temporal_issue_ids = raw_unresolved_temporal_issue_ids

        temporal_handoff = {
            "unresolved_temporal_issue_count": unresolved_temporal_issue_count,
            "unresolved_temporal_issue_ids": unresolved_temporal_issue_ids,
            "chronology_task_count": int(packet_summary.get("temporal_gap_task_count", 0) or 0),
            "event_ids": _dedupe_text_values(event_ids) or raw_event_ids,
            "temporal_fact_ids": _dedupe_text_values(temporal_fact_ids) or raw_temporal_fact_ids,
            "temporal_relation_ids": _dedupe_text_values(temporal_relation_ids) or raw_temporal_relation_ids,
            "timeline_issue_ids": _dedupe_text_values(timeline_issue_ids) or raw_temporal_issue_ids or summary_issue_ids,
            "temporal_issue_ids": _dedupe_text_values(temporal_issue_ids) or raw_temporal_issue_ids or summary_issue_ids,
            "temporal_proof_bundle_ids": _dedupe_text_values(temporal_proof_bundle_ids),
            "temporal_proof_objectives": _dedupe_text_values(temporal_proof_objectives),
        }
        normalized_timeline_anchor_ids = _dedupe_text_values(timeline_anchor_ids) or raw_timeline_anchor_ids
        if normalized_timeline_anchor_ids:
            temporal_handoff["timeline_anchor_ids"] = normalized_timeline_anchor_ids
        normalized_missing_temporal_predicates = _dedupe_text_values(missing_temporal_predicates) or summary_missing_temporal_predicates
        if normalized_missing_temporal_predicates:
            temporal_handoff["missing_temporal_predicates"] = normalized_missing_temporal_predicates
        normalized_required_provenance_kinds = _dedupe_text_values(required_provenance_kinds) or summary_required_provenance_kinds
        if normalized_required_provenance_kinds:
            temporal_handoff["required_provenance_kinds"] = normalized_required_provenance_kinds
        if not temporal_handoff["unresolved_temporal_issue_count"] and not any(
            temporal_handoff.get(key)
            for key in (
                "unresolved_temporal_issue_ids",
                "event_ids",
                "temporal_fact_ids",
                "temporal_relation_ids",
                "timeline_anchor_ids",
                "timeline_issue_ids",
                "temporal_issue_ids",
                "missing_temporal_predicates",
                "required_provenance_kinds",
                "temporal_proof_bundle_ids",
                "temporal_proof_objectives",
            )
        ):
            return {}
        return temporal_handoff

    def _resolve_email_timeline_handoff(
        self,
        *,
        draft: Dict[str, Any],
        document_optimization: Optional[Dict[str, Any]],
    ) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        for candidate in (
            optimization_report.get("email_timeline_handoff"),
            (draft.get("source_context") or {}).get("email_timeline_handoff")
            if isinstance(draft.get("source_context"), dict)
            else None,
        ):
            if isinstance(candidate, dict) and candidate:
                return deepcopy(candidate)
        return {}

    def _resolve_email_authority_enrichment(
        self,
        *,
        draft: Dict[str, Any],
        document_optimization: Optional[Dict[str, Any]],
    ) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        for candidate in (
            optimization_report.get("email_authority_enrichment"),
            (draft.get("source_context") or {}).get("email_authority_enrichment")
            if isinstance(draft.get("source_context"), dict)
            else None,
        ):
            if isinstance(candidate, dict) and candidate:
                return deepcopy(candidate)
        return {}

    def _merge_claim_support_temporal_handoff(
        self,
        base_handoff: Dict[str, Any],
        email_timeline_handoff: Dict[str, Any],
    ) -> Dict[str, Any]:
        merged = dict(base_handoff or {}) if isinstance(base_handoff, dict) else {}
        email_packet = (
            email_timeline_handoff.get("claim_support_temporal_handoff")
            if isinstance(email_timeline_handoff.get("claim_support_temporal_handoff"), dict)
            else {}
        )
        if not email_packet:
            return merged

        list_keys = (
            "unresolved_temporal_issue_ids",
            "event_ids",
            "temporal_fact_ids",
            "temporal_relation_ids",
            "timeline_anchor_ids",
            "timeline_issue_ids",
            "temporal_issue_ids",
            "temporal_proof_bundle_ids",
            "temporal_proof_objectives",
            "missing_temporal_predicates",
            "required_provenance_kinds",
        )
        for key in list_keys:
            merged_values = _dedupe_text_values(list(merged.get(key) or []) + list(email_packet.get(key) or []))
            if merged_values:
                merged[key] = merged_values

        for key in ("unresolved_temporal_issue_count", "chronology_task_count"):
            merged[key] = max(int(merged.get(key) or 0), int(email_packet.get(key) or 0))

        for key in ("timeline_anchor_count", "event_count"):
            if email_packet.get(key) not in (None, ""):
                merged[key] = int(email_packet.get(key) or 0)

        if isinstance(email_packet.get("topic_summary"), dict) and email_packet.get("topic_summary"):
            merged["email_topic_summary"] = deepcopy(email_packet.get("topic_summary") or {})

        contract_version = str(email_packet.get("contract_version") or "").strip()
        if contract_version:
            merged.setdefault("contract_version", contract_version)
        return merged

    def _build_email_timeline_fact_entries(
        self,
        email_timeline_handoff: Dict[str, Any],
        *,
        limit: int = 6,
        exclude_predicate_families: Optional[Iterable[str]] = None,
    ) -> List[Dict[str, Any]]:
        canonical_facts = (
            list(email_timeline_handoff.get("canonical_facts") or [])
            if isinstance(email_timeline_handoff, dict)
            else []
        )
        default_claim_type = str(email_timeline_handoff.get("claim_type") or "").strip()
        default_claim_element_id = str(email_timeline_handoff.get("claim_element_id") or "").strip()
        excluded_predicates = {
            str(value or "").strip().lower()
            for value in (exclude_predicate_families or [])
            if str(value or "").strip()
        }
        entries: List[Dict[str, Any]] = []
        seen = set()
        for fact in canonical_facts:
            if not isinstance(fact, dict):
                continue
            predicate_family = str(fact.get("predicate_family") or "").strip().lower()
            if predicate_family and predicate_family in excluded_predicates:
                continue
            text = re.sub(r"\s+", " ", str(fact.get("text") or "").strip())
            if len(text) < 12:
                continue
            key = text.lower()
            if key in seen:
                continue
            seen.add(key)
            entries.append(
                {
                    "text": text,
                    "fact_ids": _normalize_identifier_list(
                        [fact.get("fact_id")] + list(fact.get("related_fact_ids") or [])
                    ),
                    "source_artifact_ids": _normalize_identifier_list(
                        list(fact.get("source_artifact_ids") or [])
                    ),
                    "claim_types": _normalize_identifier_list(
                        list(fact.get("claim_types") or [])
                        + ([default_claim_type] if default_claim_type else [])
                    ),
                    "claim_element_ids": _normalize_identifier_list(
                        list(fact.get("element_tags") or [])
                        + ([default_claim_element_id] if default_claim_element_id else [])
                        + ([fact.get("predicate_family")] if fact.get("predicate_family") else [])
                    ),
                    "source_kind": str(fact.get("source_kind") or "email_timeline_candidate").strip()
                    or "email_timeline_candidate",
                    "source_ref": str(fact.get("source_ref") or fact.get("fact_id") or "").strip() or None,
                }
            )
            if len(entries) >= limit:
                break
        return entries

    def _load_email_timeline_message_metadata(self, source_ref: Any) -> Dict[str, Any]:
        source_text = str(source_ref or "").strip()
        if not source_text:
            return {}
        source_path = Path(source_text)
        candidates: List[Path] = []
        if source_path.suffix.lower() == ".json":
            candidates.append(source_path)
        if source_path.name == "message.eml":
            candidates.append(source_path.with_name("message.json"))
        for candidate in candidates:
            try:
                if candidate.exists():
                    payload = json.loads(candidate.read_text(encoding="utf-8"))
                    if isinstance(payload, dict):
                        return payload
            except Exception:
                continue
        return {}

    def _build_email_timeline_narrative_text(
        self,
        predicate_family: str,
        facts: List[Dict[str, Any]],
    ) -> str:
        ordered_facts = sorted(
            [fact for fact in facts if isinstance(fact, dict)],
            key=lambda fact: str(
                (fact.get("temporal_context") or {}).get("sortable_date")
                or fact.get("event_date_or_range")
                or ""
            ),
        )
        if not ordered_facts:
            return ""

        def _date_text(fact: Dict[str, Any]) -> str:
            return _format_timeline_date(
                (fact.get("temporal_context") or {}).get("start_date")
                or fact.get("event_date_or_range")
            )

        first_date = _date_text(ordered_facts[0])
        last_date = _date_text(ordered_facts[-1])
        clackamas_participants = _unique_preserving_order(
            [
                str(email).strip().lower()
                for fact in ordered_facts
                for email in list(fact.get("participants") or [])
                if str(email).strip().lower().endswith("@clackamas.us")
            ]
        )
        participant_phrase = _join_chronology_segments(clackamas_participants[:4]) if clackamas_participants else ""

        if predicate_family == "fraud_household" and first_date and last_date:
            return (
                f"Beginning on {first_date} and continuing through {last_date}, Plaintiff emailed "
                "Clackamas County Housing Authority staff in the "
                "'Allegations of Fraud - JC Household' chain, "
                f"including {participant_phrase or 'multiple Clackamas housing staff members'}."
            )
        if predicate_family == "additional_information" and first_date and last_date:
            return (
                f"From {first_date} through {last_date}, Clackamas housing staff sent "
                "'Additional Information Needed' emails requesting additional housing-program information from Plaintiff."
            )
        if predicate_family == "cortez_case" and first_date:
            return (
                f"On {first_date}, Plaintiff emailed Clackamas personnel regarding the "
                "'Jane Kay Cortez vs Solomon Samuel Barber' matter."
            )
        if predicate_family == "hcv_orientation" and first_date:
            attachment_names = _unique_preserving_order(
                [
                    str(attachment.get("filename") or "").strip()
                    for fact in ordered_facts
                    for attachment in list(
                        self._load_email_timeline_message_metadata(fact.get("source_ref")).get("attachments") or []
                    )
                    if isinstance(attachment, dict) and str(attachment.get("filename") or "").strip()
                ]
            )
            denial_attachment = next(
                (name for name in attachment_names if "denial" in name.lower() or "ra-" in name.lower()),
                "",
            )
            attachment_clause = (
                f" One of those Clackamas emails included the attachment '{denial_attachment}'."
                if denial_attachment
                else ""
            )
            return (
                f"On {first_date}, Clackamas housing staff emailed Plaintiff in the 'HCV Orientation' thread "
                f"involving {participant_phrase or 'Clackamas voucher staff'}."
                f"{attachment_clause}"
            )
        return ""

    def _build_email_timeline_narrative_fact_entries(
        self,
        email_timeline_handoff: Dict[str, Any],
        *,
        limit: int = 4,
    ) -> List[Dict[str, Any]]:
        canonical_facts = (
            list(email_timeline_handoff.get("canonical_facts") or [])
            if isinstance(email_timeline_handoff, dict)
            else []
        )
        grouped: Dict[str, List[Dict[str, Any]]] = {}
        for fact in canonical_facts:
            if not isinstance(fact, dict):
                continue
            predicate_family = str(fact.get("predicate_family") or "").strip().lower()
            if not predicate_family:
                continue
            grouped.setdefault(predicate_family, []).append(fact)

        entries: List[Dict[str, Any]] = []
        for predicate_family in ("hcv_orientation", "additional_information", "cortez_case", "fraud_household"):
            facts = grouped.get(predicate_family) or []
            text = self._build_email_timeline_narrative_text(predicate_family, facts)
            if not text:
                continue
            entries.append(
                {
                    "text": text,
                    "fact_ids": _normalize_identifier_list([fact.get("fact_id") for fact in facts]),
                    "source_artifact_ids": _normalize_identifier_list(
                        [fact.get("source_ref") for fact in facts if fact.get("source_ref")]
                    ),
                    "claim_types": _normalize_identifier_list(
                        [claim_type for fact in facts for claim_type in list(fact.get("claim_types") or [])]
                    ),
                    "claim_element_ids": _normalize_identifier_list(
                        [predicate_family]
                        + [element_id for fact in facts for element_id in list(fact.get("element_tags") or [])]
                    ),
                    "source_kind": "email_timeline_narrative",
                    "source_ref": str(facts[0].get("source_ref") or "").strip() or None,
                }
            )
            if len(entries) >= limit:
                break
        return entries

    def _build_email_timeline_chronology_lines(
        self,
        email_timeline_handoff: Dict[str, Any],
        *,
        limit: int = 6,
    ) -> List[str]:
        canonical_facts = (
            list(email_timeline_handoff.get("canonical_facts") or [])
            if isinstance(email_timeline_handoff, dict)
            else []
        )
        lines: List[str] = []
        seen = set()
        for fact in canonical_facts:
            if not isinstance(fact, dict):
                continue
            temporal_context = fact.get("temporal_context") if isinstance(fact.get("temporal_context"), dict) else {}
            event_date = _format_timeline_date(
                temporal_context.get("start_date")
                or fact.get("event_date_or_range")
            )
            text = re.sub(r"\s+", " ", str(fact.get("text") or "").strip())
            if len(text) < 12 or not event_date:
                continue
            line = f"On {event_date}, {text}"
            if not line.endswith((".", "?", "!")):
                line = f"{line}."
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            lines.append(line)
            if len(lines) >= limit:
                break
        return lines

    def _enrich_draft_with_email_timeline_handoff(
        self,
        draft: Dict[str, Any],
        *,
        email_timeline_handoff: Dict[str, Any],
    ) -> Dict[str, Any]:
        if not isinstance(draft, dict) or not isinstance(email_timeline_handoff, dict) or not email_timeline_handoff:
            return draft

        narrative_entries = self._build_email_timeline_narrative_fact_entries(email_timeline_handoff)
        narrative_predicates = _normalize_identifier_list(
            [
                element_id
                for entry in narrative_entries
                for element_id in _coerce_list(entry.get("claim_element_ids"))
                if str(element_id or "").strip().lower()
                in {"hcv_orientation", "additional_information", "cortez_case", "fraud_household"}
            ]
        )
        raw_limit = 0 if narrative_entries else 6
        email_fact_entries = narrative_entries + self._build_email_timeline_fact_entries(
            email_timeline_handoff,
            limit=raw_limit,
            exclude_predicate_families=narrative_predicates,
        )
        email_chronology_lines = self._build_email_timeline_chronology_lines(email_timeline_handoff)
        if not email_fact_entries and not email_chronology_lines:
            return draft

        updated = dict(draft)
        existing_fact_entries = [
            dict(entry)
            for entry in _coerce_list(updated.get("summary_of_fact_entries"))
            if isinstance(entry, dict)
        ]
        merged_fact_entries: List[Dict[str, Any]] = []
        seen_fact_texts = set()
        for entry in email_fact_entries + existing_fact_entries:
            text = re.sub(r"\s+", " ", str(entry.get("text") or "").strip())
            if len(text) < 12:
                continue
            key = text.lower()
            if key in seen_fact_texts:
                continue
            seen_fact_texts.add(key)
            normalized_entry = dict(entry)
            normalized_entry["text"] = text
            normalized_entry["fact_ids"] = _normalize_identifier_list(normalized_entry.get("fact_ids") or [])
            normalized_entry["source_artifact_ids"] = _normalize_identifier_list(
                normalized_entry.get("source_artifact_ids") or []
            )
            normalized_entry["claim_types"] = _normalize_identifier_list(normalized_entry.get("claim_types") or [])
            normalized_entry["claim_element_ids"] = _normalize_identifier_list(
                normalized_entry.get("claim_element_ids") or []
            )
            if normalized_entry.get("source_ref") is not None:
                source_ref = str(normalized_entry.get("source_ref") or "").strip()
                normalized_entry["source_ref"] = source_ref or None
            merged_fact_entries.append(normalized_entry)
            if len(merged_fact_entries) >= 14:
                break

        if narrative_entries:
            filtered_fact_entries = [
                entry for entry in merged_fact_entries if not self._is_low_signal_email_summary_entry(entry)
            ]
            if filtered_fact_entries:
                merged_fact_entries = filtered_fact_entries

        if merged_fact_entries:
            updated["summary_of_fact_entries"] = merged_fact_entries
            updated["summary_of_facts"] = [
                str(entry.get("text") or "").strip()
                for entry in merged_fact_entries
                if str(entry.get("text") or "").strip()
            ]

        chronology_lines = _unique_preserving_order(
            list(self._normalize_text_lines(updated.get("anchored_chronology_summary", [])))
            + email_chronology_lines
        )
        if chronology_lines:
            updated["anchored_chronology_summary"] = chronology_lines[:8]

        claims_for_relief = (
            list(updated.get("claims_for_relief") or [])
            if isinstance(updated.get("claims_for_relief"), list)
            else []
        )
        if merged_fact_entries:
            updated["factual_allegation_entries"] = self._build_factual_allegation_entries(
                summary_fact_entries=merged_fact_entries,
                claims_for_relief=claims_for_relief,
            )
            updated["factual_allegations"] = self._build_factual_allegations(
                summary_of_facts=updated.get("summary_of_facts", []),
                claims_for_relief=claims_for_relief,
            )
        return updated

    def _build_email_authority_summary_lines(
        self,
        email_authority_enrichment: Dict[str, Any],
        *,
        authority_limit: int = 5,
        guidance_limit: int = 3,
    ) -> List[str]:
        if not isinstance(email_authority_enrichment, dict) or not email_authority_enrichment:
            return []
        summary = (
            dict(email_authority_enrichment.get("summary") or {})
            if isinstance(email_authority_enrichment.get("summary"), dict)
            else {}
        )
        total_counts = (
            dict(summary.get("total_counts") or {})
            if isinstance(summary.get("total_counts"), dict)
            else {}
        )
        statutes = int(total_counts.get("statutes") or 0)
        regulations = int(total_counts.get("regulations") or 0)
        case_law = int(total_counts.get("case_law") or 0)
        guidance = int(total_counts.get("state_web_archives") or 0) + int(total_counts.get("web_archives") or 0)
        lines: List[str] = []
        if any(value > 0 for value in (statutes, regulations, case_law, guidance)):
            count_parts = []
            if statutes:
                count_parts.append(f"{statutes} statutes")
            if regulations:
                count_parts.append(f"{regulations} regulations")
            if case_law:
                count_parts.append(f"{case_law} cases")
            if guidance:
                count_parts.append(f"{guidance} agency or local guidance sources")
            lines.append(
                "Email-aligned authority review identified "
                f"{_join_chronology_segments(count_parts)} supporting the Clackamas chronology."
            )

        def _authority_priority(item: Dict[str, Any]) -> tuple[int, int, str]:
            authority_type = str(item.get("authority_type") or "").strip().lower()
            citation = str(item.get("citation") or "").strip()
            title = str(item.get("title") or "").strip()
            source_url = str(item.get("source_url") or "").strip().lower()
            priority = 0
            if citation.startswith("ORS "):
                priority = 6
            elif "clackamas" in title.lower() or "clackamas" in source_url:
                priority = 5
            elif citation.startswith("24 C.F.R."):
                priority = 4
            elif citation.startswith("42 U.S.C."):
                priority = 3
            elif authority_type in {"guidance", "agency_guidance"}:
                priority = 2
            return (priority, len(citation or title), citation or title)

        authority_names: List[str] = []
        recommended_authorities = [
            dict(item)
            for item in list(email_authority_enrichment.get("recommended_authorities") or [])
            if isinstance(item, dict)
        ]
        recommended_authorities.sort(key=_authority_priority, reverse=True)
        for item in recommended_authorities:
            if not isinstance(item, dict):
                continue
            label = str(item.get("citation") or item.get("title") or "").strip()
            if label and label not in authority_names:
                authority_names.append(label)
            if len(authority_names) >= authority_limit:
                break
        if authority_names:
            lines.append(
                "Key authorities include "
                f"{_join_chronology_segments(authority_names)}."
            )

        guidance_names: List[str] = []
        seen = set()
        for query_result in list(email_authority_enrichment.get("query_results") or []):
            if not isinstance(query_result, dict):
                continue
            for row in list(query_result.get("state_web_archives") or []):
                if not isinstance(row, dict):
                    continue
                title = str(row.get("title") or row.get("citation") or "").strip()
                if not title or title in seen:
                    continue
                seen.add(title)
                guidance_names.append(title)
                if len(guidance_names) >= guidance_limit:
                    break
            if len(guidance_names) >= guidance_limit:
                break
        if guidance_names:
            lines.append(
                "Relevant agency and local guidance includes "
                f"{_join_chronology_segments(guidance_names)}."
            )
        return lines[:4]

    def _enrich_draft_with_email_authority_enrichment(
        self,
        draft: Dict[str, Any],
        *,
        email_authority_enrichment: Dict[str, Any],
    ) -> Dict[str, Any]:
        if not isinstance(draft, dict) or not isinstance(email_authority_enrichment, dict) or not email_authority_enrichment:
            return draft
        summary_lines = self._build_email_authority_summary_lines(email_authority_enrichment)
        if not summary_lines:
            return draft
        updated = dict(draft)
        updated["email_authority_summary_lines"] = summary_lines
        return updated

    def _collect_email_authority_citations(self, email_authority_enrichment: Dict[str, Any]) -> List[str]:
        citations: List[str] = []
        seen = set()
        for item in list(email_authority_enrichment.get("recommended_authorities") or []):
            if not isinstance(item, dict):
                continue
            citation = str(item.get("citation") or item.get("title") or "").strip()
            if citation and citation not in seen:
                seen.add(citation)
                citations.append(citation)
        for query_result in list(email_authority_enrichment.get("query_results") or []):
            if not isinstance(query_result, dict):
                continue
            for bucket in ("statutes", "regulations", "case_law"):
                for row in list((query_result.get("results") or {}).get(bucket) or []):
                    if not isinstance(row, dict):
                        continue
                    citation = str(row.get("citation") or row.get("title") or "").strip()
                    if citation and citation not in seen:
                        seen.add(citation)
                        citations.append(citation)
        return citations

    def _build_claim_authority_lines_from_email_enrichment(
        self,
        claim_type: str,
        *,
        email_authority_enrichment: Dict[str, Any],
    ) -> List[str]:
        citations = self._collect_email_authority_citations(email_authority_enrichment)
        normalized = normalize_claim_type(claim_type or "")

        def _available(preferred: List[str]) -> List[str]:
            return [citation for citation in preferred if citation in citations]

        if normalized == "due_process_failure":
            preferred = _available(["24 C.F.R. § 982.555", "24 C.F.R. § 982.552", "24 C.F.R. § 982.551"])
            if preferred:
                return [f"Authority support for this count includes {_join_chronology_segments(preferred)}."]
        if normalized == "housing_discrimination":
            preferred = _available(["ORS 659A.145", "42 U.S.C. § 3604(f)(3)(B)", "HUD/DOJ Joint Statement (May 17, 2004)"])
            if preferred:
                return [f"Authority support for this count includes {_join_chronology_segments(preferred)}."]
        if normalized == "retaliation":
            preferred = _available(["ORS 659A.145", "42 U.S.C. § 3604(f)(3)(B)", "24 C.F.R. § 982.555"])
            if preferred:
                return [f"Authority support for this count includes {_join_chronology_segments(preferred)}."]
        return []

    def _infer_claim_types_from_email_support(
        self,
        *,
        email_timeline_handoff: Dict[str, Any],
        email_authority_enrichment: Dict[str, Any],
    ) -> List[str]:
        topic_summary = (
            dict((email_timeline_handoff.get("claim_support_temporal_handoff") or {}).get("topic_summary") or {})
            if isinstance((email_timeline_handoff.get("claim_support_temporal_handoff") or {}), dict)
            else {}
        )
        topics = {str(key).strip() for key in topic_summary.keys() if str(key).strip()}
        citations = set(self._collect_email_authority_citations(email_authority_enrichment))
        inferred: List[str] = []

        if topics & {"hcv_orientation", "fraud_household", "additional_information"} or citations & {
            "24 C.F.R. § 982.555",
            "24 C.F.R. § 982.552",
            "24 C.F.R. § 982.551",
        }:
            inferred.append("due_process_failure")
        if topics & {"cortez_case", "hcv_orientation"} or citations & {
            "ORS 659A.145",
            "42 U.S.C. § 3604(f)(3)(B)",
        }:
            inferred.append("housing_discrimination")
        if str(email_timeline_handoff.get("claim_type") or "").strip().lower() == "retaliation":
            inferred.append("retaliation")

        return _unique_preserving_order(inferred)

    def _specialize_generic_claims_from_email_support(
        self,
        draft: Dict[str, Any],
        *,
        email_timeline_handoff: Dict[str, Any],
        email_authority_enrichment: Dict[str, Any],
    ) -> Dict[str, Any]:
        if not isinstance(draft, dict):
            return draft
        claims = [dict(item) for item in list(draft.get("claims_for_relief") or []) if isinstance(item, dict)]
        if len(claims) != 1:
            return draft
        base_claim = claims[0]
        if normalize_claim_type(base_claim.get("claim_type") or "") != "general_civil_action":
            return draft

        inferred_claim_types = self._infer_claim_types_from_email_support(
            email_timeline_handoff=email_timeline_handoff,
            email_authority_enrichment=email_authority_enrichment,
        )
        if not inferred_claim_types:
            return draft

        specialized_claims: List[Dict[str, Any]] = []
        base_facts = [
            str(item or "").strip()
            for item in list(base_claim.get("supporting_facts") or [])
            if str(item or "").strip() and not self._is_generic_claim_support_text(item)
        ]
        summary_facts = [
            line
            for line in self._normalize_text_lines(draft.get("summary_of_facts", []))
            if not self._is_generic_claim_support_text(line)
        ]
        merged_claim_facts = _unique_preserving_order(base_facts + summary_facts)[:8]
        for claim_type in inferred_claim_types[:3]:
            specialized = deepcopy(base_claim)
            specialized["claim_type"] = claim_type
            specialized["count_title"] = self._humanize_claim_title(claim_type, merged_claim_facts or base_facts)
            if merged_claim_facts:
                specialized["supporting_facts"] = merged_claim_facts
            specialized["legal_standards"] = (
                self._build_claim_legal_standard_fallbacks(claim_type)
                + self._build_claim_authority_lines_from_email_enrichment(
                    claim_type,
                    email_authority_enrichment=email_authority_enrichment,
                )
            )
            specialized_claims.append(specialized)

        updated = dict(draft)
        updated["claims_for_relief"] = sorted(specialized_claims, key=self._claim_order_score)
        return updated

    def _build_chronology_blocker_summary(
        self,
        *,
        intake_case_summary: Any,
        claim_support_temporal_handoff: Any,
    ) -> Dict[str, Any]:
        summary = intake_case_summary if isinstance(intake_case_summary, dict) else {}
        packet_summary = summary.get("claim_support_packet_summary")
        packet_summary = packet_summary if isinstance(packet_summary, dict) else {}
        alignment_task_summary = summary.get("alignment_task_summary")
        alignment_task_summary = alignment_task_summary if isinstance(alignment_task_summary, dict) else {}
        handoff = dict(claim_support_temporal_handoff or {}) if isinstance(claim_support_temporal_handoff, dict) else {}

        unresolved_issue_ids = _dedupe_text_values(
            packet_summary.get("claim_support_unresolved_temporal_issue_ids")
            or handoff.get("unresolved_temporal_issue_ids")
            or []
        )
        unresolved_issue_count = int(
            packet_summary.get("claim_support_unresolved_temporal_issue_count")
            or handoff.get("unresolved_temporal_issue_count")
            or len(unresolved_issue_ids)
            or 0
        )
        temporal_gap_task_count = int(
            packet_summary.get("temporal_gap_task_count")
            or alignment_task_summary.get("temporal_gap_task_count")
            or 0
        )
        proof_readiness_score = float(packet_summary.get("proof_readiness_score", 0.0) or 0.0)
        chronology_blocked = bool(unresolved_issue_count > 0 or temporal_gap_task_count > 0)
        if not chronology_blocked and proof_readiness_score <= 0.0 and not unresolved_issue_ids:
            return {}

        summary_parts: List[str] = []
        if temporal_gap_task_count > 0:
            task_label = "task" if temporal_gap_task_count == 1 else "tasks"
            summary_parts.append(f"{temporal_gap_task_count} pending chronology gap {task_label}")
        if unresolved_issue_count > 0:
            issue_label = "issue" if unresolved_issue_count == 1 else "issues"
            summary_parts.append(f"{unresolved_issue_count} unresolved temporal {issue_label}")
        summary_text = (
            f"Chronology blockers remain: {'; '.join(summary_parts)}."
            if summary_parts
            else "No chronology blockers currently reduce proof readiness."
        )

        return {
            "chronology_blocked": chronology_blocked,
            "proof_readiness_score": round(proof_readiness_score, 3),
            "temporal_gap_task_count": temporal_gap_task_count,
            "unresolved_temporal_issue_count": unresolved_issue_count,
            "unresolved_temporal_issue_ids": unresolved_issue_ids,
            "summary": summary_text,
        }

    def _build_claim_reasoning_review(self, document_optimization: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        optimization_review = optimization_report.get("claim_reasoning_review")
        if isinstance(optimization_review, dict) and optimization_review:
            return deepcopy(optimization_review)

        intake_case_summary = optimization_report.get("intake_case_summary")
        if not isinstance(intake_case_summary, dict) or not intake_case_summary:
            intake_case_summary = build_intake_case_review_summary(self.mediator)
        if not isinstance(intake_case_summary, dict) or not intake_case_summary:
            return {}

        existing_review = intake_case_summary.get("claim_reasoning_review")
        if isinstance(existing_review, dict) and existing_review:
            return deepcopy(existing_review)

        validation_getter = getattr(self.mediator, "get_claim_support_validation", None)
        if not callable(validation_getter):
            return {}

        mediator_state = getattr(self.mediator, "state", None)
        legal_classification = getattr(mediator_state, "legal_classification", {})
        legal_classification = legal_classification if isinstance(legal_classification, dict) else {}
        legal_classification_claim_types = legal_classification.get("claim_types")
        legal_classification_claim_types = (
            legal_classification_claim_types
            if isinstance(legal_classification_claim_types, list)
            else []
        )
        claim_types = _unique_preserving_order(
            [
                *[
                    str((claim or {}).get("claim_type") or "").strip()
                    for claim in (intake_case_summary.get("candidate_claims") or [])
                    if isinstance(claim, dict)
                ],
                *[
                    str(item).strip()
                    for item in legal_classification_claim_types
                ],
            ]
        )
        if not claim_types:
            return {}

        review_by_claim: Dict[str, Any] = {}
        for claim_type in claim_types:
            try:
                validation_payload = validation_getter(claim_type=claim_type)
            except Exception:
                continue
            if not isinstance(validation_payload, dict):
                continue
            validation_claims = validation_payload.get("claims")
            validation_claims = validation_claims if isinstance(validation_claims, dict) else {}
            validation_claim = validation_claims.get(claim_type)
            if not isinstance(validation_claim, dict) or not validation_claim:
                normalized_claim_type = normalize_claim_type(claim_type)
                validation_claim = validation_claims.get(normalized_claim_type)
            if not isinstance(validation_claim, dict) or not validation_claim:
                continue
            claim_review = summarize_claim_reasoning_review(validation_claim)
            flagged_elements = claim_review.get("flagged_elements")
            if isinstance(flagged_elements, list):
                for flagged_element in flagged_elements:
                    if not isinstance(flagged_element, dict):
                        continue
                    theorem_export_metadata = flagged_element.get("proof_artifact_theorem_export_metadata")
                    if isinstance(theorem_export_metadata, dict) and theorem_export_metadata:
                        continue
                    fallback_metadata = _build_claim_reasoning_theorem_export_metadata(
                        intake_case_summary,
                        claim_type=claim_type,
                        claim_element_id=str(flagged_element.get("element_id") or ""),
                    )
                    if fallback_metadata:
                        flagged_element["proof_artifact_theorem_export_metadata"] = fallback_metadata
            if claim_review:
                review_by_claim[claim_type] = claim_review
        return review_by_claim

    def _adapt_formal_complaint_to_package_draft(self, formal_complaint: Dict[str, Any]) -> Dict[str, Any]:
        caption = formal_complaint.get("caption", {}) if isinstance(formal_complaint.get("caption"), dict) else {}
        claims_for_relief = []
        for claim in _coerce_list(formal_complaint.get("legal_claims")):
            if not isinstance(claim, dict):
                continue
            normalized_missing_elements = _unique_preserving_order(
                _extract_text_candidates(claim.get("missing_requirements"))
            )
            claims_for_relief.append(
                {
                    "claim_type": claim.get("claim_type") or claim.get("claim_name") or claim.get("title") or "Claim",
                    "count_title": claim.get("claim_name") or claim.get("title") or "Claim",
                    "legal_standards": _unique_preserving_order(
                        [claim.get("legal_standard", "")]
                        + [
                            f"{item.get('citation')} - {item.get('element')}"
                            if item.get("citation")
                            else str(item.get("element") or "")
                            for item in _coerce_list(claim.get("legal_standard_elements"))
                            if isinstance(item, dict) and (item.get("element") or item.get("citation"))
                        ]
                    ),
                    "supporting_facts": _unique_preserving_order(_extract_text_candidates(claim.get("supporting_facts"))),
                    "missing_elements": normalized_missing_elements,
                    "partially_supported_elements": [],
                    "support_summary": {
                        "elements_satisfied": claim.get("elements_satisfied", ""),
                        "authority_count": len(_coerce_list(claim.get("supporting_authorities"))),
                        "temporal_gap_hint_count": sum(
                            1
                            for item in normalized_missing_elements
                            if str(item or "").strip().lower().startswith("chronology gap")
                        ),
                    },
                    "supporting_exhibits": [
                        {
                            "label": exhibit.get("label"),
                            "title": exhibit.get("title"),
                            "link": exhibit.get("reference") or exhibit.get("source_url") or exhibit.get("link"),
                        }
                        for exhibit in _coerce_list(claim.get("supporting_exhibits"))
                        if isinstance(exhibit, dict)
                    ],
                }
            )

        legal_standards = []
        for standard in _coerce_list(formal_complaint.get("legal_standards")):
            if isinstance(standard, dict):
                claim_name = str(standard.get("claim_name") or standard.get("claim_type") or "").strip()
                body = str(standard.get("standard") or "").strip()
                citations = ", ".join(_unique_preserving_order(_extract_text_candidates(standard.get("citations"))))
                if claim_name and body and citations:
                    legal_standards.append(f"{claim_name}: {body} ({citations})")
                elif claim_name and body:
                    legal_standards.append(f"{claim_name}: {body}")
                elif body:
                    legal_standards.append(body)
            else:
                text = str(standard or "").strip()
                if text:
                    legal_standards.append(text)

        exhibits = []
        for exhibit in _coerce_list(formal_complaint.get("exhibits")):
            if not isinstance(exhibit, dict):
                continue
            exhibits.append(
                {
                    "label": exhibit.get("label"),
                    "title": exhibit.get("title") or exhibit.get("description") or "Supporting exhibit",
                    "claim_type": exhibit.get("claim_type"),
                    "kind": exhibit.get("kind") or "evidence",
                    "link": exhibit.get("reference") or exhibit.get("source_url") or exhibit.get("link") or "",
                    "source_ref": exhibit.get("cid") or exhibit.get("reference") or "",
                    "summary": exhibit.get("summary") or exhibit.get("description") or "",
                }
            )

        nature_of_action = formal_complaint.get("nature_of_action")
        if isinstance(nature_of_action, str):
            nature_of_action = [nature_of_action]

        factual_allegations = _unique_preserving_order(
            _extract_text_candidates(formal_complaint.get("factual_allegations") or formal_complaint.get("summary_of_facts"))
        )
        if not factual_allegations:
            factual_allegations = self._build_factual_allegations(
                summary_of_facts=_extract_text_candidates(formal_complaint.get("summary_of_facts")),
                claims_for_relief=claims_for_relief,
            )

        draft = {
            "court_header": formal_complaint.get("court_header", ""),
            "case_caption": {
                "plaintiffs": _coerce_list(formal_complaint.get("parties", {}).get("plaintiffs", [])) if isinstance(formal_complaint.get("parties"), dict) else [],
                "defendants": _coerce_list(formal_complaint.get("parties", {}).get("defendants", [])) if isinstance(formal_complaint.get("parties"), dict) else [],
                "case_number": caption.get("case_number") or formal_complaint.get("case_number") or "________________",
                "county": caption.get("county_line") or ((formal_complaint.get("caption") or {}).get("county_line") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "lead_case_number": caption.get("lead_case_number") or ((formal_complaint.get("caption") or {}).get("lead_case_number") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "related_case_number": caption.get("related_case_number") or ((formal_complaint.get("caption") or {}).get("related_case_number") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "assigned_judge": caption.get("assigned_judge") or ((formal_complaint.get("caption") or {}).get("assigned_judge") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "courtroom": caption.get("courtroom") or ((formal_complaint.get("caption") or {}).get("courtroom") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "jury_demand_notice": caption.get("jury_demand_notice") or ((formal_complaint.get("caption") or {}).get("jury_demand_notice") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "document_title": "COMPLAINT",
            },
            "title": formal_complaint.get("title") or caption.get("case_title") or "Complaint",
            "nature_of_action": _unique_preserving_order(_extract_text_candidates(nature_of_action)),
            "parties": formal_complaint.get("parties", {}),
            "jurisdiction_statement": formal_complaint.get("jurisdiction_statement", ""),
            "venue_statement": formal_complaint.get("venue_statement", ""),
            "factual_allegations": factual_allegations,
            "summary_of_facts": _unique_preserving_order(_extract_text_candidates(formal_complaint.get("summary_of_facts") or formal_complaint.get("factual_allegations"))),
            "anchored_chronology_summary": _unique_preserving_order(
                _extract_text_candidates(formal_complaint.get("anchored_chronology_summary"))
            ),
            "claims_for_relief": claims_for_relief,
            "legal_standards": _unique_preserving_order(legal_standards),
            "requested_relief": _unique_preserving_order(_extract_text_candidates(formal_complaint.get("requested_relief") or formal_complaint.get("prayer_for_relief"))),
            "jury_demand": formal_complaint.get("jury_demand", {}),
            "exhibits": exhibits,
            "signature_block": formal_complaint.get("signature_block", {}),
            "verification": formal_complaint.get("verification", {}),
            "certificate_of_service": formal_complaint.get("certificate_of_service", {}),
            "source_context": {
                "generated_at": formal_complaint.get("generated_at") or _utcnow().isoformat(),
                "district": formal_complaint.get("district") or caption.get("district") or "",
                "jurisdiction": formal_complaint.get("jurisdiction", "unknown"),
            },
        }
        self._annotate_claim_temporal_gap_hints(draft)
        self._attach_allegation_references(draft)
        self._annotate_case_caption_display(draft)
        built_affidavit = self._build_affidavit(draft)
        existing_affidavit = formal_complaint.get("affidavit", {}) if isinstance(formal_complaint.get("affidavit"), dict) else {}
        draft["affidavit"] = {**built_affidavit, **existing_affidavit}
        rendered_draft_text = self._render_draft_text(draft)
        supplied_draft_text = str(formal_complaint.get("draft_text") or "").strip()
        expected_case_line = (
            f"{draft['case_caption'].get('case_number_label', 'Civil Action No.')} "
            f"{draft['case_caption'].get('case_number', '________________')}"
        )
        draft["draft_text"] = (
            supplied_draft_text
            if supplied_draft_text and expected_case_line in supplied_draft_text
            else rendered_draft_text
        )
        return draft

    def _format_county_for_header(self, county: Optional[str]) -> str:
        county_text = str(county or "").strip().upper()
        if not county_text:
            return ""
        if county_text.startswith("COUNTY OF "):
            return county_text
        if county_text.endswith(" COUNTY"):
            return f"COUNTY OF {county_text[:-7].strip()}"
        return f"COUNTY OF {county_text}"

    def _annotate_case_caption_display(self, draft: Dict[str, Any]) -> None:
        caption = draft.get("case_caption")
        if not isinstance(caption, dict):
            return
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        jurisdiction = str(source_context.get("jurisdiction") or "").strip()
        forum_type = self._infer_forum_type(
            classification={"jurisdiction": jurisdiction},
            court_name=str(draft.get("court_header") or ""),
        )
        caption["forum_type"] = forum_type
        caption["case_number_label"] = caption.get("case_number_label") or (
            "Case No." if forum_type == "state" else "Civil Action No."
        )
        caption["lead_case_number_label"] = caption.get("lead_case_number_label") or (
            "Related Proceeding No." if forum_type == "state" else "Lead Case No."
        )
        caption["related_case_number_label"] = caption.get("related_case_number_label") or (
            "Coordination No." if forum_type == "state" else "Related Case No."
        )
        caption["assigned_judge_label"] = caption.get("assigned_judge_label") or (
            "Judicial Officer" if forum_type == "state" else "Assigned Judge"
        )
        caption["courtroom_label"] = caption.get("courtroom_label") or (
            "Department" if forum_type == "state" else "Courtroom"
        )
        plaintiff_names = caption.get("plaintiffs") if isinstance(caption.get("plaintiffs"), list) else []
        defendant_names = caption.get("defendants") if isinstance(caption.get("defendants"), list) else []
        caption["plaintiff_caption_label"] = caption.get("plaintiff_caption_label") or (
            "Plaintiff" if len(plaintiff_names) == 1 else "Plaintiffs"
        )
        caption["defendant_caption_label"] = caption.get("defendant_caption_label") or (
            "Defendant" if len(defendant_names) == 1 else "Defendants"
        )
        caption["caption_party_lines"] = caption.get("caption_party_lines") or self._build_caption_party_lines(caption)

    def _build_caption_party_lines(self, caption: Dict[str, Any]) -> List[str]:
        plaintiffs = caption.get("plaintiffs") if isinstance(caption.get("plaintiffs"), list) else []
        defendants = caption.get("defendants") if isinstance(caption.get("defendants"), list) else []
        plaintiff_names = [str(name).strip() for name in plaintiffs if str(name).strip()] or ["Plaintiff"]
        defendant_names = [str(name).strip() for name in defendants if str(name).strip()] or ["Defendant"]
        plaintiff_label = str(
            caption.get("plaintiff_caption_label")
            or ("Plaintiff" if len(plaintiff_names) == 1 else "Plaintiffs")
        ).strip()
        defendant_label = str(
            caption.get("defendant_caption_label")
            or ("Defendant" if len(defendant_names) == 1 else "Defendants")
        ).strip()
        return [
            f"{'\n'.join(plaintiff_names)}, {plaintiff_label},",
            "v.",
            f"{'\n'.join(defendant_names)}, {defendant_label}.",
        ]

    def _resolve_draft_forum_type(self, draft: Dict[str, Any]) -> str:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        forum_type = str(caption.get("forum_type") or "").strip().lower()
        if forum_type:
            return forum_type
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        return self._infer_forum_type(
            classification={"jurisdiction": source_context.get("jurisdiction")},
            court_name=str(draft.get("court_header") or ""),
        )

    def _build_party_section_lines(
        self,
        *,
        plaintiffs: List[str],
        defendants: List[str],
        forum_type: str,
    ) -> List[str]:
        plaintiff_names = [str(name).strip() for name in _coerce_list(plaintiffs) if str(name).strip()] or ["Plaintiff"]
        defendant_names = [str(name).strip() for name in _coerce_list(defendants) if str(name).strip()] or ["Defendant"]
        plaintiff_label = "Plaintiff" if len(plaintiff_names) == 1 else "Plaintiffs"
        defendant_label = "Defendant" if len(defendant_names) == 1 else "Defendants"
        plaintiff_names_text = ", ".join(plaintiff_names)
        defendant_names_text = ", ".join(defendant_names)
        if forum_type == "state":
            plaintiff_verb = "is" if len(plaintiff_names) == 1 else "are"
            defendant_verb = "is" if len(defendant_names) == 1 else "are"
            return [
                f"{plaintiff_label} {plaintiff_names_text} {plaintiff_verb} a party bringing this civil action in this Court.",
                f"{defendant_label} {defendant_names_text} {defendant_verb} named as the party from whom relief is sought.",
            ]
        return [
            f"{plaintiff_label}: {plaintiff_names_text}.",
            f"{defendant_label}: {defendant_names_text}.",
        ]

    def _build_jurisdiction_statement(
        self,
        *,
        classification: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        court_name: str,
    ) -> str:
        forum_type = self._infer_forum_type(classification=classification, court_name=court_name)
        first_citation = next(
            (
                str(statute.get("citation") or "").strip()
                for statute in statutes
                if isinstance(statute, dict) and statute.get("citation")
            ),
            "",
        )
        if forum_type == "federal":
            if first_citation:
                return (
                    "This Court has subject-matter jurisdiction under federal law, including "
                    f"{first_citation}, because Plaintiff alleges violations arising under the laws of the United States."
                )
            return "This Court has subject-matter jurisdiction under 28 U.S.C. § 1331 because Plaintiff alleges claims arising under federal law."
        if forum_type == "state":
            if first_citation:
                return (
                    "This Court has subject-matter jurisdiction because Plaintiff asserts claims arising under "
                    f"the governing state law, including {first_citation}, and seeks relief within this Court's authority."
                )
            return (
                "This Court has subject-matter jurisdiction because Plaintiff asserts claims arising under the "
                "governing state law and seeks relief within this Court's authority."
            )
        return "This Court has subject-matter jurisdiction because the claims arise under the governing law identified in this pleading."

    def _build_venue_statement(
        self,
        *,
        district: str,
        county: Optional[str],
        division: Optional[str],
        classification: Dict[str, Any],
        court_name: str,
    ) -> str:
        district_text = str(district or "").strip()
        county_text = str(county or "").strip()
        division_text = str(division or "").strip()
        forum_type = self._infer_forum_type(classification=classification, court_name=court_name)
        if forum_type == "state" and county_text:
            return (
                "Venue is proper in this Court because a substantial part of the events or omissions giving rise "
                f"to these claims occurred in {county_text}."
            )
        if forum_type == "federal" and district_text and division_text:
            return (
                f"Venue is proper in the {division_text} Division of the {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
            )
        if forum_type == "federal" and district_text:
            return (
                f"Venue is proper in the {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
            )
        if forum_type == "state" and district_text and division_text:
            return (
                "Venue is proper in this Court because a substantial part of the events or omissions giving rise "
                f"to these claims occurred in {division_text}, {district_text}."
            )
        if forum_type == "state" and district_text:
            return (
                "Venue is proper in this Court because a substantial part of the events or omissions giving rise "
                f"to these claims occurred in {district_text}."
            )
        return "Venue is proper in this Court because a substantial part of the events or omissions giving rise to these claims occurred in this judicial district."

    def _render_draft_text(self, draft: Dict[str, Any]) -> str:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        plaintiff_list = parties.get("plaintiffs", []) or caption.get("plaintiffs", []) or ["Plaintiff"]
        defendant_list = parties.get("defendants", []) or caption.get("defendants", []) or ["Defendant"]
        plaintiffs = ", ".join(plaintiff_list)
        defendants = ", ".join(defendant_list)
        forum_type = self._resolve_draft_forum_type(draft)
        caption_party_lines = caption.get("caption_party_lines") if isinstance(caption.get("caption_party_lines"), list) else self._build_caption_party_lines(caption)
        party_section_lines = self._build_party_section_lines(
            plaintiffs=plaintiff_list,
            defendants=defendant_list,
            forum_type=forum_type,
        )
        case_number_label = str(caption.get("case_number_label") or "Civil Action No.")
        lead_case_number_label = str(caption.get("lead_case_number_label") or "Lead Case No.")
        related_case_number_label = str(caption.get("related_case_number_label") or "Related Case No.")
        assigned_judge_label = str(caption.get("assigned_judge_label") or "Assigned Judge")
        courtroom_label = str(caption.get("courtroom_label") or "Courtroom")
        lines = [
            str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"),
            *([str(caption.get("county"))] if caption.get("county") else []),
            "",
            *caption_party_lines,
            f"{case_number_label} {caption.get('case_number', '________________')}",
            *([f"{lead_case_number_label} {caption.get('lead_case_number')}"] if caption.get('lead_case_number') else []),
            *([f"{related_case_number_label} {caption.get('related_case_number')}"] if caption.get('related_case_number') else []),
            *([f"{assigned_judge_label}: {caption.get('assigned_judge')}"] if caption.get('assigned_judge') else []),
            *([f"{courtroom_label}: {caption.get('courtroom')}"] if caption.get('courtroom') else []),
            "",
            str(caption.get("document_title") or "COMPLAINT"),
            *([str(caption.get("jury_demand_notice"))] if caption.get("jury_demand_notice") else []),
            "",
            "NATURE OF THE ACTION",
        ]
        lines.extend(self._normalize_text_lines(draft.get("nature_of_action", [])))
        lines.extend([
            "",
            "PARTIES",
            *party_section_lines,
            "",
            "JURISDICTION AND VENUE",
        ])
        if draft.get("jurisdiction_statement"):
            lines.append(str(draft["jurisdiction_statement"]))
        if draft.get("venue_statement"):
            lines.append(str(draft["venue_statement"]))
        lines.extend(["", "FACTUAL ALLEGATIONS"])
        lines.extend(self._grouped_allegation_text_lines(draft))
        chronology_lines = self._normalize_text_lines(draft.get("anchored_chronology_summary", []))
        if chronology_lines:
            lines.extend(["", "ANCHORED CHRONOLOGY"])
            lines.extend(self._numbered_lines(chronology_lines))
        authority_summary_lines = self._normalize_text_lines(draft.get("email_authority_summary_lines", []))
        if authority_summary_lines:
            lines.extend(["", "EMAIL-ALIGNED AUTHORITY SUPPORT"])
            lines.extend(self._bulletize_lines(authority_summary_lines))
        claims = draft.get("claims_for_relief", []) if isinstance(draft.get("claims_for_relief"), list) else []
        if claims:
            lines.extend(["", "CLAIMS FOR RELIEF"])
        for index, claim in enumerate(claims, start=1):
            lines.extend([
                "",
                f"COUNT {_roman(index)} - {claim.get('count_title', claim.get('claim_type', 'Claim'))}",
            ])
            lines.extend(self._build_claim_render_lines(claim))
            missing = self._normalize_text_lines(claim.get("missing_elements", []))
            if missing:
                lines.append("Open Support Gaps:")
                lines.extend([f"- {line}" for line in missing])
        lines.extend(["", "REQUESTED RELIEF"])
        if forum_type == "state":
            lines.append("Wherefore, Plaintiff prays for judgment against Defendant as follows:")
        lines.extend(self._numbered_lines(draft.get("requested_relief", [])))
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            lines.extend(["", str(jury_demand.get("title") or "JURY DEMAND").upper()])
            if jury_demand.get("text"):
                lines.append(str(jury_demand.get("text")))
        exhibits = draft.get("exhibits", []) if isinstance(draft.get("exhibits"), list) else []
        if exhibits:
            lines.extend(["", "EXHIBITS"])
            for exhibit in exhibits:
                if not isinstance(exhibit, dict):
                    continue
                text = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
                if exhibit.get("link"):
                    text = f"{text} ({exhibit['link']})"
                lines.append(text)
                if exhibit.get("summary"):
                    lines.append(f"  {exhibit['summary']}")
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            lines.extend([
                "",
                str(verification.get("title") or "Verification").upper(),
                str(verification.get("text") or ""),
                str(verification.get("dated") or ""),
                str(verification.get("signature_line") or ""),
            ])
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            lines.extend([
                "",
                str(certificate_of_service.get("title") or "Certificate of Service").upper(),
                str(certificate_of_service.get("text") or ""),
                str(certificate_of_service.get("dated") or ""),
                str(certificate_of_service.get("signature_line") or ""),
            ])
        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            lines.extend([
                "",
                str(affidavit.get("title") or "AFFIDAVIT IN SUPPORT OF COMPLAINT"),
            ])
            lines.extend(str(line) for line in _coerce_list(affidavit.get("venue_lines")) if str(line or "").strip())
            lines.extend([
                "",
                str(affidavit.get("intro") or ""),
                str(affidavit.get("knowledge_graph_note") or ""),
                "",
                "Affiant states as follows:",
            ])
            lines.extend(self._numbered_lines(affidavit.get("facts", [])))
            supporting_exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
            if supporting_exhibits:
                lines.extend(["", "AFFIDAVIT SUPPORTING EXHIBITS"])
                for exhibit in supporting_exhibits:
                    if not isinstance(exhibit, dict):
                        continue
                    exhibit_text = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
                    if exhibit.get("link"):
                        exhibit_text = f"{exhibit_text} ({exhibit['link']})"
                    lines.append(exhibit_text)
            lines.extend([
                "",
                str(affidavit.get("dated") or ""),
                str(affidavit.get("signature_line") or ""),
                str(affidavit.get("jurat") or ""),
            ])
            lines.extend(str(line) for line in _coerce_list(affidavit.get("notary_block")) if str(line or "").strip())
        lines.extend(["", *self._build_signature_section_lines(signature_block, forum_type)])
        return "\n".join(line for line in lines if line is not None)

    def _build_claim_render_lines(self, claim: Dict[str, Any]) -> List[str]:
        if not isinstance(claim, dict):
            return []
        rendered: List[str] = []
        incorporated_clause = self._format_incorporated_reference_clause(
            claim.get("allegation_references", []),
            claim.get("supporting_exhibits", []),
        )
        if incorporated_clause:
            rendered.append(incorporated_clause)

        legal_standards = self._normalize_text_lines(claim.get("legal_standards", []))
        supporting_facts = self._render_claim_supporting_facts(claim)
        claim_type = normalize_claim_type(claim.get("claim_type") or "")

        standard_lead_in = "Plaintiff alleges that"
        if claim_type == "due_process_failure":
            standard_lead_in = "Before enforcing a final adverse housing decision, Defendant was required to"
        if claim_type == "housing_discrimination":
            standard_paragraph = self._compose_housing_standard_paragraph(legal_standards[:3])
        else:
            standard_paragraph = self._compose_count_paragraph(
                legal_standards[:3],
                lead_in=standard_lead_in,
                mode="sentences",
            )
        if standard_paragraph:
            rendered.append(standard_paragraph)

        authority_paragraph = self._compose_authority_paragraph(
            legal_standards[3:],
            claim_type=claim_type,
        )
        if authority_paragraph:
            rendered.append(authority_paragraph)

        support_paragraph = self._compose_count_paragraph(
            supporting_facts,
            lead_in="The pleaded facts further show that",
            mode="sentences",
        )
        if support_paragraph:
            rendered.append(support_paragraph)
        return rendered

    def _render_claim_supporting_facts(self, claim: Dict[str, Any]) -> List[str]:
        supporting_facts = self._normalize_text_lines(claim.get("supporting_facts", []))
        supporting_entries = self._align_entries_to_lines(
            claim.get("supporting_fact_entries"),
            supporting_facts,
        )
        if not supporting_entries:
            return supporting_facts

        prioritized: List[str] = []
        seen = set()
        for entry in supporting_entries:
            if not isinstance(entry, dict):
                continue
            text = str(entry.get("text") or "").strip()
            if not text:
                continue
            lowered = text.lower()
            source_kind = str(entry.get("source_kind") or "").strip().lower()
            has_fact_ids = bool(_normalize_identifier_list(entry.get("fact_ids") or []))
            if not has_fact_ids and supporting_entries:
                continue
            if source_kind in {"uploaded_evidence_fact", "claim_chronology_support"} or _contains_date_anchor(text):
                if lowered not in seen:
                    prioritized.append(text)
                    seen.add(lowered)
        for entry in supporting_entries:
            if not isinstance(entry, dict):
                continue
            text = str(entry.get("text") or "").strip()
            if not text:
                continue
            lowered = text.lower()
            has_fact_ids = bool(_normalize_identifier_list(entry.get("fact_ids") or []))
            if lowered in seen or not has_fact_ids:
                continue
            if "required before enforcement of that adverse action" in lowered or "before a final adverse housing decision was enforced" in lowered:
                prioritized.append(text)
                seen.add(lowered)
        for entry in supporting_entries:
            if not isinstance(entry, dict):
                continue
            text = str(entry.get("text") or "").strip()
            if not text:
                continue
            lowered = text.lower()
            if lowered in seen:
                continue
            prioritized.append(text)
            seen.add(lowered)
        return prioritized or supporting_facts

    def _compose_authority_paragraph(self, lines: List[str], *, claim_type: str) -> str:
        normalized = [str(line or "").strip().rstrip(".") for line in lines if str(line or "").strip()]
        if not normalized:
            return ""
        text = " ".join(normalized)
        if claim_type == "due_process_failure":
            citation_match = re.search(r"(24\s+C\.F\.R\.\s*982\.555)", text)
            citation = citation_match.group(1) if citation_match else "24 C.F.R. 982.555"
            return (
                f"Federal housing regulations, including {citation}, required written notice and an opportunity "
                "for informal review before a final adverse housing decision was enforced."
            )
        if len(normalized) == 1 and normalized[0].lower().startswith("authority support for this count includes"):
            sentence = normalized[0].strip()
            return sentence if sentence.endswith((".", "?", "!")) else f"{sentence}."
        return self._compose_count_paragraph(normalized, lead_in="Authority for this count includes")

    def _compose_housing_standard_paragraph(self, lines: List[str]) -> str:
        normalized = [str(line or "").strip().rstrip(".") for line in lines if str(line or "").strip()]
        if not normalized:
            return ""
        cleaned = [
            self._normalize_count_paragraph_fragment(
                line,
                index=index,
                lead_in="Plaintiff alleges that",
            )
            for index, line in enumerate(normalized)
        ]
        cleaned = [line for line in cleaned if line]
        if not cleaned:
            return ""
        sentences: List[str] = []
        first = self._normalize_housing_standard_fragment(cleaned[0], first=True)
        if first:
            sentences.append(f"{first}.")
        for fragment in cleaned[1:]:
            promoted = self._normalize_housing_standard_fragment(fragment, first=False)
            if promoted:
                sentences.append(f"{promoted}.")
        return " ".join(sentences)

    def _normalize_housing_standard_fragment(self, fragment: str, *, first: bool) -> str:
        text = str(fragment or "").strip()
        if not text:
            return ""
        lowered = text.lower()
        if lowered.startswith("defendant denied, limited, or otherwise interfered with housing assistance"):
            return (
                "Defendant denied, limited, or otherwise interfered with Plaintiff's housing assistance "
                "or related housing rights"
            )
        if lowered.startswith("the challenged housing action was unlawful because"):
            return "That housing action was unlawful because" + text[len("the challenged housing action was unlawful because"):]
        if lowered.startswith("the requested relief addresses the resulting denial of housing opportunity"):
            return (
                "That unlawful housing decision caused the resulting denial of housing opportunity"
                + text[len("the requested relief addresses the resulting denial of housing opportunity"):]
            )
        if first:
            return self._promote_count_sentence_fragment(text)
        return self._promote_count_sentence_fragment(text)

    def _compose_count_paragraph(self, lines: List[str], *, lead_in: str, mode: str = "joined") -> str:
        normalized = [str(line or "").strip().rstrip(".") for line in lines if str(line or "").strip()]
        if not normalized:
            return ""
        cleaned: List[str] = []
        for index, line in enumerate(normalized):
            line = self._normalize_count_paragraph_fragment(
                line,
                index=index,
                lead_in=lead_in,
            )
            cleaned.append(line)
        if mode == "sentences":
            return self._compose_count_sentence_paragraph(cleaned, lead_in=lead_in)
        joined = _join_chronology_segments(cleaned)
        return f"{lead_in} {joined}." if joined else ""

    def _compose_count_sentence_paragraph(self, lines: List[str], *, lead_in: str) -> str:
        fragments = [self._trim_count_sentence_connector(line) for line in lines if str(line or "").strip()]
        fragments = [fragment for fragment in fragments if fragment]
        if not fragments:
            return ""
        sentences = [f"{lead_in} {fragments[0]}."]
        for fragment in fragments[1:]:
            promoted = self._promote_count_sentence_fragment(fragment)
            if promoted:
                sentences.append(f"{promoted}.")
        return " ".join(sentences)

    def _trim_count_sentence_connector(self, fragment: str) -> str:
        text = str(fragment or "").strip()
        return re.sub(r"^(?:and|but)\s+", "", text, flags=re.IGNORECASE)

    def _promote_count_sentence_fragment(self, fragment: str) -> str:
        text = str(fragment or "").strip()
        if not text:
            return ""
        lowered = text.lower()
        if lowered.startswith("the requested relief addresses the deprivation of housing benefits and review rights"):
            return (
                "That failure of notice and process caused the deprivation of housing benefits and review rights"
                + text[len("the requested relief addresses the deprivation of housing benefits and review rights"):]
            )
        promotions = (
            ("plaintiff ", "Plaintiff "),
            ("defendant ", "Defendant "),
            ("the ", "The "),
            ("on ", "On "),
        )
        for prefix, replacement in promotions:
            if lowered.startswith(prefix):
                return replacement + text[len(prefix):]
        return text[0].upper() + text[1:] if text[:1].islower() else text

    def _normalize_count_paragraph_fragment(
        self,
        line: str,
        *,
        index: int,
        lead_in: str,
    ) -> str:
        fragment = str(line or "").strip().rstrip(".")
        if not fragment:
            return ""
        lowered_fragment = fragment.lower()
        lowered_lead_in = lead_in.lower()
        if index == 0 and lowered_fragment.startswith(lowered_lead_in + " "):
            fragment = fragment[len(lead_in):].strip()
            lowered_fragment = fragment.lower()
        if (
            index == 0
            and lead_in.endswith(",")
            and lowered_fragment.startswith("before enforcing a final adverse housing decision,")
        ):
            prefix = "Before enforcing a final adverse housing decision,"
            fragment = fragment[len(prefix):].strip() if fragment.startswith(prefix) else fragment
            lowered_fragment = fragment.lower()

        prefix_replacements = (
            ("plaintiff alleges that ", ""),
            ("plaintiff further alleges that ", ""),
            ("plaintiff seeks relief for ", "the requested relief addresses "),
        )
        for prefix, replacement in prefix_replacements:
            if lowered_fragment.startswith(prefix):
                fragment = replacement + fragment[len(prefix):]
                lowered_fragment = fragment.lower()
                break

        if index > 0:
            if fragment.startswith("Plaintiff "):
                fragment = "plaintiff " + fragment[len("Plaintiff "):]
            elif fragment.startswith("Defendant "):
                fragment = "defendant " + fragment[len("Defendant "):]
            elif fragment.startswith("HACC's "):
                fragment = "HACC's " + fragment[len("HACC's "):]
            elif fragment.startswith("HACC "):
                fragment = "HACC " + fragment[len("HACC "):]
            elif fragment.startswith("24 C.F.R."):
                fragment = fragment
            elif fragment.startswith("On "):
                fragment = "on " + fragment[len("On "):]
            elif fragment and fragment[0].isupper():
                fragment = fragment[0].lower() + fragment[1:]
        elif index == 0 and fragment.startswith("On ") and (
            lead_in.startswith("The chronology and governing policy further show that")
            or lead_in.startswith("The record further shows that")
            or lead_in.startswith("These facts further show that")
            or lead_in.startswith("The pleaded facts further show that")
        ):
            fragment = "on " + fragment[len("On "):]

        if ": " in fragment and lead_in.startswith("Authority for this count includes"):
            head, tail = fragment.split(": ", 1)
            if head and tail:
                fragment = f"{head} {tail[0].lower() + tail[1:]}" if tail[:1].isupper() else f"{head} {tail}"

        return fragment

    def _normalize_text_lines(self, values: Any) -> List[str]:
        normalized = []
        for value in _unique_preserving_order(_extract_text_candidates(values)):
            text = re.sub(r"\s+", " ", value).strip()
            if text:
                normalized.append(text)
        return normalized

    def _apply_document_drafting_focus(
        self,
        draft: Dict[str, Any],
        *,
        document_drafting_next_action: Dict[str, Any],
        document_grounding_improvement_next_action: Dict[str, Any],
    ) -> Dict[str, Any]:
        action = (
            document_drafting_next_action
            if isinstance(document_drafting_next_action, dict)
            else {}
        )
        action_code = str(action.get("action") or "").strip().lower()
        focus_source = "document_drafting_next_action"
        if action_code != "realign_document_drafting":
            grounding_action = (
                document_grounding_improvement_next_action
                if isinstance(document_grounding_improvement_next_action, dict)
                else {}
            )
            grounding_action_code = str(grounding_action.get("action") or "").strip().lower()
            if grounding_action_code == "retarget_document_grounding":
                action = dict(grounding_action)
                suggested_claim_element_id = str(action.get("suggested_claim_element_id") or "").strip()
                if suggested_claim_element_id:
                    action["claim_element_id"] = suggested_claim_element_id
                action_code = grounding_action_code
                focus_source = "document_grounding_improvement_next_action"
            else:
                return draft
        if action_code not in {"realign_document_drafting", "retarget_document_grounding"}:
            return draft

        focused_draft = deepcopy(draft)
        focus_section = str(action.get("focus_section") or "").strip()
        claim_element_id = str(action.get("claim_element_id") or "").strip()
        preferred_support_kind = str(action.get("preferred_support_kind") or "").strip()
        executed_claim_element_id = str(action.get("executed_claim_element_id") or "").strip()
        original_claim_element_id = str(
            action.get("original_claim_element_id")
            or document_grounding_improvement_next_action.get("claim_element_id")
            or ""
        ).strip()
        focus_metadata = self._build_document_focus_metadata(
            action=action,
            focus_source=focus_source,
            claim_element_id=claim_element_id,
            preferred_support_kind=preferred_support_kind,
            original_claim_element_id=original_claim_element_id,
        )

        if focus_section in {"factual_allegations", "summary_of_facts"}:
            focus_claim_type = self._resolve_document_focus_claim_type(
                focused_draft,
                claim_element_id=claim_element_id,
                preferred_support_kind=preferred_support_kind,
            )
            if isinstance(focused_draft.get("summary_of_facts"), list):
                summary_entries = self._prioritize_document_focus_entries(
                    self._align_entries_to_lines(
                        focused_draft.get("summary_of_fact_entries"),
                        focused_draft.get("summary_of_facts") or [],
                    ),
                    claim_type=focus_claim_type,
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                )
                summary_entries = self._annotate_document_focus_entries(
                    summary_entries,
                    focus_metadata=focus_metadata,
                )
                focused_draft["summary_of_fact_entries"] = summary_entries
                focused_draft["summary_of_facts"] = [entry.get("text") for entry in summary_entries if entry.get("text")]
            if isinstance(focused_draft.get("factual_allegations"), list):
                allegation_entries = self._prioritize_document_focus_entries(
                    self._align_entries_to_lines(
                        focused_draft.get("factual_allegation_entries"),
                        focused_draft.get("factual_allegations") or [],
                    ),
                    claim_type=focus_claim_type,
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                )
                allegation_entries = self._annotate_document_focus_entries(
                    allegation_entries,
                    focus_metadata=focus_metadata,
                )
                focused_draft["factual_allegation_entries"] = allegation_entries
                focused_draft["factual_allegations"] = [entry.get("text") for entry in allegation_entries if entry.get("text")]

        if focus_section == "claims_for_relief" and isinstance(focused_draft.get("claims_for_relief"), list):
            prioritized_claims: List[Dict[str, Any]] = []
            for claim in focused_draft.get("claims_for_relief") or []:
                if not isinstance(claim, dict):
                    continue
                updated_claim = deepcopy(claim)
                supporting_fact_entries = self._prioritize_document_focus_entries(
                    self._align_entries_to_lines(
                        updated_claim.get("supporting_fact_entries"),
                        updated_claim.get("supporting_facts") or [],
                    ),
                    claim_type=str(updated_claim.get("claim_type") or "").strip(),
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                )
                supporting_fact_entries = self._annotate_document_focus_entries(
                    supporting_fact_entries,
                    focus_metadata=focus_metadata,
                )
                updated_claim["supporting_fact_entries"] = supporting_fact_entries
                updated_claim["supporting_facts"] = [entry.get("text") for entry in supporting_fact_entries if entry.get("text")]
                prioritized_claims.append(updated_claim)
            prioritized_claims.sort(
                key=lambda claim: (
                    -self._document_focus_claim_score(
                        claim,
                        claim_element_id=claim_element_id,
                        preferred_support_kind=preferred_support_kind,
                    ),
                    str(claim.get("count_title") or claim.get("claim_type") or ""),
                )
            )
            focused_draft["claims_for_relief"] = prioritized_claims

        focused_draft["document_drafting_focus_section"] = focus_section
        focused_draft["document_drafting_focus_claim_element_id"] = claim_element_id
        focused_draft["document_drafting_focus_source"] = focus_source
        if preferred_support_kind:
            focused_draft["document_drafting_focus_support_kind"] = preferred_support_kind
        if executed_claim_element_id:
            focused_draft["document_drafting_executed_claim_element_id"] = executed_claim_element_id
        return focused_draft

    def _build_document_focus_metadata(
        self,
        *,
        action: Dict[str, Any],
        focus_source: str,
        claim_element_id: str,
        preferred_support_kind: str,
        original_claim_element_id: str,
    ) -> Dict[str, Any]:
        action_code = str(action.get("action") or "").strip().lower()
        metadata = {
            "focus_source": focus_source,
            "action": action_code,
            "target_claim_element_id": claim_element_id,
            "preferred_support_kind": preferred_support_kind,
        }
        if original_claim_element_id and original_claim_element_id != claim_element_id:
            metadata["original_claim_element_id"] = original_claim_element_id
        return metadata

    def _annotate_document_focus_entries(
        self,
        entries: List[Dict[str, Any]],
        *,
        focus_metadata: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        annotated: List[Dict[str, Any]] = []
        normalized_focus_metadata = dict(focus_metadata or {})
        for index, entry in enumerate(entries, start=1):
            if not isinstance(entry, dict):
                continue
            updated_entry = dict(entry)
            if normalized_focus_metadata:
                updated_entry["document_focus"] = dict(normalized_focus_metadata)
                updated_entry["document_focus_priority_rank"] = index
            annotated.append(updated_entry)
        return annotated

    def _document_focus_claim_score(
        self,
        claim: Dict[str, Any],
        *,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> int:
        if not isinstance(claim, dict):
            return 0
        claim_type = str(claim.get("claim_type") or "").strip()
        packet_element = self._claim_support_packet_element(
            claim_type=claim_type,
            claim_element_id=claim_element_id,
        )
        score = 0
        if packet_element:
            score += 12
            support_status = str(packet_element.get("support_status") or "").strip().lower()
            if support_status == "supported":
                score += 4
            elif support_status == "partially_supported":
                score += 3
            elif support_status == "unsupported":
                score += 2
            elif support_status == "contradicted":
                score += 1
        candidate_entries = self._align_entries_to_lines(
            claim.get("supporting_fact_entries"),
            claim.get("supporting_facts") or [],
        )
        candidate_lines = [entry.get("text") for entry in candidate_entries if entry.get("text")] + self._normalize_text_lines(
            claim.get("missing_elements") or []
        )
        return score + max(
            (
                self._document_focus_entry_score(
                    candidate_entries[index],
                    claim_type=claim_type,
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                )
                for index in range(len(candidate_entries))
            ),
            default=0,
        ) + max(
            (
                self._document_focus_line_score(
                    line,
                    claim_type=claim_type,
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                )
                for line in candidate_lines
            ),
            default=0,
        )

    def _resolve_document_focus_claim_type(
        self,
        draft: Dict[str, Any],
        *,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> str:
        claims = draft.get("claims_for_relief")
        if not isinstance(claims, list):
            return ""
        scored_claims = [
            (
                self._document_focus_claim_score(
                    claim,
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                ),
                str(claim.get("claim_type") or "").strip(),
            )
            for claim in claims
            if isinstance(claim, dict) and str(claim.get("claim_type") or "").strip()
        ]
        scored_claims = [item for item in scored_claims if item[1]]
        if not scored_claims:
            return ""
        scored_claims.sort(key=lambda item: (-item[0], item[1]))
        return scored_claims[0][1]

    def _claim_support_packet_element(
        self,
        *,
        claim_type: str,
        claim_element_id: str,
    ) -> Dict[str, Any]:
        normalized_claim_type = str(claim_type or "").strip()
        normalized_element_id = str(claim_element_id or "").strip().lower()
        if not normalized_claim_type or not normalized_element_id:
            return {}
        phase_manager = getattr(self.mediator, "phase_manager", None)
        if phase_manager is None:
            return {}
        packets = phase_manager.get_phase_data(ComplaintPhase.EVIDENCE, "claim_support_packets")
        if not isinstance(packets, dict):
            return {}
        packet = packets.get(normalized_claim_type)
        if not isinstance(packet, dict):
            return {}
        for element in packet.get("elements") or []:
            if not isinstance(element, dict):
                continue
            if str(element.get("element_id") or "").strip().lower() == normalized_element_id:
                return element
        return {}

    def _prioritize_document_focus_lines(
        self,
        lines: List[str],
        *,
        claim_type: str,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> List[str]:
        entries = [{"text": line} for line in self._normalize_text_lines(lines)]
        return [
            str(entry.get("text") or "").strip()
            for entry in self._prioritize_document_focus_entries(
                entries,
                claim_type=claim_type,
                claim_element_id=claim_element_id,
                preferred_support_kind=preferred_support_kind,
            )
            if str(entry.get("text") or "").strip()
        ]

    def _prioritize_document_focus_entries(
        self,
        entries: List[Dict[str, Any]],
        *,
        claim_type: str,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> List[Dict[str, Any]]:
        indexed_entries = list(enumerate([dict(entry) for entry in entries if isinstance(entry, dict)]))
        indexed_entries.sort(
            key=lambda item: (
                -self._document_focus_entry_score(
                    item[1],
                    claim_type=claim_type,
                    claim_element_id=claim_element_id,
                    preferred_support_kind=preferred_support_kind,
                ),
                item[0],
            )
        )
        return [entry for _, entry in indexed_entries]

    def _document_focus_entry_score(
        self,
        entry: Dict[str, Any],
        *,
        claim_type: str,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> int:
        if not isinstance(entry, dict):
            return 0
        score = self._document_focus_line_score(
            str(entry.get("text") or ""),
            claim_type=claim_type,
            claim_element_id=claim_element_id,
            preferred_support_kind=preferred_support_kind,
        )
        packet_element = self._claim_support_packet_element(
            claim_type=claim_type,
            claim_element_id=claim_element_id,
        )
        entry_fact_ids = set(_normalize_identifier_list(entry.get("fact_ids") or []))
        packet_fact_ids = set(
            _normalize_identifier_list(
                list(packet_element.get("canonical_fact_ids") or [])
                + list(packet_element.get("fact_ids") or [])
            )
        )
        if entry_fact_ids and packet_fact_ids and entry_fact_ids & packet_fact_ids:
            score += 30
        entry_element_ids = {
            value.strip().lower()
            for value in _normalize_identifier_list(entry.get("claim_element_ids") or [])
            if value
        }
        if claim_element_id and claim_element_id.strip().lower() in entry_element_ids:
            score += 12
        support_trace_ids = set(_normalize_identifier_list(entry.get("support_trace_ids") or []))
        packet_trace_ids = {
            str(trace.get("source_ref") or "").strip()
            for trace in _coerce_list(packet_element.get("support_traces"))
            if isinstance(trace, dict) and str(trace.get("source_ref") or "").strip()
        }
        if support_trace_ids and packet_trace_ids and support_trace_ids & packet_trace_ids:
            score += 10
        artifact_ids = set(_normalize_identifier_list(entry.get("source_artifact_ids") or []))
        packet_artifact_ids = set(
            _normalize_identifier_list(
                list(packet_element.get("supporting_artifact_ids") or [])
                + list(packet_element.get("supporting_testimony_ids") or [])
            )
        )
        if artifact_ids and packet_artifact_ids and artifact_ids & packet_artifact_ids:
            score += 8
        return score

    def _document_focus_line_score(
        self,
        line: str,
        *,
        claim_type: str,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> int:
        text = str(line or "").strip().lower()
        if not text:
            return 0
        score = 0
        for token in self._document_focus_tokens(
            claim_type=claim_type,
            claim_element_id=claim_element_id,
            preferred_support_kind=preferred_support_kind,
        ):
            if token and token in text:
                score += max(2, len(token.split()))
        return score

    def _document_focus_tokens(
        self,
        *,
        claim_type: str,
        claim_element_id: str,
        preferred_support_kind: str,
    ) -> List[str]:
        normalized_claim_type = normalize_claim_type(claim_type)
        normalized_element = str(claim_element_id or "").strip().lower()
        registry_element = registry_element_for_claim_type(normalized_claim_type, normalized_element)
        packet_element = self._claim_support_packet_element(
            claim_type=claim_type,
            claim_element_id=claim_element_id,
        )
        tokens = [token for token in normalized_element.replace("_", " ").split() if token]
        tokens.extend(
            term
            for term in str(registry_element.get("label") or normalized_element).replace("/", " ").lower().split()
            if term and len(term) > 2
        )
        tokens.extend(
            str(keyword or "").strip().lower()
            for keyword in list(registry_element.get("keywords") or [])
            if str(keyword or "").strip()
        )
        tokens.extend(
            str(role or "").strip().replace("_", " ").lower()
            for role in list(registry_element.get("actor_roles") or [])
            if str(role or "").strip()
        )
        tokens.extend(
            str(kind or "").strip().replace("_", " ").lower()
            for kind in list(registry_element.get("evidence_classes") or [])
            if str(kind or "").strip()
        )
        tokens.extend(
            str(kind or "").strip().replace("_", " ").lower()
            for kind in list(packet_element.get("preferred_evidence_classes") or [])
            if str(kind or "").strip()
        )
        tokens.extend(
            str(bundle or "").strip().replace("_", " ").lower()
            for bundle in list(packet_element.get("required_fact_bundle") or [])
            if str(bundle or "").strip()
        )
        tokens.extend(
            str(bundle or "").strip().replace("_", " ").lower()
            for bundle in list(packet_element.get("missing_fact_bundle") or [])
            if str(bundle or "").strip()
        )
        tokens.extend(
            str(packet_element.get("element_text") or "").strip().lower().split()
            if str(packet_element.get("element_text") or "").strip()
            else []
        )
        alias_map = {
            "protected_activity": [
                "protected activity",
                "reported",
                "report",
                "complained",
                "complaint",
                "grievance",
                "hearing request",
                "review request",
            ],
            "causation": [
                "because",
                "after",
                "in retaliation",
                "retaliatory",
                "in response",
                "shortly after",
                "days after",
                "weeks after",
            ],
            "adverse_action": [
                "adverse action",
                "terminated",
                "termination",
                "denied",
                "denial",
                "suspended",
                "evicted",
                "eviction",
            ],
            "notice": ["notice", "written notice"],
            "hearing": ["hearing", "informal review", "grievance hearing"],
            "response": ["response", "decision", "responded on"],
            "timeline": ["before", "after", "date", "on "],
        }
        tokens.extend(alias_map.get(normalized_element, []))
        support_kind = str(preferred_support_kind or "").strip().lower().replace("_", " ")
        if support_kind:
            tokens.append(support_kind)
        return _unique_preserving_order(token for token in tokens if token)

    def _split_allegation_fragments(self, value: Any) -> List[str]:
        text = re.sub(r"\s+", " ", str(value or "")).strip(" -;")
        if not text:
            return []
        if ": " in text:
            prefix, suffix = text.split(": ", 1)
            prefix_lower = prefix.strip().lower()
            if (
                prefix.strip().endswith("?")
                or prefix_lower.startswith(("what ", "when ", "where ", "why ", "how ", "who ", "describe ", "explain "))
                or prefix_lower in {"what happened", "what relief do you want"}
            ):
                text = suffix.strip()
        parts = [
            part.strip(" -;")
            for part in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
            if part.strip(" -;")
        ]
        return parts or [text]

    def _formalize_allegation_fragment(self, value: Any) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip(" -;")
        if not text:
            return ""
        replacements = (
            (r"^i was\b", "Plaintiff was"),
            (r"^i am\b", "Plaintiff is"),
            (r"^i need\b", "Plaintiff needs"),
            (r"^i needed\b", "Plaintiff needed"),
            (r"^i lost\b", "Plaintiff lost"),
            (r"^i asked\b", "Plaintiff asked"),
            (r"^i reported\b", "Plaintiff reported"),
            (r"^i complained\b", "Plaintiff complained"),
            (r"^i informed\b", "Plaintiff informed"),
            (r"^i notified\b", "Plaintiff notified"),
            (r"^i requested\b", "Plaintiff requested"),
            (r"^i sought\b", "Plaintiff sought"),
            (r"^i experienced\b", "Plaintiff experienced"),
            (r"^i suffered\b", "Plaintiff suffered"),
            (r"^i told\b", "Plaintiff told"),
            (r"^they\b", "Defendant"),
        )
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        clause_replacements = (
            (r"([,;]\s+)i was\b", r"\1Plaintiff was"),
            (r"([,;]\s+)i am\b", r"\1Plaintiff is"),
            (r"([,;]\s+)i need\b", r"\1Plaintiff needs"),
            (r"([,;]\s+)i needed\b", r"\1Plaintiff needed"),
            (r"([,;]\s+)i lost\b", r"\1Plaintiff lost"),
            (r"([,;]\s+)i asked\b", r"\1Plaintiff asked"),
            (r"([,;]\s+)i reported\b", r"\1Plaintiff reported"),
            (r"([,;]\s+)i complained\b", r"\1Plaintiff complained"),
            (r"([,;]\s+)i requested\b", r"\1Plaintiff requested"),
            (r"([,;]\s+)i informed\b", r"\1Plaintiff informed"),
            (r"([,;]\s+)i notified\b", r"\1Plaintiff notified"),
            (r"([,;]\s+)i suffered\b", r"\1Plaintiff suffered"),
            (r"([,;]\s+)i experienced\b", r"\1Plaintiff experienced"),
            (r"([,;]\s+)i told\b", r"\1Plaintiff told"),
            (r"(\band\s+)i was\b", r"\1Plaintiff was"),
            (r"(\band\s+)i am\b", r"\1Plaintiff is"),
            (r"(\band\s+)i need\b", r"\1Plaintiff needs"),
            (r"(\band\s+)i needed\b", r"\1Plaintiff needed"),
            (r"(\band\s+)i lost\b", r"\1Plaintiff lost"),
            (r"(\band\s+)i asked\b", r"\1Plaintiff asked"),
            (r"(\band\s+)i reported\b", r"\1Plaintiff reported"),
            (r"(\band\s+)i complained\b", r"\1Plaintiff complained"),
            (r"(\band\s+)i requested\b", r"\1Plaintiff requested"),
            (r"(\band\s+)i informed\b", r"\1Plaintiff informed"),
            (r"(\band\s+)i notified\b", r"\1Plaintiff notified"),
            (r"(\band\s+)i suffered\b", r"\1Plaintiff suffered"),
            (r"(\band\s+)i experienced\b", r"\1Plaintiff experienced"),
            (r"(\band\s+)i told\b", r"\1Plaintiff told"),
            (r"(\bafter\s+)i was\b", r"\1Plaintiff was"),
            (r"(\bafter\s+)i am\b", r"\1Plaintiff is"),
            (r"(\bafter\s+)i need\b", r"\1Plaintiff needs"),
            (r"(\bafter\s+)i needed\b", r"\1Plaintiff needed"),
            (r"(\bafter\s+)i lost\b", r"\1Plaintiff lost"),
            (r"(\bafter\s+)i asked\b", r"\1Plaintiff asked"),
            (r"(\bafter\s+)i reported\b", r"\1Plaintiff reported"),
            (r"(\bafter\s+)i complained\b", r"\1Plaintiff complained"),
            (r"(\bafter\s+)i requested\b", r"\1Plaintiff requested"),
            (r"(\bafter\s+)i informed\b", r"\1Plaintiff informed"),
            (r"(\bafter\s+)i notified\b", r"\1Plaintiff notified"),
            (r"(\bafter\s+)i suffered\b", r"\1Plaintiff suffered"),
            (r"(\bafter\s+)i experienced\b", r"\1Plaintiff experienced"),
            (r"(\bafter\s+)i told\b", r"\1Plaintiff told"),
            (r"(\bthat\s+)i am\b", r"\1Plaintiff is"),
            (r"(\bthat\s+)i need\b", r"\1Plaintiff needs"),
            (r"(\bthat\s+)i needed\b", r"\1Plaintiff needed"),
            (r"(\bthat\s+)i asked\b", r"\1Plaintiff asked"),
            (r"(\bthat\s+)i complained\b", r"\1Plaintiff complained"),
            (r"(\bthat\s+)i requested\b", r"\1Plaintiff requested"),
            (r"(\bthat\s+)i told\b", r"\1Plaintiff told"),
        )
        for pattern, replacement in clause_replacements:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        text = re.sub(r"\bmy\b", "Plaintiff's", text, flags=re.IGNORECASE)
        text = re.sub(r"\bmine\b", "Plaintiff's", text, flags=re.IGNORECASE)
        text = re.sub(r"\bme\b", "Plaintiff", text, flags=re.IGNORECASE)
        text = re.sub(r"\blost Plaintiff's pay and benefits\b", "lost pay and benefits", text, flags=re.IGNORECASE)
        text = re.sub(r"\blost Plaintiff's (pay|wages|salary|income|benefits)\b", r"lost \1", text, flags=re.IGNORECASE)
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        if len(text) < 12:
            return ""
        return text if text.endswith((".", "?", "!")) else f"{text}."

    def _is_factual_allegation_candidate(self, value: Any) -> bool:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            return False
        lowered = text.lower()
        if re.match(r"^(as to [^,]+, )?plaintiff (seeks|requests|asks|demands)\b", lowered):
            return False
        if lowered.startswith(("requested relief", "relief requested", "element supported:")):
            return False
        if lowered.startswith(("evidence shows facts supporting", "the intake record describes facts supporting")):
            return False
        if "the strongest supporting material is" in lowered:
            return False
        if re.match(r"^(as to [^,]+, )?(title\s+[ivxlcdm0-9]+\b|\d+\s+u\.s\.c\.|\d+\s+c\.f\.r\.|[a-z]{2,6}\.\s+gov\.\s+code\b)", lowered):
            return False
        if len(text) > 360:
            return False
        if not re.search(
            r"\b(was|were|is|are|reported|complained|terminated|fired|retaliated|denied|refused|told|informed|notified|requested|sought|seek|experienced|suffered|lost|made|engaged|opposed|filed|submitted|sent|issued|signed|emailed|wrote|received|occurred|happened|subjected|demoted|suspended|disciplined|reduced|exchanged|used|included|includes|requires|required|must|states|describes|shows|reflects|discusses|provides)\b",
            lowered,
        ):
            return False
        return True

    def _is_generic_claim_support_text(self, value: Any) -> bool:
        lowered = re.sub(r"\s+", " ", str(value or "")).strip().lower()
        return (
            lowered.startswith(("evidence shows facts supporting", "the intake record describes facts supporting"))
            or "the strongest supporting material is" in lowered
            or lowered.startswith("for this question, the strongest supporting material is")
            or lowered == "additional factual development is required before filing."
        )

    def _is_priority_email_timeline_allegation(self, value: Any) -> bool:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text or not _contains_date_anchor(text):
            return False
        lowered = text.lower()
        priority_markers = (
            "hcv orientation",
            "additional information needed",
            "jane kay cortez",
            "jc household",
            "clackamas housing staff",
            "clackamas county housing authority staff",
            "clackamas personnel",
            "cortez-j-ra-denial",
        )
        return any(marker in lowered for marker in priority_markers)

    def _is_low_signal_email_summary_entry(self, entry: Dict[str, Any]) -> bool:
        if not isinstance(entry, dict):
            return False
        source_kind = str(entry.get("source_kind") or "").strip().lower()
        if source_kind not in {"email_timeline_candidate", "email_timeline_fact"}:
            return False
        text = re.sub(r"\s+", " ", str(entry.get("text") or "")).strip()
        if not text:
            return False
        lowered = text.lower()
        if self._is_priority_email_timeline_allegation(text):
            return False
        if "denial" in lowered or "ra-" in lowered:
            return False
        return True

    def _expand_allegation_sources(self, values: Any, *, limit: Optional[int] = None) -> List[str]:
        expanded: List[str] = []
        for value in _extract_text_candidates(values):
            for fragment in self._split_allegation_fragments(value):
                sentence = self._formalize_allegation_fragment(fragment)
                if not sentence or not self._is_factual_allegation_candidate(sentence):
                    continue
                expanded.append(sentence)
        unique = _unique_preserving_order(expanded)
        return unique[:limit] if limit is not None else unique

    def _merge_email_attachment_follow_on_allegations(self, allegations: List[str]) -> List[str]:
        merged: List[str] = []
        for item in allegations:
            text = str(item or "").strip()
            if not text:
                continue
            lowered = text.lower()
            if (
                merged
                and lowered.startswith("one of those clackamas emails included the attachment ")
                and "hcv orientation" in merged[-1].lower()
            ):
                prior = merged[-1].rstrip(".!?")
                attachment_clause = text[0].lower() + text[1:] if text[:1].isupper() else text
                merged[-1] = f"{prior}, and {attachment_clause.rstrip('.!?')}."
                continue
            merged.append(text)
        return merged

    def _extract_uploaded_evidence_text_candidates(self, value: Any, *, limit: int = 3) -> List[str]:
        text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
        if not text.strip():
            return []
        cleaned_source = re.sub(r"\[(.*?)\]\([^)]+\)", r"\1", text)
        cleaned_source = re.sub(r"`{1,3}", "", cleaned_source)
        cleaned_source = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned_source)
        cleaned_source = re.sub(r"__(.*?)__", r"\1", cleaned_source)
        cleaned_source = cleaned_source.replace("•", "\n")
        raw_segments = re.split(r"\n+|(?<=[.!?])\s+", cleaned_source)
        priority_markers = (
            "hacc",
            "hearing",
            "appeal",
            "notice",
            "grievance",
            "review",
            "request",
            "retaliat",
            "adverse",
            "deny",
            "denial",
            "termination",
            "voucher",
            "discrimin",
            "accommodation",
            "tenant",
            "policy",
        )
        ignored_prefixes = (
            "source:",
            "usage:",
            "use for:",
            "solution:",
            "issue:",
            "quick start",
            "integration pattern",
            "file structure",
            "python ",
            "from ",
            "import ",
            "def ",
            "class ",
            "return ",
            "print(",
        )

        scored_segments: List[tuple[int, str]] = []
        seen = set()
        for segment in raw_segments:
            candidate = re.sub(r"\s+", " ", str(segment or "")).strip(" -*#>\t")
            if len(candidate) < 24 or len(candidate) > 360:
                continue
            lowered = candidate.lower()
            if lowered.startswith(ignored_prefixes):
                continue
            if candidate.count("{") + candidate.count("}") >= 2 or candidate.count("=") >= 3:
                continue
            if "```" in candidate or "::" in candidate:
                continue
            marker_hits = sum(1 for marker in priority_markers if marker in lowered)
            if marker_hits <= 0:
                continue
            key = candidate.lower()
            if key in seen:
                continue
            seen.add(key)
            scored_segments.append((marker_hits, candidate))

        scored_segments.sort(key=lambda item: (-item[0], len(item[1])))
        return [candidate for _, candidate in scored_segments[:limit]]

    def _synthesize_narrative_allegations(self, allegations: List[str]) -> List[str]:
        cleaned = [str(item).strip() for item in allegations if str(item).strip()]
        if not cleaned:
            return []

        def _normalize_adverse_clause(clause: str) -> str:
            text = str(clause or "").strip().rstrip(".!?")
            if re.match(r"^(after|following)\b", text, flags=re.IGNORECASE) and "," in text:
                text = text.split(",", 1)[1].strip()
            return text

        def _normalize_harm_clause(clause: str) -> str:
            text = str(clause or "").strip().rstrip(".!?")
            text = re.sub(r",?\s+as a result$", "", text, flags=re.IGNORECASE)
            text = re.sub(r",?\s+as a direct result$", "", text, flags=re.IGNORECASE)
            return text.strip()

        def _pick(pattern: str, *, require_plaintiff: bool = False) -> str:
            for item in cleaned:
                lowered = item.lower()
                if require_plaintiff and "plaintiff" not in lowered:
                    continue
                if re.search(pattern, lowered):
                    return item.rstrip(".!?")
            return ""

        report_clause = _pick(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", require_plaintiff=True)
        adverse_clause = _pick(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b")
        harm_clause = _pick(r"\blost (pay|wages|salary|income|benefits)\b|\b(suffered|experienced)\b", require_plaintiff=True)
        harm_already_tied_to_adverse_action = any(
            re.search(r"\b(lost|suffered|experienced)\b", item.lower())
            and re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", item.lower())
            for item in cleaned
        )

        synthesized: List[str] = []
        if report_clause and adverse_clause:
            synthesized.append(f"After {report_clause}, {_normalize_adverse_clause(adverse_clause)}.")
        if harm_clause and not harm_already_tied_to_adverse_action:
            normalized_harm_clause = _normalize_harm_clause(harm_clause)
            loss_match = re.search(r"\blost ([^.]+)", normalized_harm_clause, flags=re.IGNORECASE)
            if loss_match:
                synthesized.append(f"As a direct result of Defendant's conduct, Plaintiff lost {loss_match.group(1).strip()}." )
        return _unique_preserving_order(synthesized)

    def _prune_subsumed_narrative_clauses(self, allegations: List[str]) -> List[str]:
        cleaned = [str(item).strip() for item in allegations if str(item).strip()]
        if not cleaned:
            return []

        def _pick(pattern: str, *, require_plaintiff: bool = False) -> str:
            for item in cleaned:
                lowered = item.lower()
                if require_plaintiff and "plaintiff" not in lowered:
                    continue
                if re.search(pattern, lowered):
                    return item.strip()
            return ""

        report_clause = _pick(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", require_plaintiff=True)
        adverse_clause = _pick(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b")
        has_harm_tied_to_adverse_action = any(
            re.search(r"\b(lost|suffered|experienced)\b", item.lower())
            and re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", item.lower())
            for item in cleaned
        )
        consumed = {item.lower() for item in (report_clause, adverse_clause) if item}
        if has_harm_tied_to_adverse_action:
            combined_clause = _pick(
                r"\b(reported|complained|opposed|informed|notified|told|requested)\b.*\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b"
                r"|\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b.*\b(reported|complained|opposed|informed|notified|told|requested)\b",
                require_plaintiff=True,
            )
            if combined_clause:
                consumed.add(combined_clause.lower())
        return [item for item in cleaned if item.lower() not in consumed]

    def _prune_near_duplicate_allegations(self, allegations: List[str]) -> List[str]:
        def _tokens(value: str) -> set[str]:
            scrubbed = re.sub(r"\(see exhibit [^)]+\)", "", value, flags=re.IGNORECASE)
            return {
                token
                for token in re.split(r"\W+", scrubbed.lower())
                if len(token) >= 4 and token not in {"plaintiff", "defendant", "exhibit", "after", "those", "this", "that"}
            }

        def _categories(value: str) -> set[str]:
            lowered = value.lower()
            flags = set()
            if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
                flags.add("report")
            if re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied|removed|stripped)\b", lowered) or re.search(r"\b(end(?:ed|ing))\b[^.]{0,40}\bemployment\b", lowered):
                flags.add("adverse")
            if re.search(r"\b(lost|suffered|experienced|benefits|wages|salary|income|opportunities)\b", lowered):
                flags.add("harm")
            if re.search(r"\b(notice|informal review|review protections?|hearing|grievance|appeal|due process)\b", lowered):
                flags.add("policy_process")
            return flags

        def _features(value: str) -> set[str]:
            lowered = value.lower()
            flags = set()
            if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
                flags.add("report")
            if re.search(r"\b(human resources|hr)\b", lowered):
                flags.add("hr")
            if re.search(r"\bregional management|management\b", lowered):
                flags.add("management")
            if re.search(r"\b(key|major)\s+accounts?\b|\b(accounts?)\b[^.]{0,20}\b(removed|stripped|taken away)\b|\b(removed|stripped|took away)\b[^.]{0,20}\baccounts?\b", lowered):
                flags.add("accounts")
            if re.search(r"\bovertime\b", lowered):
                flags.add("overtime")
            if re.search(r"\bshift(s)?\b", lowered):
                flags.add("shifts")
            if re.search(r"\b(absences?|attendance|treatment-related absences?)\b", lowered):
                flags.add("absences")
            if re.search(r"\b(disciplined|discipline|wrote me up|write-up|write up)\b", lowered):
                flags.add("discipline")
            if re.search(r"\b(accommodation|accommodate|light duty|schedule flexibility|medical restrictions?|doctor-imposed restrictions?)\b", lowered):
                flags.add("accommodation")
            if re.search(r"\b(restrictions?|light duty|schedule flexibility)\b", lowered):
                flags.add("restrictions")
            if re.search(r"\b(terminated|fired)\b", lowered) or re.search(r"\b(end(?:ed|ing))\b[^.]{0,40}\bemployment\b", lowered):
                flags.add("termination")
            if re.search(r"\b(wages|pay|salary|income|benefits)\b", lowered):
                flags.add("economic_harm")
            if re.search(r"\b(career opportunities|future opportunities|opportunities)\b", lowered):
                flags.add("opportunities")
            if re.search(r"\bnotice to the applicant\b|\bwritten notice\b", lowered):
                flags.add("notice")
            if re.search(r"\binformal review\b|\breview protections?\b|\bhearing\b|\bgrievance\b|\bappeal\b", lowered):
                flags.add("review_process")
            if re.search(r"\b(denial notice|denied assistance|denying assistance|loss of assistance|terminated assistance|termination)\b", lowered):
                flags.add("denial_action")
            if re.search(r"\bwithout providing\b|\bfailed to provide\b|\bshould have accompanied\b", lowered):
                flags.add("process_violation")
            return flags

        kept: List[str] = []
        for candidate in allegations:
            candidate_tokens = _tokens(candidate)
            candidate_categories = _categories(candidate)
            candidate_features = _features(candidate)
            candidate_is_chronology = candidate.lower().startswith("the chronology shows that ")
            skip = False
            for index, existing in enumerate(kept):
                existing_tokens = _tokens(existing)
                existing_categories = _categories(existing)
                existing_features = _features(existing)
                existing_is_chronology = existing.lower().startswith("the chronology shows that ")
                if not candidate_tokens or not existing_tokens:
                    continue
                if candidate_is_chronology and not existing_is_chronology:
                    continue
                if existing_is_chronology and not candidate_is_chronology:
                    continue
                if (
                    "without the notice and review protections described in the governing process" in candidate.lower()
                    and "without the notice and review protections described in the governing process" not in existing.lower()
                    and "denial notice" in candidate.lower()
                    and "denial notice" in existing.lower()
                    and _contains_date_anchor(candidate)
                    and _contains_date_anchor(existing)
                ):
                    kept[index] = candidate
                    skip = True
                    break
                if not (candidate_categories & existing_categories):
                    continue
                overlap = len(candidate_tokens & existing_tokens) / max(1, min(len(candidate_tokens), len(existing_tokens)))
                shared_features = candidate_features & existing_features
                if overlap >= 0.7:
                    skip = True
                    break
                if "adverse" in candidate_categories and "adverse" in existing_categories and len(shared_features) >= 3:
                    skip = True
                    break
                if "report" in candidate_categories and "report" in existing_categories and "accommodation" in shared_features and len(shared_features) >= 2:
                    skip = True
                    break
                if "policy_process" in candidate_categories and "policy_process" in existing_categories:
                    if (
                        "denial_action" in candidate_features
                        and "denial_action" in existing_features
                        and "process_violation" in candidate_features
                        and "process_violation" not in existing_features
                    ):
                        kept[index] = candidate
                        skip = True
                        break
                    if "denial_action" in candidate_features ^ existing_features:
                        continue
                    if (
                        "process_violation" in candidate_features
                        and "process_violation" not in existing_features
                        and len(shared_features & {"notice", "review_process"}) >= 1
                    ):
                        kept[index] = candidate
                        skip = True
                        break
                    if "process_violation" in shared_features and len(shared_features & {"notice", "review_process", "process_violation"}) >= 2:
                        skip = True
                        break
                    if overlap >= 0.55 and len(shared_features & {"notice", "review_process"}) >= 1:
                        skip = True
                        break
            if not skip:
                kept.append(candidate)
        return kept

    def _is_near_duplicate_allegation(self, candidate: str, existing: List[str]) -> bool:
        if not candidate:
            return False
        pruned = self._prune_near_duplicate_allegations([*existing, candidate])
        return len(pruned) == len(existing)

    def _infer_actor_label_from_allegations(self, allegations: List[str]) -> str:
        combined = " ".join(str(item or "") for item in allegations)
        lowered = combined.lower()
        if "hacc" in lowered:
            return "HACC"
        if any(
            token in lowered
            for token in (
                "housing authority",
                "voucher",
                "denying assistance",
                "notice to the applicant",
                "informal review",
                "grievance",
                "appeal rights",
            )
        ):
            return "the Housing Authority"
        return "Defendant"

    def _actor_possessive_label(self, actor_label: str) -> str:
        text = str(actor_label or "").strip()
        if not text:
            return "Defendant's"
        if text.endswith("s"):
            return f"{text}'"
        return f"{text}'s"

    def _build_concrete_sequence_allegation(self, allegations: List[str]) -> str:
        if not allegations:
            return ""
        dated_events: Dict[str, tuple[str, str]] = {}
        for item in allegations:
            text = str(item or "").strip()
            lowered = text.lower()
            if not text or not _contains_date_anchor(text):
                continue
            date_text, clause = _extract_dated_event_clause(text)
            if not date_text or not clause:
                continue
            if any(
                marker in lowered
                for marker in (
                    "grievance request",
                    "requested a grievance hearing",
                    "requested review",
                    "review request",
                    "submitted a grievance",
                    "informal review request",
                )
            ):
                dated_events.setdefault("review_request", (date_text, clause))
                continue
            if any(
                marker in lowered
                for marker in (
                    "review decision",
                    "hearing outcome",
                    "final decision",
                    "issued the review decision",
                    "issued the decision",
                )
            ):
                dated_events.setdefault("review_decision", (date_text, clause))
                continue
            if (
                "denial notice" in lowered
                or ("written notice" in lowered and "decision denying assistance" in lowered)
                or ("sent plaintiff" in lowered and "notice" in lowered)
            ):
                dated_events.setdefault("notice", (date_text, clause))

        ordered_segments = []
        for key in ("notice", "review_request", "review_decision"):
            if key not in dated_events:
                continue
            date_text, clause = dated_events[key]
            ordered_segments.append(f"on {date_text}, {clause}")
        if len(ordered_segments) < 2:
            return ""
        return f"The chronology shows that {_join_chronology_segments(ordered_segments)}."

    def _build_missing_detail_allegations(
        self,
        allegations: List[str],
        source_allegations: Optional[List[str]] = None,
    ) -> List[str]:
        source_pool = [str(item or "").strip() for item in (source_allegations or allegations) if str(item or "").strip()]
        if not allegations and not source_pool:
            return []
        priority_email_count = sum(1 for line in source_pool if self._is_priority_email_timeline_allegation(line))
        if priority_email_count >= 3:
            concrete_sequence = self._build_concrete_sequence_allegation(source_pool)
            return [concrete_sequence] if concrete_sequence else []
        combined = " ".join(source_pool or [str(item or "") for item in allegations])
        lowered = combined.lower()
        references_notice = "notice" in lowered
        references_review = any(token in lowered for token in ("informal review", "hearing", "grievance", "appeal"))
        actor_label = self._infer_actor_label_from_allegations(source_pool or allegations)

        fallbacks: List[str] = []
        if not any(_contains_hearing_timing_marker(line) for line in source_pool):
            if references_review:
                fallbacks.append(
                    f"The present record indicates that Plaintiff sought review of the challenged action, but the exact date of {self._actor_possessive_label(actor_label)} hearing or review request remains to be confirmed."
                )
            else:
                fallbacks.append(
                    "The present record indicates that Plaintiff challenged the adverse action, but the exact date of that request for review remains to be confirmed."
                )
        if not any(_contains_response_date_marker(line) for line in source_pool):
            if references_notice and references_review:
                fallbacks.append(
                    f"The current record further indicates deficiencies in {actor_label}'s notice and review process, although the exact dates of the notice, response, and final decision have not yet been confirmed."
                )
            elif references_notice:
                fallbacks.append(
                    f"The current record indicates that {actor_label} issued or should have issued written notice, but the exact notice and response dates have not yet been confirmed."
                )
        if not any(_contains_staff_identity_marker(line) for line in source_pool):
            if references_notice or references_review:
                fallbacks.append(
                    f"The present record does not yet identify by name the official at {actor_label} who issued the notice or handled Plaintiff's request for review."
                )
            else:
                fallbacks.append(
                    f"The present record does not yet identify by name the responsible {actor_label} decisionmaker."
                )
        if not any(_contains_sequence_timing_marker(line) for line in source_pool):
            concrete_sequence = self._build_concrete_sequence_allegation(source_pool)
            if concrete_sequence:
                fallbacks.append(concrete_sequence)
            else:
                fallbacks.append(
                    "The current chronology indicates that the protected activity, challenged notice, and resulting adverse treatment occurred in close sequence, although the precise intervals still require confirmation."
                )
        return fallbacks

    def _build_policy_process_allegations(
        self,
        allegations: List[str],
        source_allegations: Optional[List[str]] = None,
    ) -> List[str]:
        source_pool = [str(item or "").strip() for item in (source_allegations or allegations) if str(item or "").strip()]
        if not allegations and not source_pool:
            return []
        combined = " ".join(source_pool or [str(item or "") for item in allegations])
        lowered = combined.lower()
        references_notice = "written notice" in lowered or "notice to the applicant" in lowered or "notice" in lowered
        references_review = any(token in lowered for token in ("informal review", "grievance", "appeal", "hearing"))
        references_denial = any(
            token in lowered
            for token in ("denying assistance", "denial notice", "denied assistance", "loss of assistance", "termination", "voucher")
        )
        actor_label = self._infer_actor_label_from_allegations(source_pool or allegations)

        case_specific_adverse_clause = ""
        for text in source_pool:
            lowered_text = text.lower()
            if not _contains_date_anchor(text):
                continue
            if any(
                token in lowered_text
                for token in ("denial notice", "denied assistance", "denying assistance", "loss of assistance", "termination")
            ):
                case_specific_adverse_clause = text.rstrip(".!?")
                break

        adverse_action_phrase = "denied Plaintiff housing assistance, or maintained the challenged loss of assistance"
        if "denial notice" in lowered:
            adverse_action_phrase = "issued a denial notice and denied Plaintiff housing assistance"
        elif "denying assistance" in lowered or "denied assistance" in lowered:
            adverse_action_phrase = "issued or maintained a denial of assistance"
        elif "loss of assistance" in lowered:
            adverse_action_phrase = "maintained the challenged loss of assistance"
        elif "termination" in lowered:
            adverse_action_phrase = "terminated or threatened to terminate Plaintiff's assistance"

        synthesized: List[str] = []
        if references_notice and not any(
            "without providing the prompt written notice required" in str(item).lower()
            for item in allegations
        ) and not any(
            "notice to the applicant requires prompt written notice" in str(item).lower()
            or "requires prompt written notice of a decision denying assistance" in str(item).lower()
            for item in allegations
        ):
            synthesized.append(
                f"Plaintiff alleges that {actor_label} took or maintained the challenged adverse housing action without providing the prompt written notice required for a decision denying assistance."
            )
        if references_review and not any(
            "failed to provide the informal review" in str(item).lower()
            or "failed to provide the grievance" in str(item).lower()
            or "failed to provide the appeal" in str(item).lower()
            for item in allegations
        ):
            synthesized.append(
                f"Plaintiff further alleges that {actor_label} failed to provide the informal review, grievance, or appeal process that should have accompanied the challenged action."
            )
        if references_denial and not any(
            "without the notice and review protections described in the governing process" in str(item).lower()
            or "without providing the prompt written notice required" in str(item).lower()
            for item in allegations
        ):
            if case_specific_adverse_clause:
                lead_in = "Plaintiff further alleges that"
                normalized_clause = case_specific_adverse_clause
                if re.match(r"^(On|By|After|Before|During)\b", case_specific_adverse_clause):
                    normalized_clause = case_specific_adverse_clause[0].lower() + case_specific_adverse_clause[1:]
                synthesized.append(
                    f"{lead_in} {normalized_clause}, without the notice and review protections described in the governing process."
                )
            else:
                synthesized.append(
                    f"Plaintiff further alleges that {actor_label} {adverse_action_phrase} without the notice and review protections described in the governing process."
                )
        return synthesized

    def _build_email_timeline_issue_allegations(
        self,
        allegations: List[str],
        source_allegations: Optional[List[str]] = None,
    ) -> List[str]:
        source_pool = [str(item or "").strip() for item in (source_allegations or allegations) if str(item or "").strip()]
        if not source_pool:
            return []
        priority_lines = [line for line in source_pool if self._is_priority_email_timeline_allegation(line)]
        if not priority_lines:
            return []

        combined = " ".join(priority_lines).lower()
        has_hcv_orientation = "hcv orientation" in combined
        has_additional_information = "additional information needed" in combined
        has_cortez_case = "jane kay cortez" in combined
        has_ra_denial = "cortez-j-ra-denial" in combined or "ra-denial" in combined or "ra denial" in combined
        has_fraud_chain = "jc household" in combined or "allegations of fraud" in combined

        allegations_out: List[str] = []
        if has_hcv_orientation and has_ra_denial:
            allegations_out.append(
                "Plaintiff alleges that, by March 26, 2026, Clackamas housing staff were circulating an HCV orientation email chain that included the RA-denial attachment, reflecting maintenance of the challenged denial of housing assistance."
            )
        elif has_hcv_orientation:
            allegations_out.append(
                "Plaintiff alleges that, by March 26, 2026, Clackamas housing staff were still handling Plaintiff's housing matter through the HCV Orientation email chain while maintaining the challenged adverse housing action."
            )

        if has_additional_information:
            allegations_out.append(
                "Plaintiff further alleges that Clackamas housing staff used the 'Additional Information Needed' thread between February 9 and March 10, 2026 to demand additional housing-program information while the challenged denial or loss of assistance remained in place."
            )

        if has_cortez_case and (has_hcv_orientation or has_additional_information):
            allegations_out.append(
                "Plaintiff alleges that the March 10, 2026 Cortez case emails and the March 26, 2026 HCV Orientation emails show Clackamas personnel were actively handling Plaintiff's housing matter without first providing the notice and review protections required before enforcing an adverse housing decision."
            )

        if has_ra_denial:
            allegations_out.append(
                "Plaintiff further alleges that the RA-denial attachment and related HCV communications support Plaintiff's claim that Defendant denied housing-related benefits or accommodations on discriminatory grounds."
            )
        elif has_fraud_chain and has_hcv_orientation:
            allegations_out.append(
                "Plaintiff alleges that the December 2025 through March 2026 Clackamas email sequence shows a continuing course of adverse housing treatment rather than an isolated communication."
            )

        return _unique_preserving_order(allegations_out)

    def _prune_redundant_policy_rule_allegations(self, allegations: List[str]) -> List[str]:
        if not allegations:
            return []
        lowered_all = [str(item or "").lower() for item in allegations]
        has_notice_process_allegation = any(
            "without the notice and review protections described in the governing process" in item
            or "failed to provide the informal review" in item
            for item in lowered_all
        )
        has_chronology = any(item.startswith("the chronology shows that ") for item in lowered_all)
        if not (has_notice_process_allegation and has_chronology):
            return allegations

        pruned: List[str] = []
        for item in allegations:
            lowered = str(item or "").lower()
            if (
                "notice to the applicant requires prompt written notice of a decision denying assistance" in lowered
            ):
                continue
            pruned.append(item)
        return pruned or allegations

    def _build_factual_allegations(
        self,
        *,
        summary_of_facts: Any,
        claims_for_relief: List[Dict[str, Any]],
    ) -> List[str]:
        base_allegations = self._merge_email_attachment_follow_on_allegations(
            list(self._expand_allegation_sources(summary_of_facts, limit=14))
        )
        priority_email_allegations = [
            item for item in base_allegations if self._is_priority_email_timeline_allegation(item)
        ]
        allegations = list(priority_email_allegations)
        for item in self._synthesize_narrative_allegations(base_allegations):
            if item.lower() not in {entry.lower() for entry in allegations}:
                allegations.append(item)
        for item in self._prune_subsumed_narrative_clauses(base_allegations):
            if item.lower() not in {entry.lower() for entry in allegations}:
                allegations.append(item)
        seen = {entry.lower() for entry in allegations}

        for claim in _coerce_list(claims_for_relief):
            if not isinstance(claim, dict):
                continue
            count_title = str(claim.get("count_title") or claim.get("claim_type") or "Claim").strip()
            for fact in self._expand_allegation_sources(claim.get("supporting_facts", []), limit=10):
                if not fact:
                    continue
                if self._is_priority_email_timeline_allegation(fact):
                    continue
                if str(fact or "").strip().lower().startswith("one of those clackamas emails included the attachment "):
                    continue
                if self._is_near_duplicate_allegation(fact, allegations):
                    continue
                prefixed_fact = fact
                if count_title and not fact.lower().startswith("as to ") and fact.lower() not in seen:
                    lowered = fact
                    if not re.match(r"^(Plaintiff|Defendant)\b", fact):
                        lowered = fact[0].lower() + fact[1:] if len(fact) > 1 and fact[0].isalpha() else fact
                    prefixed_fact = f"As to {count_title}, {lowered}"
                    if not prefixed_fact.endswith((".", "?", "!")):
                        prefixed_fact = f"{prefixed_fact}."
                key = prefixed_fact.lower()
                if key in seen:
                    continue
                seen.add(key)
                allegations.append(prefixed_fact)
                if len(allegations) >= 24:
                    return self._prune_near_duplicate_allegations(allegations)

        pruned = self._prune_near_duplicate_allegations(allegations)
        for item in self._build_email_timeline_issue_allegations(pruned, base_allegations):
            if item.lower() not in {entry.lower() for entry in pruned}:
                pruned.append(item)
        pruned.extend(self._build_policy_process_allegations(pruned, base_allegations))
        pruned.extend(self._build_missing_detail_allegations(pruned, base_allegations))
        pruned = self._prune_near_duplicate_allegations(pruned)
        pruned = self._prune_redundant_policy_rule_allegations(pruned)
        return pruned[:24] or ["Additional factual development is required before filing."]

    def _build_factual_allegation_entries(
        self,
        *,
        summary_fact_entries: List[Dict[str, Any]],
        claims_for_relief: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        allegation_texts = self._build_factual_allegations(
            summary_of_facts=[entry.get("text") for entry in summary_fact_entries if isinstance(entry, dict)],
            claims_for_relief=claims_for_relief,
        )
        source_entries: List[Dict[str, Any]] = []
        source_entries.extend([dict(entry) for entry in summary_fact_entries if isinstance(entry, dict)])
        for claim in _coerce_list(claims_for_relief):
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or "").strip()
            for entry in _coerce_list(claim.get("supporting_fact_entries")):
                if not isinstance(entry, dict):
                    continue
                claim_entry = dict(entry)
                if claim_type:
                    claim_entry["claim_types"] = _normalize_identifier_list(
                        list(claim_entry.get("claim_types") or []) + [claim_type]
                    )
                source_entries.append(claim_entry)

        entries: List[Dict[str, Any]] = []
        for text in allegation_texts:
            matched_entries = self._match_document_source_entries(text, source_entries)
            if "failed to provide the informal review" in str(text or "").lower():
                matched_entries = self._enrich_review_process_fact_matches(
                    matched_entries=matched_entries,
                    source_entries=source_entries,
                )
            if str(text or "").lower().startswith("the chronology shows that "):
                matched_entries = self._enrich_chronology_fact_matches(
                    matched_entries=matched_entries,
                    source_entries=source_entries,
                )
            entries.append(
                {
                    "text": text,
                    "fact_ids": _normalize_identifier_list(
                        [
                            fact_id
                            for entry in matched_entries
                            for fact_id in _coerce_list(entry.get("fact_ids"))
                        ]
                    ),
                    "source_artifact_ids": _normalize_identifier_list(
                        [
                            artifact_id
                            for entry in matched_entries
                            for artifact_id in _coerce_list(entry.get("source_artifact_ids"))
                        ]
                    ),
                    "claim_types": _normalize_identifier_list(
                        [
                            claim_type
                            for entry in matched_entries
                            for claim_type in _coerce_list(entry.get("claim_types"))
                        ]
                    ),
                    "claim_element_ids": _normalize_identifier_list(
                        [
                            claim_element_id
                            for entry in matched_entries
                            for claim_element_id in _coerce_list(entry.get("claim_element_ids"))
                        ]
                    ),
                    "support_trace_ids": _normalize_identifier_list(
                        [
                            trace_id
                            for entry in matched_entries
                            for trace_id in _coerce_list(entry.get("support_trace_ids"))
                        ]
                    ),
                    "source_kind": "factual_allegation",
                }
            )
        return entries

    def _enrich_review_process_fact_matches(
        self,
        *,
        matched_entries: List[Dict[str, Any]],
        source_entries: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        def _entry_key(entry: Dict[str, Any]) -> tuple[tuple[str, ...], str]:
            return (
                tuple(_normalize_identifier_list(entry.get("fact_ids"))),
                str(entry.get("text") or ""),
            )

        def _is_requirement_fact(entry: Dict[str, Any]) -> bool:
            text = str(entry.get("text") or "").lower()
            return bool(_coerce_list(entry.get("fact_ids"))) and (
                ("require" in text or "required" in text)
                and ("review" in text or "grievance" in text or "appeal" in text)
            )

        def _is_event_fact(entry: Dict[str, Any]) -> bool:
            text = str(entry.get("text") or "").lower()
            if not _coerce_list(entry.get("fact_ids")):
                return False
            if "require" in text or "required" in text:
                return False
            return (
                "grievance request" in text
                or "review decision" in text
                or ("submitted" in text and "grievance" in text)
                or ("issued" in text and "review decision" in text)
            )

        enriched = [dict(entry) for entry in matched_entries if isinstance(entry, dict)]
        seen_keys = {_entry_key(entry) for entry in enriched}

        requirement_entry = next((entry for entry in enriched if _is_requirement_fact(entry)), None)
        event_entry = next((entry for entry in enriched if _is_event_fact(entry)), None)

        if requirement_entry is None:
            for entry in source_entries:
                if isinstance(entry, dict) and _is_requirement_fact(entry):
                    candidate_key = _entry_key(entry)
                    if candidate_key not in seen_keys:
                        enriched.append(dict(entry))
                        seen_keys.add(candidate_key)
                    requirement_entry = dict(entry)
                    break
        if event_entry is None:
            for entry in source_entries:
                if isinstance(entry, dict) and _is_event_fact(entry):
                    candidate_key = _entry_key(entry)
                    if candidate_key not in seen_keys:
                        enriched.append(dict(entry))
                        seen_keys.add(candidate_key)
                    event_entry = dict(entry)
                    break
        if event_entry is None:
            for entry in source_entries:
                if not isinstance(entry, dict) or not _coerce_list(entry.get("fact_ids")):
                    continue
                text = str(entry.get("text") or "").lower()
                if "grievance request" not in text and "review decision" not in text:
                    continue
                candidate_key = _entry_key(entry)
                if candidate_key not in seen_keys:
                    enriched.append(dict(entry))
                    seen_keys.add(candidate_key)
                event_entry = dict(entry)
                break
        has_grievance_request = any(
            "grievance request" in str(entry.get("text") or "").lower()
            or ("submitted" in str(entry.get("text") or "").lower() and "grievance" in str(entry.get("text") or "").lower())
            for entry in enriched
            if isinstance(entry, dict)
        )
        if not has_grievance_request:
            for entry in source_entries:
                if not isinstance(entry, dict) or not _coerce_list(entry.get("fact_ids")):
                    continue
                text = str(entry.get("text") or "").lower()
                if "grievance request" not in text and not ("submitted" in text and "grievance" in text):
                    continue
                candidate_key = _entry_key(entry)
                if candidate_key not in seen_keys:
                    enriched.append(dict(entry))
                    seen_keys.add(candidate_key)
                break
        return enriched

    def _enrich_chronology_fact_matches(
        self,
        *,
        matched_entries: List[Dict[str, Any]],
        source_entries: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        def _entry_key(entry: Dict[str, Any]) -> tuple[tuple[str, ...], str]:
            return (
                tuple(_normalize_identifier_list(entry.get("fact_ids"))),
                str(entry.get("text") or ""),
            )

        def _entry_category(entry: Dict[str, Any]) -> str:
            text = str(entry.get("text") or "").lower()
            if not _coerce_list(entry.get("fact_ids")):
                return ""
            if any(
                marker in text
                for marker in (
                    "grievance request",
                    "requested a grievance hearing",
                    "requested review",
                    "review request",
                    "submitted a grievance",
                    "informal review request",
                )
            ):
                return "review_request"
            if any(
                marker in text
                for marker in (
                    "review decision",
                    "hearing outcome",
                    "final decision",
                    "issued the review decision",
                    "issued the decision",
                )
            ):
                return "review_decision"
            if (
                "denial notice" in text
                or ("written notice" in text and "decision denying assistance" in text)
                or ("sent plaintiff" in text and "notice" in text)
            ):
                return "notice"
            return ""

        enriched = [dict(entry) for entry in matched_entries if isinstance(entry, dict)]
        seen_keys = {_entry_key(entry) for entry in enriched}
        present_categories = {
            category
            for entry in enriched
            if (category := _entry_category(entry))
        }

        for category in ("notice", "review_request", "review_decision"):
            if category in present_categories:
                continue
            for entry in source_entries:
                if not isinstance(entry, dict):
                    continue
                if _entry_category(entry) != category:
                    continue
                candidate_key = _entry_key(entry)
                if candidate_key not in seen_keys:
                    enriched.append(dict(entry))
                    seen_keys.add(candidate_key)
                present_categories.add(category)
                break
        return enriched

    def _attach_allegation_references(self, draft: Dict[str, Any]) -> None:
        allegation_lines = self._normalize_text_lines(
            draft.get("factual_allegations") or draft.get("summary_of_facts", [])
        )
        allegation_entries = self._align_entries_to_lines(
            draft.get("factual_allegation_entries"),
            allegation_lines,
        )
        paragraph_entries = [
            {
                "number": index,
                "text": text,
                "fact_ids": allegation_entries[index - 1].get("fact_ids", []),
                "source_artifact_ids": allegation_entries[index - 1].get("source_artifact_ids", []),
                "claim_types": allegation_entries[index - 1].get("claim_types", []),
                "claim_element_ids": allegation_entries[index - 1].get("claim_element_ids", []),
                "support_trace_ids": allegation_entries[index - 1].get("support_trace_ids", []),
                "source_kind": allegation_entries[index - 1].get("source_kind"),
                "document_focus": (
                    dict(allegation_entries[index - 1].get("document_focus") or {})
                    if isinstance(allegation_entries[index - 1].get("document_focus"), dict)
                    else {}
                ),
                "document_focus_priority_rank": allegation_entries[index - 1].get("document_focus_priority_rank"),
            }
            for index, text in enumerate(allegation_lines, start=1)
        ]
        paragraph_entries = self._annotate_entries_with_exhibits(
            paragraph_entries,
            draft.get("exhibits") if isinstance(draft.get("exhibits"), list) else [],
        )
        allegation_entries = [
            {
                **dict(allegation_entries[index]),
                "text": paragraph_entries[index].get("text"),
                "exhibit_label": paragraph_entries[index].get("exhibit_label"),
            }
            for index in range(min(len(allegation_entries), len(paragraph_entries)))
        ]
        allegation_lines = [str(entry.get("text") or "").strip() for entry in paragraph_entries if str(entry.get("text") or "").strip()]
        draft["factual_allegations"] = allegation_lines
        draft["factual_allegation_entries"] = allegation_entries
        draft["factual_allegation_paragraphs"] = paragraph_entries
        draft["factual_allegation_groups"] = self._build_factual_allegation_groups(paragraph_entries)

        claims = draft.get("claims_for_relief") if isinstance(draft.get("claims_for_relief"), list) else []
        for claim in claims:
            if not isinstance(claim, dict):
                continue
            claim["supporting_fact_provenance"] = self._build_claim_supporting_fact_provenance(
                claim=claim,
            )
            claim["allegation_references"] = self._select_allegation_references_for_claim(
                claim=claim,
                allegation_paragraphs=paragraph_entries,
            )
        draft["document_provenance_summary"] = self._build_document_provenance_summary(draft)

    def _build_factual_allegation_groups(self, allegation_paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        ordered_titles = [
            "Clackamas Email Chronology",
            "Protected Activity and Complaints",
            "Adverse Action and Retaliatory Conduct",
            "Damages and Resulting Harm",
            "Additional Factual Support",
        ]
        groups: Dict[str, List[Dict[str, Any]]] = {title: [] for title in ordered_titles}

        for paragraph in allegation_paragraphs:
            if not isinstance(paragraph, dict):
                continue
            title = self._classify_factual_allegation_group(paragraph)
            groups[title].append(paragraph)

        return [
            {"title": title, "paragraphs": self._order_factual_group_paragraphs(groups[title])}
            for title in ordered_titles
            if groups[title]
        ]

    def _classify_factual_allegation_group(self, paragraph: Dict[str, Any]) -> str:
        text = str(paragraph.get("text") or "").strip()
        lowered = text.lower()
        if self._is_priority_email_timeline_allegation(text):
            return "Clackamas Email Chronology"
        if re.search(
            r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied|denial notice|loss of assistance|review decision|adverse action)\b",
            lowered,
        ):
            return "Adverse Action and Retaliatory Conduct"
        if re.search(r"\b(lost|damages|harm|injur|suffered|experienced|benefits|wages|salary|income)\b", lowered):
            return "Damages and Resulting Harm"
        if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
            return "Protected Activity and Complaints"
        return "Additional Factual Support"

    def _order_factual_group_paragraphs(self, paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        def _score(paragraph: Dict[str, Any]) -> tuple[int, int, int, int, str]:
            text = str(paragraph.get("text") or "").strip()
            lowered = text.lower()
            fact_backed = 0 if _normalize_identifier_list(paragraph.get("fact_ids") or []) else 1
            dated = 0 if _contains_date_anchor(text) else 1
            adverse_specific = 1
            if any(
                marker in lowered
                for marker in (
                    "denial notice",
                    "review decision",
                    "loss of assistance",
                    "adverse action",
                    "denied plaintiff housing assistance",
                )
            ):
                adverse_specific = 0
            number = int(paragraph.get("number", 0) or 0)
            return (fact_backed, dated, adverse_specific, number, lowered)

        return sorted(
            [dict(paragraph) for paragraph in paragraphs if isinstance(paragraph, dict)],
            key=_score,
        )

    def _build_claim_supporting_fact_provenance(
        self,
        *,
        claim: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        supporting_entries = self._align_entries_to_lines(
            claim.get("supporting_fact_entries"),
            claim.get("supporting_facts") or [],
        )
        provenance_rows: List[Dict[str, Any]] = []
        for index, entry in enumerate(supporting_entries, start=1):
            text = str(entry.get("text") or "").strip()
            if not text:
                continue
            provenance_rows.append(
                {
                    "number": index,
                    "text": text,
                    "fact_ids": _normalize_identifier_list(entry.get("fact_ids") or []),
                    "source_artifact_ids": _normalize_identifier_list(entry.get("source_artifact_ids") or []),
                    "claim_types": _normalize_identifier_list(entry.get("claim_types") or []),
                    "claim_element_ids": _normalize_identifier_list(entry.get("claim_element_ids") or []),
                    "support_trace_ids": _normalize_identifier_list(entry.get("support_trace_ids") or []),
                    "source_kind": str(entry.get("source_kind") or "").strip() or None,
                    "exhibit_label": str(entry.get("exhibit_label") or "").strip() or None,
                    "document_focus": (
                        dict(entry.get("document_focus") or {})
                        if isinstance(entry.get("document_focus"), dict)
                        else {}
                    ),
                    "document_focus_priority_rank": entry.get("document_focus_priority_rank"),
                }
            )
        return provenance_rows

    def _build_document_provenance_summary(self, draft: Dict[str, Any]) -> Dict[str, Any]:
        factual_paragraphs = [
            paragraph
            for paragraph in _coerce_list(draft.get("factual_allegation_paragraphs"))
            if isinstance(paragraph, dict)
        ]
        claims = [
            claim
            for claim in _coerce_list(draft.get("claims_for_relief"))
            if isinstance(claim, dict)
        ]
        claim_rows = []
        total_claim_support_rows = 0
        total_claim_support_fact_backed_rows = 0
        total_claim_support_exhibit_backed_rows = 0
        focused_entry_count = 0
        focus_source_counts: Dict[str, int] = {}
        all_fact_ids: List[str] = []
        all_artifact_ids: List[str] = []
        for claim in claims:
            provenance_rows = [
                row
                for row in _coerce_list(claim.get("supporting_fact_provenance"))
                if isinstance(row, dict)
            ]
            fact_backed_count = sum(1 for row in provenance_rows if _normalize_identifier_list(row.get("fact_ids") or []))
            artifact_backed_count = sum(1 for row in provenance_rows if _normalize_identifier_list(row.get("source_artifact_ids") or []))
            exhibit_backed_count = sum(1 for row in provenance_rows if str(row.get("exhibit_label") or "").strip())
            total_claim_support_rows += len(provenance_rows)
            total_claim_support_fact_backed_rows += fact_backed_count
            total_claim_support_exhibit_backed_rows += exhibit_backed_count
            for row in provenance_rows:
                all_fact_ids.extend(_normalize_identifier_list(row.get("fact_ids") or []))
                all_artifact_ids.extend(_normalize_identifier_list(row.get("source_artifact_ids") or []))
                focus_metadata = row.get("document_focus") if isinstance(row.get("document_focus"), dict) else {}
                focus_source = str(focus_metadata.get("focus_source") or "").strip()
                if focus_source:
                    focused_entry_count += 1
                    focus_source_counts[focus_source] = focus_source_counts.get(focus_source, 0) + 1
            claim_rows.append(
                {
                    "claim_type": str(claim.get("claim_type") or "").strip(),
                    "supporting_fact_count": len(provenance_rows),
                    "fact_backed_supporting_fact_count": fact_backed_count,
                    "artifact_backed_supporting_fact_count": artifact_backed_count,
                    "exhibit_backed_supporting_fact_count": exhibit_backed_count,
                    "fact_ids": _normalize_identifier_list(
                        [
                            fact_id
                            for row in provenance_rows
                            for fact_id in _coerce_list(row.get("fact_ids"))
                        ]
                    ),
                    "source_artifact_ids": _normalize_identifier_list(
                        [
                            artifact_id
                            for row in provenance_rows
                            for artifact_id in _coerce_list(row.get("source_artifact_ids"))
                        ]
                    ),
                }
            )

        factual_fact_backed_count = sum(
            1 for paragraph in factual_paragraphs if _normalize_identifier_list(paragraph.get("fact_ids") or [])
        )
        factual_exhibit_backed_count = sum(
            1 for paragraph in factual_paragraphs if str(paragraph.get("exhibit_label") or "").strip()
        )
        for paragraph in factual_paragraphs:
            all_fact_ids.extend(_normalize_identifier_list(paragraph.get("fact_ids") or []))
            all_artifact_ids.extend(_normalize_identifier_list(paragraph.get("source_artifact_ids") or []))
            focus_metadata = paragraph.get("document_focus") if isinstance(paragraph.get("document_focus"), dict) else {}
            focus_source = str(focus_metadata.get("focus_source") or "").strip()
            if focus_source:
                focused_entry_count += 1
                focus_source_counts[focus_source] = focus_source_counts.get(focus_source, 0) + 1

        summary_entries = [
            entry for entry in _coerce_list(draft.get("summary_of_fact_entries")) if isinstance(entry, dict)
        ]
        summary_fact_backed_count = sum(1 for entry in summary_entries if _normalize_identifier_list(entry.get("fact_ids") or []))
        summary_exhibit_backed_count = sum(1 for entry in summary_entries if str(entry.get("exhibit_label") or "").strip())
        for entry in summary_entries:
            all_fact_ids.extend(_normalize_identifier_list(entry.get("fact_ids") or []))
            all_artifact_ids.extend(_normalize_identifier_list(entry.get("source_artifact_ids") or []))
            focus_metadata = entry.get("document_focus") if isinstance(entry.get("document_focus"), dict) else {}
            focus_source = str(focus_metadata.get("focus_source") or "").strip()
            if focus_source:
                focused_entry_count += 1
                focus_source_counts[focus_source] = focus_source_counts.get(focus_source, 0) + 1

        summary_fact_ratio = (float(summary_fact_backed_count) / float(len(summary_entries))) if summary_entries else 0.0
        allegation_ratio = (float(factual_fact_backed_count) / float(len(factual_paragraphs))) if factual_paragraphs else 0.0
        claim_ratio = (float(total_claim_support_fact_backed_rows) / float(total_claim_support_rows)) if total_claim_support_rows else 0.0
        fact_backed_ratio = (summary_fact_ratio + allegation_ratio + claim_ratio) / 3.0
        return {
            "summary_fact_count": len(summary_entries),
            "summary_fact_backed_count": summary_fact_backed_count,
            "summary_fact_exhibit_backed_count": summary_exhibit_backed_count,
            "factual_allegation_paragraph_count": len(factual_paragraphs),
            "factual_allegation_fact_backed_count": factual_fact_backed_count,
            "factual_allegation_exhibit_backed_count": factual_exhibit_backed_count,
            "claim_count": len(claims),
            "claim_supporting_fact_count": total_claim_support_rows,
            "claim_supporting_fact_backed_count": total_claim_support_fact_backed_rows,
            "claim_supporting_fact_exhibit_backed_count": total_claim_support_exhibit_backed_rows,
            "fact_id_count": len(_normalize_identifier_list(all_fact_ids)),
            "source_artifact_id_count": len(_normalize_identifier_list(all_artifact_ids)),
            "fact_backed_ratio": round(fact_backed_ratio, 4),
            "low_grounding_flag": bool(fact_backed_ratio < 0.6),
            "focused_entry_count": focused_entry_count,
            "focus_source_counts": focus_source_counts,
            "claims": claim_rows,
        }

    def _grouped_allegation_text_lines(self, draft: Dict[str, Any]) -> List[str]:
        groups = draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else []
        if not groups:
            return self._numbered_lines(draft.get("factual_allegations") or draft.get("summary_of_facts", []))

        lines: List[str] = []
        for group in groups:
            if not isinstance(group, dict):
                continue
            title = str(group.get("title") or "").strip()
            paragraphs = group.get("paragraphs") if isinstance(group.get("paragraphs"), list) else []
            if not paragraphs:
                continue
            if title:
                lines.append(self._format_group_heading(title, paragraphs))
            for paragraph in paragraphs:
                if not isinstance(paragraph, dict):
                    continue
                number = paragraph.get("number")
                text = str(paragraph.get("text") or "").strip()
                if text:
                    lines.append(f"{number}. {text}" if number else text)
        return lines

    def _format_group_heading(self, title: str, paragraphs: List[Dict[str, Any]]) -> str:
        heading = str(title or "").strip().upper()
        if not heading:
            return ""
        primary_exhibit = self._select_primary_group_exhibit_label(paragraphs)
        if primary_exhibit:
            return f"{heading} ({primary_exhibit})"
        return heading

    def _select_primary_group_exhibit_label(self, paragraphs: List[Dict[str, Any]]) -> str:
        label_counts: Dict[str, int] = {}
        for paragraph in paragraphs:
            if not isinstance(paragraph, dict):
                continue
            label = str(paragraph.get("exhibit_label") or "").strip()
            if not label:
                continue
            label_counts[label] = label_counts.get(label, 0) + 1
        if not label_counts:
            return ""
        label, count = max(label_counts.items(), key=lambda item: (item[1], item[0]))
        if count < 2 and len(paragraphs) > 1:
            return ""
        return label

    def _select_allegation_references_for_claim(
        self,
        *,
        claim: Dict[str, Any],
        allegation_paragraphs: List[Dict[str, Any]],
    ) -> List[int]:
        references: List[int] = []
        supporting_facts = self._normalize_text_lines(claim.get("supporting_facts", []))
        supporting_entries = self._align_entries_to_lines(
            claim.get("supporting_fact_entries"),
            supporting_facts,
        )
        count_title = str(claim.get("count_title") or claim.get("claim_type") or "").strip().lower()

        for fact, entry in zip(supporting_facts, supporting_entries):
            fact_tokens = self._text_tokens(fact)
            if not fact_tokens:
                continue
            best_number: Optional[int] = None
            best_score = 0
            fact_lower = fact.lower()
            fact_ids = set(_normalize_identifier_list(entry.get("fact_ids") or []))
            for paragraph in allegation_paragraphs:
                if not isinstance(paragraph, dict):
                    continue
                paragraph_text = str(paragraph.get("text") or "").strip()
                paragraph_lower = paragraph_text.lower()
                paragraph_tokens = self._text_tokens(paragraph_text)
                score = len(fact_tokens & paragraph_tokens)
                if fact_lower in paragraph_lower:
                    score += 100
                paragraph_fact_ids = set(_normalize_identifier_list(paragraph.get("fact_ids") or []))
                if fact_ids and paragraph_fact_ids and fact_ids & paragraph_fact_ids:
                    score += 50
                if count_title and count_title in paragraph_lower:
                    score += 5
                if score > best_score:
                    best_score = score
                    best_number = int(paragraph.get("number", 0) or 0)
            if best_number and best_number not in references:
                references.append(best_number)
                if len(references) >= 6:
                    break

        if references:
            return references

        fallback = []
        for paragraph in allegation_paragraphs:
            paragraph_text = str(paragraph.get("text") or "").lower()
            if count_title and count_title in paragraph_text:
                fallback.append(int(paragraph.get("number", 0) or 0))
        return fallback[:4]

    def _format_paragraph_reference_clause(self, references: Any) -> str:
        values = []
        for value in _coerce_list(references):
            try:
                number = int(value)
            except (TypeError, ValueError):
                continue
            if number > 0 and number not in values:
                values.append(number)
        if not values:
            return ""
        citation = self._format_paragraph_citation(values)
        return f"Plaintiff repeats and realleges {citation} as if fully set forth herein."

    def _format_incorporated_reference_clause(self, references: Any, exhibits: Any) -> str:
        paragraph_citation = self._format_paragraph_citation(references)
        exhibit_phrase = self._format_exhibit_reference_phrase(exhibits)
        if paragraph_citation and exhibit_phrase:
            return (
                f"Plaintiff repeats and realleges {paragraph_citation} and incorporates {exhibit_phrase} "
                "as if fully set forth herein."
            )
        if paragraph_citation:
            return f"Plaintiff repeats and realleges {paragraph_citation} as if fully set forth herein."
        if exhibit_phrase:
            return f"Plaintiff incorporates {exhibit_phrase} as if fully set forth herein."
        return ""

    def _format_paragraph_citation(self, references: Any) -> str:
        values = []
        for value in _coerce_list(references):
            try:
                number = int(value)
            except (TypeError, ValueError):
                continue
            if number > 0 and number not in values:
                values.append(number)
        if not values:
            return ""
        values.sort()
        ranges: List[str] = []
        range_start = values[0]
        range_end = values[0]
        for number in values[1:]:
            if number == range_end + 1:
                range_end = number
                continue
            ranges.append(self._format_paragraph_range(range_start, range_end))
            range_start = number
            range_end = number
        ranges.append(self._format_paragraph_range(range_start, range_end))
        marker = "¶" if len(values) == 1 else "¶¶"
        return f"{marker} {', '.join(ranges)}"

    def _format_exhibit_reference_phrase(self, exhibits: Any) -> str:
        exhibit_list = [exhibit for exhibit in _coerce_list(exhibits) if isinstance(exhibit, dict)]
        include_titles_for_all = len(exhibit_list) <= 2
        exhibit_entries: List[str] = []
        for exhibit in exhibit_list:
            label = str(exhibit.get("label") or "").strip()
            if not label:
                continue
            title = str(exhibit.get("title") or "").strip()
            kind = str(exhibit.get("kind") or "").strip().lower()
            entry = label
            if title and (kind == "evidence" or include_titles_for_all):
                entry = f"{label} ({title})"
            if entry not in exhibit_entries:
                exhibit_entries.append(entry)
        if not exhibit_entries:
            return ""
        if len(exhibit_entries) == 1:
            return exhibit_entries[0]
        if len(exhibit_entries) == 2:
            return f"{exhibit_entries[0]} and {exhibit_entries[1]}"
        return f"{', '.join(exhibit_entries[:-1])}, and {exhibit_entries[-1]}"

    def _format_paragraph_range(self, start: int, end: int) -> str:
        return str(start) if start == end else f"{start}-{end}"

    def _numbered_lines(self, values: Any) -> List[str]:
        return [f"{index}. {line}" for index, line in enumerate(self._normalize_text_lines(values), start=1)]

    def _bulletize_lines(self, values: Any) -> List[str]:
        return [f"- {line}" for line in self._normalize_text_lines(values)]

    def _build_signature_block(
        self,
        plaintiffs: List[str],
        *,
        signer_name: Optional[str] = None,
        signer_title: Optional[str] = None,
        signer_firm: Optional[str] = None,
        signer_bar_number: Optional[str] = None,
        signer_contact: Optional[str] = None,
        additional_signers: Optional[List[Dict[str, str]]] = None,
        signature_date: Optional[str] = None,
    ) -> Dict[str, Any]:
        plaintiff_name = str(signer_name or "").strip() or (plaintiffs or ["Plaintiff"])[0]
        return {
            "name": plaintiff_name,
            "signature_line": f"/s/ {plaintiff_name}",
            "title": str(signer_title or "").strip() or "Plaintiff, Pro Se",
            "firm": str(signer_firm or "").strip() or "",
            "bar_number": str(signer_bar_number or "").strip(),
            "contact": str(signer_contact or "").strip() or "Mailing address, telephone number, and email address to be completed before filing.",
            "additional_signers": self._normalize_additional_signers(additional_signers),
            "dated": self._format_dated_line("Dated", signature_date),
        }

    def _normalize_additional_signers(self, values: Any) -> List[Dict[str, str]]:
        normalized: List[Dict[str, str]] = []
        seen: set[tuple[str, str, str, str, str]] = set()
        for item in _coerce_list(values):
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or item.get("signer_name") or "").strip()
            title = str(item.get("title") or item.get("signer_title") or "").strip()
            firm = str(item.get("firm") or item.get("signer_firm") or "").strip()
            bar_number = str(item.get("bar_number") or item.get("signer_bar_number") or "").strip()
            contact = str(item.get("contact") or item.get("signer_contact") or "").strip()
            if not any((name, title, firm, bar_number, contact)):
                continue
            key = (name, title, firm, bar_number, contact)
            if key in seen:
                continue
            seen.add(key)
            normalized.append(
                {
                    "name": name or "Additional Counsel",
                    "signature_line": f"/s/ {name}" if name else "",
                    "title": title,
                    "firm": firm,
                    "bar_number": bar_number,
                    "contact": contact,
                }
            )
        return normalized

    def _signature_block_lines(self, signature_block: Dict[str, Any], *, include_dated: bool = True) -> List[str]:
        lines: List[str] = [
            str(signature_block.get("signature_line") or "/s/ Plaintiff"),
            str(signature_block.get("name") or "Plaintiff"),
        ]
        for key in ("title", "firm"):
            if signature_block.get(key):
                lines.append(str(signature_block[key]))
        if signature_block.get("bar_number"):
            lines.append(f"Bar No. {signature_block['bar_number']}")
        if signature_block.get("contact"):
            lines.append(str(signature_block["contact"]))
        for signer in _coerce_list(signature_block.get("additional_signers")):
            if not isinstance(signer, dict):
                continue
            lines.append("")
            if signer.get("signature_line"):
                lines.append(str(signer["signature_line"]))
            lines.append(str(signer.get("name") or "Additional Counsel"))
            for key in ("title", "firm"):
                if signer.get(key):
                    lines.append(str(signer[key]))
            if signer.get("bar_number"):
                lines.append(f"Bar No. {signer['bar_number']}")
            if signer.get("contact"):
                lines.append(str(signer["contact"]))
        if include_dated and signature_block.get("dated"):
            lines.append(str(signature_block["dated"]))
        return lines

    def _build_signature_section_lines(self, signature_block: Dict[str, Any], forum_type: str) -> List[str]:
        if forum_type == "state":
            lines: List[str] = []
            if signature_block.get("dated"):
                lines.append(str(signature_block["dated"]))
            lines.extend(["", "Respectfully submitted,", *self._signature_block_lines(signature_block, include_dated=False)])
            return lines
        return ["Respectfully submitted,", *self._signature_block_lines(signature_block)]

    def _build_jury_demand(
        self,
        *,
        jury_demand: Optional[bool] = None,
        jury_demand_text: Optional[str] = None,
    ) -> Dict[str, str]:
        text = str(jury_demand_text or "").strip()
        if text:
            return {
                "title": "Jury Demand",
                "text": text if text.endswith((".", "?", "!")) else f"{text}.",
            }
        if jury_demand:
            return {
                "title": "Jury Demand",
                "text": "Plaintiff demands a trial by jury on all issues so triable.",
            }
        return {}

    def _build_verification(
        self,
        plaintiffs: List[str],
        *,
        declarant_name: Optional[str] = None,
        signer_name: Optional[str] = None,
        verification_date: Optional[str] = None,
        jurisdiction: Optional[str] = None,
    ) -> Dict[str, str]:
        plaintiff_name = str(declarant_name or "").strip() or str(signer_name or "").strip() or (plaintiffs or ["Plaintiff"])[0]
        is_state = str(jurisdiction or "").strip().lower() == "state"
        return {
            "title": "Verification",
            "text": (
                f"I, {plaintiff_name}, verify that I have reviewed this Complaint and know its contents. "
                "The facts stated in this Complaint are true of my own knowledge, except as to those matters "
                "stated on information and belief, and as to those matters I believe them to be true."
                if is_state
                else (
                    f"I, {plaintiff_name}, declare under penalty of perjury that I have reviewed this Complaint "
                    "and that the factual allegations stated in it are true and correct to the best of my knowledge, "
                    "information, and belief."
                )
            ),
            "dated": self._format_dated_line("Verified on" if is_state else "Executed on", verification_date),
            "signature_line": f"/s/ {plaintiff_name}",
        }

    def _build_certificate_of_service(
        self,
        plaintiffs: List[str],
        defendants: List[str],
        *,
        signer_name: Optional[str] = None,
        service_method: Optional[str] = None,
        service_recipients: Optional[List[str]] = None,
        service_recipient_details: Optional[List[Dict[str, str]]] = None,
        service_date: Optional[str] = None,
        jurisdiction: Optional[str] = None,
    ) -> Dict[str, Any]:
        plaintiff_name = str(signer_name or "").strip() or (plaintiffs or ["Plaintiff"])[0]
        recipient_details = self._normalize_service_recipient_details(service_recipient_details)
        detail_recipients = [detail["recipient"] for detail in recipient_details if detail.get("recipient")]
        recipients_list = _unique_preserving_order([str(item or "").strip() for item in _coerce_list(service_recipients)] + detail_recipients) or defendants or ["all defendants"]
        recipients = ", ".join(recipients_list)
        method_text = str(service_method or "").strip() or "a method authorized by the applicable rules of civil procedure"
        detail_lines = [self._format_service_recipient_detail(detail) for detail in recipient_details]
        is_state = str(jurisdiction or "").strip().lower() == "state"
        return {
            "title": "Proof of Service" if is_state else "Certificate of Service",
            "text": (
                ("I declare that a true and correct copy of this Complaint will be served promptly after filing on the following recipients."
                if is_state
                else "I certify that a true and correct copy of this Complaint will be served promptly after filing on the following recipients.")
                if detail_lines
                else (("I declare that a true and correct copy of this Complaint will be served on "
                if is_state else "I certify that a true and correct copy of this Complaint will be served on ")
                + f"{recipients} using {method_text} promptly after filing.")
            ),
            "recipients": recipients_list,
            "recipient_details": recipient_details,
            "detail_lines": detail_lines,
            "dated": self._format_dated_line("Service date", service_date),
            "signature_line": f"/s/ {plaintiff_name}",
        }

    def _normalize_service_recipient_details(self, values: Any) -> List[Dict[str, str]]:
        details: List[Dict[str, str]] = []
        seen = set()
        for item in _coerce_list(values):
            if not isinstance(item, dict):
                continue
            detail = {
                "recipient": str(item.get("recipient") or "").strip(),
                "method": str(item.get("method") or "").strip(),
                "address": str(item.get("address") or "").strip(),
                "notes": str(item.get("notes") or "").strip(),
            }
            if not any(detail.values()):
                continue
            key = (detail["recipient"], detail["method"], detail["address"], detail["notes"])
            if key in seen:
                continue
            seen.add(key)
            details.append(detail)
        return details

    def _format_service_recipient_detail(self, detail: Dict[str, str]) -> str:
        segments = [detail.get("recipient") or "Recipient"]
        if detail.get("method"):
            segments.append(f"Method: {detail['method']}")
        if detail.get("address"):
            segments.append(f"Address: {detail['address']}")
        if detail.get("notes"):
            segments.append(f"Notes: {detail['notes']}")
        return " | ".join(segment for segment in segments if segment)

    def _format_dated_line(self, label: str, value: Optional[str]) -> str:
        cleaned = str(value or "").strip()
        return f"{label}: {cleaned}" if cleaned else f"{label}: __________________"

    def render_artifacts(
        self,
        draft: Dict[str, Any],
        *,
        output_dir: Optional[str],
        output_formats: List[str],
    ) -> Dict[str, Dict[str, Any]]:
        output_root = Path(output_dir).expanduser() if output_dir else DEFAULT_OUTPUT_DIR
        output_root.mkdir(parents=True, exist_ok=True)
        timestamp = _utcnow().strftime("%Y%m%dT%H%M%SZ")
        file_stem = f"{_slugify(draft.get('title') or 'complaint')}-{timestamp}"
        artifacts: Dict[str, Dict[str, Any]] = {}

        for output_format in output_formats:
            if output_format == "packet":
                continue
            path = self._artifact_path(output_root, file_stem, output_format)
            if output_format == "docx":
                self._render_docx(draft, path)
                affidavit_path = self._artifact_path(output_root, file_stem, output_format, document_kind="affidavit")
                self._render_affidavit_docx(draft, affidavit_path)
                artifacts["affidavit_docx"] = {
                    "path": str(affidavit_path),
                    "filename": affidavit_path.name,
                    "size_bytes": affidavit_path.stat().st_size,
                }
            elif output_format == "pdf":
                self._render_pdf(draft, path)
                affidavit_path = self._artifact_path(output_root, file_stem, output_format, document_kind="affidavit")
                self._render_affidavit_pdf(draft, affidavit_path)
                artifacts["affidavit_pdf"] = {
                    "path": str(affidavit_path),
                    "filename": affidavit_path.name,
                    "size_bytes": affidavit_path.stat().st_size,
                }
            elif output_format == "txt":
                self._render_txt(draft, path)
                affidavit_path = self._artifact_path(output_root, file_stem, output_format, document_kind="affidavit")
                self._render_affidavit_txt(draft, affidavit_path)
                artifacts["affidavit_txt"] = {
                    "path": str(affidavit_path),
                    "filename": affidavit_path.name,
                    "size_bytes": affidavit_path.stat().st_size,
                }
            elif output_format == "checklist":
                self._render_checklist_txt(draft, path)
            artifacts[output_format] = {
                "path": str(path),
                "filename": path.name,
                "size_bytes": path.stat().st_size,
            }

        if "packet" in output_formats:
            path = self._artifact_path(output_root, file_stem, "packet")
            self._render_packet_json(draft, path, artifacts=artifacts)
            artifacts["packet"] = {
                "path": str(path),
                "filename": path.name,
                "size_bytes": path.stat().st_size,
            }

        return artifacts

    def _resolve_user_id(self, user_id: Optional[str]) -> str:
        if user_id:
            return user_id
        state = getattr(self.mediator, "state", None)
        return (
            getattr(state, "username", None)
            or getattr(state, "hashed_username", None)
            or "anonymous"
        )

    def _normalize_formats(self, output_formats: Optional[List[str]]) -> List[str]:
        values = output_formats or ["docx", "pdf"]
        normalized = []
        for value in values:
            current = str(value or "").strip().lower()
            if current in {"docx", "pdf", "txt", "checklist", "packet"} and current not in normalized:
                normalized.append(current)
        return normalized or ["docx", "pdf"]

    def _build_affidavit_overrides(
        self,
        *,
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        normalized_override_facts = []
        for value in affidavit_facts or []:
            cleaned = self._sanitize_affidavit_fact(value)
            if cleaned:
                normalized_override_facts.append(cleaned)
        normalized_supporting_exhibits = []
        for exhibit in _coerce_list(affidavit_supporting_exhibits):
            if not isinstance(exhibit, dict):
                continue
            normalized = {
                "label": str(exhibit.get("label") or "Exhibit").strip(),
                "title": str(exhibit.get("title") or exhibit.get("summary") or "Supporting exhibit").strip(),
                "link": str(exhibit.get("link") or exhibit.get("reference") or "").strip(),
                "summary": str(exhibit.get("summary") or "").strip(),
            }
            if any(normalized.values()):
                normalized_supporting_exhibits.append(normalized)
        return {
            "title": str(affidavit_title or "").strip() or None,
            "intro": str(affidavit_intro or "").strip() or None,
            "facts": normalized_override_facts,
            "supporting_exhibits": normalized_supporting_exhibits,
            "include_complaint_exhibits": affidavit_include_complaint_exhibits,
            "venue_lines": self._normalize_text_lines(affidavit_venue_lines or []),
            "jurat": str(affidavit_jurat or "").strip() or None,
            "notary_block": self._normalize_text_lines(affidavit_notary_block or []),
        }

    def _build_affidavit(self, draft: Dict[str, Any]) -> Dict[str, Any]:
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        affidavit_overrides = draft.get("affidavit_overrides", {}) if isinstance(draft.get("affidavit_overrides"), dict) else {}
        declarant_name = self._derive_affidavit_declarant_name(draft)
        is_state = self._resolve_draft_forum_type(draft) == "state"
        exhibits = []
        for exhibit in _coerce_list(draft.get("exhibits")):
            if not isinstance(exhibit, dict):
                continue
            exhibits.append(
                {
                    "label": str(exhibit.get("label") or "Exhibit").strip(),
                    "title": str(exhibit.get("title") or exhibit.get("summary") or "Supporting exhibit").strip(),
                    "link": str(exhibit.get("link") or exhibit.get("reference") or "").strip(),
                    "summary": str(exhibit.get("summary") or "").strip(),
                }
            )
        return {
            "title": str(affidavit_overrides.get("title") or f"AFFIDAVIT OF {declarant_name.upper()} IN SUPPORT OF COMPLAINT"),
            "declarant_name": declarant_name,
            "intro": str(
                affidavit_overrides.get("intro")
                or (
                    (
                        f"I, {declarant_name}, being duly sworn, state that I am competent to testify to the matters stated below, "
                        "that these statements are based on my personal knowledge and the complaint intake knowledge graph assembled from the facts, records, and exhibits provided in support of this action, and that the following facts are true and correct."
                    )
                    if is_state
                    else (
                        f"I, {declarant_name}, declare under penalty of perjury that I am competent to testify to the matters stated below, "
                        "that these statements are based on my personal knowledge and the complaint intake knowledge graph assembled from the facts, records, and exhibits provided in support of this action, and that the following facts are true and correct."
                    )
                )
            ),
            "knowledge_graph_note": "This affidavit is generated from the complaint intake knowledge graph and supporting records rather than a turn-by-turn chat transcript.",
            "venue_lines": list(affidavit_overrides.get("venue_lines") or self._build_affidavit_venue_lines(draft)),
            "facts": list(affidavit_overrides.get("facts") or self._collect_affidavit_facts(draft)),
            "supporting_exhibits": list(
                affidavit_overrides.get("supporting_exhibits")
                or ([] if affidavit_overrides.get("include_complaint_exhibits") is False else exhibits)
            ),
            "dated": str(verification.get("dated") or signature_block.get("dated") or self._format_dated_line("Verified on" if is_state else "Executed on", None)),
            "signature_line": str(verification.get("signature_line") or signature_block.get("signature_line") or f"/s/ {declarant_name}"),
            "jurat": str(
                affidavit_overrides.get("jurat")
                or (
                    f"Subscribed and sworn to before me on __________________ by {declarant_name}."
                    if is_state
                    else f"Subscribed and sworn to (or affirmed) before me on __________________ by {declarant_name}."
                )
            ),
            "notary_block": list(
                affidavit_overrides.get("notary_block")
                or [
                    "__________________________________",
                    "Notary Public",
                    "My commission expires: __________________",
                ]
            ),
            "case_number": str(case_caption.get("case_number") or "________________"),
        }

    def _derive_affidavit_declarant_name(self, draft: Dict[str, Any]) -> str:
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        signature_line = str(verification.get("signature_line") or signature_block.get("signature_line") or "").strip()
        if signature_line.startswith("/s/ "):
            return signature_line[4:].strip() or str(signature_block.get("name") or "Plaintiff")
        plaintiffs = [str(name).strip() for name in _coerce_list(parties.get("plaintiffs")) if str(name).strip()]
        return str(signature_block.get("name") or (plaintiffs[0] if plaintiffs else "Plaintiff")).strip() or "Plaintiff"

    def _build_affidavit_venue_lines(self, draft: Dict[str, Any]) -> List[str]:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        county = str(caption.get("county") or "").strip()
        district = str(source_context.get("district") or "").strip()
        jurisdiction = str(source_context.get("jurisdiction") or caption.get("forum_type") or "").strip().lower()
        lines: List[str] = []
        if district:
            lines.append(f"State/District: {district}")
        elif jurisdiction == "federal":
            lines.append("State/District: __________________")
        if county:
            lines.append(f"County: {county.title()}")
        elif jurisdiction == "state":
            lines.append("County: __________________")
        return lines or ["Venue: __________________"]

    def _collect_affidavit_facts(self, draft: Dict[str, Any]) -> List[str]:
        candidates: List[str] = []
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        plaintiffs = [str(name).strip() for name in _coerce_list(parties.get("plaintiffs")) if str(name).strip()]
        if plaintiffs:
            candidates.append(f"I am {plaintiffs[0]}, the plaintiff in this action.")
        candidates.extend(self._normalize_text_lines(draft.get("factual_allegations", [])))

        facts: List[str] = []
        seen: set[str] = set()
        for candidate in candidates:
            cleaned = self._sanitize_affidavit_fact(candidate)
            if not cleaned:
                continue
            key = cleaned.lower()
            if key in seen:
                continue
            seen.add(key)
            facts.append(cleaned)
            if len(facts) >= 12:
                break
        return facts or ["Additional fact development is required before the affidavit can be finalized."]

    def _sanitize_affidavit_fact(self, value: str) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            return ""
        text = re.sub(r"^As to [^,]+,\s*", "", text, flags=re.IGNORECASE)
        if ": " in text:
            prefix, suffix = text.split(": ", 1)
            prefix_lower = prefix.strip().lower()
            if (
                prefix.strip().endswith("?")
                or prefix_lower.startswith(("what ", "when ", "where ", "why ", "how ", "who ", "describe ", "explain "))
                or prefix_lower in {"what happened", "what relief do you want"}
            ):
                text = suffix.strip()
        lowered = text.lower()
        if lowered.startswith("plaintiff repeats and realleges"):
            return ""
        if not self._is_factual_allegation_candidate(text) and not lowered.startswith("i am "):
            return ""
        if len(text) < 12:
            return ""
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        if text[-1] not in ".!?":
            text = f"{text}."
        return text

    def _render_txt(self, draft: Dict[str, Any], path: Path) -> None:
        path.write_text(str(draft.get("draft_text") or self._render_draft_text(draft)), encoding="utf-8")

    def _render_affidavit_txt(self, draft: Dict[str, Any], path: Path) -> None:
        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else self._build_affidavit(draft)
        path.write_text(self._render_affidavit_text(draft, affidavit), encoding="utf-8")

    def _render_checklist_txt(self, draft: Dict[str, Any], path: Path) -> None:
        checklist = draft.get("filing_checklist") if isinstance(draft.get("filing_checklist"), list) else []
        title = str(draft.get("title") or draft.get("case_caption", {}).get("document_title") or "Complaint").strip()
        lines = [
            f"PRE-FILING CHECKLIST: {title}",
            "",
        ]
        if not checklist:
            lines.append("No pre-filing checklist items were generated.")
        else:
            for index, item in enumerate(checklist, start=1):
                if not isinstance(item, dict):
                    continue
                scope = str(item.get("scope") or "item").strip().upper()
                title_text = str(item.get("title") or "Checklist Item").strip()
                status = str(item.get("status") or "ready").strip().upper()
                summary = str(item.get("summary") or "").strip()
                detail = str(item.get("detail") or "").strip()
                review_url = str(item.get("review_url") or "").strip()
                chip_labels = _extract_text_candidates(item.get("chip_labels"))
                claim_missing_temporal_predicates = _extract_text_candidates(item.get("claim_missing_temporal_predicates"))
                claim_required_provenance_kinds = _extract_text_candidates(item.get("claim_required_provenance_kinds"))
                claim_unresolved_temporal_issue_count = int(item.get("claim_unresolved_temporal_issue_count") or 0)
                lines.append(f"{index}. [{status}] {scope}: {title_text}")
                if summary:
                    lines.append(f"   Summary: {summary}")
                if detail:
                    lines.append(f"   Detail: {detail}")
                if chip_labels:
                    lines.append(f"   Signals: {' | '.join(chip_labels)}")
                if claim_unresolved_temporal_issue_count > 0:
                    lines.append(f"   Unresolved temporal issues: {claim_unresolved_temporal_issue_count}")
                if claim_missing_temporal_predicates:
                    lines.append(
                        f"   Missing temporal predicates: {'; '.join(claim_missing_temporal_predicates)}"
                    )
                if claim_required_provenance_kinds:
                    lines.append(
                        "   Required provenance kinds: "
                        + ", ".join(_humanize_checklist_label(value) or str(value) for value in claim_required_provenance_kinds)
                    )
                if review_url:
                    lines.append(f"   Review URL: {review_url}")
                lines.append("")
        path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    def _artifact_path(self, output_root: Path, file_stem: str, output_format: str, document_kind: str = "complaint") -> Path:
        suffix = "-affidavit" if document_kind == "affidavit" else ""
        if output_format == "checklist":
            return output_root / f"{file_stem}{suffix}-checklist.txt"
        if output_format == "packet":
            return output_root / f"{file_stem}-packet.json"
        return output_root / f"{file_stem}{suffix}.{output_format}"

    def _render_packet_json(
        self,
        draft: Dict[str, Any],
        path: Path,
        *,
        artifacts: Dict[str, Dict[str, Any]],
    ) -> None:
        payload = self._build_filing_packet_payload(draft, artifacts=artifacts)
        path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")

    def _build_filing_packet_payload(
        self,
        draft: Dict[str, Any],
        *,
        artifacts: Dict[str, Dict[str, Any]],
    ) -> Dict[str, Any]:
        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        packet_artifacts = {
            key: {
                "filename": value.get("filename"),
                "path": value.get("path"),
                "size_bytes": value.get("size_bytes"),
            }
            for key, value in artifacts.items()
            if isinstance(value, dict)
        }
        return {
            "title": draft.get("title"),
            "court_header": draft.get("court_header"),
            "generated_at": source_context.get("generated_at") or _utcnow().isoformat(),
            "claim_support_temporal_handoff": dict(source_context.get("claim_support_temporal_handoff") or {}) if isinstance(source_context.get("claim_support_temporal_handoff"), dict) else {},
            "email_timeline_handoff": deepcopy(source_context.get("email_timeline_handoff") or {}) if isinstance(source_context.get("email_timeline_handoff"), dict) else {},
            "email_authority_enrichment": deepcopy(source_context.get("email_authority_enrichment") or {}) if isinstance(source_context.get("email_authority_enrichment"), dict) else {},
            "claim_reasoning_review": deepcopy(source_context.get("claim_reasoning_review") or {}) if isinstance(source_context.get("claim_reasoning_review"), dict) else {},
            "chronology_blocker_summary": deepcopy(source_context.get("chronology_blocker_summary") or {}) if isinstance(source_context.get("chronology_blocker_summary"), dict) else {},
            "source_context": source_context,
            "case_caption": {
                "plaintiffs": case_caption.get("plaintiffs", []),
                "defendants": case_caption.get("defendants", []),
                "case_number": case_caption.get("case_number"),
                "document_title": case_caption.get("document_title"),
                "jury_demand_notice": case_caption.get("jury_demand_notice"),
            },
            "sections": {
                "nature_of_action": draft.get("nature_of_action", []),
                "summary_of_facts": draft.get("summary_of_facts", []),
                "factual_allegations": draft.get("factual_allegations", []),
                "anchored_chronology_summary": draft.get("anchored_chronology_summary", []),
                "email_authority_summary_lines": draft.get("email_authority_summary_lines", []),
                "claims_for_relief": draft.get("claims_for_relief", []),
                "legal_standards": draft.get("legal_standards", []),
                "requested_relief": draft.get("requested_relief", []),
            },
            "affidavit": draft.get("affidavit", {}),
            "verification": draft.get("verification", {}),
            "certificate_of_service": draft.get("certificate_of_service", {}),
            "exhibits": draft.get("exhibits", []),
            "filing_checklist": draft.get("filing_checklist", []),
            "drafting_readiness": draft.get("drafting_readiness", {}),
            "artifacts": packet_artifacts,
        }

    def _render_affidavit_text(self, draft: Dict[str, Any], affidavit: Dict[str, Any]) -> str:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        caption_party_lines = caption.get("caption_party_lines") if isinstance(caption.get("caption_party_lines"), list) else self._build_caption_party_lines(caption)
        lines = [
            str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"),
            *([str(caption.get("county"))] if caption.get("county") else []),
            "",
            *caption_party_lines,
            f"{caption.get('case_number_label', 'Civil Action No.')} {caption.get('case_number', '________________')}",
            "",
            str(affidavit.get("title") or "AFFIDAVIT IN SUPPORT OF COMPLAINT"),
            *[str(line) for line in _coerce_list(affidavit.get("venue_lines")) if str(line or "").strip()],
            "",
            str(affidavit.get("intro") or ""),
            str(affidavit.get("knowledge_graph_note") or ""),
            "",
            "Affiant states as follows:",
            *self._numbered_lines(affidavit.get("facts", [])),
        ]
        exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
        if exhibits:
            lines.extend(["", "SUPPORTING EXHIBITS"])
            for exhibit in exhibits:
                if not isinstance(exhibit, dict):
                    continue
                exhibit_text = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
                if exhibit.get("link"):
                    exhibit_text = f"{exhibit_text} ({exhibit['link']})"
                lines.append(exhibit_text)
        lines.extend(["", str(affidavit.get("dated") or ""), str(affidavit.get("signature_line") or ""), str(affidavit.get("jurat") or "")])
        lines.extend(str(line) for line in _coerce_list(affidavit.get("notary_block")) if str(line or "").strip())
        return "\n".join(line for line in lines if line is not None)

    def _render_affidavit_docx(self, draft: Dict[str, Any], path: Path) -> None:
        from docx import Document

        document = Document()
        for line in self._render_affidavit_text(
            draft,
            draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else self._build_affidavit(draft),
        ).split("\n"):
            document.add_paragraph(line)
        document.save(path)

    def _render_affidavit_pdf(self, draft: Dict[str, Any], path: Path) -> None:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        styles = getSampleStyleSheet()
        story = []
        for line in self._render_affidavit_text(
            draft,
            draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else self._build_affidavit(draft),
        ).split("\n"):
            story.append(Paragraph(escape(line or "&nbsp;"), styles["Normal"]))
            story.append(Spacer(1, 4))
        doc = SimpleDocTemplate(
            str(path),
            pagesize=LETTER,
            topMargin=inch,
            bottomMargin=inch,
            leftMargin=inch,
            rightMargin=inch,
        )
        doc.build(story)

    def _get_existing_formal_complaint(self) -> Dict[str, Any]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        if phase_manager is None:
            return {}
        existing = _safe_call(
            phase_manager,
            "get_phase_data",
            ComplaintPhase.FORMALIZATION,
            "formal_complaint",
        )
        if isinstance(existing, dict) and existing:
            return existing
        return {}

    def _derive_claim_types(
        self,
        generated_complaint: Dict[str, Any],
        classification: Dict[str, Any],
        support_claims: Dict[str, Any],
        requirements: Dict[str, Any],
        *,
        user_id: Optional[str] = None,
    ) -> List[str]:
        claim_names = []
        claim_names.extend(_coerce_list(classification.get("claim_types")))
        claim_names.extend(list(support_claims.keys()))
        claim_names.extend(list(requirements.keys()))
        if user_id:
            for evidence in _coerce_list(_safe_call(self.mediator, "get_user_evidence", user_id=user_id) or []):
                if not isinstance(evidence, dict):
                    continue
                claim_names.extend(_coerce_list(evidence.get("claim_type")))
                metadata = evidence.get("metadata") if isinstance(evidence.get("metadata"), dict) else {}
                claim_names.extend(_coerce_list(metadata.get("claim_type")))
                claim_names.extend(_coerce_list(metadata.get("claim_types")))
        for claim in _coerce_list(generated_complaint.get("legal_claims")):
            if isinstance(claim, dict):
                claim_names.append(claim.get("title"))
        return _unique_preserving_order(claim_names) or ["General civil action"]

    def _derive_parties(
        self,
        generated_complaint: Dict[str, Any],
        *,
        plaintiff_names: Optional[List[str]],
        defendant_names: Optional[List[str]],
    ) -> tuple[List[str], List[str]]:
        parties = generated_complaint.get("parties", {}) if isinstance(generated_complaint, dict) else {}
        plaintiffs = _unique_preserving_order(
            list(plaintiff_names or []) + list(parties.get("plaintiffs", []) or [])
        ) or ["Plaintiff"]
        defendants = _unique_preserving_order(
            list(defendant_names or []) + list(parties.get("defendants", []) or [])
        ) or ["Defendant"]
        return plaintiffs, defendants

    def _derive_title(self, plaintiffs: List[str], defendants: List[str]) -> str:
        return f"{plaintiffs[0]} v. {defendants[0]}"

    def _build_court_header(
        self,
        *,
        court_name: str,
        district: str,
        county: Optional[str],
        division: Optional[str],
        override: Optional[str],
    ) -> str:
        if override:
            return override.strip().upper()
        court = str(court_name or "United States District Court").strip().upper()
        parts = [f"IN THE {court}"]
        forum_type = self._infer_forum_type(classification={}, court_name=court_name)
        county_text = self._format_county_for_header(county)
        if county_text and forum_type == "state":
            parts.append(f"FOR THE {county_text}")
        elif district:
            parts.append(f"FOR THE {str(district).strip().upper()}")
        if division:
            parts.append(str(division).strip().upper())
        return " ".join(parts)

    def _infer_forum_type(
        self,
        *,
        classification: Dict[str, Any],
        court_name: str,
    ) -> str:
        jurisdiction = str(classification.get("jurisdiction") or "").strip().lower()
        if jurisdiction in {"federal", "us", "united states"}:
            return "federal"
        if jurisdiction in {"state", "state court", "county", "local"}:
            return "state"

        court_name_text = str(court_name or "").strip().lower()
        if "united states" in court_name_text or "u.s." in court_name_text:
            return "federal"
        if any(
            marker in court_name_text
            for marker in ("superior court", "circuit court", "common pleas", "state of", "county")
        ):
            return "state"
        return "unknown"

    def _build_nature_of_action(
        self,
        *,
        claim_types: List[str],
        classification: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        court_name: str,
    ) -> List[str]:
        claim_phrase = ", ".join(claim_types)
        legal_areas = ", ".join(_coerce_list(classification.get("legal_areas")))
        forum_type = self._infer_forum_type(classification=classification, court_name=court_name)
        jurisdiction = str(classification.get("jurisdiction") or "").strip().lower()
        if not jurisdiction or jurisdiction == "unknown":
            if forum_type == "federal":
                jurisdiction = "federal"
            elif forum_type == "state":
                jurisdiction = "state"
            else:
                jurisdiction = "the applicable court"
        statute_refs = _unique_preserving_order(
            [s.get("citation") for s in statutes if isinstance(s, dict) and s.get("citation")]
        )
        if forum_type == "federal":
            paragraphs = [
                (
                    "This is a civil action arising under federal law and the facts disclosed during the "
                    f"complaint intake process. Plaintiff seeks relief for {claim_phrase} within {jurisdiction} jurisdiction."
                )
            ]
        elif forum_type == "state":
            paragraphs = [
                (
                    "This is a civil action brought in state court arising from the facts disclosed during "
                    f"the complaint intake process. Plaintiff seeks relief for {claim_phrase} under the governing state law."
                )
            ]
        else:
            paragraphs = [
                (
                    "This is a civil action arising from the facts disclosed during the complaint intake "
                    f"process. Plaintiff seeks relief for {claim_phrase} within {jurisdiction} jurisdiction."
                )
            ]
        if legal_areas:
            paragraphs.append(
                f"The action implicates the following areas of law: {legal_areas}."
            )
        if statute_refs:
            paragraphs.append(
                "The draft relies on the following principal legal authorities: "
                f"{', '.join(statute_refs[:5])}."
            )
        return paragraphs

    def _collect_general_facts(
        self,
        generated_complaint: Dict[str, Any],
        classification: Dict[str, Any],
        state: Any,
    ) -> List[str]:
        facts: List[str] = []
        for allegation in _coerce_list(generated_complaint.get("factual_allegations")):
            facts.extend(_extract_text_candidates(allegation))
        facts.extend(_extract_text_candidates(classification.get("key_facts")))
        for inquiry in _coerce_list(getattr(state, "inquiries", []) if state is not None else []):
            if isinstance(inquiry, dict):
                question = str(inquiry.get("question") or "").strip()
                answer = str(inquiry.get("answer") or "").strip()
                if answer:
                    if question:
                        facts.append(f"{question}: {answer}")
                    else:
                        facts.append(answer)
        complaint_text = getattr(state, "complaint", None) if state is not None else None
        original_text = getattr(state, "original_complaint", None) if state is not None else None
        if isinstance(complaint_text, dict):
            explicit_facts = _extract_text_candidates(complaint_text.get("facts"))
            if explicit_facts:
                facts.extend(explicit_facts)
            else:
                facts.extend(_extract_text_candidates(complaint_text.get("summary") or complaint_text))
        elif complaint_text:
            facts.extend(_extract_text_candidates(complaint_text))
        elif original_text:
            facts.extend(_extract_text_candidates(original_text))

        normalized = []
        for item in _unique_preserving_order(facts):
            text = re.sub(r"\s+", " ", item).strip()
            if len(text) < 12:
                continue
            normalized.append(text)
            if len(normalized) >= 12:
                break
        return normalized or ["Additional factual development is required before filing."]

    def _build_summary_fact_entries(
        self,
        *,
        user_id: str = "",
        generated_complaint: Dict[str, Any],
        classification: Dict[str, Any],
        state: Any,
    ) -> List[Dict[str, Any]]:
        entries: List[Dict[str, Any]] = []
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        canonical_facts = (
            list(intake_case_file.get("canonical_facts") or [])
            if isinstance(intake_case_file, dict)
            else []
        )
        for fact in canonical_facts:
            if not isinstance(fact, dict):
                continue
            text = re.sub(r"\s+", " ", str(fact.get("text") or "").strip())
            if len(text) < 12:
                continue
            entries.append(
                {
                    "text": text,
                    "fact_ids": _normalize_identifier_list(
                        [fact.get("fact_id")] + list(fact.get("related_fact_ids") or [])
                    ),
                    "source_artifact_ids": _normalize_identifier_list(
                        list(fact.get("source_artifact_ids") or [])
                        + ([fact.get("source_artifact_id")] if fact.get("source_artifact_id") else [])
                    ),
                    "claim_types": _normalize_identifier_list(fact.get("claim_types") or []),
                    "claim_element_ids": _normalize_identifier_list(
                        fact.get("element_tags")
                        or fact.get("claim_element_ids")
                        or ([] if not fact.get("predicate_family") else [fact.get("predicate_family")])
                    ),
                    "source_kind": str(fact.get("source_kind") or "canonical_fact").strip() or "canonical_fact",
                    "source_ref": str(fact.get("source_ref") or fact.get("fact_id") or "").strip() or None,
                }
            )

        entries.extend(self._collect_uploaded_evidence_fact_entries(user_id=user_id, limit=8))

        for allegation in _coerce_list(generated_complaint.get("factual_allegations")):
            for text in _extract_text_candidates(allegation):
                entries.append({"text": text, "source_kind": "generated_complaint"})
        for text in _extract_text_candidates(classification.get("key_facts")):
            entries.append({"text": text, "source_kind": "classification"})
        for inquiry in _coerce_list(getattr(state, "inquiries", []) if state is not None else []):
            if not isinstance(inquiry, dict):
                continue
            question = str(inquiry.get("question") or "").strip()
            answer = str(inquiry.get("answer") or "").strip()
            if answer:
                entries.append(
                    {
                        "text": f"{question}: {answer}" if question else answer,
                        "source_kind": "inquiry",
                        "source_ref": question or None,
                    }
                )

        complaint_text = getattr(state, "complaint", None) if state is not None else None
        original_text = getattr(state, "original_complaint", None) if state is not None else None
        if isinstance(complaint_text, dict):
            explicit_facts = _extract_text_candidates(complaint_text.get("facts"))
            if explicit_facts:
                for text in explicit_facts:
                    entries.append({"text": text, "source_kind": "complaint"})
            else:
                for text in _extract_text_candidates(complaint_text.get("summary") or complaint_text):
                    entries.append({"text": text, "source_kind": "complaint"})
        elif complaint_text:
            for text in _extract_text_candidates(complaint_text):
                entries.append({"text": text, "source_kind": "complaint"})
        elif original_text:
            for text in _extract_text_candidates(original_text):
                entries.append({"text": text, "source_kind": "original_complaint"})

        normalized: List[Dict[str, Any]] = []
        seen = set()
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            text = re.sub(r"\s+", " ", str(entry.get("text") or "").strip())
            if len(text) < 12:
                continue
            key = text.lower()
            if key in seen:
                continue
            seen.add(key)
            normalized.append(
                {
                    "text": text,
                    "fact_ids": _normalize_identifier_list(entry.get("fact_ids") or []),
                    "source_artifact_ids": _normalize_identifier_list(entry.get("source_artifact_ids") or []),
                    "claim_types": _normalize_identifier_list(entry.get("claim_types") or []),
                    "claim_element_ids": _normalize_identifier_list(entry.get("claim_element_ids") or []),
                    "source_kind": str(entry.get("source_kind") or "").strip() or None,
                    "source_ref": str(entry.get("source_ref") or "").strip() or None,
                }
            )
            if len(normalized) >= 12:
                break
        return normalized or [{"text": "Additional factual development is required before filing.", "source_kind": "fallback"}]

    def _collect_uploaded_evidence_fact_entries(
        self,
        *,
        user_id: str,
        claim_type: Optional[str] = None,
        limit: int = 8,
    ) -> List[Dict[str, Any]]:
        evidence_records = _safe_call(self.mediator, "get_user_evidence", user_id=user_id) or []
        normalized_claim_type = str(claim_type or "").strip().lower()
        entries: List[Dict[str, Any]] = []

        for record in _coerce_list(evidence_records):
            if not isinstance(record, dict):
                continue
            record_claim_type = str(record.get("claim_type") or "").strip().lower()
            metadata = record.get("metadata") if isinstance(record.get("metadata"), dict) else {}
            record_claim_types = _normalize_identifier_list(
                ([record.get("claim_type")] if record.get("claim_type") else [])
                + list(metadata.get("claim_types") or [])
            )
            if normalized_claim_type and record_claim_types:
                normalized_record_claim_types = {str(value or "").strip().lower() for value in record_claim_types if str(value or "").strip()}
                if normalized_claim_type not in normalized_record_claim_types:
                    continue
            evidence_id = record.get("id")
            source_artifact_ids = _normalize_identifier_list(
                [record.get("cid"), evidence_id]
                + list(record.get("source_artifact_ids") or [])
                + ([record.get("source_artifact_id")] if record.get("source_artifact_id") else [])
            )
            claim_types = record_claim_types
            fact_rows = _safe_call(self.mediator, "get_evidence_facts", evidence_id=evidence_id) or []

            for row in _coerce_list(fact_rows):
                if not isinstance(row, dict):
                    continue
                for text in _extract_text_candidates(row):
                    for fragment in self._extract_uploaded_evidence_text_candidates(text, limit=3):
                        cleaned = re.sub(r"\s+", " ", str(fragment or "").strip())
                        if len(cleaned) < 12:
                            continue
                        entries.append(
                            {
                                "text": cleaned,
                                "fact_ids": _normalize_identifier_list([row.get("fact_id")] + list(row.get("fact_ids") or [])),
                                "source_artifact_ids": source_artifact_ids,
                                "claim_types": claim_types,
                                "claim_element_ids": _normalize_identifier_list(
                                    [row.get("claim_element_id"), row.get("element_id")] + list(row.get("element_tags") or [])
                                ),
                                "support_trace_ids": _normalize_identifier_list(
                                    [
                                        trace.get("source_ref")
                                        for trace in _coerce_list(row.get("support_traces"))
                                        if isinstance(trace, dict) and trace.get("source_ref")
                                    ]
                                ),
                                "source_kind": "uploaded_evidence_fact",
                                "source_ref": str(row.get("source_ref") or record.get("cid") or evidence_id or "").strip() or None,
                            }
                        )

            preview_texts = _extract_text_candidates(
                {
                    "parsed_text_preview": record.get("parsed_text_preview"),
                    "description": record.get("description"),
                    "summary": metadata.get("summary") if isinstance(metadata, dict) else "",
                }
            )
            for text in preview_texts:
                for fragment in self._extract_uploaded_evidence_text_candidates(text, limit=2):
                    cleaned = re.sub(r"\s+", " ", str(fragment or "").strip())
                    if len(cleaned) < 12:
                        continue
                    entries.append(
                        {
                            "text": cleaned,
                            "fact_ids": [],
                            "source_artifact_ids": source_artifact_ids,
                            "claim_types": claim_types,
                            "claim_element_ids": [],
                            "support_trace_ids": [],
                            "source_kind": "uploaded_evidence_preview",
                            "source_ref": str(record.get("cid") or evidence_id or "").strip() or None,
                        }
                    )

        normalized: List[Dict[str, Any]] = []
        seen = set()
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            text = re.sub(r"\s+", " ", str(entry.get("text") or "").strip())
            if len(text) < 12 or self._is_generic_claim_support_text(text):
                continue
            key = text.lower()
            if key in seen:
                continue
            seen.add(key)
            normalized.append(
                {
                    "text": text,
                    "fact_ids": _normalize_identifier_list(entry.get("fact_ids") or []),
                    "source_artifact_ids": _normalize_identifier_list(entry.get("source_artifact_ids") or []),
                    "claim_types": _normalize_identifier_list(entry.get("claim_types") or []),
                    "claim_element_ids": _normalize_identifier_list(entry.get("claim_element_ids") or []),
                    "support_trace_ids": _normalize_identifier_list(entry.get("support_trace_ids") or []),
                    "source_kind": str(entry.get("source_kind") or "").strip() or "uploaded_evidence_fact",
                    "source_ref": str(entry.get("source_ref") or "").strip() or None,
                }
            )
            if len(normalized) >= limit:
                break
        return normalized

    def _build_claims_for_relief(
        self,
        *,
        user_id: str,
        claim_types: List[str],
        requirements: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        support_claims: Dict[str, Any],
        exhibits: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        claims: List[Dict[str, Any]] = []
        for claim_type in claim_types:
            support_claim = support_claims.get(claim_type, {}) if isinstance(support_claims, dict) else {}
            overview = _safe_call(
                self.mediator,
                "get_claim_overview",
                claim_type=claim_type,
                user_id=user_id,
                required_support_kinds=["evidence", "authority"],
            ) or {}
            overview_claim = overview.get("claims", {}).get(claim_type, {}) if isinstance(overview, dict) else {}
            claim_fact_entries = self._collect_claim_fact_entries(claim_type, user_id, support_claim)
            related_exhibits = self._select_related_exhibits_for_claim(
                claim_type=claim_type,
                exhibits=exhibits,
                claim_fact_entries=claim_fact_entries,
            )
            claim_fact_entries = self._annotate_entries_with_exhibits(claim_fact_entries, related_exhibits)
            claim_fact_entries = self._order_claim_fact_entries(claim_fact_entries)
            claim_fact_entries = self._prune_redundant_claim_fact_entries(claim_fact_entries)
            claim_facts = [str(entry.get("text") or "").strip() for entry in claim_fact_entries if str(entry.get("text") or "").strip()]
            source_context = self._extract_support_source_context_counts(support_claim)
            claims.append(
                {
                    "claim_type": claim_type,
                    "count_title": self._humanize_claim_title(claim_type, claim_facts),
                    "legal_standards": self._build_claim_legal_standards(
                        claim_type=claim_type,
                        requirements=requirements,
                        statutes=statutes,
                        support_claim=support_claim,
                        related_exhibits=related_exhibits,
                    ),
                    "supporting_facts": claim_facts,
                    "supporting_fact_entries": claim_fact_entries,
                    "missing_elements": self._extract_overview_elements(overview_claim.get("missing")),
                    "partially_supported_elements": self._extract_overview_elements(
                        overview_claim.get("partially_supported")
                    ),
                    "support_summary": {
                        "total_elements": support_claim.get("total_elements", 0),
                        "covered_elements": support_claim.get("covered_elements", 0),
                        "uncovered_elements": support_claim.get("uncovered_elements", 0),
                        "support_by_kind": support_claim.get("support_by_kind", {}),
                        "support_by_source": source_context["support_by_source"],
                        "source_family_counts": source_context["source_family_counts"],
                        "record_scope_counts": source_context["record_scope_counts"],
                        "artifact_family_counts": source_context["artifact_family_counts"],
                        "corpus_family_counts": source_context["corpus_family_counts"],
                        "content_origin_counts": source_context["content_origin_counts"],
                    },
                    "supporting_exhibits": [
                        {
                            "label": exhibit.get("label"),
                            "title": exhibit.get("title"),
                            "kind": exhibit.get("kind"),
                            "link": exhibit.get("link"),
                        }
                        for exhibit in related_exhibits[:8]
                    ],
                }
            )
        claims.sort(key=self._claim_order_score)
        return claims

    def _order_claim_fact_entries(self, entries: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        def _entry_score(entry: Dict[str, Any]) -> tuple[int, int, int, str]:
            if not isinstance(entry, dict):
                return (9, 9, 9, "")
            text = str(entry.get("text") or "").strip()
            lowered = text.lower()
            source_kind = str(entry.get("source_kind") or "").strip().lower()
            source_priority = {
                "uploaded_evidence_fact": 0,
                "claim_chronology_support": 1,
                "claim_support_fact": 2,
                "claim_support_link": 3,
            }.get(source_kind, 8)
            process_priority = 8
            if "denial notice" in lowered or ("sent plaintiff" in lowered and "notice" in lowered):
                process_priority = 0
            elif "grievance request" in lowered or ("submitted" in lowered and "grievance" in lowered):
                process_priority = 1
            elif "review decision" in lowered or ("issued" in lowered and "review decision" in lowered):
                process_priority = 2
            elif "notice to the applicant requires" in lowered:
                process_priority = 3
            elif "informal review requires" in lowered:
                process_priority = 4
            elif "element supported:" in lowered:
                process_priority = 5
            elif "requires written notice and an opportunity for informal review" in lowered:
                process_priority = 6
            elif "informal review for denial of assistance" in lowered:
                process_priority = 7
            return (
                source_priority,
                process_priority,
                0 if _contains_date_anchor(text) else 1,
                text,
            )

        return sorted(
            [dict(entry) for entry in entries if isinstance(entry, dict)],
            key=_entry_score,
        )

    def _prune_redundant_claim_fact_entries(self, entries: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        ordered_entries = [dict(entry) for entry in entries if isinstance(entry, dict)]
        if not ordered_entries:
            return []

        has_notice_requirement = any(
            "notice to the applicant requires" in str(entry.get("text") or "").lower()
            for entry in ordered_entries
        )
        has_review_requirement = any(
            "scheduling an informal review requires" in str(entry.get("text") or "").lower()
            for entry in ordered_entries
        )
        has_authority_notice_review = any(
            "requires written notice and an opportunity for informal review" in str(entry.get("text") or "").lower()
            for entry in ordered_entries
        )

        pruned: List[Dict[str, Any]] = []
        for entry in ordered_entries:
            text = str(entry.get("text") or "").strip()
            lowered = text.lower()
            source_kind = str(entry.get("source_kind") or "").strip().lower()
            if lowered.startswith("element supported: required notice and review process"):
                if has_notice_requirement and has_review_requirement:
                    continue
            if "informal review for denial of assistance" in lowered:
                if has_authority_notice_review:
                    continue
            if source_kind == "uploaded_evidence_preview":
                if "chronology" in lowered or "denial notice and review" in lowered:
                    continue
            pruned.append(entry)
        return self._merge_claim_support_policy_entries(pruned)

    def _merge_claim_support_policy_entries(self, entries: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        ordered_entries = [dict(entry) for entry in entries if isinstance(entry, dict)]
        if not ordered_entries:
            return []

        notice_entry = next(
            (
                entry
                for entry in ordered_entries
                if "notice to the applicant requires" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        review_entry = next(
            (
                entry
                for entry in ordered_entries
                if "scheduling an informal review requires" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        authority_entry = next(
            (
                entry
                for entry in ordered_entries
                if "requires written notice and an opportunity for informal review" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        adverse_action_entry = next(
            (
                entry
                for entry in ordered_entries
                if "denial notice" in str(entry.get("text") or "").lower()
                or "loss of assistance" in str(entry.get("text") or "").lower()
                or "challenged adverse housing action" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        if not (notice_entry and review_entry and authority_entry):
            return self._merge_housing_policy_entries(ordered_entries)

        merged_source_entries = [
            item for item in (adverse_action_entry, notice_entry, review_entry, authority_entry) if isinstance(item, dict)
        ]
        merged_entry = {
            "text": (
                "HACC policy required prompt written notice of a decision denying assistance and a written "
                "opportunity to request informal review before a final adverse housing decision was enforced "
                f"({self._merge_support_exhibit_labels(notice_entry, review_entry, authority_entry)})."
            ),
            "fact_ids": _normalize_identifier_list(
                [
                    fact_id
                    for source_entry in merged_source_entries
                    for fact_id in list(source_entry.get("fact_ids") or [])
                ]
            ),
            "source_artifact_ids": _normalize_identifier_list(
                [
                    artifact_id
                    for source_entry in merged_source_entries
                    for artifact_id in list(source_entry.get("source_artifact_ids") or [])
                ]
            ),
            "claim_types": _normalize_identifier_list(
                [
                    claim_type
                    for source_entry in merged_source_entries
                    for claim_type in list(source_entry.get("claim_types") or [])
                ]
            ),
            "claim_element_ids": _normalize_identifier_list(
                [
                    element_id
                    for source_entry in merged_source_entries
                    for element_id in list(source_entry.get("claim_element_ids") or [])
                ]
            ),
            "support_trace_ids": _normalize_identifier_list(
                [
                    trace_id
                    for source_entry in merged_source_entries
                    for trace_id in list(source_entry.get("support_trace_ids") or [])
                ]
            ),
            "source_kind": "claim_support_merged",
            "exhibit_label": self._merge_support_exhibit_labels(notice_entry, review_entry, authority_entry),
        }

        merged_entries: List[Dict[str, Any]] = []
        merged_inserted = False
        for entry in ordered_entries:
            text = str(entry.get("text") or "").lower()
            if (
                "notice to the applicant requires" in text
                or "scheduling an informal review requires" in text
                or "requires written notice and an opportunity for informal review" in text
            ):
                if not merged_inserted:
                    merged_entries.append(merged_entry)
                    merged_inserted = True
                continue
            merged_entries.append(entry)
        return merged_entries

    def _merge_housing_policy_entries(self, entries: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        ordered_entries = [dict(entry) for entry in entries if isinstance(entry, dict)]
        if not ordered_entries:
            return []

        notice_entry = next(
            (
                entry
                for entry in ordered_entries
                if "notice to the applicant requires" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        review_entry = next(
            (
                entry
                for entry in ordered_entries
                if "scheduling an informal review requires" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        adverse_action_entry = next(
            (
                entry
                for entry in ordered_entries
                if "denial notice" in str(entry.get("text") or "").lower()
                or "denied or maintained the denial of housing assistance" in str(entry.get("text") or "").lower()
                or "loss of assistance" in str(entry.get("text") or "").lower()
            ),
            None,
        )
        if not (notice_entry and review_entry):
            return ordered_entries

        merged_source_entries = [item for item in (adverse_action_entry, notice_entry, review_entry) if isinstance(item, dict)]
        merged_entry = {
            "text": (
                "HACC wrongfully denied or maintained the denial of housing assistance without the written notice "
                "and informal review "
                f"required before enforcement of that adverse action ({self._merge_support_exhibit_labels(notice_entry, review_entry)})."
            ),
            "fact_ids": _normalize_identifier_list(
                [
                    fact_id
                    for source_entry in merged_source_entries
                    for fact_id in list(source_entry.get("fact_ids") or [])
                ]
            ),
            "source_artifact_ids": _normalize_identifier_list(
                [
                    artifact_id
                    for source_entry in merged_source_entries
                    for artifact_id in list(source_entry.get("source_artifact_ids") or [])
                ]
            ),
            "claim_types": _normalize_identifier_list(
                [
                    claim_type
                    for source_entry in merged_source_entries
                    for claim_type in list(source_entry.get("claim_types") or [])
                ]
            ),
            "claim_element_ids": _normalize_identifier_list(
                [
                    element_id
                    for source_entry in merged_source_entries
                    for element_id in list(source_entry.get("claim_element_ids") or [])
                ]
            ),
            "support_trace_ids": _normalize_identifier_list(
                [
                    trace_id
                    for source_entry in merged_source_entries
                    for trace_id in list(source_entry.get("support_trace_ids") or [])
                ]
            ),
            "source_kind": "claim_support_merged",
            "exhibit_label": self._merge_support_exhibit_labels(notice_entry, review_entry),
        }

        merged_entries: List[Dict[str, Any]] = []
        merged_inserted = False
        for entry in ordered_entries:
            text = str(entry.get("text") or "").lower()
            if (
                "notice to the applicant requires" in text
                or "scheduling an informal review requires" in text
            ):
                if not merged_inserted:
                    merged_entries.append(merged_entry)
                    merged_inserted = True
                continue
            merged_entries.append(entry)
        return merged_entries

    def _merge_support_exhibit_labels(self, *entries: Dict[str, Any]) -> str:
        labels = [
            str(entry.get("exhibit_label") or "").strip()
            for entry in entries
            if isinstance(entry, dict) and str(entry.get("exhibit_label") or "").strip()
        ]
        labels = _unique_preserving_order(labels)
        if not labels:
            return ""
        if len(labels) == 1:
            return f"See {labels[0]}"
        return "See " + " and ".join(labels[:2])

    def _select_related_exhibits_for_claim(
        self,
        *,
        claim_type: str,
        exhibits: List[Dict[str, Any]],
        claim_fact_entries: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        artifact_ids = {
            str(artifact_id)
            for entry in _coerce_list(claim_fact_entries)
            if isinstance(entry, dict)
            for artifact_id in _coerce_list(entry.get("source_artifact_ids"))
            if str(artifact_id or "").strip()
        }
        related: List[Dict[str, Any]] = []
        seen_keys = set()
        for exhibit in exhibits:
            if not isinstance(exhibit, dict):
                continue
            exhibit_claim_type = str(exhibit.get("claim_type") or "").strip()
            exhibit_source_ref = str(exhibit.get("source_ref") or "").strip()
            include = (
                not exhibit_claim_type
                or exhibit_claim_type == claim_type
                or (exhibit_source_ref and exhibit_source_ref in artifact_ids)
            )
            if not include:
                continue
            key = (
                str(exhibit.get("label") or ""),
                str(exhibit.get("source_ref") or ""),
                str(exhibit.get("kind") or ""),
            )
            if key in seen_keys:
                continue
            seen_keys.add(key)
            related.append(exhibit)
        return related

    def _claim_order_score(self, claim: Dict[str, Any]) -> tuple[int, int, int, str]:
        if not isinstance(claim, dict):
            return (9, 999, 999, "")
        claim_type = normalize_claim_type(claim.get("claim_type") or "")
        priority_map = {
            "due_process_failure": 0,
            "housing_discrimination": 1,
            "retaliation": 2,
            "disability_discrimination": 3,
        }
        support_summary = claim.get("support_summary") if isinstance(claim.get("support_summary"), dict) else {}
        authority_count = len(_coerce_list(claim.get("supporting_exhibits")))
        covered = int(support_summary.get("covered_elements", 0) or 0)
        uncovered = int(support_summary.get("uncovered_elements", 0) or 0)
        return (
            priority_map.get(claim_type, 8),
            -covered,
            uncovered - authority_count,
            str(claim.get("count_title") or claim_type),
        )

    def _humanize_claim_title(self, claim_type: str, claim_facts: Optional[List[str]] = None) -> str:
        normalized = normalize_claim_type(claim_type or "")
        combined_facts = " ".join(str(item or "") for item in (claim_facts or []))
        lowered = combined_facts.lower()
        if normalized == "housing_discrimination":
            if any(
                marker in lowered
                for marker in (
                    "denial notice",
                    "denied assistance",
                    "denying assistance",
                    "loss of assistance",
                    "voucher",
                    "informal review",
                    "grievance",
                )
            ):
                return "Housing Discrimination and Wrongful Denial of Assistance"
            return "Housing Discrimination"
        if normalized == "due_process_failure":
            return "Denial of Required Notice and Informal Review"
        if normalized == "retaliation":
            return "Retaliation"
        if normalized == "disability_discrimination":
            return "Disability Discrimination"
        label = str(claim_type or "").strip().replace("_", " ")
        return label.title() if label else "Claim"

    def _build_claim_specific_relief(self, *, claim_types: List[str], facts: List[str]) -> List[str]:
        combined = " ".join(str(item or "") for item in facts).lower()
        normalized_claim_types = {normalize_claim_type(item) for item in claim_types if str(item or "").strip()}
        relief: List[str] = []
        housing_process_context = bool(
            normalized_claim_types & {"housing_discrimination", "due_process_failure"}
        ) and any(
            marker in combined
            for marker in (
                "denial notice",
                "written notice",
                "denied assistance",
                "denying assistance",
                "loss of assistance",
                "informal review",
                "grievance",
                "review decision",
            )
        )
        if housing_process_context:
            relief.extend(
                [
                    "Declaratory relief that Defendant's challenged denial or loss of housing assistance was imposed without the notice and review protections required by law.",
                    "Injunctive relief requiring Defendant to rescind or stay the challenged denial or loss of housing assistance unless and until lawful notice and review procedures are provided.",
                    "Injunctive relief requiring Defendant to provide the informal review, grievance hearing, appeal, or other process required before any final adverse housing decision is enforced.",
                ]
            )
        return relief

    def _extract_support_source_context_counts(self, support_claim: Dict[str, Any]) -> Dict[str, Dict[str, int]]:
        packet_summary = (
            support_claim.get("support_packet_summary", {})
            if isinstance(support_claim, dict) and isinstance(support_claim.get("support_packet_summary"), dict)
            else {}
        )

        def _normalized_counts(key: str) -> Dict[str, int]:
            primary = support_claim.get(key, {}) if isinstance(support_claim, dict) else {}
            fallback = packet_summary.get(key, {})
            source = primary if isinstance(primary, dict) and primary else fallback
            if not isinstance(source, dict):
                return {}
            counts: Dict[str, int] = {}
            for label, value in source.items():
                normalized_label = str(label or "").strip()
                if not normalized_label:
                    continue
                count = int(value or 0)
                if count <= 0:
                    continue
                counts[normalized_label] = count
            return counts

        return {
            "support_by_source": _normalized_counts("support_by_source"),
            "source_family_counts": _normalized_counts("source_family_counts"),
            "record_scope_counts": _normalized_counts("record_scope_counts"),
            "artifact_family_counts": _normalized_counts("artifact_family_counts"),
            "corpus_family_counts": _normalized_counts("corpus_family_counts"),
            "content_origin_counts": _normalized_counts("content_origin_counts"),
        }

    def _collect_claim_facts(
        self,
        claim_type: str,
        user_id: str,
        support_claim: Dict[str, Any],
    ) -> List[str]:
        return [
            str(entry.get("text") or "").strip()
            for entry in self._collect_claim_fact_entries(claim_type, user_id, support_claim)
            if str(entry.get("text") or "").strip()
        ]

    def _collect_claim_fact_entries(
        self,
        claim_type: str,
        user_id: str,
        support_claim: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        facts: List[str] = []
        entries: List[Dict[str, Any]] = []
        fact_rows = _safe_call(
            self.mediator,
            "get_claim_support_facts",
            claim_type=claim_type,
            user_id=user_id,
        ) or []
        for row in _coerce_list(fact_rows):
            row_texts = _extract_text_candidates(row)
            facts.extend(row_texts)
            if isinstance(row, dict):
                fact_ids = _normalize_identifier_list([row.get("fact_id")] + list(row.get("fact_ids") or []))
                source_artifact_ids = _normalize_identifier_list(
                    list(row.get("source_artifact_ids") or [])
                    + ([row.get("source_artifact_id")] if row.get("source_artifact_id") else [])
                )
                claim_element_ids = _normalize_identifier_list(
                    [row.get("claim_element_id"), row.get("element_id")] + list(row.get("element_tags") or [])
                )
                support_trace_ids = _normalize_identifier_list(
                    list(row.get("support_trace_ids") or [])
                    + [
                        trace.get("source_ref")
                        for trace in _coerce_list(row.get("support_traces"))
                        if isinstance(trace, dict) and trace.get("source_ref")
                    ]
                )
                for text in row_texts:
                    entries.append(
                        {
                            "text": text,
                            "fact_ids": fact_ids,
                            "source_artifact_ids": source_artifact_ids,
                            "claim_types": [claim_type] if claim_type else [],
                            "claim_element_ids": claim_element_ids,
                            "support_trace_ids": support_trace_ids,
                            "source_kind": "claim_support_fact",
                            "source_ref": str(row.get("source_ref") or row.get("fact_id") or "").strip() or None,
                        }
                    )

        for element in _coerce_list(support_claim.get("elements") if isinstance(support_claim, dict) else []):
            if not isinstance(element, dict):
                continue
            element_text = str(element.get("element_text") or element.get("claim_element") or "").strip()
            if element_text:
                facts.append(f"Element supported: {element_text}")
            for link in _coerce_list(element.get("links")):
                if isinstance(link, dict):
                    link_texts = _extract_text_candidates(link)
                    facts.extend(link_texts)
                    for text in link_texts:
                        entries.append(
                            {
                                "text": text,
                                "fact_ids": _normalize_identifier_list(
                                    [link.get("fact_id")] + list(link.get("fact_ids") or []) + list(element.get("fact_ids") or [])
                                ),
                                "source_artifact_ids": _normalize_identifier_list(
                                    list(link.get("source_artifact_ids") or [])
                                    + ([link.get("source_artifact_id")] if link.get("source_artifact_id") else [])
                                    + list(element.get("supporting_artifact_ids") or [])
                                ),
                                "claim_types": [claim_type] if claim_type else [],
                                "claim_element_ids": _normalize_identifier_list(
                                    [element.get("element_id"), link.get("claim_element_id"), link.get("element_id")]
                                ),
                                "support_trace_ids": _normalize_identifier_list(
                                    [
                                        trace.get("source_ref")
                                        for trace in _coerce_list(element.get("support_traces"))
                                        if isinstance(trace, dict) and trace.get("source_ref")
                                    ]
                                ),
                                "source_kind": "claim_support_link",
                                "source_ref": str(link.get("source_ref") or link.get("fact_id") or "").strip() or None,
                            }
                        )

        evidence_entries = self._collect_uploaded_evidence_fact_entries(
            user_id=user_id,
            claim_type=claim_type,
            limit=6,
        )
        for entry in evidence_entries:
            if not isinstance(entry, dict):
                continue
            text = str(entry.get("text") or "").strip()
            if text:
                facts.append(text)
            entries.append(dict(entry))

        normalized = []
        for item in _unique_preserving_order(facts):
            text = re.sub(r"\s+", " ", item).strip()
            if len(text) < 10 or self._is_generic_claim_support_text(text):
                continue
            normalized.append(text)
            if len(normalized) >= 8:
                break
        chronology_support = self._build_claim_chronology_support(claim_type=claim_type, claim_name=claim_type.title())
        combined = _unique_preserving_order(chronology_support + normalized)
        if not combined:
            combined = [f"The intake record describes facts supporting the {claim_type} claim."]
        entry_index = {
            str(entry.get("text") or "").strip().lower(): entry
            for entry in entries
            if isinstance(entry, dict) and str(entry.get("text") or "").strip()
        }
        result: List[Dict[str, Any]] = []
        for text in combined:
            key = str(text or "").strip().lower()
            entry = dict(entry_index.get(key) or {})
            entry["text"] = text
            entry["fact_ids"] = _normalize_identifier_list(entry.get("fact_ids") or [])
            entry["source_artifact_ids"] = _normalize_identifier_list(entry.get("source_artifact_ids") or [])
            entry["claim_types"] = _normalize_identifier_list(entry.get("claim_types") or [claim_type])
            entry["claim_element_ids"] = _normalize_identifier_list(entry.get("claim_element_ids") or [])
            entry["support_trace_ids"] = _normalize_identifier_list(entry.get("support_trace_ids") or [])
            if not entry.get("source_kind"):
                entry["source_kind"] = "claim_chronology_support" if text in chronology_support else "claim_support_fact"
            result.append(entry)
        return self._enrich_claim_support_entries(result, entries)

    def _enrich_claim_support_entries(
        self,
        entries: List[Dict[str, Any]],
        source_entries: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        enriched: List[Dict[str, Any]] = []
        evidence_pool = [
            dict(entry)
            for entry in source_entries
            if isinstance(entry, dict) and _coerce_list(entry.get("fact_ids"))
        ]
        for entry in entries:
            current = dict(entry) if isinstance(entry, dict) else {}
            if current.get("fact_ids"):
                enriched.append(current)
                continue
            text = str(current.get("text") or "").strip()
            if not text:
                enriched.append(current)
                continue
            matched_entries = self._match_document_source_entries(text, evidence_pool, limit=4)
            if "informal review" in text.lower() or "required notice and review process" in text.lower():
                matched_entries = self._enrich_review_process_fact_matches(
                    matched_entries=matched_entries,
                    source_entries=evidence_pool,
                )
                matched_entries = self._narrow_claim_process_fact_matches(
                    text=text,
                    matched_entries=matched_entries,
                )
            current["fact_ids"] = _normalize_identifier_list(
                [
                    fact_id
                    for matched in matched_entries
                    for fact_id in _coerce_list(matched.get("fact_ids"))
                ]
            )
            current["source_artifact_ids"] = _normalize_identifier_list(
                list(current.get("source_artifact_ids") or [])
                + [
                    artifact_id
                    for matched in matched_entries
                    for artifact_id in _coerce_list(matched.get("source_artifact_ids"))
                ]
            )
            current["claim_types"] = _normalize_identifier_list(
                list(current.get("claim_types") or [])
                + [
                    claim_type
                    for matched in matched_entries
                    for claim_type in _coerce_list(matched.get("claim_types"))
                ]
            )
            current["claim_element_ids"] = _normalize_identifier_list(
                list(current.get("claim_element_ids") or [])
                + [
                    claim_element_id
                    for matched in matched_entries
                    for claim_element_id in _coerce_list(matched.get("claim_element_ids"))
                ]
            )
            current["support_trace_ids"] = _normalize_identifier_list(
                list(current.get("support_trace_ids") or [])
                + [
                    trace_id
                    for matched in matched_entries
                    for trace_id in _coerce_list(matched.get("support_trace_ids"))
                ]
            )
            enriched.append(current)
        return enriched

    def _narrow_claim_process_fact_matches(
        self,
        *,
        text: str,
        matched_entries: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        text_lower = str(text or "").lower()

        def _entry_category(entry: Dict[str, Any]) -> str:
            candidate = str(entry.get("text") or "").lower()
            if not _coerce_list(entry.get("fact_ids")):
                return ""
            if "notice to the applicant requires" in candidate or (
                "written notice" in candidate and "decision denying assistance" in candidate
            ):
                return "notice_requirement"
            if "informal review requires" in candidate or (
                "require" in candidate and "review" in candidate
            ):
                return "review_requirement"
            if "grievance request" in candidate or ("submitted" in candidate and "grievance" in candidate):
                return "review_request"
            if "review decision" in candidate or ("issued" in candidate and "review decision" in candidate):
                return "review_decision"
            if "denial notice" in candidate or ("sent plaintiff" in candidate and "notice" in candidate):
                return "notice_event"
            return ""

        desired_categories: List[str]
        if "required notice and review process" in text_lower:
            desired_categories = ["notice_requirement", "review_requirement", "review_request"]
        elif "requires written notice and an opportunity for informal review" in text_lower:
            desired_categories = ["notice_requirement", "review_requirement"]
        elif "informal review for denial of assistance" in text_lower:
            desired_categories = ["review_requirement", "review_request", "review_decision"]
        else:
            desired_categories = ["notice_requirement", "review_requirement", "review_request"]

        selected: List[Dict[str, Any]] = []
        seen = set()
        for category in desired_categories:
            for entry in matched_entries:
                entry_category = _entry_category(entry)
                if entry_category != category:
                    continue
                key = (
                    tuple(_normalize_identifier_list(entry.get("fact_ids"))),
                    str(entry.get("text") or ""),
                )
                if key in seen:
                    continue
                selected.append(dict(entry))
                seen.add(key)
                break
        return selected or matched_entries

    def _build_claim_chronology_support(self, *, claim_type: str, claim_name: str, limit: int = 2) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        if not isinstance(intake_case_file, dict):
            return []

        facts = [dict(item) for item in list(intake_case_file.get("canonical_facts") or []) if isinstance(item, dict)]
        relations = [dict(item) for item in list(intake_case_file.get("timeline_relations") or []) if isinstance(item, dict)]
        if not facts or not relations:
            return []

        combined = " ".join([str(claim_type or ""), str(claim_name or "")]).strip().lower()
        if any(token in combined for token in ("retaliat", "reprisal", "protected activity")):
            focus_families = {"protected_activity", "adverse_action", "causation"}
        elif any(token in combined for token in ("due process", "grievance", "hearing", "appeal", "review", "notice")):
            focus_families = {"notice_chain", "hearing_process", "response_timeline", "adverse_action"}
        elif any(token in combined for token in ("accommodation", "disabil", "fair housing", "discrimination", "termination", "denial")):
            focus_families = {"adverse_action", "response_timeline", "notice_chain"}
        else:
            focus_families = set()

        fact_by_id = {
            str(fact.get("fact_id") or "").strip(): fact
            for fact in facts
            if str(fact.get("fact_id") or "").strip()
        }
        relation_records = []
        for relation in relations:
            if str(relation.get("relation_type") or "").strip().lower() != "before":
                continue
            source_id = str(relation.get("source_fact_id") or "").strip()
            target_id = str(relation.get("target_fact_id") or "").strip()
            source_fact = fact_by_id.get(source_id)
            target_fact = fact_by_id.get(target_id)
            if not source_fact or not target_fact:
                continue
            source_date = _format_timeline_date((source_fact.get("temporal_context") or {}).get("start_date") or relation.get("source_start_date"))
            target_date = _format_timeline_date((target_fact.get("temporal_context") or {}).get("start_date") or relation.get("target_start_date"))
            if not source_date or not target_date:
                continue
            relation_records.append(
                {
                    "key": (source_id, target_id),
                    "source_id": source_id,
                    "target_id": target_id,
                    "source_fact": source_fact,
                    "target_fact": target_fact,
                    "source_date": source_date,
                    "target_date": target_date,
                    "source_family": str(source_fact.get("predicate_family") or "").strip().lower(),
                    "target_family": str(target_fact.get("predicate_family") or "").strip().lower(),
                }
            )
        if not relation_records:
            return []

        filtered_records = [
            record for record in relation_records
            if not focus_families or ({record['source_family'], record['target_family']} & focus_families)
        ]
        if not filtered_records:
            filtered_records = relation_records

        outgoing: Dict[str, List[Dict[str, Any]]] = {}
        incoming_count: Dict[str, int] = {}
        for record in filtered_records:
            outgoing.setdefault(record["source_id"], []).append(record)
            incoming_count[record["target_id"]] = incoming_count.get(record["target_id"], 0) + 1
            incoming_count.setdefault(record["source_id"], incoming_count.get(record["source_id"], 0))

        lines: List[str] = []
        fallback_lines: List[str] = []
        seen = set()
        used_keys = set()
        for record in filtered_records:
            if len(lines) >= limit:
                break
            if record["key"] in used_keys:
                continue
            if incoming_count.get(record["source_id"], 0) != 0 or len(outgoing.get(record["source_id"], [])) != 1:
                continue
            chain = [record]
            next_id = record["target_id"]
            temp_used = {record["key"]}
            while len(outgoing.get(next_id, [])) == 1 and incoming_count.get(next_id, 0) == 1:
                next_record = outgoing[next_id][0]
                if next_record["key"] in temp_used:
                    break
                chain.append(next_record)
                temp_used.add(next_record["key"])
                next_id = next_record["target_id"]
            if len(chain) < 2:
                continue
            segments = [
                f"{_chronology_fact_label(chain[0]['source_fact']).lower()} on {chain[0]['source_date']}"
            ]
            segments.extend(
                f"{_chronology_fact_label(item['target_fact']).lower()} on {item['target_date']}"
                for item in chain
            )
            line = f"The chronology shows {_join_chronology_segments(segments)} in sequence."
            last_target = chain[-1]["target_fact"]
            target_context = last_target.get("temporal_context") if isinstance(last_target.get("temporal_context"), dict) else {}
            if target_context.get("derived_from_relative_anchor"):
                relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
                if relative_markers:
                    line = line.rstrip(".") + f". The later date is derived from reported timing ({relative_markers[0]})."
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            used_keys.update(temp_used)
            lines.append(line)

        for record in filtered_records:
            if record["key"] in used_keys:
                continue
            source_label = _chronology_fact_label(record["source_fact"])
            target_label = _chronology_fact_label(record["target_fact"]).lower()
            line = f"The chronology shows {source_label.lower()} on {record['source_date']} before {target_label} on {record['target_date']}."
            target_context = record["target_fact"].get("temporal_context") if isinstance(record["target_fact"].get("temporal_context"), dict) else {}
            if target_context.get("derived_from_relative_anchor"):
                relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
                if relative_markers:
                    line = line.rstrip(".") + f". The later date is derived from reported timing ({relative_markers[0]})."
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            fallback_lines.append(line)
        if lines:
            return lines
        if fallback_lines:
            return fallback_lines[:1]
        return []

    def _claim_temporal_gap_focus(self, claim_type: str, claim_name: str) -> Dict[str, set[str]]:
        combined = " ".join([str(claim_type or ""), str(claim_name or "")]).strip().lower()
        issue_families = {"timeline"}
        element_tags = {"timeline"}
        objectives = {"timeline", "exact_dates"}
        if any(token in combined for token in ("retaliat", "reprisal", "protected activity")):
            issue_families.update({"causation", "protected_activity", "adverse_action"})
            element_tags.update({"causation", "protected_activity", "adverse_action"})
            objectives.update({"causation_link", "causation_sequence", "anchor_adverse_action"})
        if any(token in combined for token in ("due process", "grievance", "hearing", "appeal", "review", "notice")):
            issue_families.update({"notice_chain", "hearing_process", "response_timeline"})
            element_tags.update({"notice", "hearing", "appeal", "response", "review"})
            objectives.update({"anchor_appeal_rights", "hearing_request_timing", "response_dates"})
        if any(token in combined for token in ("accommodation", "disabil", "discrimination", "termination", "denial")):
            issue_families.update({"adverse_action", "notice_chain", "response_timeline"})
            element_tags.update({"adverse_action", "response", "notice"})
            objectives.update({"response_dates", "anchor_adverse_action"})
        return {
            "issue_families": issue_families,
            "element_tags": element_tags,
            "objectives": objectives,
        }

    def _build_claim_temporal_gap_hints(self, claim_type: str, claim_name: str, *, limit: int = 3) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        if not isinstance(intake_case_file, dict):
            return []

        focus = self._claim_temporal_gap_focus(claim_type, claim_name)
        normalized_claim_type = str(claim_type or "").strip().lower()
        hints: List[str] = []

        for issue in _coerce_list(intake_case_file.get("temporal_issue_registry")):
            if not isinstance(issue, dict):
                continue
            status = str(issue.get("status") or "open").strip().lower()
            if status not in {"open", "blocking", "warning"}:
                continue
            issue_claim_types = {str(item).strip().lower() for item in _coerce_list(issue.get("claim_types")) if str(item).strip()}
            issue_element_tags = {str(item).strip().lower() for item in _coerce_list(issue.get("element_tags")) if str(item).strip()}
            if issue_claim_types:
                if normalized_claim_type not in issue_claim_types:
                    continue
            elif issue_element_tags and not (issue_element_tags & focus["element_tags"]):
                continue
            summary = str(issue.get("summary") or "").strip()
            if summary:
                hints.append(f"Chronology gap: {summary}")

        blocker_summary = intake_case_file.get("blocker_follow_up_summary") if isinstance(intake_case_file.get("blocker_follow_up_summary"), dict) else {}
        for blocker in _coerce_list(blocker_summary.get("blocking_items")):
            if not isinstance(blocker, dict):
                continue
            issue_family = str(blocker.get("issue_family") or "").strip().lower()
            primary_objective = str(blocker.get("primary_objective") or "").strip().lower()
            blocker_objectives = {str(item).strip().lower() for item in _coerce_list(blocker.get("blocker_objectives")) if str(item).strip()}
            matched_objectives = ({primary_objective} if primary_objective else set()) | blocker_objectives
            if issue_family:
                if issue_family not in focus["issue_families"] and not (matched_objectives & focus["objectives"]):
                    continue
            elif not (matched_objectives & focus["objectives"]):
                continue
            reason = str(blocker.get("reason") or "").strip()
            if reason:
                hints.append(f"Chronology gap: {reason}")

        return _unique_preserving_order(hints)[:limit]

    def _annotate_claim_temporal_gap_hints(self, draft: Dict[str, Any]) -> None:
        claims = draft.get("claims_for_relief") if isinstance(draft.get("claims_for_relief"), list) else []
        for claim in claims:
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or claim.get("count_title") or "").strip()
            claim_name = str(claim.get("count_title") or claim.get("claim_type") or "").strip()
            hints = self._build_claim_temporal_gap_hints(claim_type, claim_name)
            if not hints:
                continue
            claim["missing_elements"] = _unique_preserving_order(
                list(_coerce_list(claim.get("missing_elements"))) + hints
            )
            support_summary = claim.get("support_summary") if isinstance(claim.get("support_summary"), dict) else {}
            claim["support_summary"] = {
                **support_summary,
                "temporal_gap_hint_count": len(hints),
            }

    def _build_claim_legal_standards(
        self,
        *,
        claim_type: str,
        requirements: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        support_claim: Optional[Dict[str, Any]] = None,
        related_exhibits: Optional[List[Dict[str, Any]]] = None,
    ) -> List[str]:
        standards = _unique_preserving_order(_extract_text_candidates(requirements.get(claim_type, [])))
        explicit_requirement_standards = bool(standards)
        related_statutes = self._select_statutes_for_claim(claim_type, statutes)
        for statute in related_statutes:
            citation = statute.get("citation")
            title = statute.get("title")
            relevance = statute.get("relevance")
            parts = [part for part in [citation, title, relevance] if part]
            if parts:
                standards.append(" - ".join(parts))
        if not explicit_requirement_standards:
            standards.extend(self._build_claim_legal_standard_fallbacks(claim_type))
        standards.extend(
            self._build_authority_backed_standard_lines(
                claim_type=claim_type,
                support_claim=support_claim or {},
                related_exhibits=related_exhibits or [],
            )
        )
        if standards:
            return _unique_preserving_order(standards)
        return self._build_claim_legal_standard_fallbacks(claim_type)

    def _build_authority_backed_standard_lines(
        self,
        *,
        claim_type: str,
        support_claim: Dict[str, Any],
        related_exhibits: List[Dict[str, Any]],
    ) -> List[str]:
        lines: List[str] = []
        seen = set()
        for exhibit in related_exhibits:
            if not isinstance(exhibit, dict):
                continue
            if str(exhibit.get("kind") or "").strip().lower() != "authority":
                continue
            title = str(exhibit.get("title") or "").strip()
            summary = str(exhibit.get("summary") or "").strip()
            if not title and not summary:
                continue
            if summary:
                line = f"{title}: {summary}" if title else summary
            else:
                line = title
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            lines.append(line)
            if len(lines) >= 2:
                return lines

        for element in _coerce_list((support_claim or {}).get("elements")):
            if not isinstance(element, dict):
                continue
            for link in _coerce_list(element.get("links")):
                if not isinstance(link, dict):
                    continue
                if str(link.get("support_kind") or "").strip().lower() != "authority":
                    continue
                citation = str(link.get("citation") or link.get("support_label") or link.get("title") or "").strip()
                relevance = str(link.get("relevance") or link.get("description") or "").strip()
                if not citation and not relevance:
                    continue
                line = f"{citation} - {relevance}" if citation and relevance else (citation or relevance)
                key = line.lower()
                if key in seen:
                    continue
                seen.add(key)
                lines.append(line)
                if len(lines) >= 2:
                    return lines
        return lines

    def _build_claim_legal_standard_fallbacks(self, claim_type: str) -> List[str]:
        normalized = normalize_claim_type(claim_type or "")
        if normalized == "housing_discrimination":
            return [
                "Plaintiff alleges that Defendant denied, limited, or otherwise interfered with housing assistance or related housing rights.",
                "Plaintiff further alleges that the challenged housing action was unlawful because Defendant enforced or maintained the denial of housing assistance without the notice, review, and fair treatment required by governing housing law and program rules.",
                "Plaintiff seeks relief for the resulting denial of housing opportunity, loss of assistance, and related harms caused by that unlawful housing decision.",
            ]
        if normalized == "due_process_failure":
            return [
                "Before enforcing a final adverse housing decision, Defendant was required to provide the written notice, review opportunity, hearing, grievance, appeal, or comparable process required by law or program rules.",
                "Plaintiff alleges that Defendant failed to provide the required written notice and meaningful review process before or while imposing the challenged denial or loss of assistance.",
                "Plaintiff seeks relief for the deprivation of housing benefits and review rights caused by that failure of notice and process.",
            ]
        if normalized == "retaliation":
            return [
                "Plaintiff alleges that Plaintiff engaged in protected activity and that Defendant thereafter took materially adverse action.",
                "Plaintiff further alleges that the adverse action was motivated, at least in part, by Plaintiff's protected conduct.",
            ]
        return [f"Plaintiff must prove the elements of {claim_type} under the applicable law."]

    def _build_legal_standards_summary(
        self,
        *,
        statutes: List[Dict[str, Any]],
        requirements: Dict[str, Any],
    ) -> List[str]:
        summary = []
        for claim_type, elements in requirements.items():
            summary.append(
                f"{claim_type.title()}: {', '.join(_unique_preserving_order(_extract_text_candidates(elements))[:4])}"
            )
        for statute in statutes[:5]:
            if isinstance(statute, dict):
                parts = [statute.get("citation"), statute.get("title"), statute.get("relevance")]
                text = " - ".join([part for part in parts if part])
                if text:
                    summary.append(text)
        return _unique_preserving_order(summary)

    def _safe_mediator_dict(self, method_name: str, **kwargs: Any) -> Dict[str, Any]:
        method = getattr(self.mediator, method_name, None)
        if not callable(method):
            return {}
        try:
            result = method(**kwargs)
        except Exception:
            return {}
        return result if isinstance(result, dict) else {}

    def _extract_blocker_follow_up_signals(self, optimization_report: Dict[str, Any]) -> Dict[str, Any]:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        intake_case_summary = (
            report.get("intake_case_summary")
            if isinstance(report.get("intake_case_summary"), dict)
            else build_intake_case_review_summary(self.mediator)
        )
        intake_case_summary = intake_case_summary if isinstance(intake_case_summary, dict) else {}
        blocker_follow_up_summary = (
            intake_case_summary.get("blocker_follow_up_summary")
            if isinstance(intake_case_summary.get("blocker_follow_up_summary"), dict)
            else {}
        )
        blocker_items = [
            dict(item)
            for item in list(blocker_follow_up_summary.get("blocking_items") or [])
            if isinstance(item, dict)
        ]
        open_items = [
            dict(item)
            for item in list(intake_case_summary.get("open_items") or [])
            if isinstance(item, dict) and str(item.get("kind") or "").strip().lower() == "blocker_follow_up"
        ]
        combined_items = blocker_items + open_items
        issue_types = _dedupe_text_values(
            str(item.get("issue_type") or item.get("type") or item.get("gap_type") or "").strip().lower()
            for item in combined_items
        )
        extraction_targets = _dedupe_text_values(
            list(blocker_follow_up_summary.get("extraction_targets") or [])
            + [
                target
                for item in combined_items
                for target in list(item.get("extraction_targets") or [])
            ]
        )
        workflow_phases = _dedupe_text_values(
            list(blocker_follow_up_summary.get("workflow_phases") or [])
            + [str(item.get("workflow_phase") or "").strip() for item in combined_items]
        )
        follow_up_questions = _dedupe_text_values(
            str(item.get("suggested_question") or item.get("next_question_template") or "").strip()
            for item in combined_items
        )

        def _item_has_confirmation_placeholder(item: Dict[str, Any]) -> bool:
            return any(
                _contains_confirmation_placeholder(value)
                for value in (
                    item.get("summary"),
                    item.get("reason"),
                    item.get("suggested_question"),
                    item.get("next_question_template"),
                    item.get("description"),
                )
            )

        placeholder_items = [
            item
            for item in combined_items
            if "confirmation_placeholder" in str(item.get("issue_type") or "").strip().lower()
            or _item_has_confirmation_placeholder(item)
        ]
        decision_maker_items = [
            item
            for item in combined_items
            if "decision_maker" in str(item.get("issue_type") or "").strip().lower()
            or any(
                target in {"decision_maker", "actor_name", "actor_role"}
                for target in list(item.get("extraction_targets") or [])
            )
        ]
        causation_items = [
            item
            for item in combined_items
            if (
                "causation" in str(item.get("issue_type") or "").strip().lower()
                or "retaliation_missing_sequence" in str(item.get("issue_type") or "").strip().lower()
                or "retaliation_missing_decision_maker" in str(item.get("issue_type") or "").strip().lower()
                or "causation_link" in list(item.get("extraction_targets") or [])
            )
        ]
        document_anchor_items = [
            item
            for item in combined_items
            if any(
                target in {"document_type", "document_date", "document_owner", "evidence_record", "verification_source"}
                for target in list(item.get("extraction_targets") or [])
            )
            or "document_anchor" in str(item.get("issue_type") or "").strip().lower()
        ]

        return {
            "blocker_count": int(blocker_follow_up_summary.get("blocking_item_count") or len(blocker_items) or 0),
            "issue_types": issue_types,
            "extraction_targets": extraction_targets,
            "workflow_phases": workflow_phases,
            "follow_up_questions": follow_up_questions,
            "confirmation_placeholder_count": len(placeholder_items),
            "decision_maker_probe_count": len(decision_maker_items),
            "causation_probe_count": len(causation_items),
            "document_anchor_probe_count": len(document_anchor_items),
            "needs_confirmation_follow_up": bool(placeholder_items),
            "needs_decision_maker_follow_up": bool(decision_maker_items),
            "needs_causation_follow_up": bool(causation_items),
            "needs_document_anchor_follow_up": bool(document_anchor_items),
        }

    def _extract_actor_critic_priority_metrics(self, optimization_report: Dict[str, Any]) -> Dict[str, float]:
        if not isinstance(optimization_report, dict):
            return {}
        final_review = optimization_report.get("final_review") if isinstance(optimization_report.get("final_review"), dict) else {}
        dimension_scores = final_review.get("dimension_scores") if isinstance(final_review.get("dimension_scores"), dict) else {}
        adversarial_batch = (
            optimization_report.get("adversarial_batch")
            if isinstance(optimization_report.get("adversarial_batch"), dict)
            else {}
        )
        metric_candidates = [
            optimization_report.get("actor_critic_metrics"),
            optimization_report.get("adversarial_batch_metrics"),
            optimization_report.get("latest_adversarial_batch_metrics"),
            optimization_report.get("priority_metrics"),
            optimization_report.get("baseline_metrics"),
            optimization_report.get("metrics"),
            adversarial_batch.get("metrics") if isinstance(adversarial_batch, dict) else {},
            (optimization_report.get("actor_critic_optimizer") or {}).get("metrics")
            if isinstance(optimization_report.get("actor_critic_optimizer"), dict)
            else {},
        ]
        normalized: Dict[str, float] = {}
        aliases = {
            "empathy": "empathy",
            "empathy_avg": "empathy",
            "avg_empathy": "empathy",
            "question_quality": "question_quality",
            "question_quality_avg": "question_quality",
            "avg_question_quality": "question_quality",
            "question_quality_score": "question_quality",
            "information_extraction": "information_extraction",
            "information_extraction_avg": "information_extraction",
            "avg_information_extraction": "information_extraction",
            "coverage": "coverage",
            "coverage_avg": "coverage",
            "avg_coverage": "coverage",
            "efficiency": "efficiency",
            "efficiency_avg": "efficiency",
            "avg_efficiency": "efficiency",
        }
        for candidate in metric_candidates:
            if not isinstance(candidate, dict):
                continue
            for key, value in candidate.items():
                metric_name = aliases.get(str(key).strip().lower())
                if not metric_name:
                    continue
                normalized[metric_name] = max(0.0, min(1.0, _safe_float(value, DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS.get(metric_name, 0.0))))
        if "question_quality" not in normalized:
            normalized["question_quality"] = max(
                _safe_float(dimension_scores.get("coherence"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["question_quality"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["question_quality"],
            )
        if "coverage" not in normalized:
            normalized["coverage"] = max(
                _safe_float(dimension_scores.get("completeness"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["coverage"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["coverage"],
            )
        if "information_extraction" not in normalized:
            normalized["information_extraction"] = max(
                _safe_float(dimension_scores.get("grounding"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["information_extraction"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["information_extraction"],
            )
        if "efficiency" not in normalized:
            normalized["efficiency"] = max(
                _safe_float(dimension_scores.get("procedural"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["efficiency"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["efficiency"],
            )
        if "empathy" not in normalized:
            normalized["empathy"] = DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["empathy"]
        return normalized

    def _extract_actor_critic_priority_value(self, optimization_report: Dict[str, Any]) -> int:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        priority_candidates = [
            report.get("priority"),
            report.get("actor_critic_priority"),
            (report.get("actor_critic_optimizer") or {}).get("priority")
            if isinstance(report.get("actor_critic_optimizer"), dict)
            else None,
        ]
        for candidate in priority_candidates:
            if candidate is None:
                continue
            try:
                return max(1, min(100, int(candidate)))
            except Exception:
                continue
        return 70

    def _resolve_router_backed_question_quality(self, optimization_report: Dict[str, Any]) -> bool:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        router_status = dict(report.get("router_status") or {})
        router_usage = dict(report.get("router_usage") or {})
        actor_critic_optimizer = (
            report.get("actor_critic_optimizer")
            if isinstance(report.get("actor_critic_optimizer"), dict)
            else {}
        )
        adversarial_batch = report.get("adversarial_batch") if isinstance(report.get("adversarial_batch"), dict) else {}
        candidate_flags = [
            str(router_status.get("llm_router") or "").strip().lower() == "available",
            _coerce_bool(router_status.get("llm_router_available"), default=False),
            _coerce_bool(router_status.get("available"), default=False),
            _coerce_bool(router_usage.get("llm_router_available"), default=False),
            _coerce_bool(router_usage.get("router_backed_question_quality"), default=False),
            _coerce_bool(report.get("router_backed_question_quality"), default=False),
            _coerce_bool(actor_critic_optimizer.get("router_backed_question_quality"), default=False),
            _coerce_bool(adversarial_batch.get("router_backed_question_quality"), default=False),
        ]
        return any(candidate_flags)

    def _extract_actor_critic_guidance(self, optimization_report: Dict[str, Any]) -> Dict[str, Any]:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        optimization_method = str(report.get("optimization_method") or "").strip().lower()
        method_name = str(report.get("method") or "").strip().lower()
        if "actor_critic" not in optimization_method and "actor_critic" not in method_name:
            return {}

        provided_order = [
            str(item).strip()
            for item in list(report.get("phase_focus_order") or [])
            if str(item).strip()
        ]
        phase_focus_order = [name for name in ACTOR_CRITIC_PHASE_FOCUS_ORDER if name in provided_order] or list(ACTOR_CRITIC_PHASE_FOCUS_ORDER)
        phase_focus_order.extend(name for name in provided_order if name not in phase_focus_order)
        final_review = report.get("final_review") if isinstance(report.get("final_review"), dict) else {}
        section_scores = final_review.get("section_scores") if isinstance(final_review.get("section_scores"), dict) else {}

        metrics = self._extract_actor_critic_priority_metrics(report)
        priority = self._extract_actor_critic_priority_value(report)
        intake_score = _safe_float(section_scores.get("intake_questioning"), 0.0)
        priority_findings = _dedupe_text_values(
            _extract_latest_adversarial_priority_findings(
                {
                    "priorities": report.get("priorities"),
                    "priority_findings": report.get("priority_findings"),
                    "latest_adversarial_batch": report.get("latest_adversarial_batch"),
                    "adversarial_batch": report.get("adversarial_batch"),
                    "latest_batch_priorities": report.get("latest_batch_priorities"),
                }
            )
        )
        if intake_score < 0.85 or priority >= 70:
            forced_order = [name for name in ACTOR_CRITIC_PHASE_FOCUS_ORDER if name in phase_focus_order]
            forced_order.extend(name for name in phase_focus_order if name not in forced_order)
            phase_focus_order = forced_order

        return {
            "optimization_method": optimization_method or "actor_critic",
            "phase_focus_order": phase_focus_order,
            "priority": priority,
            "metrics": metrics,
            "intake_questioning_score": intake_score,
            "router_backed_question_quality": self._resolve_router_backed_question_quality(report),
            "priority_findings": priority_findings,
            "needs_chronology_closure": _has_chronology_gap_priority(priority_findings),
            "needs_decision_document_precision": _has_decision_or_document_precision_priority(priority_findings),
        }

    def _extract_adversarial_session_flow_signals(self, optimization_report: Dict[str, Any]) -> Dict[str, Any]:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        if not report:
            return {
                "available": False,
                "successful_session_count": 0,
                "session_count": 0,
                "assessment_blocked": False,
            }

        adversarial_batch = report.get("adversarial_batch") if isinstance(report.get("adversarial_batch"), dict) else {}
        actor_critic_optimizer = (
            report.get("actor_critic_optimizer")
            if isinstance(report.get("actor_critic_optimizer"), dict)
            else {}
        )
        review_metadata = report.get("review_metadata") if isinstance(report.get("review_metadata"), dict) else {}

        def _first_count(candidates: List[Any], *, default: int = 0) -> int:
            for candidate in candidates:
                if candidate is None:
                    continue
                try:
                    if isinstance(candidate, list):
                        return max(0, len(candidate))
                    return max(0, int(candidate))
                except Exception:
                    continue
            return default

        successful_session_count = _first_count(
            [
                report.get("successful_session_count"),
                report.get("successful_sessions"),
                adversarial_batch.get("successful_session_count"),
                adversarial_batch.get("successful_sessions"),
                actor_critic_optimizer.get("successful_session_count"),
                actor_critic_optimizer.get("successful_sessions"),
                review_metadata.get("successful_session_count"),
                review_metadata.get("successful_sessions"),
                report.get("accepted_iterations"),
            ]
        )
        session_count = _first_count(
            [
                report.get("session_count"),
                report.get("total_session_count"),
                report.get("adversarial_session_count"),
                adversarial_batch.get("session_count"),
                adversarial_batch.get("total_session_count"),
                adversarial_batch.get("adversarial_session_count"),
                adversarial_batch.get("sessions"),
                actor_critic_optimizer.get("session_count"),
                actor_critic_optimizer.get("sessions"),
                report.get("section_history"),
            ]
        )
        has_adversarial_evidence = any(
            (
                bool(adversarial_batch),
                bool(report.get("latest_adversarial_batch")),
                bool(report.get("latest_batch_priorities")),
                bool(report.get("priority_findings")),
                bool(report.get("adversarial_batch_metrics")),
            )
        )
        available = bool(report)
        assessment_blocked = available and successful_session_count <= 0 and (
            session_count > 0 or has_adversarial_evidence
        )
        return {
            "available": available,
            "successful_session_count": successful_session_count,
            "session_count": session_count,
            "assessment_blocked": assessment_blocked,
        }

    def _refresh_drafting_readiness_workflow_warnings(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        workflow_phase_plan: Dict[str, Any],
    ) -> None:
        if not isinstance(drafting_readiness, dict):
            return
        all_existing_warnings = [
            dict(item) for item in list(drafting_readiness.get("warnings") or []) if isinstance(item, dict)
        ]
        prior_workflow_warning_count = sum(
            1
            for item in all_existing_warnings
            if str(item.get("code") or "").strip().lower().startswith("workflow_")
        )
        existing_warnings = [
            item
            for item in all_existing_warnings
            if not str(item.get("code") or "").strip().lower().startswith("workflow_")
        ]
        workflow_warnings = self._build_workflow_phase_warning_entries(workflow_phase_plan)
        combined_warnings = existing_warnings + workflow_warnings
        base_warning_count = max(
            0,
            int(drafting_readiness.get("warning_count") or 0) - prior_workflow_warning_count,
        )
        if combined_warnings:
            drafting_readiness["warnings"] = combined_warnings
            drafting_readiness["warning_count"] = base_warning_count + len(workflow_warnings)
            for warning in workflow_warnings:
                drafting_readiness["status"] = _merge_status(
                    str(drafting_readiness.get("status") or "ready"),
                    str(warning.get("severity") or "ready"),
                )
        elif "warnings" in drafting_readiness:
            drafting_readiness.pop("warnings", None)
            drafting_readiness["warning_count"] = base_warning_count

    def _build_runtime_workflow_phase_plan(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        document_optimization: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        if phase_manager is None and not optimization_report:
            return {}

        phases: Dict[str, Dict[str, Any]] = {}
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        adversarial_flow_signals = self._extract_adversarial_session_flow_signals(optimization_report)
        graph_phase = self._build_graph_analysis_phase_guidance(
            phase_manager,
            document_optimization=optimization_report,
        )
        if graph_phase:
            phases["graph_analysis"] = graph_phase

        document_phase = self._build_document_generation_phase_guidance(
            drafting_readiness=drafting_readiness,
            document_optimization=optimization_report,
        )
        graph_status = str(graph_phase.get("status") or "ready").strip().lower() if graph_phase else "ready"
        graph_signals = (
            dict(graph_phase.get("signals") or {})
            if isinstance(graph_phase.get("signals"), dict)
            else {}
        )
        readiness_graph_signals = (
            dict(drafting_readiness.get("graph_completeness_signals") or {})
            if isinstance(drafting_readiness.get("graph_completeness_signals"), dict)
            else {}
        )
        unresolved_factual_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_factual_gaps") or [])
            if str(item).strip()
        ]
        unresolved_legal_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_legal_gaps") or [])
            if str(item).strip()
        ]
        uncovered_intake_objectives = [
            str(item).strip()
            for item in list(drafting_readiness.get("uncovered_intake_objectives") or [])
            if str(item).strip()
        ]
        missing_required_intake_objectives = [
            str(item).strip()
            for item in list(drafting_readiness.get("missing_required_intake_objectives") or [])
            if str(item).strip()
        ]
        structured_handoff_signals = (
            dict(drafting_readiness.get("structured_intake_handoff_signals") or {})
            if isinstance(drafting_readiness.get("structured_intake_handoff_signals"), dict)
            else {}
        )
        structured_handoff_gap_count = int(structured_handoff_signals.get("gap_count", 0) or 0)
        weak_complaint_types = [
            str(item).strip()
            for item in list(drafting_readiness.get("weak_complaint_types") or [])
            if str(item).strip()
        ]
        weak_evidence_modalities = [
            str(item).strip()
            for item in list(drafting_readiness.get("weak_evidence_modalities") or [])
            if str(item).strip()
        ]
        targeted_weak_complaint_types = [
            item
            for item in weak_complaint_types
            if item.lower() in {"housing_discrimination", "hacc_research_engine"}
        ]
        targeted_weak_evidence_modalities = [
            item
            for item in weak_evidence_modalities
            if item.lower() in {"policy_document", "file_evidence"}
        ]
        graph_remaining_gap_count = max(
            int(graph_signals.get("remaining_gap_count", 0) or 0),
            int(graph_signals.get("current_gap_count", 0) or 0),
            int(readiness_graph_signals.get("remaining_gap_count", 0) or 0),
            int(readiness_graph_signals.get("current_gap_count", 0) or 0),
        )
        graph_gate_active = (
            graph_status != "ready"
            or graph_remaining_gap_count > 0
            or not _coerce_bool(
                readiness_graph_signals.get("knowledge_graph_available", graph_signals.get("knowledge_graph_available", True)),
                default=True,
            )
            or not _coerce_bool(
                readiness_graph_signals.get("dependency_graph_available", graph_signals.get("dependency_graph_available", True)),
                default=True,
            )
            )
        if document_phase:
            updated_document_phase = dict(document_phase)
            actions = [str(item).strip() for item in list(updated_document_phase.get("recommended_actions") or []) if str(item).strip()]
            handoff_gap_active = bool(
                unresolved_factual_gaps
                or unresolved_legal_gaps
                or uncovered_intake_objectives
                or missing_required_intake_objectives
                or structured_handoff_gap_count > 0
                or targeted_weak_complaint_types
                or targeted_weak_evidence_modalities
            )
            if graph_gate_active:
                gate_status = "blocked"
                updated_document_phase["status"] = _merge_status(
                    str(updated_document_phase.get("status") or "ready"),
                    gate_status,
                )
                summary = str(updated_document_phase.get("summary") or "").strip()
                gate_summary = (
                    "Document generation is gated on graph completeness and should not be treated as final until graph blockers are resolved."
                )
                updated_document_phase["summary"] = f"{summary} {gate_summary}".strip()
                actions.append(
                    "Resolve graph completeness blockers (knowledge/dependency graph availability and unresolved graph gaps) before formalization."
                )
            if handoff_gap_active:
                updated_document_phase["status"] = _merge_status(
                    str(updated_document_phase.get("status") or "ready"),
                    "warning",
                )
                summary = str(updated_document_phase.get("summary") or "").strip()
                handoff_summary = (
                    "Drafting handoff still has unresolved factual/legal/objective gaps that should be closed before formalization."
                )
                updated_document_phase["summary"] = f"{summary} {handoff_summary}".strip()
            if unresolved_factual_gaps:
                actions.append(
                    "Close unresolved factual gaps before formalization: "
                    + "; ".join(unresolved_factual_gaps[:3])
                )
            if unresolved_legal_gaps:
                actions.append(
                    "Close unresolved legal gaps before formalization: "
                    + "; ".join(unresolved_legal_gaps[:3])
                )
            if uncovered_intake_objectives or missing_required_intake_objectives:
                actions.append(
                    "Resolve uncovered intake objectives before formalization: "
                    + ", ".join((missing_required_intake_objectives or uncovered_intake_objectives)[:4])
                )
            if structured_handoff_gap_count > 0:
                actions.append(
                    "Promote structured intake facts, date/actor anchors, and evidence references directly into summary-of-facts and claim-support paragraphs before document optimization runs."
                )
            if targeted_weak_complaint_types:
                actions.append(
                    "Generalize drafting quality for weak complaint types before formalization: "
                    + ", ".join(targeted_weak_complaint_types)
                )
            if targeted_weak_evidence_modalities:
                actions.append(
                    "Strengthen allegation support for weak evidence modalities before formalization: "
                    + ", ".join(targeted_weak_evidence_modalities)
                )
            if bool(adversarial_flow_signals.get("assessment_blocked")):
                gate_status = "blocked" if str(drafting_readiness.get("phase_status") or "").strip().lower() == "critical" else "warning"
                updated_document_phase["status"] = _merge_status(
                    str(updated_document_phase.get("status") or "ready"),
                    gate_status,
                )
                summary = str(updated_document_phase.get("summary") or "").strip()
                session_summary = (
                    "No successful adversarial sessions were available to assess drafting handoff quality."
                )
                updated_document_phase["summary"] = f"{summary} {session_summary}".strip()
                actions.append(
                    "Restore a stable adversarial session flow before tuning document-generation handoffs."
                )
            updated_document_phase["recommended_actions"] = _dedupe_text_values(actions)
            signals = dict(updated_document_phase.get("signals") or {})
            signals["gate_on_graph_completeness"] = bool(graph_gate_active)
            signals["graph_phase_status"] = graph_status
            signals["graph_remaining_gap_count"] = int(graph_remaining_gap_count)
            signals["drafting_coverage"] = _safe_float(drafting_readiness.get("coverage"), 0.0)
            signals["drafting_phase_status"] = str(drafting_readiness.get("phase_status") or "ready").strip().lower() or "ready"
            signals["unresolved_factual_gap_count"] = len(unresolved_factual_gaps)
            signals["unresolved_legal_gap_count"] = len(unresolved_legal_gaps)
            signals["uncovered_intake_objective_count"] = len(uncovered_intake_objectives)
            signals["missing_required_intake_objective_count"] = len(missing_required_intake_objectives)
            signals["structured_intake_handoff_gap_count"] = int(structured_handoff_gap_count)
            signals["uncovered_intake_objectives"] = uncovered_intake_objectives[:8]
            signals["unresolved_factual_gaps"] = unresolved_factual_gaps[:6]
            signals["unresolved_legal_gaps"] = unresolved_legal_gaps[:6]
            signals["weak_complaint_types"] = weak_complaint_types
            signals["weak_evidence_modalities"] = weak_evidence_modalities
            signals["targeted_weak_complaint_types"] = targeted_weak_complaint_types
            signals["targeted_weak_evidence_modalities"] = targeted_weak_evidence_modalities
            signals["drafting_handoff_gap_active"] = bool(handoff_gap_active)
            signals["drafting_handoff_gap_count"] = (
                len(unresolved_factual_gaps)
                + len(unresolved_legal_gaps)
                + len(uncovered_intake_objectives)
                + len(missing_required_intake_objectives)
                + int(structured_handoff_gap_count > 0)
                + len(targeted_weak_complaint_types)
                + len(targeted_weak_evidence_modalities)
            )
            signals["ready_for_formalization"] = not bool(
                graph_gate_active
                or unresolved_factual_gaps
                or unresolved_legal_gaps
                or uncovered_intake_objectives
                or missing_required_intake_objectives
                or structured_handoff_gap_count > 0
                or targeted_weak_complaint_types
                or targeted_weak_evidence_modalities
            )
            signals["adversarial_session_flow_available"] = bool(adversarial_flow_signals.get("available"))
            signals["adversarial_session_count"] = int(adversarial_flow_signals.get("session_count") or 0)
            signals["adversarial_successful_session_count"] = int(
                adversarial_flow_signals.get("successful_session_count") or 0
            )
            signals["adversarial_session_flow_stable"] = not bool(adversarial_flow_signals.get("assessment_blocked"))
            updated_document_phase["signals"] = signals
            phases["document_generation"] = updated_document_phase
        intake_phase = self._build_intake_questioning_phase_guidance(
            drafting_readiness=drafting_readiness,
            document_optimization=optimization_report,
        )
        if intake_phase:
            phases["intake_questioning"] = intake_phase

        plan = build_workflow_phase_plan(phases)
        if not plan:
            return {}
        preferred_order = (
            list(actor_critic_guidance.get("phase_focus_order") or [])
            if actor_critic_guidance
            else ["graph_analysis", "document_generation", "intake_questioning"]
        )
        if bool(adversarial_flow_signals.get("assessment_blocked")):
            preferred_order = ["graph_analysis", "intake_questioning", "document_generation"]
        ordered = [name for name in preferred_order if name in phases]
        ordered.extend(
            name
            for name in list(plan.get("recommended_order") or [])
            if name in phases and name not in ordered
        )
        plan["recommended_order"] = ordered
        if actor_critic_guidance:
            plan["actor_critic_guidance"] = actor_critic_guidance
        return plan

    def _build_graph_analysis_phase_guidance(
        self,
        phase_manager: Any,
        *,
        document_optimization: Dict[str, Any],
    ) -> Dict[str, Any]:
        phase = build_graph_analysis_phase_guidance(phase_manager, audience="drafting")
        if not phase:
            return {}
        updated = dict(phase)
        signals = dict(updated.get("signals") or {})
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        temporal_handoff = (
            optimization_report.get("claim_support_temporal_handoff")
            if isinstance(optimization_report.get("claim_support_temporal_handoff"), dict)
            else {}
        )
        unresolved_temporal_count = int(temporal_handoff.get("unresolved_temporal_issue_count", 0) or 0)
        chronology_tasks = int(temporal_handoff.get("chronology_task_count", 0) or 0)
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        blocker_signals = self._extract_blocker_follow_up_signals(optimization_report)
        needs_chronology_closure = bool(actor_critic_guidance.get("needs_chronology_closure")) or bool(
            blocker_signals.get("needs_causation_follow_up")
        )
        if unresolved_temporal_count > 0 or chronology_tasks > 0:
            updated["status"] = "warning" if str(updated.get("status") or "").lower() == "ready" else updated.get("status")
            summary = str(updated.get("summary") or "").strip()
            suffix = (
                f" Temporal graph alignment still has {unresolved_temporal_count} unresolved chronology issue(s) "
                f"across {chronology_tasks} chronology task(s)."
            )
            updated["summary"] = f"{summary}{suffix}".strip()
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Resolve chronology edges for protected activity, hearing/review requests, response dates, and adverse-action outcomes before finalizing the complaint timeline."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        if needs_chronology_closure:
            updated["status"] = "warning" if str(updated.get("status") or "").lower() == "ready" else updated.get("status")
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Close critical chronology gaps by confirming exact date anchors for protected activity, each notice/response event, and each adverse-action step in sequence."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        signals["unresolved_temporal_issue_count"] = unresolved_temporal_count
        signals["chronology_task_count"] = chronology_tasks
        signals["needs_chronology_closure"] = needs_chronology_closure
        signals["decision_maker_probe_count"] = int(blocker_signals.get("decision_maker_probe_count", 0) or 0)
        signals["document_anchor_probe_count"] = int(blocker_signals.get("document_anchor_probe_count", 0) or 0)
        updated["signals"] = signals
        return updated

    def _build_document_generation_phase_guidance(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        document_optimization: Dict[str, Any],
    ) -> Dict[str, Any]:
        phase = build_drafting_document_generation_phase_guidance(
            drafting_readiness=drafting_readiness,
            document_optimization=document_optimization,
        )
        if not phase:
            return {}
        updated = dict(phase)
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        blocker_signals = self._extract_blocker_follow_up_signals(optimization_report)
        final_review = optimization_report.get("final_review") if isinstance(optimization_report.get("final_review"), dict) else {}
        section_scores = final_review.get("section_scores") if isinstance(final_review.get("section_scores"), dict) else {}
        intake_score = float(section_scores.get("intake_questioning") or 0.0)
        if intake_score < 0.8:
            updated["status"] = "warning" if str(updated.get("status") or "").lower() == "ready" else updated.get("status")
            summary = str(updated.get("summary") or "").strip()
            updated["summary"] = (
                f"{summary} Document generation should preserve patchability while improving fact sequencing for retaliation causation and adverse-action chronology."
            ).strip()
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Strengthen factual paragraphs so each adverse action is paired with exact date anchors, named/titled staff actors, hearing-request timing, response dates, and causation sequencing."
            )
            actions.append(
                "Preserve patchability by keeping one material fact per sentence and explicitly anchoring each sentence to dates, actors, and response events extracted during intake questioning."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        if actor_critic_guidance and bool(actor_critic_guidance.get("router_backed_question_quality")):
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Use router-backed drafting passes to convert unresolved intake objectives into concrete chronology-aligned allegations and claim-support paragraphs."
            )
            actions.append(
                "Preserve patchability by emitting single-sentence allegation units with explicit date, actor/title, and source anchors for each router-refined paragraph."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        if bool(actor_critic_guidance.get("needs_decision_document_precision")) or bool(
            blocker_signals.get("needs_document_anchor_follow_up")
        ):
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Increase precision in adverse-action allegations by naming the decision-maker (or known title), the specific decision communicated, and the controlling documentary artifact for each step."
            )
            actions.append(
                "Keep each allegation patchable by isolating one decision event per sentence with explicit fields for date, actor/title, adverse action detail, and source document anchor."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        signals = dict(updated.get("signals") or {})
        signals["intake_questioning_score"] = intake_score
        if actor_critic_guidance:
            signals["actor_critic_priority"] = int(actor_critic_guidance.get("priority") or 70)
            signals["router_backed_question_quality"] = bool(actor_critic_guidance.get("router_backed_question_quality"))
            signals["needs_decision_document_precision"] = bool(
                actor_critic_guidance.get("needs_decision_document_precision")
            )
        signals["decision_maker_probe_count"] = int(blocker_signals.get("decision_maker_probe_count", 0) or 0)
        signals["document_anchor_probe_count"] = int(blocker_signals.get("document_anchor_probe_count", 0) or 0)
        updated["signals"] = signals
        return updated

    def _build_intake_questioning_phase_guidance(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        document_optimization: Dict[str, Any],
    ) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        actor_critic_metrics = dict(actor_critic_guidance.get("metrics") or {}) if actor_critic_guidance else {}
        blocker_signals = self._extract_blocker_follow_up_signals(optimization_report)
        needs_chronology_closure = bool(actor_critic_guidance.get("needs_chronology_closure")) or bool(
            blocker_signals.get("needs_causation_follow_up")
        )
        needs_decision_document_precision = bool(actor_critic_guidance.get("needs_decision_document_precision")) or bool(
            blocker_signals.get("needs_decision_maker_follow_up")
        ) or bool(blocker_signals.get("needs_document_anchor_follow_up"))
        intake_status = (
            optimization_report.get("intake_status")
            if isinstance(optimization_report.get("intake_status"), dict)
            else build_intake_status_summary(self.mediator)
        )
        intake_handoff = (
            intake_status.get("intake_summary_handoff")
            if isinstance(intake_status.get("intake_summary_handoff"), dict)
            else {}
        )
        confirmation = (
            intake_handoff.get("complainant_summary_confirmation")
            if isinstance(intake_handoff.get("complainant_summary_confirmation"), dict)
            else {}
        )
        confirmation_snapshot = (
            confirmation.get("confirmed_summary_snapshot")
            if isinstance(confirmation.get("confirmed_summary_snapshot"), dict)
            else {}
        )
        intake_priority_summary = (
            confirmation_snapshot.get("adversarial_intake_priority_summary")
            if isinstance(confirmation_snapshot.get("adversarial_intake_priority_summary"), dict)
            else {}
        )
        objective_aliases = {
            "staff_names": "staff_names_titles",
            "staff_titles": "staff_names_titles",
            "hearing_timing": "hearing_request_timing",
            "response_timing": "response_dates",
            "causation": "causation_link",
            "adverse_action": "anchor_adverse_action",
            "appeal_rights": "anchor_appeal_rights",
        }

        def _normalize_objective(value: Any) -> str:
            text = str(value or "").strip().lower()
            if not text:
                return ""
            return objective_aliases.get(text, text)

        uncovered_objectives = _dedupe_text_values(
            _normalize_objective(item)
            for item in (intake_priority_summary.get("uncovered_objectives") or [])
            if _normalize_objective(item)
        )
        objective_question_counts = {
            _normalize_objective(key): int(value or 0)
            for key, value in dict(intake_priority_summary.get("objective_question_counts") or {}).items()
            if _normalize_objective(key)
        }
        required_objectives = (
            "timeline",
            "actors",
            "staff_names_titles",
            "causation_link",
            "anchor_adverse_action",
            "anchor_appeal_rights",
            "hearing_request_timing",
            "response_dates",
        )
        missing_required = [
            objective
            for objective in required_objectives
            if objective in uncovered_objectives or int(objective_question_counts.get(objective, 0)) <= 0
        ]
        objective_actions = {
            "timeline": "Capture exact dates for complaint activity, notices, hearing/review requests, and adverse action outcomes.",
            "actors": "Identify who at HACC made, communicated, and carried out each key decision.",
            "staff_names_titles": "Capture each HACC staff member name and title, or best-known title when name is unknown.",
            "causation_link": "Document direct causation facts linking protected activity to adverse action.",
            "anchor_adverse_action": "Confirm the exact denial, termination, or threatened loss of assistance and its communication date.",
            "anchor_appeal_rights": "Confirm whether written notice, informal review, grievance hearing, and appeal rights were provided, requested, denied, or ignored.",
            "hearing_request_timing": "Capture when hearing/review was requested and whether the request timing was acknowledged.",
            "response_dates": "Capture exact response dates for HACC notices, hearing/review responses, and final decision outcomes.",
        }
        if missing_required:
            status = "warning"
            summary = (
                "Intake questioning still needs closure on key blockers (exact dates, staff names/titles, hearing request timing, response dates, "
                "and causation links between protected activity and adverse treatment)."
            )
            actions = _dedupe_text_values(
                objective_actions.get(objective, "")
                for objective in missing_required
                if objective_actions.get(objective, "")
            )
        else:
            status = "ready"
            summary = (
                "Intake questioning currently includes timeline anchors, staff identification, hearing/response timing, and causation probes."
            )
            actions = []

        empathy_score = _safe_float(actor_critic_metrics.get("empathy"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["empathy"])
        question_quality_score = _safe_float(actor_critic_metrics.get("question_quality"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["question_quality"])
        information_extraction_score = _safe_float(actor_critic_metrics.get("information_extraction"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["information_extraction"])
        coverage_score = _safe_float(actor_critic_metrics.get("coverage"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["coverage"])
        efficiency_score = _safe_float(actor_critic_metrics.get("efficiency"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["efficiency"])
        router_backed_question_quality = bool(actor_critic_guidance.get("router_backed_question_quality")) if actor_critic_guidance else False

        if empathy_score < 0.5:
            actions.append(
                "Lead each intake follow-up with a short empathy frame before asking for chronology or decision details."
            )
            actions.append(
                "Use an empathy-forward transition format ('I hear the impact this had. To keep your record accurate...') before requesting exact dates or staff identifiers."
            )
        if question_quality_score < 0.7:
            actions.append(
                "Upgrade question quality by asking one targeted objective per question with explicit date, actor/title, and decision-anchor prompts."
            )
            actions.append(
                "Require each follow-up question to include one objective tag plus a concrete answer target (date, actor/title, notice/response event, or quoted decision language)."
            )
        if information_extraction_score < 0.7:
            actions.append(
                "Improve information extraction by capturing exact event sequence, communication channel, and quoted decision language where available."
            )
        if coverage_score < 0.7:
            actions.append(
                "Increase coverage by ensuring every required objective has at least one answered question and no zero-count objective remains."
            )
        if efficiency_score < 0.75:
            actions.append(
                "Improve intake efficiency by removing multi-part questions and collapsing duplicates to the shortest sufficient prompt."
            )
        if router_backed_question_quality:
            actions.append(
                "Use router-backed question generation to refine prompts for specificity while preserving patchability and objective coverage."
            )
            if question_quality_score < 0.75:
                actions.append(
                    "Route low-quality follow-ups through the llm router with an objective-specific schema so each prompt asks for one verifiable fact and one temporal anchor."
                )
            if information_extraction_score < 0.75:
                actions.append(
                    "Use router-backed extraction checks to normalize each answer into date, actor/title, event, and source fields before drafting updates."
                )
        if needs_chronology_closure:
            actions.append(
                "Close chronology gaps with follow-up questions that require exact dates, elapsed response timing, and event sequence ordering before drafting updates."
            )
            if router_backed_question_quality:
                actions.append(
                    "Use router-backed single-objective prompts to ask one chronology gap at a time (event date, response date, or sequence delta) and reject answers without explicit date anchors."
                )
        if needs_decision_document_precision:
            actions.append(
                "Ask decision-precision follow-ups that capture who made each adverse decision, their title, the exact decision communicated, and the document or notice that records it."
            )
            if router_backed_question_quality:
                actions.append(
                    "Use router-backed extraction checks to normalize adverse-action answers into decision-maker, decision detail, communication channel, document artifact, and date fields."
                )
        actions.append(
            "Keep follow-up prompts patchable by using short, single-objective question templates that can be edited independently without changing global intake flow."
        )
        actions = _dedupe_text_values(actions)
        if actions and status == "ready":
            status = "warning"
        if actions and actor_critic_guidance:
            summary = (
                f"{summary} Actor-critic optimization highlights low empathy/question quality/extraction/coverage/efficiency signals that should be addressed before final drafting."
            ).strip()

        return {
            "priority": 2,
            "status": status,
            "summary": summary,
            "signals": {
                "uncovered_objectives": uncovered_objectives,
                "objective_question_counts": objective_question_counts,
                "missing_required_objectives": missing_required,
                "drafting_status": str(drafting_readiness.get("status") or ""),
                "actor_critic_priority": int(actor_critic_guidance.get("priority") or 70) if actor_critic_guidance else None,
                "actor_critic_metrics": actor_critic_metrics if actor_critic_guidance else {},
                "router_backed_question_quality": router_backed_question_quality,
                "blocker_follow_up_signals": blocker_signals,
                "needs_chronology_closure": needs_chronology_closure,
                "needs_decision_document_precision": needs_decision_document_precision,
                "priority_findings": list(actor_critic_guidance.get("priority_findings") or []) if actor_critic_guidance else [],
            },
            "recommended_actions": actions,
        }

    def _build_workflow_phase_warning_entries(self, workflow_phase_plan: Dict[str, Any]) -> List[Dict[str, Any]]:
        return build_workflow_phase_warning_entries(workflow_phase_plan)

    def _build_formalization_gate_payload(self, drafting_readiness: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(drafting_readiness, dict):
            return {}
        phase_status = str(drafting_readiness.get("phase_status") or drafting_readiness.get("status") or "ready").strip().lower() or "ready"
        blockers = [
            str(item).strip()
            for item in list(drafting_readiness.get("blockers") or [])
            if str(item).strip()
        ]
        unresolved_factual_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_factual_gaps") or [])
            if str(item).strip()
        ]
        unresolved_legal_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_legal_gaps") or [])
            if str(item).strip()
        ]
        return {
            "ready_for_formalization": phase_status == "ready" and not blockers,
            "phase_status": phase_status,
            "coverage": _safe_float(drafting_readiness.get("coverage"), 0.0),
            "blockers": blockers,
            "unresolved_factual_gaps": unresolved_factual_gaps[:6],
            "unresolved_legal_gaps": unresolved_legal_gaps[:6],
            "weak_complaint_types": list(drafting_readiness.get("weak_complaint_types") or []),
            "weak_evidence_modalities": list(drafting_readiness.get("weak_evidence_modalities") or []),
            "document_fact_backed_ratio": _safe_float(drafting_readiness.get("document_fact_backed_ratio"), 0.0),
            "document_low_grounding_flag": bool(drafting_readiness.get("document_low_grounding_flag")),
        }

    def _build_graph_completeness_signals(self, phase_manager: Any) -> Dict[str, Any]:
        if phase_manager is None:
            return {
                "status": "warning",
                "knowledge_graph_available": False,
                "dependency_graph_available": False,
                "remaining_gap_count": 0,
                "current_gap_count": 0,
                "knowledge_graph_enhanced": False,
            }
        graph_phase = build_graph_analysis_phase_guidance(phase_manager, audience="drafting")
        signals = dict(graph_phase.get("signals") or {}) if isinstance(graph_phase, dict) else {}
        return {
            "status": str(graph_phase.get("status") or "warning").strip().lower() if isinstance(graph_phase, dict) else "warning",
            "knowledge_graph_available": _coerce_bool(signals.get("knowledge_graph_available"), default=False),
            "dependency_graph_available": _coerce_bool(signals.get("dependency_graph_available"), default=False),
            "remaining_gap_count": int(signals.get("remaining_gap_count", 0) or 0),
            "current_gap_count": int(signals.get("current_gap_count", 0) or 0),
            "knowledge_graph_enhanced": _coerce_bool(signals.get("knowledge_graph_enhanced"), default=False),
        }

    def _collect_evidence_modality_signals(
        self,
        *,
        draft: Dict[str, Any],
        claim_readiness: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        policy_count = 0
        file_count = 0

        for claim in claim_readiness:
            if not isinstance(claim, dict):
                continue
            for key, value in dict(claim.get("artifact_family_counts") or {}).items():
                count = int(value or 0)
                token = str(key or "").strip().lower()
                if count <= 0:
                    continue
                if "policy" in token:
                    policy_count += count
                if any(marker in token for marker in ("file", "upload", "document", "pdf", "exhibit", "record")):
                    file_count += count
            for key, value in dict(claim.get("support_by_source") or {}).items():
                count = int(value or 0)
                token = str(key or "").strip().lower()
                if count <= 0:
                    continue
                if "policy_document" in token:
                    policy_count += count
                if "file_evidence" in token:
                    file_count += count

        for exhibit in _coerce_list(draft.get("exhibits")):
            if not isinstance(exhibit, dict):
                continue
            type_text = " ".join(
                str(exhibit.get(key) or "").strip().lower()
                for key in ("type", "evidence_type", "source_type", "title", "description", "link", "path")
            )
            if any(marker in type_text for marker in ("policy", "administrative plan", "acop", "grievance")):
                policy_count += 1
            if any(marker in type_text for marker in ("file", "upload", ".pdf", "exhibit", "record", "document")):
                file_count += 1

        weak_modalities: List[str] = []
        if policy_count <= 0:
            weak_modalities.append("policy_document")
        if file_count <= 0:
            weak_modalities.append("file_evidence")
        return {
            "modalities": {
                "policy_document": int(policy_count),
                "file_evidence": int(file_count),
            },
            "weak_modalities": weak_modalities,
        }

    def _collect_unresolved_readiness_gaps(
        self,
        *,
        claim_readiness: List[Dict[str, Any]],
        sections: Dict[str, Dict[str, Any]],
        graph_signals: Dict[str, Any],
    ) -> Dict[str, List[str]]:
        factual_gaps: List[str] = []
        legal_gaps: List[str] = []

        for claim in claim_readiness:
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or "claim").strip()
            for warning in _coerce_list(claim.get("warnings")):
                if not isinstance(warning, dict):
                    continue
                code = str(warning.get("code") or "").strip().lower()
                message = " ".join(str(warning.get("message") or "").split()).strip()
                if not message:
                    continue
                if code in {"proof_gaps_present", "unresolved_elements", "claim_contradicted"}:
                    factual_gaps.append(f"{claim_type}: {message}")
                elif code in {"adverse_authority_present", "authority_reliability_uncertain"}:
                    legal_gaps.append(f"{claim_type}: {message}")
                else:
                    factual_gaps.append(f"{claim_type}: {message}")

        for section_key, section in sections.items():
            if not isinstance(section, dict):
                continue
            section_name = str(section.get("title") or section_key or "section").strip()
            for warning in _coerce_list(section.get("warnings")):
                if not isinstance(warning, dict):
                    continue
                code = str(warning.get("code") or "").strip().lower()
                message = " ".join(str(warning.get("message") or "").split()).strip()
                if not message:
                    continue
                if code in {"procedural_prerequisites_identified", "jurisdiction_or_venue_incomplete"}:
                    legal_gaps.append(f"{section_name}: {message}")
                else:
                    factual_gaps.append(f"{section_name}: {message}")

        graph_status = str(graph_signals.get("status") or "ready").strip().lower()
        graph_gap_count = max(
            int(graph_signals.get("remaining_gap_count", 0) or 0),
            int(graph_signals.get("current_gap_count", 0) or 0),
        )
        if graph_status != "ready":
            factual_gaps.append(
                "Graph analysis is not ready and chronology/cross-claim dependencies should be closed before formalization."
            )
        if graph_gap_count > 0:
            factual_gaps.append(
                f"Graph analysis still has {graph_gap_count} unresolved intake gap(s) that affect drafting completeness."
            )

        return {
            "factual_gaps": _dedupe_text_values(factual_gaps)[:6],
            "legal_gaps": _dedupe_text_values(legal_gaps)[:6],
        }

    def _build_drafting_readiness(
        self,
        *,
        user_id: str,
        draft: Dict[str, Any],
    ) -> Dict[str, Any]:
        support_summary = self._safe_mediator_dict("summarize_claim_support", user_id=user_id)
        gap_summary = self._safe_mediator_dict("get_claim_support_gaps", user_id=user_id)
        validation_summary = self._safe_mediator_dict("get_claim_support_validation", user_id=user_id)

        support_claims = support_summary.get("claims", {}) if isinstance(support_summary.get("claims"), dict) else {}
        gap_claims = gap_summary.get("claims", {}) if isinstance(gap_summary.get("claims"), dict) else {}
        validation_claims = validation_summary.get("claims", {}) if isinstance(validation_summary.get("claims"), dict) else {}

        claim_types = _unique_preserving_order(
            _extract_text_candidates((draft.get("source_context") or {}).get("claim_types"))
            + list(support_claims.keys())
            + list(validation_claims.keys())
            + [
                str(claim.get("claim_type") or "").strip()
                for claim in _coerce_list(draft.get("claims_for_relief"))
                if isinstance(claim, dict)
            ]
        )

        claim_readiness: List[Dict[str, Any]] = []
        aggregate_warning_count = 0
        overall_status = "ready"
        draft_claims_by_type = {
            str(claim.get("claim_type") or "").strip(): claim
            for claim in _coerce_list(draft.get("claims_for_relief"))
            if isinstance(claim, dict) and str(claim.get("claim_type") or "").strip()
        }

        for claim_type in claim_types:
            support_claim = support_claims.get(claim_type, {}) if isinstance(support_claims.get(claim_type), dict) else {}
            gap_claim = gap_claims.get(claim_type, {}) if isinstance(gap_claims.get(claim_type), dict) else {}
            validation_claim = validation_claims.get(claim_type, {}) if isinstance(validation_claims.get(claim_type), dict) else {}
            if not isinstance(validation_claim, dict) or not validation_claim:
                claim_validation_payload = self._safe_mediator_dict(
                    "get_claim_support_validation",
                    claim_type=claim_type,
                    user_id=user_id,
                )
                claim_validation_claims = (
                    claim_validation_payload.get("claims", {})
                    if isinstance(claim_validation_payload.get("claims"), dict)
                    else {}
                )
                validation_claim = (
                    claim_validation_claims.get(claim_type, {})
                    if isinstance(claim_validation_claims.get(claim_type), dict)
                    else {}
                )
            draft_claim = draft_claims_by_type.get(claim_type, {}) if isinstance(draft_claims_by_type.get(claim_type), dict) else {}
            overview_payload = self._safe_mediator_dict(
                "get_claim_overview",
                claim_type=claim_type,
                user_id=user_id,
                required_support_kinds=["evidence", "authority"],
            )
            overview_claim = overview_payload.get("claims", {}).get(claim_type, {}) if isinstance(overview_payload.get("claims"), dict) else {}
            treatment_summary = support_claim.get("authority_treatment_summary", {}) if isinstance(support_claim.get("authority_treatment_summary"), dict) else {}
            rule_summary = support_claim.get("authority_rule_candidate_summary", {}) if isinstance(support_claim.get("authority_rule_candidate_summary"), dict) else {}
            source_context = self._extract_support_source_context_counts(support_claim)
            claim_reasoning_summary = summarize_claim_reasoning_review(validation_claim) if validation_claim else {}
            claim_temporal_issue_count = int(claim_reasoning_summary.get("claim_temporal_issue_count", 0) or 0)
            claim_unresolved_temporal_issue_count = int(
                claim_reasoning_summary.get("claim_unresolved_temporal_issue_count", 0) or 0
            )
            claim_resolved_temporal_issue_count = int(
                claim_reasoning_summary.get("claim_resolved_temporal_issue_count", 0) or 0
            )
            claim_temporal_issue_ids = _extract_text_candidates(claim_reasoning_summary.get("claim_temporal_issue_ids"))
            claim_missing_temporal_predicates = _extract_text_candidates(
                claim_reasoning_summary.get("claim_missing_temporal_predicates")
            )
            claim_required_provenance_kinds = _extract_text_candidates(
                claim_reasoning_summary.get("claim_required_provenance_kinds")
            )

            claim_status = "ready"
            warnings: List[Dict[str, Any]] = []

            validation_status = str(validation_claim.get("validation_status") or "")
            if validation_status == "contradicted":
                claim_status = _merge_status(claim_status, "blocked")
                warnings.append(
                    {
                        "code": "claim_contradicted",
                        "severity": "blocked",
                        "message": f"{claim_type.title()} has contradiction signals that should be resolved before filing.",
                    }
                )
            elif validation_status in {"missing", "incomplete"}:
                claim_status = _merge_status(claim_status, "warning")

            if int(validation_claim.get("proof_gap_count", 0) or 0) > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "proof_gaps_present",
                        "severity": "warning",
                        "message": f"{claim_type.title()} still has proof or failed-premise gaps.",
                    }
                )

            temporal_gap_hint_count = int(
                (draft_claim.get("support_summary") or {}).get("temporal_gap_hint_count") or 0
            )
            if temporal_gap_hint_count <= 0:
                temporal_gap_hint_count = sum(
                    1
                    for item in _coerce_list(draft_claim.get("missing_elements"))
                    if str(item or "").strip().lower().startswith("chronology gap")
                )
            temporal_gap_hint_count = max(
                temporal_gap_hint_count,
                claim_unresolved_temporal_issue_count,
                len(claim_missing_temporal_predicates),
            )
            if temporal_gap_hint_count > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "chronology_gaps_present",
                        "severity": "warning",
                        "message": f"{claim_type.title()} still has {temporal_gap_hint_count} chronology gap(s) that should be resolved before filing.",
                    }
                )

            if int(treatment_summary.get("adverse_authority_link_count", 0) or 0) > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "adverse_authority_present",
                        "severity": "warning",
                        "message": f"{claim_type.title()} includes adverse or limiting authority that should be reviewed before relying on it in the draft.",
                    }
                )

            uncertain_authority_count = int(treatment_summary.get("uncertain_authority_link_count", 0) or 0)
            uncertain_treatment_types = sorted(
                str(name)
                for name in (treatment_summary.get("treatment_type_counts", {}) or {}).keys()
                if str(name) in {"questioned", "limits", "superseded", "good_law_unconfirmed"}
            )
            if uncertain_authority_count > 0 or uncertain_treatment_types:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "authority_reliability_uncertain",
                        "severity": "warning",
                        "message": f"{claim_type.title()} has authority support with unresolved treatment or good-law uncertainty.",
                    }
                )

            unresolved_elements = int(gap_claim.get("unresolved_count", 0) or 0)
            if unresolved_elements == 0:
                unresolved_elements = len(_coerce_list(overview_claim.get("missing"))) + len(_coerce_list(overview_claim.get("partially_supported")))
            if unresolved_elements > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "unresolved_elements",
                        "severity": "warning",
                        "message": f"{claim_type.title()} still has {unresolved_elements} unresolved claim element(s).",
                    }
                )

            claim_entry = {
                "claim_type": claim_type,
                "status": claim_status,
                "validation_status": validation_status or ("supported" if claim_status == "ready" else "incomplete"),
                "covered_elements": int(support_claim.get("covered_elements", 0) or 0),
                "total_elements": int(support_claim.get("total_elements", 0) or 0),
                "unresolved_element_count": unresolved_elements,
                "proof_gap_count": int(validation_claim.get("proof_gap_count", 0) or 0),
                "temporal_gap_hint_count": temporal_gap_hint_count,
                "contradiction_candidate_count": int(validation_claim.get("contradiction_candidate_count", 0) or 0),
                "claim_temporal_issue_count": claim_temporal_issue_count,
                "claim_unresolved_temporal_issue_count": claim_unresolved_temporal_issue_count,
                "claim_resolved_temporal_issue_count": claim_resolved_temporal_issue_count,
                "claim_temporal_issue_ids": claim_temporal_issue_ids,
                "claim_missing_temporal_predicates": claim_missing_temporal_predicates,
                "claim_required_provenance_kinds": claim_required_provenance_kinds,
                "support_by_kind": support_claim.get("support_by_kind", {}),
                "support_by_source": source_context["support_by_source"],
                "source_family_counts": source_context["source_family_counts"],
                "record_scope_counts": source_context["record_scope_counts"],
                "artifact_family_counts": source_context["artifact_family_counts"],
                "corpus_family_counts": source_context["corpus_family_counts"],
                "content_origin_counts": source_context["content_origin_counts"],
                "authority_treatment_summary": treatment_summary,
                "authority_rule_candidate_summary": rule_summary,
                "warnings": warnings,
            }
            claim_entry["chip_labels"] = _build_claim_checklist_chip_labels(claim_entry)
            aggregate_warning_count += len(warnings)
            overall_status = _merge_status(overall_status, claim_status)
            claim_readiness.append(claim_entry)

        claims_section_status = "ready"
        for claim_entry in claim_readiness:
            claims_section_status = _merge_status(claims_section_status, claim_entry.get("status", "ready"))

        total_fact_count = sum(int(claim.get("total_facts", 0) or 0) for claim in support_claims.values() if isinstance(claim, dict))
        if total_fact_count <= 0:
            total_fact_count = sum(
                len(self._normalize_text_lines(claim.get("supporting_facts", [])))
                for claim in _coerce_list(draft.get("claims_for_relief"))
                if isinstance(claim, dict)
            )
        summary_fact_count = len(self._normalize_text_lines(draft.get("summary_of_facts", [])))
        exhibits = _coerce_list(draft.get("exhibits"))
        relief_items = self._normalize_text_lines(draft.get("requested_relief", []))
        document_provenance_summary = (
            dict(draft.get("document_provenance_summary") or {})
            if isinstance(draft.get("document_provenance_summary"), dict)
            else {}
        )
        document_fact_backed_ratio = _safe_float(document_provenance_summary.get("fact_backed_ratio"), 0.0)
        document_low_grounding_flag = bool(document_provenance_summary.get("low_grounding_flag"))

        sections: Dict[str, Dict[str, Any]] = {}

        facts_status = "ready" if total_fact_count > 0 and summary_fact_count > 0 else "warning"
        facts_warnings: List[Dict[str, Any]] = []
        if facts_status != "ready":
            facts_warnings.append(
                {
                    "code": "fact_support_thin",
                    "severity": "warning",
                    "message": "The factual allegations section has limited fact-backed support and should be reviewed before filing.",
                }
            )
        sections["summary_of_facts"] = {
            "title": "Summary of Facts",
            "status": facts_status,
            "metrics": {
                "summary_fact_count": summary_fact_count,
                "support_fact_count": total_fact_count,
                "document_fact_backed_ratio": round(document_fact_backed_ratio, 4),
            },
            "warnings": facts_warnings,
        }
        if document_low_grounding_flag:
            sections["summary_of_facts"]["status"] = _merge_status(
                str(sections["summary_of_facts"].get("status") or "ready"),
                "warning",
            )
            sections["summary_of_facts"].setdefault("warnings", []).append(
                {
                    "code": "document_provenance_grounding_thin",
                    "severity": "warning",
                    "message": "The current draft is not grounded enough in canonical facts or artifact-backed support rows.",
                }
            )
            aggregate_warning_count += 1

        jurisdiction_status = "ready" if draft.get("jurisdiction_statement") and draft.get("venue_statement") else "warning"
        jurisdiction_warnings: List[Dict[str, Any]] = []
        procedural_rule_count = sum(
            int((entry.get("authority_rule_candidate_summary", {}).get("rule_type_counts", {}) or {}).get("procedural_prerequisite", 0) or 0)
            for entry in claim_readiness
            if isinstance(entry, dict)
        )
        if jurisdiction_status != "ready":
            jurisdiction_warnings.append(
                {
                    "code": "jurisdiction_or_venue_incomplete",
                    "severity": "warning",
                    "message": "Jurisdiction or venue language is incomplete and should be confirmed before export.",
                }
            )
        if procedural_rule_count > 0:
            jurisdiction_status = _merge_status(jurisdiction_status, "warning")
            jurisdiction_warnings.append(
                {
                    "code": "procedural_prerequisites_identified",
                    "severity": "warning",
                    "message": "Authority-derived procedural prerequisites were identified and should be checked against the current facts before filing.",
                }
            )
        sections["jurisdiction_and_venue"] = {
            "title": "Jurisdiction and Venue",
            "status": jurisdiction_status,
            "metrics": {
                "procedural_rule_count": procedural_rule_count,
            },
            "warnings": jurisdiction_warnings,
        }

        sections["claims_for_relief"] = {
            "title": "Claims for Relief",
            "status": claims_section_status,
            "metrics": {
                "claim_count": len(claim_readiness),
                "blocked_claim_count": len([entry for entry in claim_readiness if entry.get("status") == "blocked"]),
                "warning_claim_count": len([entry for entry in claim_readiness if entry.get("status") == "warning"]),
            },
            "warnings": [
                warning
                for entry in claim_readiness
                for warning in entry.get("warnings", [])
                if isinstance(warning, dict)
            ],
        }

        exhibits_status = "ready" if exhibits else "warning"
        exhibits_warnings: List[Dict[str, Any]] = []
        if not exhibits:
            exhibits_warnings.append(
                {
                    "code": "no_exhibits",
                    "severity": "warning",
                    "message": "No exhibits are currently attached to the draft package.",
                }
            )
        sections["exhibits"] = {
            "title": "Exhibits",
            "status": exhibits_status,
            "metrics": {
                "exhibit_count": len(exhibits),
            },
            "warnings": exhibits_warnings,
        }

        relief_status = "ready" if relief_items else "warning"
        relief_warnings: List[Dict[str, Any]] = []
        if not relief_items:
            relief_warnings.append(
                {
                    "code": "requested_relief_missing",
                    "severity": "warning",
                    "message": "Requested relief should be confirmed before filing.",
                }
            )
        sections["requested_relief"] = {
            "title": "Requested Relief",
            "status": relief_status,
            "metrics": {
                "requested_relief_count": len(relief_items),
            },
            "warnings": relief_warnings,
        }

        for section in sections.values():
            overall_status = _merge_status(overall_status, str(section.get("status") or "ready"))
            aggregate_warning_count += len(section.get("warnings", []) or [])

        graph_signals = self._build_graph_completeness_signals(getattr(self.mediator, "phase_manager", None))
        gap_signals = self._collect_unresolved_readiness_gaps(
            claim_readiness=claim_readiness,
            sections=sections,
            graph_signals=graph_signals,
        )
        intake_status = build_intake_status_summary(self.mediator)
        intake_handoff = (
            intake_status.get("intake_summary_handoff")
            if isinstance(intake_status, dict) and isinstance(intake_status.get("intake_summary_handoff"), dict)
            else {}
        )
        confirmation = (
            intake_handoff.get("complainant_summary_confirmation")
            if isinstance(intake_handoff.get("complainant_summary_confirmation"), dict)
            else {}
        )
        confirmation_snapshot = (
            confirmation.get("confirmed_summary_snapshot")
            if isinstance(confirmation.get("confirmed_summary_snapshot"), dict)
            else {}
        )
        intake_priority_summary = (
            confirmation_snapshot.get("adversarial_intake_priority_summary")
            if isinstance(confirmation_snapshot.get("adversarial_intake_priority_summary"), dict)
            else {}
        )
        objective_aliases = {
            "staff_names": "staff_names_titles",
            "staff_titles": "staff_names_titles",
            "hearing_timing": "hearing_request_timing",
            "response_timing": "response_dates",
            "causation": "causation_link",
            "adverse_action": "anchor_adverse_action",
            "appeal_rights": "anchor_appeal_rights",
        }

        def _normalize_objective(value: Any) -> str:
            objective_text = str(value or "").strip().lower()
            if not objective_text:
                return ""
            return objective_aliases.get(objective_text, objective_text)

        uncovered_intake_objectives = _dedupe_text_values(
            _normalize_objective(item)
            for item in list(intake_priority_summary.get("uncovered_objectives") or [])
            if _normalize_objective(item)
        )
        objective_question_counts = {
            _normalize_objective(key): int(value or 0)
            for key, value in dict(intake_priority_summary.get("objective_question_counts") or {}).items()
            if _normalize_objective(key)
        }
        required_intake_objectives = (
            "timeline",
            "actors",
            "staff_names_titles",
            "causation_link",
            "anchor_adverse_action",
            "anchor_appeal_rights",
            "hearing_request_timing",
            "response_dates",
        )
        missing_required_intake_objectives = [
            objective
            for objective in required_intake_objectives
            if objective in uncovered_intake_objectives or int(objective_question_counts.get(objective, 0)) <= 0
        ]

        summary_fact_lines = self._normalize_text_lines(draft.get("summary_of_facts", []))
        claim_support_lines = [
            line
            for claim in _coerce_list(draft.get("claims_for_relief"))
            if isinstance(claim, dict)
            for line in self._normalize_text_lines(claim.get("supporting_facts", []))
        ]

        def _has_date_anchor(text: str) -> bool:
            candidate = str(text or "").strip().lower()
            if not candidate:
                return False
            return bool(
                re.search(r"\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?\b", candidate)
                or re.search(r"\b\d{4}-\d{2}-\d{2}\b", candidate)
                or re.search(r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\b", candidate)
            )

        def _has_actor_anchor(text: str) -> bool:
            candidate = str(text or "").strip().lower()
            if not candidate:
                return False
            return bool(
                re.search(r"\b(manager|director|coordinator|officer|supervisor|caseworker|staff|agent|representative|administrator)\b", candidate)
                or re.search(r"\b(hacc|housing authority)\b", candidate)
            )

        def _has_evidence_anchor(text: str) -> bool:
            candidate = str(text or "").strip().lower()
            if not candidate:
                return False
            return bool(
                re.search(r"\b(exhibit|policy|document|notice|letter|email|record|file|pdf|attachment|cid|http)\b", candidate)
            )

        summary_date_anchor_count = sum(1 for line in summary_fact_lines if _has_date_anchor(line))
        summary_actor_anchor_count = sum(1 for line in summary_fact_lines if _has_actor_anchor(line))
        summary_evidence_anchor_count = sum(1 for line in summary_fact_lines if _has_evidence_anchor(line))
        claim_support_date_anchor_count = sum(1 for line in claim_support_lines if _has_date_anchor(line))
        claim_support_actor_anchor_count = sum(1 for line in claim_support_lines if _has_actor_anchor(line))
        claim_support_evidence_anchor_count = sum(1 for line in claim_support_lines if _has_evidence_anchor(line))
        structured_handoff_gaps: List[str] = []
        if summary_fact_lines and summary_date_anchor_count <= 0:
            structured_handoff_gaps.append("summary_of_facts_missing_date_anchors")
        if summary_fact_lines and summary_actor_anchor_count <= 0:
            structured_handoff_gaps.append("summary_of_facts_missing_actor_anchors")
        if summary_fact_lines and summary_evidence_anchor_count <= 0:
            structured_handoff_gaps.append("summary_of_facts_missing_evidence_anchors")
        if claim_support_lines and claim_support_date_anchor_count <= 0:
            structured_handoff_gaps.append("claim_support_missing_date_anchors")
        if claim_support_lines and claim_support_actor_anchor_count <= 0:
            structured_handoff_gaps.append("claim_support_missing_actor_anchors")
        if claim_support_lines and claim_support_evidence_anchor_count <= 0:
            structured_handoff_gaps.append("claim_support_missing_evidence_anchors")
        structured_handoff_signals = {
            "summary_fact_count": len(summary_fact_lines),
            "claim_support_fact_count": len(claim_support_lines),
            "summary_date_anchor_count": int(summary_date_anchor_count),
            "summary_actor_anchor_count": int(summary_actor_anchor_count),
            "summary_evidence_anchor_count": int(summary_evidence_anchor_count),
            "claim_support_date_anchor_count": int(claim_support_date_anchor_count),
            "claim_support_actor_anchor_count": int(claim_support_actor_anchor_count),
            "claim_support_evidence_anchor_count": int(claim_support_evidence_anchor_count),
            "gap_codes": _dedupe_text_values(structured_handoff_gaps),
            "gap_count": len(_dedupe_text_values(structured_handoff_gaps)),
        }
        unresolved_factual_gaps = list(gap_signals.get("factual_gaps") or [])
        unresolved_legal_gaps = list(gap_signals.get("legal_gaps") or [])
        if uncovered_intake_objectives:
            unresolved_factual_gaps.insert(
                0,
                "Intake coverage remains incomplete for one or more required objectives; chronology/actor/notice details should be closed before formalization.",
            )
        if missing_required_intake_objectives:
            unresolved_factual_gaps.insert(
                0,
                "Required intake objectives remain uncovered: " + ", ".join(missing_required_intake_objectives[:5]),
            )
        if int(structured_handoff_signals.get("gap_count", 0) or 0) > 0:
            unresolved_factual_gaps.append(
                "Structured drafting handoff is incomplete; promote intake facts, date/actor anchors, and evidence references into summary-of-facts and claim-support before optimization."
            )
        unresolved_factual_gaps = _dedupe_text_values(unresolved_factual_gaps)[:6]
        evidence_modality_signals = self._collect_evidence_modality_signals(
            draft=draft,
            claim_readiness=claim_readiness,
        )
        weak_evidence_modalities = list(evidence_modality_signals.get("weak_modalities") or [])
        weak_complaint_types = [
            claim_type
            for claim_type in claim_types
            if str(claim_type or "").strip().lower() in {"housing_discrimination", "hacc_research_engine"}
        ]
        targeted_weak_evidence_modalities = [
            modality
            for modality in weak_evidence_modalities
            if str(modality or "").strip().lower() in {"policy_document", "file_evidence"}
        ]
        policy_document_count = int(
            dict(evidence_modality_signals.get("modalities") or {}).get("policy_document", 0) or 0
        )
        file_evidence_count = int(
            dict(evidence_modality_signals.get("modalities") or {}).get("file_evidence", 0) or 0
        )
        total_elements = sum(int(entry.get("total_elements", 0) or 0) for entry in claim_readiness if isinstance(entry, dict))
        covered_elements = sum(int(entry.get("covered_elements", 0) or 0) for entry in claim_readiness if isinstance(entry, dict))
        claim_coverage = (
            (float(covered_elements) / float(total_elements))
            if total_elements > 0
            else (1.0 if summary_fact_count > 0 else 0.0)
        )
        ready_section_count = sum(
            1
            for section in sections.values()
            if isinstance(section, dict) and str(section.get("status") or "ready").strip().lower() == "ready"
        )
        section_coverage = (float(ready_section_count) / float(len(sections))) if sections else 0.0
        graph_gap_count = max(
            int(graph_signals.get("remaining_gap_count", 0) or 0),
            int(graph_signals.get("current_gap_count", 0) or 0),
        )
        graph_coverage = 1.0
        if str(graph_signals.get("status") or "ready").strip().lower() != "ready":
            graph_coverage -= 0.05
        if graph_gap_count > 0:
            graph_coverage -= min(0.25, 0.02 * float(graph_gap_count))
        if not _coerce_bool(graph_signals.get("knowledge_graph_enhanced"), default=True):
            graph_coverage -= 0.05
        graph_coverage = max(0.0, min(1.0, graph_coverage))
        graph_gate_active = (
            str(graph_signals.get("status") or "ready").strip().lower() != "ready"
            or graph_gap_count > 0
            or not _coerce_bool(graph_signals.get("knowledge_graph_available"), default=True)
            or not _coerce_bool(graph_signals.get("dependency_graph_available"), default=True)
        )
        coverage = max(
            0.0,
            min(
                1.0,
                round(
                    (0.5 * claim_coverage) + (0.25 * section_coverage) + (0.25 * graph_coverage),
                    3,
                ),
            ),
        )
        phase_status = str(overall_status or "ready").strip().lower() or "ready"
        if str(graph_signals.get("status") or "ready").strip().lower() != "ready":
            phase_status = _merge_status(phase_status, "warning")
        if graph_gate_active:
            phase_status = _merge_status(phase_status, "blocked")
        if (
            coverage < 0.98
            or unresolved_factual_gaps
            or unresolved_legal_gaps
            or missing_required_intake_objectives
            or targeted_weak_evidence_modalities
            or weak_complaint_types
            or document_low_grounding_flag
        ):
            phase_status = _merge_status(phase_status, "warning")
        if coverage <= 0.05 and (
            graph_gap_count > 0
            or not claim_readiness
            or unresolved_factual_gaps
            or unresolved_legal_gaps
        ):
            phase_status = "critical"
        if document_fact_backed_ratio < 0.25 and document_low_grounding_flag:
            phase_status = _merge_status(phase_status, "blocked")

        blockers: List[str] = []
        if graph_gate_active:
            blockers.append("graph_analysis_not_ready")
        if phase_status in {"warning", "blocked", "critical"}:
            blockers.append("document_generation_not_ready")
        if phase_status == "critical":
            blockers.append("document_generation_critical")
        if unresolved_factual_gaps:
            blockers.append("unresolved_factual_gaps_not_closed")
        if unresolved_legal_gaps:
            blockers.append("unresolved_legal_gaps_not_closed")
        if uncovered_intake_objectives or missing_required_intake_objectives:
            blockers.append("uncovered_intake_objectives")
        if int(structured_handoff_signals.get("gap_count", 0) or 0) > 0:
            blockers.append("structured_intake_handoff_incomplete")
        if weak_complaint_types:
            blockers.append("weak_complaint_type_generalization_needed")
        if targeted_weak_evidence_modalities:
            blockers.append("weak_evidence_modality_support_needed")
        if document_low_grounding_flag:
            blockers.append("document_provenance_grounding_needed")
        blockers = _dedupe_text_values(blockers)

        readiness_payload = {
            "status": overall_status,
            "claim_types": claim_types,
            "warning_count": aggregate_warning_count,
            "claims": claim_readiness,
            "sections": sections,
            "coverage": coverage,
            "phase_status": phase_status,
            "blockers": blockers,
            "unresolved_factual_gaps": unresolved_factual_gaps,
            "unresolved_legal_gaps": unresolved_legal_gaps,
            "uncovered_intake_objectives": uncovered_intake_objectives,
            "missing_required_intake_objectives": missing_required_intake_objectives,
            "objective_question_counts": objective_question_counts,
            "structured_intake_handoff_signals": structured_handoff_signals,
            "weak_complaint_types": weak_complaint_types,
            "evidence_modalities": dict(evidence_modality_signals.get("modalities") or {}),
            "weak_evidence_modalities": weak_evidence_modalities,
            "graph_completeness_signals": graph_signals,
            "document_provenance_summary": document_provenance_summary,
            "document_fact_backed_ratio": round(document_fact_backed_ratio, 4),
            "document_low_grounding_flag": document_low_grounding_flag,
            "drafting_handoff": {
                "gate_on_graph_completeness": bool(graph_gate_active),
                "graph_phase_status": str(graph_signals.get("status") or "ready").strip().lower() or "ready",
                "graph_remaining_gap_count": int(graph_gap_count),
                "coverage": float(coverage),
                "ready_for_formalization": phase_status == "ready" and not blockers,
                "blockers": list(blockers),
                "summary_fact_count": len(summary_fact_lines),
                "claim_support_fact_count": len(claim_support_lines),
                "exhibit_count": len(exhibits),
                "policy_document_count": int(policy_document_count),
                "file_evidence_count": int(file_evidence_count),
                "document_fact_backed_ratio": round(document_fact_backed_ratio, 4),
                "document_low_grounding_flag": bool(document_low_grounding_flag),
                "uncovered_intake_objectives": uncovered_intake_objectives[:8],
                "missing_required_intake_objectives": missing_required_intake_objectives[:8],
                "unresolved_factual_gaps": unresolved_factual_gaps[:6],
                "unresolved_legal_gaps": unresolved_legal_gaps[:6],
                "targeted_weak_complaint_types": weak_complaint_types[:4],
                "targeted_weak_evidence_modalities": targeted_weak_evidence_modalities[:4],
                "structured_intake_handoff_signals": structured_handoff_signals,
                "formalization_readiness_snapshot": {
                    "phase_status": phase_status,
                    "coverage": float(coverage),
                    "graph_gate_active": bool(graph_gate_active),
                    "unresolved_factual_gap_count": len(unresolved_factual_gaps),
                    "unresolved_legal_gap_count": len(unresolved_legal_gaps),
                    "uncovered_intake_objective_count": len(uncovered_intake_objectives),
                    "missing_required_intake_objective_count": len(missing_required_intake_objectives),
                    "structured_handoff_gap_count": int(structured_handoff_signals.get("gap_count", 0) or 0),
                    "weak_complaint_type_count": len(weak_complaint_types),
                    "weak_evidence_modality_count": len(targeted_weak_evidence_modalities),
                    "document_fact_backed_ratio": round(document_fact_backed_ratio, 4),
                    "document_low_grounding_flag": bool(document_low_grounding_flag),
                    "ready_for_formalization": phase_status == "ready" and not blockers,
                },
            },
        }
        workflow_phase_plan = self._build_runtime_workflow_phase_plan(
            drafting_readiness=readiness_payload,
            document_optimization=None,
        )
        workflow_optimization_guidance = _build_runtime_workflow_optimization_guidance(
            mediator=self.mediator,
            drafting_readiness=readiness_payload,
            workflow_phase_plan=workflow_phase_plan,
            document_optimization=None,
        )
        workflow_warnings = self._build_workflow_phase_warning_entries(workflow_phase_plan)
        optimization_warnings = _build_workflow_optimization_warning_entries(workflow_optimization_guidance)
        if workflow_phase_plan:
            readiness_payload["workflow_phase_plan"] = workflow_phase_plan
        if workflow_optimization_guidance:
            readiness_payload["workflow_optimization_guidance"] = workflow_optimization_guidance
        if workflow_warnings:
            readiness_payload["warnings"] = workflow_warnings
            readiness_payload["warning_count"] = int(readiness_payload.get("warning_count") or 0) + len(workflow_warnings)
            for warning in workflow_warnings:
                readiness_payload["status"] = _merge_status(
                    str(readiness_payload.get("status") or "ready"),
                    str(warning.get("severity") or "ready"),
                )
        gap_warnings: List[Dict[str, Any]] = []
        if unresolved_factual_gaps:
            gap_warnings.append(
                {
                    "code": "unresolved_factual_gaps",
                    "severity": "warning",
                    "message": "Unresolved factual gaps remain and should be closed before formalization.",
                    "gaps": unresolved_factual_gaps[:5],
                }
            )
        if unresolved_legal_gaps:
            gap_warnings.append(
                {
                    "code": "unresolved_legal_gaps",
                    "severity": "warning",
                    "message": "Unresolved legal gaps remain and should be closed before formalization.",
                    "gaps": unresolved_legal_gaps[:5],
                }
            )
        if weak_complaint_types:
            gap_warnings.append(
                {
                    "code": "weak_complaint_type_generalization_needed",
                    "severity": "warning",
                    "message": "Target complaint types still need stronger drafting generalization before formalization.",
                    "claim_types": weak_complaint_types[:4],
                }
            )
        if uncovered_intake_objectives or missing_required_intake_objectives:
            gap_warnings.append(
                {
                    "code": "uncovered_intake_objectives",
                    "severity": "warning",
                    "message": "Intake objectives remain uncovered and should be resolved before formalization.",
                    "uncovered_objectives": uncovered_intake_objectives[:8],
                    "missing_required_objectives": missing_required_intake_objectives[:8],
                }
            )
        if int(structured_handoff_signals.get("gap_count", 0) or 0) > 0:
            gap_warnings.append(
                {
                    "code": "structured_intake_handoff_incomplete",
                    "severity": "warning",
                    "message": (
                        "Structured intake handoff is incomplete; promote facts, date/actor anchors, and evidence references "
                        "into summary-of-facts and claim-support generation before document optimization."
                    ),
                    "gap_codes": list(structured_handoff_signals.get("gap_codes") or [])[:8],
                }
            )
        if targeted_weak_evidence_modalities:
            gap_warnings.append(
                {
                    "code": "weak_evidence_modality_support_needed",
                    "severity": "warning",
                    "message": "Weak evidence modalities require stronger source-anchored allegations before formalization.",
                    "modalities": targeted_weak_evidence_modalities[:4],
                }
            )
        if gap_warnings:
            existing_warnings = [
                dict(item)
                for item in list(readiness_payload.get("warnings") or [])
                if isinstance(item, dict)
            ]
            readiness_payload["warnings"] = existing_warnings + gap_warnings
            readiness_payload["warning_count"] = int(readiness_payload.get("warning_count") or 0) + len(gap_warnings)
            readiness_payload["status"] = _merge_status(str(readiness_payload.get("status") or "ready"), "warning")
            readiness_payload["phase_status"] = _merge_status(str(readiness_payload.get("phase_status") or "ready"), "warning")
        if optimization_warnings:
            readiness_payload.setdefault("warnings", []).extend(optimization_warnings)
            readiness_payload["warning_count"] = int(readiness_payload.get("warning_count") or 0) + len(optimization_warnings)
            for warning in optimization_warnings:
                readiness_payload["status"] = _merge_status(
                    str(readiness_payload.get("status") or "ready"),
                    str(warning.get("severity") or "ready"),
                )

        return readiness_payload

    def _build_filing_checklist(self, drafting_readiness: Dict[str, Any]) -> List[Dict[str, Any]]:
        if not isinstance(drafting_readiness, dict):
            return []

        checklist: List[Dict[str, Any]] = []
        sections = drafting_readiness.get("sections") if isinstance(drafting_readiness.get("sections"), dict) else {}
        claims = drafting_readiness.get("claims") if isinstance(drafting_readiness.get("claims"), list) else []

        for section_key, section in sections.items():
            if not isinstance(section, dict):
                continue
            status = str(section.get("status") or "ready")
            title = str(section.get("title") or section_key or "Section").strip()
            warnings = section.get("warnings") if isinstance(section.get("warnings"), list) else []
            metrics = section.get("metrics") if isinstance(section.get("metrics"), dict) else {}
            if status == "ready":
                checklist.append(
                    {
                        "scope": "section",
                        "key": str(section_key),
                        "title": title,
                        "status": "ready",
                        "summary": f"{title} is ready for filing review.",
                        "detail": self._summarize_metrics(metrics),
                    }
                )
                continue
            primary_warning = warnings[0] if warnings and isinstance(warnings[0], dict) else {}
            checklist.append(
                {
                    "scope": "section",
                    "key": str(section_key),
                    "title": title,
                    "status": status,
                    "summary": str(primary_warning.get("message") or f"Review {title} before filing."),
                    "detail": self._summarize_metrics(metrics),
                }
            )

        for claim in claims:
            if not isinstance(claim, dict):
                continue
            status = str(claim.get("status") or "ready")
            claim_type = str(claim.get("claim_type") or "claim").strip()
            warnings = claim.get("warnings") if isinstance(claim.get("warnings"), list) else []
            chip_labels = _build_claim_checklist_chip_labels(claim)
            metrics = {
                "covered_elements": claim.get("covered_elements"),
                "total_elements": claim.get("total_elements"),
                "unresolved_element_count": claim.get("unresolved_element_count"),
                "proof_gap_count": claim.get("proof_gap_count"),
            }
            if status == "ready":
                checklist.append(
                    {
                        "scope": "claim",
                        "key": claim_type,
                        "title": claim_type.title(),
                        "status": "ready",
                        "summary": f"{claim_type.title()} is ready for filing review.",
                        "detail": self._summarize_metrics(metrics),
                        "chip_labels": chip_labels,
                        "claim_unresolved_temporal_issue_count": int(claim.get("claim_unresolved_temporal_issue_count") or 0),
                        "claim_missing_temporal_predicates": _extract_text_candidates(claim.get("claim_missing_temporal_predicates")),
                        "claim_required_provenance_kinds": _extract_text_candidates(claim.get("claim_required_provenance_kinds")),
                    }
                )
                continue
            primary_warning = warnings[0] if warnings and isinstance(warnings[0], dict) else {}
            checklist.append(
                {
                    "scope": "claim",
                    "key": claim_type,
                    "title": claim_type.title(),
                    "status": status,
                    "summary": str(primary_warning.get("message") or f"Review {claim_type.title()} before filing."),
                    "detail": self._summarize_metrics(metrics),
                    "chip_labels": chip_labels,
                    "claim_unresolved_temporal_issue_count": int(claim.get("claim_unresolved_temporal_issue_count") or 0),
                    "claim_missing_temporal_predicates": _extract_text_candidates(claim.get("claim_missing_temporal_predicates")),
                    "claim_required_provenance_kinds": _extract_text_candidates(claim.get("claim_required_provenance_kinds")),
                }
            )

        checklist.sort(key=lambda item: {"blocked": 0, "warning": 1, "ready": 2}.get(str(item.get("status")), 3))
        return checklist

    def _annotate_filing_checklist_review_links(
        self,
        *,
        filing_checklist: List[Dict[str, Any]],
        drafting_readiness: Dict[str, Any],
        user_id: Optional[str],
    ) -> None:
        if not filing_checklist or not isinstance(drafting_readiness, dict):
            return

        claim_map: Dict[str, Dict[str, Any]] = {}
        for claim in _coerce_list(drafting_readiness.get("claims")):
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or "").strip()
            if not claim_type:
                continue
            claim_map[claim_type] = {
                "review_url": self._build_review_url(user_id=user_id, claim_type=claim_type),
                "review_context": {
                    "user_id": user_id,
                    "claim_type": claim_type,
                },
                "chip_labels": list(claim.get("chip_labels") or _build_claim_checklist_chip_labels(claim)),
            }

        section_map: Dict[str, Dict[str, Any]] = {}
        for section_key, section in (drafting_readiness.get("sections") or {}).items():
            if not isinstance(section, dict):
                continue
            resolved_key = str(section_key or "").strip()
            if not resolved_key:
                continue
            section_map[resolved_key] = {
                "review_url": self._build_review_url(user_id=user_id, section=resolved_key),
                "review_context": {
                    "user_id": user_id,
                    "section": resolved_key,
                    "claim_type": None,
                },
            }

        dashboard_url = self._build_review_url(user_id=user_id)
        for item in filing_checklist:
            if not isinstance(item, dict):
                continue
            scope = str(item.get("scope") or "").strip().lower()
            key = str(item.get("key") or "").strip()
            target = None
            if scope == "claim":
                target = claim_map.get(key)
            elif scope == "section":
                target = section_map.get(key)
            if target:
                item["review_url"] = target["review_url"]
                item["review_context"] = target["review_context"]
                merged_chip_labels = _merge_chip_labels(item.get("chip_labels"), target.get("chip_labels"))
                if merged_chip_labels:
                    item["chip_labels"] = merged_chip_labels
                else:
                    item.pop("chip_labels", None)
            else:
                item["review_url"] = dashboard_url
                item["review_context"] = {"user_id": user_id}

    def _build_review_url(
        self,
        *,
        user_id: Optional[str] = None,
        claim_type: Optional[str] = None,
        section: Optional[str] = None,
    ) -> str:
        params = {}
        if user_id:
            params["user_id"] = user_id
        if claim_type:
            params["claim_type"] = claim_type
        if section:
            params["section"] = section
        query = urlencode(params)
        return f"/claim-support-review?{query}" if query else "/claim-support-review"

    def _summarize_metrics(self, metrics: Dict[str, Any]) -> str:
        parts = []
        for key, value in metrics.items():
            if value in (None, "", []):
                continue
            parts.append(f"{key.replace('_', ' ')}={value}")
            if len(parts) >= 3:
                break
        return "; ".join(parts)

    def _select_statutes_for_claim(
        self,
        claim_type: str,
        statutes: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        tokens = {token for token in re.split(r"\W+", claim_type.lower()) if token}
        scored = []
        for statute in statutes:
            if not isinstance(statute, dict):
                continue
            haystack = " ".join(
                str(statute.get(field) or "") for field in ("citation", "title", "relevance")
            ).lower()
            score = sum(1 for token in tokens if token in haystack)
            scored.append((score, statute))
        scored.sort(key=lambda item: item[0], reverse=True)
        selected = [statute for score, statute in scored if score > 0][:3]
        return selected or [statute for _, statute in scored[:3]]

    def _extract_overview_elements(self, elements: Any) -> List[str]:
        names = []
        for element in _coerce_list(elements):
            if isinstance(element, dict):
                names.extend(_extract_text_candidates(element.get("element_text") or element.get("claim_element") or element))
            else:
                names.extend(_extract_text_candidates(element))
        return _unique_preserving_order(names)

    def _extract_requested_relief_from_facts(self, facts: List[str]) -> List[str]:
        remedies = []
        for fact in facts:
            lower = fact.lower()
            if "reinstat" in lower:
                remedies.append("Reinstatement or front pay in lieu of reinstatement.")
            if "back pay" in lower or "lost wages" in lower:
                remedies.append("Back pay, front pay, and lost benefits.")
            if "injunct" in lower:
                remedies.append("Injunctive relief to prevent continuing violations.")
        return remedies

    def _collect_exhibits(
        self,
        *,
        user_id: str,
        claim_types: List[str],
        support_claims: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        exhibits: List[Dict[str, Any]] = []
        evidence_records = _safe_call(self.mediator, "get_user_evidence", user_id=user_id) or []
        for record in _coerce_list(evidence_records):
            if not isinstance(record, dict):
                continue
            claim_type = record.get("claim_type")
            if claim_type and claim_types and claim_type not in claim_types:
                continue
            summary = self._summarize_exhibit_record(record)
            exhibits.append(
                {
                    "label": f"Exhibit {chr(65 + len(exhibits))}",
                    "title": record.get("description") or record.get("type") or record.get("cid") or "Supporting exhibit",
                    "claim_type": claim_type,
                    "kind": "evidence",
                    "link": self._build_exhibit_link(record),
                    "source_ref": record.get("cid") or record.get("source_url") or "",
                    "summary": summary,
                }
            )

        for claim_type, claim_summary in (support_claims or {}).items():
            if not isinstance(claim_summary, dict):
                continue
            for element in _coerce_list(claim_summary.get("elements")):
                if not isinstance(element, dict):
                    continue
                for link in _coerce_list(element.get("links")):
                    if not isinstance(link, dict):
                        continue
                    support_kind = str(link.get("support_kind") or "").strip().lower()
                    if support_kind != "authority":
                        continue
                    link_url = self._build_exhibit_link(link)
                    title = link.get("support_label") or link.get("title") or link.get("citation") or element.get("element_text")
                    source_ref = link.get("support_ref") or link.get("citation") or link_url or ""
                    if not title and not source_ref:
                        continue
                    exhibits.append(
                        {
                            "label": f"Exhibit {chr(65 + len(exhibits))}",
                            "title": title or "Authority support",
                            "claim_type": claim_type,
                            "kind": "authority",
                            "link": link_url,
                            "source_ref": source_ref,
                            "summary": link.get("relevance") or link.get("description") or "",
                        }
                    )

        deduped: List[Dict[str, Any]] = []
        seen = set()
        for exhibit in exhibits:
            key = (exhibit.get("kind"), exhibit.get("title"), exhibit.get("source_ref"))
            if key in seen:
                continue
            seen.add(key)
            deduped.append(exhibit)
            if len(deduped) >= 20:
                break
        return deduped

    def _summarize_exhibit_record(self, record: Dict[str, Any]) -> str:
        metadata = record.get("metadata") if isinstance(record.get("metadata"), dict) else {}
        graph_summary = metadata.get("document_graph_summary") if isinstance(metadata.get("document_graph_summary"), dict) else {}
        parts: List[str] = []
        parsed_preview = str(record.get("parsed_text_preview") or "").strip()
        if parsed_preview:
            parts.append(parsed_preview)
        evidence_id = record.get("id")
        if evidence_id not in (None, ""):
            fact_rows = _safe_call(self.mediator, "get_evidence_facts", evidence_id=evidence_id) or []
            fact_lines = [
                str(item.get("text") or "").strip()
                for item in _coerce_list(fact_rows)
                if isinstance(item, dict) and str(item.get("text") or "").strip()
            ]
            if fact_lines:
                parts.append("; ".join(fact_lines[:2]))
        description = str(record.get("description") or "").strip()
        if description:
            parts.append(description)
        entity_count = int(graph_summary.get("entity_count") or 0)
        relationship_count = int(graph_summary.get("relationship_count") or 0)
        if entity_count or relationship_count:
            parts.append(f"Graph extraction: {entity_count} entities, {relationship_count} relationships.")
        return " ".join(part for part in parts if part)

    def _build_exhibit_link(self, record: Dict[str, Any]) -> str:
        source_url = str(record.get("source_url") or "").strip()
        if source_url:
            return source_url
        support_ref = str(record.get("support_ref") or record.get("cid") or "").strip()
        if support_ref.startswith("http://") or support_ref.startswith("https://"):
            return support_ref
        if support_ref:
            return f"https://ipfs.io/ipfs/{support_ref}"
        return ""

    def _annotate_lines_with_exhibits(
        self,
        lines: List[str],
        exhibits: List[Dict[str, Any]],
    ) -> List[str]:
        if not lines or not exhibits:
            return lines
        annotated: List[str] = []
        for index, line in enumerate(lines):
            exhibit = self._select_exhibit_for_line(line, exhibits)
            if exhibit is None and index == 0:
                exhibit = exhibits[0]
            annotated.append(self._append_exhibit_citation(line, exhibit))
        return annotated

    def _annotate_entries_with_exhibits(
        self,
        entries: List[Dict[str, Any]],
        exhibits: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        if not entries:
            return []
        annotated: List[Dict[str, Any]] = []
        for index, entry in enumerate(entries):
            if not isinstance(entry, dict):
                continue
            updated_entry = dict(entry)
            text = str(updated_entry.get("text") or "").strip()
            exhibit = self._select_exhibit_for_entry(updated_entry, exhibits)
            if exhibit is None and index == 0 and exhibits:
                exhibit = exhibits[0]
            updated_entry["text"] = self._append_exhibit_citation(text, exhibit)
            if exhibit is not None:
                updated_entry["exhibit_label"] = str(exhibit.get("label") or "").strip() or None
            annotated.append(updated_entry)
        return annotated

    def _select_exhibit_for_entry(
        self,
        entry: Dict[str, Any],
        exhibits: List[Dict[str, Any]],
    ) -> Optional[Dict[str, Any]]:
        if not isinstance(entry, dict):
            return None
        source_kind = str(entry.get("source_kind") or "").strip().lower()
        if source_kind in {"uploaded_evidence_fact", "claim_support_fact", "claim_chronology_support", "factual_allegation"}:
            preferred = self._select_exhibit_by_kind(exhibits, "evidence")
            if preferred is not None:
                return preferred
        if source_kind == "claim_support_link":
            preferred = self._select_exhibit_by_kind(exhibits, "authority")
            if preferred is not None:
                return preferred
        return self._select_exhibit_for_line(str(entry.get("text") or ""), exhibits)

    def _select_exhibit_by_kind(
        self,
        exhibits: List[Dict[str, Any]],
        kind: str,
    ) -> Optional[Dict[str, Any]]:
        normalized_kind = str(kind or "").strip().lower()
        for exhibit in exhibits:
            if not isinstance(exhibit, dict):
                continue
            exhibit_kind = str(exhibit.get("kind") or "").strip().lower()
            if exhibit_kind == normalized_kind:
                return exhibit
        return None

    def _align_entries_to_lines(self, entries: Any, lines: List[str]) -> List[Dict[str, Any]]:
        normalized_lines = self._normalize_text_lines(lines)
        entry_list = [dict(entry) for entry in _coerce_list(entries) if isinstance(entry, dict)]
        entry_index: Dict[str, Dict[str, Any]] = {}
        for entry in entry_list:
            text = str(entry.get("text") or "").strip()
            if text:
                entry_index[text.lower()] = entry
        aligned: List[Dict[str, Any]] = []
        for line in normalized_lines:
            matched = dict(entry_index.get(str(line).lower()) or {})
            matched["text"] = line
            matched["fact_ids"] = _normalize_identifier_list(matched.get("fact_ids") or [])
            matched["source_artifact_ids"] = _normalize_identifier_list(matched.get("source_artifact_ids") or [])
            matched["claim_types"] = _normalize_identifier_list(matched.get("claim_types") or [])
            matched["claim_element_ids"] = _normalize_identifier_list(matched.get("claim_element_ids") or [])
            matched["support_trace_ids"] = _normalize_identifier_list(matched.get("support_trace_ids") or [])
            aligned.append(matched)
        return aligned

    def _match_document_source_entries(
        self,
        line: str,
        source_entries: List[Dict[str, Any]],
        *,
        limit: int = 3,
    ) -> List[Dict[str, Any]]:
        line_text = self._strip_exhibit_citation_suffix(str(line or "").strip())
        if not line_text:
            return []
        line_lower = line_text.lower()
        line_tokens = self._text_tokens(line_text)
        process_markers = (
            "notice",
            "informal review",
            "grievance",
            "appeal",
            "review decision",
            "request for review",
            "review",
            "hearing",
            "denial notice",
        )
        line_process_markers = {marker for marker in process_markers if marker in line_lower}
        scored: List[tuple[int, Dict[str, Any]]] = []
        for entry in source_entries:
            if not isinstance(entry, dict):
                continue
            entry_text = self._strip_exhibit_citation_suffix(str(entry.get("text") or "").strip())
            if not entry_text:
                continue
            entry_lower = entry_text.lower()
            entry_tokens = self._text_tokens(entry_text)
            score = len(line_tokens & entry_tokens)
            if entry_lower in line_lower or line_lower in entry_lower:
                score += 20
            entry_process_markers = {marker for marker in process_markers if marker in entry_lower}
            if line_process_markers and entry_process_markers:
                score += 4 * len(line_process_markers & entry_process_markers)
            if "failed to provide the informal review" in line_lower and (
                "informal review" in entry_lower
                or "grievance" in entry_lower
                or "review decision" in entry_lower
                or "request for review" in entry_lower
            ):
                score += 10
            if "failed to provide" in line_lower and "review" in line_lower:
                if "require" in entry_lower and (
                    "review" in entry_lower or "grievance" in entry_lower or "appeal" in entry_lower
                ):
                    score += 8
                if "issued" in entry_lower and "review decision" in entry_lower:
                    score += 6
            if score <= 0:
                continue
            scored.append((score, entry))
        scored.sort(key=lambda item: (-item[0], str(item[1].get("text") or "")))
        selected_entries = [dict(item[1]) for item in scored[:limit]]
        if "failed to provide the informal review" in line_lower:
            selected_entries = self._augment_review_process_matches(
                selected_entries=selected_entries,
                scored_entries=scored,
                limit=limit,
            )
        return selected_entries

    def _strip_exhibit_citation_suffix(self, text: str) -> str:
        candidate = str(text or "").strip()
        if not candidate:
            return ""
        candidate = re.sub(r"\s*\(See Exhibit [^)]+\)\.?\s*$", "", candidate, flags=re.IGNORECASE)
        return candidate.strip()

    def _augment_review_process_matches(
        self,
        *,
        selected_entries: List[Dict[str, Any]],
        scored_entries: List[tuple[int, Dict[str, Any]]],
        limit: int,
    ) -> List[Dict[str, Any]]:
        if not selected_entries or not scored_entries:
            return selected_entries

        def _has_process_fact(entry: Dict[str, Any], markers: tuple[str, ...]) -> bool:
            text = str(entry.get("text") or "").lower()
            return bool(_coerce_list(entry.get("fact_ids"))) and any(marker in text for marker in markers)

        def _is_requirement_fact(entry: Dict[str, Any]) -> bool:
            return _has_process_fact(entry, ("require", "required")) and _has_process_fact(
                entry,
                ("review", "grievance", "appeal"),
            )

        def _is_event_fact(entry: Dict[str, Any]) -> bool:
            return _has_process_fact(entry, ("submitted", "issued", "request", "decision")) and _has_process_fact(
                entry,
                ("review", "grievance"),
            )

        seen_keys = {
            (
                tuple(_normalize_identifier_list(entry.get("fact_ids"))),
                str(entry.get("text") or ""),
            )
            for entry in selected_entries
        }
        sticky_entries: List[Dict[str, Any]] = []
        for entry in selected_entries:
            if _is_requirement_fact(entry) and not any(_is_requirement_fact(item) for item in sticky_entries):
                sticky_entries.append(entry)
            elif _is_event_fact(entry) and not any(_is_event_fact(item) for item in sticky_entries):
                sticky_entries.append(entry)

        for _, candidate in scored_entries:
            candidate_entry = dict(candidate)
            candidate_key = (
                tuple(_normalize_identifier_list(candidate_entry.get("fact_ids"))),
                str(candidate_entry.get("text") or ""),
            )
            if candidate_key in seen_keys:
                continue
            if _is_requirement_fact(candidate_entry) and not any(_is_requirement_fact(item) for item in sticky_entries):
                sticky_entries.append(candidate_entry)
                seen_keys.add(candidate_key)
            elif _is_event_fact(candidate_entry) and not any(_is_event_fact(item) for item in sticky_entries):
                sticky_entries.append(candidate_entry)
                seen_keys.add(candidate_key)
            if any(_is_requirement_fact(item) for item in sticky_entries) and any(
                _is_event_fact(item) for item in sticky_entries
            ):
                break

        ordered_entries: List[Dict[str, Any]] = []
        for entry in sticky_entries + selected_entries:
            candidate_key = (
                tuple(_normalize_identifier_list(entry.get("fact_ids"))),
                str(entry.get("text") or ""),
            )
            if candidate_key in {
                (
                    tuple(_normalize_identifier_list(item.get("fact_ids"))),
                    str(item.get("text") or ""),
                )
                for item in ordered_entries
            }:
                continue
            ordered_entries.append(entry)
            if len(ordered_entries) >= limit:
                break
        return ordered_entries

    def _select_exhibit_for_line(
        self,
        line: str,
        exhibits: List[Dict[str, Any]],
    ) -> Optional[Dict[str, Any]]:
        line_tokens = self._text_tokens(line)
        if not line_tokens:
            return exhibits[0] if exhibits else None

        best_match: Optional[Dict[str, Any]] = None
        best_score = 0
        lowered_line = str(line or "").lower()
        prefers_evidence = any(
            marker in lowered_line
            for marker in (
                "on march",
                "on april",
                "on may",
                "on june",
                "on july",
                "on august",
                "on september",
                "on october",
                "on november",
                "on december",
                "denial notice",
                "grievance request",
                "review decision",
                "hearing officer",
                "sent plaintiff",
            )
        ) or _contains_date_anchor(line)
        for exhibit in exhibits:
            if not isinstance(exhibit, dict):
                continue
            exhibit_kind = str(exhibit.get("kind") or "").strip().lower()
            exhibit_tokens = self._text_tokens(
                " ".join(
                    str(exhibit.get(field) or "")
                    for field in ("title", "summary", "source_ref", "claim_type")
                )
            )
            score = len(line_tokens & exhibit_tokens)
            if prefers_evidence and exhibit_kind == "evidence":
                score += 4
            if prefers_evidence and exhibit_kind == "authority":
                score -= 2
            if score > best_score:
                best_score = score
                best_match = exhibit

        return best_match if best_score > 0 else None

    def _append_exhibit_citation(
        self,
        line: str,
        exhibit: Optional[Dict[str, Any]],
    ) -> str:
        text = str(line or "").strip()
        if not text or exhibit is None:
            return text
        label = str(exhibit.get("label") or "").strip()
        if not label:
            return text
        if label.lower() in text.lower():
            return text
        punctuation = "." if text.endswith(".") else ""
        base = text[:-1] if punctuation else text
        return f"{base} (See {label}){punctuation}"

    def _text_tokens(self, value: str) -> set[str]:
        return {
            token
            for token in re.split(r"\W+", str(value or "").lower())
            if len(token) >= 4
        }

    def _render_docx(self, draft: Dict[str, Any], path: Path) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.opc.constants import RELATIONSHIP_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(draft.get("court_header", ""))
        run.bold = True
        run.font.size = Pt(12)

        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        caption_party_lines = case_caption.get("caption_party_lines") if isinstance(case_caption.get("caption_party_lines"), list) else self._build_caption_party_lines(case_caption)
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run("\n\n".join(caption_party_lines) + "\n")

        case_no = document.add_paragraph()
        case_no.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        case_no.add_run(
            f"{case_caption.get('case_number_label', 'Civil Action No.')} {case_caption.get('case_number', '________________')}"
        ).bold = True
        if case_caption.get("lead_case_number"):
            lead_case = document.add_paragraph()
            lead_case.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lead_case.add_run(
                f"{case_caption.get('lead_case_number_label', 'Lead Case No.')} {case_caption['lead_case_number']}"
            ).bold = True
        if case_caption.get("related_case_number"):
            related_case = document.add_paragraph()
            related_case.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            related_case.add_run(
                f"{case_caption.get('related_case_number_label', 'Related Case No.')} {case_caption['related_case_number']}"
            ).bold = True
        if case_caption.get("assigned_judge"):
            judge = document.add_paragraph()
            judge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            judge.add_run(
                f"{case_caption.get('assigned_judge_label', 'Assigned Judge')}: {case_caption['assigned_judge']}"
            ).bold = True
        if case_caption.get("courtroom"):
            courtroom = document.add_paragraph()
            courtroom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            courtroom.add_run(
                f"{case_caption.get('courtroom_label', 'Courtroom')}: {case_caption['courtroom']}"
            ).bold = True

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(draft.get("case_caption", {}).get("document_title", "COMPLAINT"))
        title_run.bold = True
        title_run.font.size = Pt(14)
        if draft.get("case_caption", {}).get("jury_demand_notice"):
            jury_notice = document.add_paragraph()
            jury_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            jury_notice_run = jury_notice.add_run(draft["case_caption"]["jury_demand_notice"])
            jury_notice_run.bold = True
            jury_notice_run.font.size = Pt(12)

        self._add_docx_section(document, "Nature of the Action", draft.get("nature_of_action", []))
        self._add_docx_section(
            document,
            "Parties",
            [
                f"Plaintiff: {', '.join(draft.get('parties', {}).get('plaintiffs', []))}.",
                f"Defendant: {', '.join(draft.get('parties', {}).get('defendants', []))}.",
            ],
        )
        self._add_docx_section(
            document,
            "Jurisdiction and Venue",
            [draft.get("jurisdiction_statement"), draft.get("venue_statement")],
        )
        self._add_docx_numbered_facts(document, "Summary of Facts", draft.get("summary_of_facts", []))
        self._add_docx_numbered_facts(
            document,
            "Factual Allegations",
            draft.get("factual_allegations") or draft.get("summary_of_facts", []),
            groups=draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else None,
        )
        chronology_lines = draft.get("anchored_chronology_summary", [])
        if chronology_lines:
            self._add_docx_numbered_facts(document, "Anchored Chronology", chronology_lines)

        legal_standards = draft.get("legal_standards", [])
        if legal_standards:
            self._add_docx_section(document, "Applicable Legal Standards", legal_standards)

        document.add_heading("Claims for Relief", level=1)
        for index, claim in enumerate(draft.get("claims_for_relief", []), start=1):
            document.add_heading(f"Count {_roman(index)} - {claim.get('count_title', 'Claim')}", level=2)
            self._add_docx_subsection(document, "Legal Standard", claim.get("legal_standards", []))
            incorporated_clause = self._format_incorporated_reference_clause(
                claim.get("allegation_references", []),
                claim.get("supporting_exhibits", []),
            )
            if incorporated_clause:
                self._add_docx_subsection(document, "Incorporated Support", [incorporated_clause])
            self._add_docx_subsection(document, "Claim-Specific Support", claim.get("supporting_facts", []))
            missing = claim.get("missing_elements", [])
            if missing:
                self._add_docx_subsection(document, "Open Support Gaps", missing)
            exhibits = claim.get("supporting_exhibits", [])
            if exhibits:
                document.add_paragraph("Supporting Exhibits:")
                for exhibit in exhibits:
                    paragraph = document.add_paragraph(style="List Bullet")
                    paragraph.add_run(f"{exhibit.get('label')}. {exhibit.get('title')}")
                    if exhibit.get("link"):
                        paragraph.add_run(" ")
                        self._append_docx_hyperlink(
                            paragraph,
                            exhibit["link"],
                            "Open exhibit",
                            RELATIONSHIP_TYPE,
                            OxmlElement,
                            qn,
                            RGBColor,
                        )

        self._add_docx_subsection(document, "Requested Relief", draft.get("requested_relief", []), numbered=True)
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            self._add_docx_section(document, jury_demand.get("title") or "Jury Demand", [jury_demand.get("text")])

        document.add_heading("Supporting Exhibits", level=1)
        for exhibit in draft.get("exhibits", []):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{exhibit.get('label')}. {exhibit.get('title')}")
            if exhibit.get("summary"):
                paragraph.add_run(f" - {exhibit.get('summary')}")
            if exhibit.get("link"):
                paragraph.add_run(" ")
                self._append_docx_hyperlink(
                    paragraph,
                    exhibit["link"],
                    "Open exhibit",
                    RELATIONSHIP_TYPE,
                    OxmlElement,
                    qn,
                    RGBColor,
                )

        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            self._add_docx_section(
                document,
                affidavit.get("title") or "Affidavit in Support of Complaint",
                list(_coerce_list(affidavit.get("venue_lines")))
                + [affidavit.get("intro"), affidavit.get("knowledge_graph_note")],
            )
            self._add_docx_numbered_facts(document, "Affiant States as Follows", affidavit.get("facts", []))
            supporting_exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
            if supporting_exhibits:
                document.add_heading("Affidavit Supporting Exhibits", level=2)
                for exhibit in supporting_exhibits:
                    if not isinstance(exhibit, dict):
                        continue
                    paragraph = document.add_paragraph(style="List Bullet")
                    paragraph.add_run(f"{exhibit.get('label')}. {exhibit.get('title')}")
                    if exhibit.get("link"):
                        paragraph.add_run(f" ({exhibit['link']})")
            self._add_docx_section(
                document,
                "Affidavit Execution",
                [affidavit.get("dated"), affidavit.get("signature_line"), affidavit.get("jurat"), *_coerce_list(affidavit.get("notary_block"))],
            )

        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            self._add_docx_section(
                document,
                verification.get("title") or "Verification",
                [verification.get("text"), verification.get("dated"), verification.get("signature_line")],
            )
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            self._add_docx_section(
                document,
                certificate_of_service.get("title") or "Certificate of Service",
                [certificate_of_service.get("text")]
                + _coerce_list(certificate_of_service.get("detail_lines"))
                + [certificate_of_service.get("dated"), certificate_of_service.get("signature_line")],
            )
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        self._add_docx_section(
            document,
            "Signature Block",
            self._build_signature_section_lines(signature_block, self._resolve_draft_forum_type(draft)),
        )

        document.save(path)

    def _add_docx_section(self, document: Any, title: str, paragraphs: List[str]) -> None:
        document.add_heading(title, level=1)
        for paragraph in paragraphs:
            if paragraph:
                document.add_paragraph(str(paragraph))

    def _add_docx_numbered_facts(self, document: Any, title: str, facts: List[str], groups: Optional[List[Dict[str, Any]]] = None) -> None:
        document.add_heading(title, level=1)
        if groups:
            for group in groups:
                if not isinstance(group, dict):
                    continue
                heading = str(group.get("title") or "").strip()
                paragraphs = group.get("paragraphs") if isinstance(group.get("paragraphs"), list) else []
                if heading:
                    document.add_paragraph(heading)
                for paragraph in paragraphs:
                    if not isinstance(paragraph, dict):
                        continue
                    number = paragraph.get("number")
                    text = str(paragraph.get("text") or "").strip()
                    if text:
                        document.add_paragraph(f"{number}. {text}" if number else text)
            return
        for index, fact in enumerate(facts, start=1):
            document.add_paragraph(f"{index}. {fact}")

    def _add_docx_subsection(
        self,
        document: Any,
        title: str,
        lines: List[str],
        numbered: bool = False,
    ) -> None:
        document.add_paragraph(title)
        for index, line in enumerate(lines, start=1):
            prefix = f"{index}. " if numbered else ""
            document.add_paragraph(f"{prefix}{line}", style="List Bullet")

    def _append_docx_hyperlink(
        self,
        paragraph: Any,
        url: str,
        text: str,
        relationship_type: Any,
        oxml_element: Any,
        qn: Any,
        rgb_color: Any,
    ) -> None:
        part = paragraph.part
        rel_id = part.relate_to(url, relationship_type.HYPERLINK, is_external=True)
        hyperlink = oxml_element("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)
        run = oxml_element("w:r")
        properties = oxml_element("w:rPr")
        color = oxml_element("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = oxml_element("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(color)
        properties.append(underline)
        run.append(properties)
        text_element = oxml_element("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _render_pdf(self, draft: Dict[str, Any], path: Path) -> None:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="CourtHeader",
                parent=styles["Normal"],
                fontName="Times-Bold",
                fontSize=12,
                leading=14,
                alignment=TA_CENTER,
                spaceAfter=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="Caption",
                parent=styles["Normal"],
                fontName="Times-Roman",
                fontSize=12,
                leading=14,
                alignment=TA_CENTER,
                spaceAfter=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="SectionHeading",
                parent=styles["Heading1"],
                fontName="Times-Bold",
                fontSize=13,
                leading=15,
                textColor=colors.black,
                alignment=TA_LEFT,
                spaceBefore=10,
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RightAligned",
                parent=styles["Normal"],
                fontName="Times-Bold",
                fontSize=12,
                leading=14,
                alignment=TA_RIGHT,
                spaceAfter=8,
            )
        )

        doc = SimpleDocTemplate(
            str(path),
            pagesize=LETTER,
            topMargin=inch,
            bottomMargin=inch,
            leftMargin=inch,
            rightMargin=inch,
        )
        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        caption_party_lines = case_caption.get("caption_party_lines") if isinstance(case_caption.get("caption_party_lines"), list) else self._build_caption_party_lines(case_caption)
        story = [
            Paragraph(escape(draft.get("court_header", "")), styles["CourtHeader"]),
            Paragraph(
                "<br/><br/>".join(escape(line).replace("\n", "<br/>") for line in caption_party_lines),
                styles["Caption"],
            ),
            Paragraph(
                escape(
                    f"{case_caption.get('case_number_label', 'Civil Action No.')} {case_caption.get('case_number', '________________')}"
                    + (
                        f"\n{case_caption.get('lead_case_number_label', 'Lead Case No.')} {case_caption.get('lead_case_number')}"
                        if case_caption.get('lead_case_number')
                        else ""
                    )
                    + (
                        f"\n{case_caption.get('related_case_number_label', 'Related Case No.')} {case_caption.get('related_case_number')}"
                        if case_caption.get('related_case_number')
                        else ""
                    )
                    + (
                        f"\n{case_caption.get('assigned_judge_label', 'Assigned Judge')}: {case_caption.get('assigned_judge')}"
                        if case_caption.get('assigned_judge')
                        else ""
                    )
                    + (
                        f"\n{case_caption.get('courtroom_label', 'Courtroom')}: {case_caption.get('courtroom')}"
                        if case_caption.get('courtroom')
                        else ""
                    )
                ),
                styles["RightAligned"],
            ),
            Paragraph(
                escape(draft.get("case_caption", {}).get("document_title", "COMPLAINT")),
                styles["CourtHeader"],
            ),
            *(
                [
                    Paragraph(
                        escape(draft["case_caption"]["jury_demand_notice"]),
                        styles["CourtHeader"],
                    )
                ]
                if draft.get("case_caption", {}).get("jury_demand_notice")
                else []
            ),
            Spacer(1, 8),
        ]

        self._append_pdf_section(story, styles, "Nature of the Action", draft.get("nature_of_action", []))
        self._append_pdf_section(
            story,
            styles,
            "Parties",
            [
                f"Plaintiff: {', '.join(draft.get('parties', {}).get('plaintiffs', []))}.",
                f"Defendant: {', '.join(draft.get('parties', {}).get('defendants', []))}.",
            ],
        )
        self._append_pdf_section(
            story,
            styles,
            "Jurisdiction and Venue",
            [draft.get("jurisdiction_statement"), draft.get("venue_statement")],
        )
        self._append_pdf_numbered_section(story, styles, "Summary of Facts", draft.get("summary_of_facts", []))
        self._append_pdf_numbered_section(
            story,
            styles,
            "Factual Allegations",
            draft.get("factual_allegations") or draft.get("summary_of_facts", []),
            groups=draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else None,
        )
        self._append_pdf_numbered_section(story, styles, "Anchored Chronology", draft.get("anchored_chronology_summary", []))
        self._append_pdf_section(
            story,
            styles,
            "Applicable Legal Standards",
            draft.get("legal_standards", []),
        )

        story.append(Paragraph("Claims for Relief", styles["SectionHeading"]))
        for index, claim in enumerate(draft.get("claims_for_relief", []), start=1):
            story.append(
                Paragraph(
                    escape(f"Count {_roman(index)} - {claim.get('count_title', 'Claim')}"),
                    styles["Heading2"],
                )
            )
            self._append_pdf_section(story, styles, "Legal Standard", claim.get("legal_standards", []), heading_style="Heading3")
            incorporated_clause = self._format_incorporated_reference_clause(
                claim.get("allegation_references", []),
                claim.get("supporting_exhibits", []),
            )
            if incorporated_clause:
                self._append_pdf_section(story, styles, "Incorporated Support", [incorporated_clause], heading_style="Heading3")
            self._append_pdf_section(story, styles, "Claim-Specific Support", claim.get("supporting_facts", []), heading_style="Heading3")
            if claim.get("missing_elements"):
                self._append_pdf_section(story, styles, "Open Support Gaps", claim.get("missing_elements", []), heading_style="Heading3")
            if claim.get("supporting_exhibits"):
                story.append(Paragraph("Supporting Exhibits", styles["Heading3"]))
                for exhibit in claim.get("supporting_exhibits", []):
                    story.append(
                        Paragraph(
                            self._pdf_exhibit_markup(exhibit),
                            styles["Normal"],
                        )
                    )

        self._append_pdf_numbered_section(story, styles, "Requested Relief", draft.get("requested_relief", []))
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            self._append_pdf_section(story, styles, jury_demand.get("title") or "Jury Demand", [jury_demand.get("text")])
        story.append(Paragraph("Supporting Exhibits", styles["SectionHeading"]))
        for exhibit in draft.get("exhibits", []):
            story.append(Paragraph(self._pdf_exhibit_markup(exhibit), styles["Normal"]))

        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            self._append_pdf_section(
                story,
                styles,
                affidavit.get("title") or "Affidavit in Support of Complaint",
                list(_coerce_list(affidavit.get("venue_lines"))) + [affidavit.get("intro"), affidavit.get("knowledge_graph_note")],
            )
            self._append_pdf_numbered_section(story, styles, "Affiant States as Follows", affidavit.get("facts", []))
            supporting_exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
            if supporting_exhibits:
                story.append(Paragraph("Affidavit Supporting Exhibits", styles["Heading3"]))
                for exhibit in supporting_exhibits:
                    if not isinstance(exhibit, dict):
                        continue
                    story.append(Paragraph(self._pdf_exhibit_markup(exhibit), styles["Normal"]))
            self._append_pdf_section(
                story,
                styles,
                "Affidavit Execution",
                [affidavit.get("dated"), affidavit.get("signature_line"), affidavit.get("jurat"), *_coerce_list(affidavit.get("notary_block"))],
            )

        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            self._append_pdf_section(
                story,
                styles,
                verification.get("title") or "Verification",
                [verification.get("text"), verification.get("dated"), verification.get("signature_line")],
            )
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            self._append_pdf_section(
                story,
                styles,
                certificate_of_service.get("title") or "Certificate of Service",
                [certificate_of_service.get("text")]
                + _coerce_list(certificate_of_service.get("detail_lines"))
                + [certificate_of_service.get("dated"), certificate_of_service.get("signature_line")],
            )
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        self._append_pdf_section(
            story,
            styles,
            "Signature Block",
            self._build_signature_section_lines(signature_block, self._resolve_draft_forum_type(draft)),
        )

        doc.build(story)

    def _append_pdf_section(
        self,
        story: List[Any],
        styles: Any,
        title: str,
        paragraphs: List[str],
        heading_style: str = "SectionHeading",
    ) -> None:
        from reportlab.platypus import Paragraph

        if not paragraphs:
            return
        story.append(Paragraph(escape(title), styles[heading_style]))
        for paragraph in paragraphs:
            story.append(Paragraph(escape(str(paragraph)), styles["Normal"]))

    def _append_pdf_numbered_section(
        self,
        story: List[Any],
        styles: Any,
        title: str,
        paragraphs: List[str],
        heading_style: str = "SectionHeading",
        groups: Optional[List[Dict[str, Any]]] = None,
    ) -> None:
        from reportlab.platypus import Paragraph

        if not paragraphs and not groups:
            return
        story.append(Paragraph(escape(title), styles[heading_style]))
        if groups:
            for group in groups:
                if not isinstance(group, dict):
                    continue
                group_title = str(group.get("title") or "").strip()
                entries = group.get("paragraphs") if isinstance(group.get("paragraphs"), list) else []
                if group_title:
                    story.append(Paragraph(escape(group_title), styles["Heading3"]))
                for entry in entries:
                    if not isinstance(entry, dict):
                        continue
                    number = entry.get("number")
                    text = str(entry.get("text") or "").strip()
                    if text:
                        prefix = f"{number}. " if number else ""
                        story.append(Paragraph(escape(f"{prefix}{text}"), styles["Normal"]))
            return
        for index, paragraph in enumerate(paragraphs, start=1):
            story.append(Paragraph(escape(f"{index}. {paragraph}"), styles["Normal"]))

    def _pdf_exhibit_markup(self, exhibit: Dict[str, Any]) -> str:
        title = escape(f"{exhibit.get('label')}. {exhibit.get('title')}")
        summary = escape(str(exhibit.get("summary") or ""))
        link = str(exhibit.get("link") or "").strip()
        if link:
            link_markup = f'<link href="{escape(link)}">Open exhibit</link>'
            if summary:
                return f"{title} - {summary} ({link_markup})"
            return f"{title} ({link_markup})"
        if summary:
            return f"{title} - {summary}"
        return title
