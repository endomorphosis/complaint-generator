"""Formal complaint document assembly and export helpers."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import re
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from complaint_phases import ComplaintPhase, NodeType
from lib.document_render import build_case_caption_table, build_case_caption_text_lines, build_case_detail_lines, build_caption_party_block_lines, make_page_footer_renderer

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    HAS_DOCX = True
except ImportError:  # pragma: no cover - optional dependency
    DocxDocument = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    HAS_DOCX = False

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    HAS_REPORTLAB = True
except ImportError:  # pragma: no cover - optional dependency
    colors = None
    TA_CENTER = None
    letter = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    inch = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None
    HAS_REPORTLAB = False


def _utc_now_isoformat() -> str:
    return datetime.now(timezone.utc).isoformat()


def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\r", "\n")
    lines = [" ".join(line.split()) for line in text.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _clean_sentence(value: Any) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    return text if text.endswith((".", "?", "!", ":")) else f"{text}."


def _split_allegation_fragments(value: Any) -> List[str]:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" -;")
    if not text:
        return []
    if ": " in text:
        prefix, suffix = text.split(": ", 1)
        prefix_lower = prefix.strip().lower()
        if (
            prefix.strip().endswith("?")
            or prefix_lower.startswith(("what ", "when ", "where ", "why ", "how ", "who ", "describe ", "explain "))
            or prefix_lower in {"what happened", "what relief do you want"}
        ):
            text = suffix.strip()
    parts = [
        part.strip(" -;")
        for part in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
        if part.strip(" -;")
    ]
    return parts or [text]


def _formalize_allegation_fragment(value: Any) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" -;")
    if not text:
        return ""
    replacements = (
        (r"^i was\b", "Plaintiff was"),
        (r"^i am\b", "Plaintiff is"),
        (r"^i need\b", "Plaintiff needs"),
        (r"^i needed\b", "Plaintiff needed"),
        (r"^i lost\b", "Plaintiff lost"),
        (r"^i asked\b", "Plaintiff asked"),
        (r"^i reported\b", "Plaintiff reported"),
        (r"^i complained\b", "Plaintiff complained"),
        (r"^i informed\b", "Plaintiff informed"),
        (r"^i notified\b", "Plaintiff notified"),
        (r"^i requested\b", "Plaintiff requested"),
        (r"^i sought\b", "Plaintiff sought"),
        (r"^i experienced\b", "Plaintiff experienced"),
        (r"^i suffered\b", "Plaintiff suffered"),
        (r"^i told\b", "Plaintiff told"),
        (r"^they\b", "Defendant"),
    )
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    clause_replacements = (
        (r"([,;]\s+)i was\b", r"\1Plaintiff was"),
        (r"([,;]\s+)i am\b", r"\1Plaintiff is"),
        (r"([,;]\s+)i need\b", r"\1Plaintiff needs"),
        (r"([,;]\s+)i needed\b", r"\1Plaintiff needed"),
        (r"([,;]\s+)i lost\b", r"\1Plaintiff lost"),
        (r"([,;]\s+)i asked\b", r"\1Plaintiff asked"),
        (r"([,;]\s+)i reported\b", r"\1Plaintiff reported"),
        (r"([,;]\s+)i complained\b", r"\1Plaintiff complained"),
        (r"([,;]\s+)i requested\b", r"\1Plaintiff requested"),
        (r"([,;]\s+)i informed\b", r"\1Plaintiff informed"),
        (r"([,;]\s+)i notified\b", r"\1Plaintiff notified"),
        (r"([,;]\s+)i suffered\b", r"\1Plaintiff suffered"),
        (r"([,;]\s+)i experienced\b", r"\1Plaintiff experienced"),
        (r"([,;]\s+)i told\b", r"\1Plaintiff told"),
        (r"(\band\s+)i was\b", r"\1Plaintiff was"),
        (r"(\band\s+)i am\b", r"\1Plaintiff is"),
        (r"(\band\s+)i need\b", r"\1Plaintiff needs"),
        (r"(\band\s+)i needed\b", r"\1Plaintiff needed"),
        (r"(\band\s+)i lost\b", r"\1Plaintiff lost"),
        (r"(\band\s+)i asked\b", r"\1Plaintiff asked"),
        (r"(\band\s+)i reported\b", r"\1Plaintiff reported"),
        (r"(\band\s+)i complained\b", r"\1Plaintiff complained"),
        (r"(\band\s+)i requested\b", r"\1Plaintiff requested"),
        (r"(\band\s+)i informed\b", r"\1Plaintiff informed"),
        (r"(\band\s+)i notified\b", r"\1Plaintiff notified"),
        (r"(\band\s+)i suffered\b", r"\1Plaintiff suffered"),
        (r"(\band\s+)i experienced\b", r"\1Plaintiff experienced"),
        (r"(\band\s+)i told\b", r"\1Plaintiff told"),
        (r"(\bafter\s+)i was\b", r"\1Plaintiff was"),
        (r"(\bafter\s+)i am\b", r"\1Plaintiff is"),
        (r"(\bafter\s+)i need\b", r"\1Plaintiff needs"),
        (r"(\bafter\s+)i needed\b", r"\1Plaintiff needed"),
        (r"(\bafter\s+)i lost\b", r"\1Plaintiff lost"),
        (r"(\bafter\s+)i asked\b", r"\1Plaintiff asked"),
        (r"(\bafter\s+)i reported\b", r"\1Plaintiff reported"),
        (r"(\bafter\s+)i complained\b", r"\1Plaintiff complained"),
        (r"(\bafter\s+)i requested\b", r"\1Plaintiff requested"),
        (r"(\bafter\s+)i informed\b", r"\1Plaintiff informed"),
        (r"(\bafter\s+)i notified\b", r"\1Plaintiff notified"),
        (r"(\bafter\s+)i suffered\b", r"\1Plaintiff suffered"),
        (r"(\bafter\s+)i experienced\b", r"\1Plaintiff experienced"),
        (r"(\bafter\s+)i told\b", r"\1Plaintiff told"),
        (r"(\bthat\s+)i am\b", r"\1Plaintiff is"),
        (r"(\bthat\s+)i need\b", r"\1Plaintiff needs"),
        (r"(\bthat\s+)i needed\b", r"\1Plaintiff needed"),
        (r"(\bthat\s+)i asked\b", r"\1Plaintiff asked"),
        (r"(\bthat\s+)i complained\b", r"\1Plaintiff complained"),
        (r"(\bthat\s+)i requested\b", r"\1Plaintiff requested"),
        (r"(\bthat\s+)i told\b", r"\1Plaintiff told"),
    )
    for pattern, replacement in clause_replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    text = re.sub(r"\bmy\b", "Plaintiff's", text, flags=re.IGNORECASE)
    text = re.sub(r"\bmine\b", "Plaintiff's", text, flags=re.IGNORECASE)
    text = re.sub(r"\bme\b", "Plaintiff", text, flags=re.IGNORECASE)
    text = re.sub(r"\blost Plaintiff's pay and benefits\b", "lost pay and benefits", text, flags=re.IGNORECASE)
    text = re.sub(r"\blost Plaintiff's (pay|wages|salary|income|benefits)\b", r"lost \1", text, flags=re.IGNORECASE)
    if text and text[0].islower():
        text = text[0].upper() + text[1:]
    if len(text) < 12:
        return ""
    return text if text.endswith((".", "?", "!")) else f"{text}."


def _is_factual_allegation_candidate(value: Any) -> bool:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if not text:
        return False
    lowered = text.lower()
    if re.match(r"^(as to [^,]+, )?plaintiff (seeks|requests|asks|demands)\b", lowered):
        return False
    if lowered.startswith(("requested relief", "relief requested", "element supported:")):
        return False
    if lowered.startswith(("evidence shows facts supporting", "the intake record describes facts supporting")):
        return False
    if re.match(r"^(as to [^,]+, )?(title\s+[ivxlcdm0-9]+\b|\d+\s+u\.s\.c\.|\d+\s+c\.f\.r\.|[a-z]{2,6}\.\s+gov\.\s+code\b)", lowered):
        return False
    if not re.search(
        r"\b(was|were|is|are|reported|complained|terminated|fired|retaliated|denied|refused|told|informed|notified|requested|sought|experienced|suffered|lost|made|engaged|opposed|filed|sent|emailed|wrote|received|occurred|happened|subjected|demoted|suspended|disciplined|reduced)\b",
        lowered,
    ):
        return False
    return True


def _expand_allegation_sources(value: Any, *, limit: Optional[int] = None) -> List[str]:
    expanded: List[str] = []
    for item in _listify(value):
        for fragment in _split_allegation_fragments(item):
            sentence = _formalize_allegation_fragment(fragment)
            if sentence and _is_factual_allegation_candidate(sentence):
                expanded.append(sentence)
    unique: List[str] = []
    seen = set()
    for item in expanded:
        marker = item.lower()
        if marker in seen:
            continue
        seen.add(marker)
        unique.append(item)
    return unique[:limit] if limit is not None else unique


def _synthesize_narrative_allegations(allegations: List[str]) -> List[str]:
    cleaned = [str(item).strip() for item in allegations if str(item).strip()]
    if not cleaned:
        return []

    def _normalize_adverse_clause(clause: str) -> str:
        text = str(clause or "").strip().rstrip(".!?")
        if re.match(r"^(after|following)\b", text, flags=re.IGNORECASE) and "," in text:
            text = text.split(",", 1)[1].strip()
        return text

    def _normalize_harm_clause(clause: str) -> str:
        text = str(clause or "").strip().rstrip(".!?")
        text = re.sub(r",?\s+as a result$", "", text, flags=re.IGNORECASE)
        text = re.sub(r",?\s+as a direct result$", "", text, flags=re.IGNORECASE)
        return text.strip()

    def _pick(pattern: str, *, require_plaintiff: bool = False) -> str:
        for item in cleaned:
            lowered = item.lower()
            if require_plaintiff and "plaintiff" not in lowered:
                continue
            if re.search(pattern, lowered):
                return item.rstrip(".!?")
        return ""

    report_clause = _pick(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", require_plaintiff=True)
    adverse_clause = _pick(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b")
    harm_clause = _pick(r"\blost (pay|wages|salary|income|benefits)\b|\b(suffered|experienced)\b", require_plaintiff=True)
    harm_already_tied_to_adverse_action = any(
        re.search(r"\b(lost|suffered|experienced)\b", item.lower())
        and re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", item.lower())
        for item in cleaned
    )

    synthesized: List[str] = []
    if report_clause and adverse_clause:
        synthesized.append(f"After {report_clause}, {_normalize_adverse_clause(adverse_clause)}.")
    if harm_clause and not harm_already_tied_to_adverse_action:
        normalized_harm_clause = _normalize_harm_clause(harm_clause)
        loss_match = re.search(r"\blost ([^.]+)", normalized_harm_clause, flags=re.IGNORECASE)
        if loss_match:
            synthesized.append(f"As a direct result of Defendant's conduct, Plaintiff lost {loss_match.group(1).strip()}.")
    unique: List[str] = []
    seen = set()
    for item in synthesized:
        marker = item.lower()
        if marker in seen:
            continue
        seen.add(marker)
        unique.append(item)
    return _expand_allegation_sources(unique, limit=4)


def _contains_date_anchor(value: Any) -> bool:
    text = str(value or "")
    return bool(
        re.search(
            r"\b(?:\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{1,2}(?:,\s+\d{2,4})?)\b",
            text,
            flags=re.IGNORECASE,
        )
    )


def _contains_actor_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "who at hacc",
            "caseworker",
            "housing specialist",
            "program manager",
            "hearing officer",
            "staff",
            "supervisor",
            "director",
            "coordinator",
            "name",
            "title",
        )
    )


def _contains_causation_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    if not lowered:
        return False
    return (
        any(marker in lowered for marker in ("because", "as a result", "after", "following", "in retaliation", "retaliat", "days after", "weeks after", "shortly after"))
        and any(marker in lowered for marker in ("complained", "reported", "grievance", "appeal", "protected activity", "requested accommodation"))
        and any(marker in lowered for marker in ("adverse action", "termination", "denial", "loss of assistance", "retaliat"))
    )


def _contains_hearing_timing_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "hearing request",
            "requested a hearing",
            "requested review",
            "review request",
            "informal hearing request",
            "grievance request",
        )
    ) and any(marker in lowered for marker in ("date", "on ", "after", "before", "within", "days", "weeks"))


def _contains_response_date_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "response date",
            "responded on",
            "response was",
            "review decision",
            "hearing outcome",
            "notice date",
            "decision date",
        )
    ) and _contains_date_anchor(value)


def _contains_staff_identity_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return (
        "hacc" in lowered
        and any(marker in lowered for marker in ("name", "title", "staff", "caseworker", "manager", "officer", "specialist", "director"))
    )


def _prioritized_intake_statements(intake_summary: List[Dict[str, str]]) -> List[str]:
    statements: List[str] = []
    for item in intake_summary:
        if not isinstance(item, dict):
            continue
        question = str(item.get("question") or "").strip().lower()
        answer = str(item.get("answer") or "").strip()
        if not answer:
            continue
        if any(token in question for token in ("when", "date", "timeline", "who", "which person", "because", "retaliat", "why")):
            statements.append(answer)
    expanded = _expand_allegation_sources(statements, limit=8)
    deduped: List[str] = []
    seen = set()
    for statement in expanded:
        marker = str(statement).lower()
        if not marker or marker in seen:
            continue
        seen.add(marker)
        deduped.append(statement)
    return deduped


def _prune_subsumed_narrative_clauses(allegations: List[str]) -> List[str]:
    cleaned = [str(item).strip() for item in allegations if str(item).strip()]
    if not cleaned:
        return []

    def _pick(pattern: str, *, require_plaintiff: bool = False) -> str:
        for item in cleaned:
            lowered = item.lower()
            if require_plaintiff and "plaintiff" not in lowered:
                continue
            if re.search(pattern, lowered):
                return item.strip()
        return ""

    report_clause = _pick(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", require_plaintiff=True)
    adverse_clause = _pick(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b")
    has_harm_tied_to_adverse_action = any(
        re.search(r"\b(lost|suffered|experienced)\b", item.lower())
        and re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", item.lower())
        for item in cleaned
    )
    consumed = {item.lower() for item in (report_clause, adverse_clause) if item}
    if has_harm_tied_to_adverse_action:
        combined_clause = _pick(
            r"\b(reported|complained|opposed|informed|notified|told|requested)\b.*\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b"
            r"|\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b.*\b(reported|complained|opposed|informed|notified|told|requested)\b",
            require_plaintiff=True,
        )
        if combined_clause:
            consumed.add(combined_clause.lower())
    return [item for item in cleaned if item.lower() not in consumed]


def _prune_near_duplicate_allegations(allegations: List[str]) -> List[str]:
    def _tokens(value: str) -> set[str]:
        scrubbed = re.sub(r"\(see exhibit [^)]+\)", "", value, flags=re.IGNORECASE)
        return {
            token
            for token in re.split(r"\W+", scrubbed.lower())
            if len(token) >= 4 and token not in {"plaintiff", "defendant", "exhibit", "after", "those", "this", "that"}
        }

    def _categories(value: str) -> set[str]:
        lowered = value.lower()
        flags = set()
        if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
            flags.add("report")
        if re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied|removed|stripped)\b", lowered) or re.search(r"\b(end(?:ed|ing))\b[^.]{0,40}\bemployment\b", lowered):
            flags.add("adverse")
        if re.search(r"\b(lost|suffered|experienced|benefits|wages|salary|income|opportunities)\b", lowered):
            flags.add("harm")
        return flags

    def _features(value: str) -> set[str]:
        lowered = value.lower()
        flags = set()
        if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
            flags.add("report")
        if re.search(r"\b(human resources|hr)\b", lowered):
            flags.add("hr")
        if re.search(r"\bregional management|management\b", lowered):
            flags.add("management")
        if re.search(r"\b(key|major)\s+accounts?\b|\b(accounts?)\b[^.]{0,20}\b(removed|stripped|taken away)\b|\b(removed|stripped|took away)\b[^.]{0,20}\baccounts?\b", lowered):
            flags.add("accounts")
        if re.search(r"\bovertime\b", lowered):
            flags.add("overtime")
        if re.search(r"\bshift(s)?\b", lowered):
            flags.add("shifts")
        if re.search(r"\b(absences?|attendance|treatment-related absences?)\b", lowered):
            flags.add("absences")
        if re.search(r"\b(disciplined|discipline|wrote me up|write-up|write up)\b", lowered):
            flags.add("discipline")
        if re.search(r"\b(accommodation|accommodate|light duty|schedule flexibility|medical restrictions?|doctor-imposed restrictions?)\b", lowered):
            flags.add("accommodation")
        if re.search(r"\b(restrictions?|light duty|schedule flexibility)\b", lowered):
            flags.add("restrictions")
        if re.search(r"\b(terminated|fired)\b", lowered) or re.search(r"\b(end(?:ed|ing))\b[^.]{0,40}\bemployment\b", lowered):
            flags.add("termination")
        if re.search(r"\b(wages|pay|salary|income|benefits)\b", lowered):
            flags.add("economic_harm")
        if re.search(r"\b(career opportunities|future opportunities|opportunities)\b", lowered):
            flags.add("opportunities")
        return flags

    kept: List[str] = []
    for candidate in allegations:
        candidate_tokens = _tokens(candidate)
        candidate_categories = _categories(candidate)
        candidate_features = _features(candidate)
        skip = False
        for existing in kept:
            existing_tokens = _tokens(existing)
            existing_categories = _categories(existing)
            existing_features = _features(existing)
            if not candidate_tokens or not existing_tokens:
                continue
            if not (candidate_categories & existing_categories):
                continue
            overlap = len(candidate_tokens & existing_tokens) / max(1, min(len(candidate_tokens), len(existing_tokens)))
            shared_features = candidate_features & existing_features
            if overlap >= 0.7:
                skip = True
                break
            if "adverse" in candidate_categories and "adverse" in existing_categories and len(shared_features) >= 3:
                skip = True
                break
            if "report" in candidate_categories and "report" in existing_categories and "accommodation" in shared_features and len(shared_features) >= 2:
                skip = True
                break
        if not skip:
            kept.append(candidate)
    return kept


def _build_factual_allegation_groups(allegations: Sequence[Any]) -> List[Dict[str, Any]]:
    ordered_titles = [
        "Protected Activity and Complaints",
        "Adverse Action and Retaliatory Conduct",
        "Damages and Resulting Harm",
        "Additional Factual Support",
    ]
    groups: Dict[str, List[Dict[str, Any]]] = {title: [] for title in ordered_titles}
    for index, value in enumerate(_listify(allegations), 1):
        text = _clean_sentence(value)
        if not text:
            continue
        lowered = text.lower()
        if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
            title = "Protected Activity and Complaints"
        elif re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", lowered):
            title = "Adverse Action and Retaliatory Conduct"
        elif re.search(r"\b(lost|damages|harm|injur|suffered|experienced|benefits|wages|salary|income)\b", lowered):
            title = "Damages and Resulting Harm"
        else:
            title = "Additional Factual Support"
        groups[title].append({"number": index, "text": text})
    return [
        {"title": title, "paragraphs": groups[title]}
        for title in ordered_titles
        if groups[title]
    ]


def _format_timeline_date(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        return datetime.strptime(text, "%Y-%m-%d").strftime("%B %-d, %Y")
    except ValueError:
        return text


def _chronology_fact_label(fact: Dict[str, Any]) -> str:
    event_label = str((fact if isinstance(fact, dict) else {}).get("event_label") or "").strip()
    if event_label:
        return event_label
    predicate_family = str((fact if isinstance(fact, dict) else {}).get("predicate_family") or "").strip().replace("_", " ")
    if predicate_family:
        return predicate_family.title()
    return "Event"


def _join_chronology_segments(segments: List[str]) -> str:
    if not segments:
        return ""
    if len(segments) == 1:
        return segments[0]
    if len(segments) == 2:
        return f"{segments[0]} and {segments[1]}"
    return f"{', '.join(segments[:-1])}, and {segments[-1]}"


def _build_anchored_chronology_summary(intake_case_file: Dict[str, Any], *, limit: int = 3) -> List[str]:
    facts = [dict(item) for item in list(intake_case_file.get("canonical_facts") or []) if isinstance(item, dict)]
    relations = [dict(item) for item in list(intake_case_file.get("timeline_relations") or []) if isinstance(item, dict)]
    if not facts or not relations:
        return []

    fact_by_id = {
        str(fact.get("fact_id") or "").strip(): fact
        for fact in facts
        if str(fact.get("fact_id") or "").strip()
    }
    relation_records = []
    for relation in relations:
        if str(relation.get("relation_type") or "").strip().lower() != "before":
            continue
        source_id = str(relation.get("source_fact_id") or "").strip()
        target_id = str(relation.get("target_fact_id") or "").strip()
        source_fact = fact_by_id.get(source_id)
        target_fact = fact_by_id.get(target_id)
        if not source_fact or not target_fact:
            continue
        source_date = _format_timeline_date((source_fact.get("temporal_context") or {}).get("start_date") or relation.get("source_start_date"))
        target_date = _format_timeline_date((target_fact.get("temporal_context") or {}).get("start_date") or relation.get("target_start_date"))
        if not source_date or not target_date:
            continue
        relation_records.append(
            {
                "key": (source_id, target_id),
                "source_id": source_id,
                "target_id": target_id,
                "source_fact": source_fact,
                "target_fact": target_fact,
                "source_date": source_date,
                "target_date": target_date,
            }
        )
    if not relation_records:
        return []

    outgoing: Dict[str, List[Dict[str, Any]]] = {}
    incoming_count: Dict[str, int] = {}
    for record in relation_records:
        outgoing.setdefault(record["source_id"], []).append(record)
        incoming_count[record["target_id"]] = incoming_count.get(record["target_id"], 0) + 1
        incoming_count.setdefault(record["source_id"], incoming_count.get(record["source_id"], 0))

    lines: List[str] = []
    seen = set()
    used_keys = set()
    for record in relation_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        if incoming_count.get(record["source_id"], 0) != 0 or len(outgoing.get(record["source_id"], [])) != 1:
            continue
        chain = [record]
        next_id = record["target_id"]
        temp_used = {record["key"]}
        while len(outgoing.get(next_id, [])) == 1 and incoming_count.get(next_id, 0) == 1:
            next_record = outgoing[next_id][0]
            if next_record["key"] in temp_used:
                break
            chain.append(next_record)
            temp_used.add(next_record["key"])
            next_id = next_record["target_id"]
        if len(chain) < 2:
            continue
        segments = [
            f"{_chronology_fact_label(chain[0]['source_fact'])} on {chain[0]['source_date']}"
        ]
        segments.extend(
            f"{_chronology_fact_label(item['target_fact'])} on {item['target_date']}"
            for item in chain
        )
        line = f"{_join_chronology_segments(segments)} occurred in sequence."
        last_target = chain[-1]["target_fact"]
        target_context = last_target.get("temporal_context") if isinstance(last_target.get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f" The later date is derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        used_keys.update(temp_used)
        lines.append(line)

    for record in relation_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        source_label = _chronology_fact_label(record["source_fact"])
        target_label = _chronology_fact_label(record["target_fact"]).lower()
        line = f"{source_label} on {record['source_date']} preceded {target_label} on {record['target_date']}."
        target_context = record["target_fact"].get("temporal_context") if isinstance(record["target_fact"].get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f" The later date is derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        lines.append(line)
    return lines


def _claim_chronology_focus_families(claim_type: str, claim_name: str) -> set[str]:
    combined = " ".join([str(claim_type or ""), str(claim_name or "")]).strip().lower()
    if any(token in combined for token in ("retaliat", "reprisal", "protected activity")):
        return {"protected_activity", "adverse_action", "causation"}
    if any(token in combined for token in ("due process", "grievance", "hearing", "appeal", "review", "notice")):
        return {"notice_chain", "hearing_process", "response_timeline", "adverse_action"}
    if any(token in combined for token in ("accommodation", "disabil", "fair housing", "discrimination", "termination", "denial")):
        return {"adverse_action", "response_timeline", "notice_chain"}
    return set()


def _build_claim_chronology_support(
    intake_case_file: Dict[str, Any],
    *,
    claim_type: str,
    claim_name: str,
    limit: int = 2,
) -> List[str]:
    facts = [dict(item) for item in list(intake_case_file.get("canonical_facts") or []) if isinstance(item, dict)]
    relations = [dict(item) for item in list(intake_case_file.get("timeline_relations") or []) if isinstance(item, dict)]
    if not facts or not relations:
        return []

    fact_by_id = {
        str(fact.get("fact_id") or "").strip(): fact
        for fact in facts
        if str(fact.get("fact_id") or "").strip()
    }
    focus_families = _claim_chronology_focus_families(claim_type, claim_name)
    relation_records = []
    for relation in relations:
        if str(relation.get("relation_type") or "").strip().lower() != "before":
            continue
        source_id = str(relation.get("source_fact_id") or "").strip()
        target_id = str(relation.get("target_fact_id") or "").strip()
        source_fact = fact_by_id.get(source_id)
        target_fact = fact_by_id.get(target_id)
        if not source_fact or not target_fact:
            continue
        source_date = _format_timeline_date((source_fact.get("temporal_context") or {}).get("start_date") or relation.get("source_start_date"))
        target_date = _format_timeline_date((target_fact.get("temporal_context") or {}).get("start_date") or relation.get("target_start_date"))
        if not source_date or not target_date:
            continue
        relation_records.append(
            {
                "key": (source_id, target_id),
                "source_id": source_id,
                "target_id": target_id,
                "source_fact": source_fact,
                "target_fact": target_fact,
                "source_date": source_date,
                "target_date": target_date,
                "source_family": str(source_fact.get("predicate_family") or "").strip().lower(),
                "target_family": str(target_fact.get("predicate_family") or "").strip().lower(),
            }
        )
    if not relation_records:
        return []

    filtered_records = [
        record for record in relation_records
        if not focus_families or ({record['source_family'], record['target_family']} & focus_families)
    ]
    if not filtered_records:
        filtered_records = relation_records

    outgoing: Dict[str, List[Dict[str, Any]]] = {}
    incoming_count: Dict[str, int] = {}
    for record in filtered_records:
        outgoing.setdefault(record["source_id"], []).append(record)
        incoming_count[record["target_id"]] = incoming_count.get(record["target_id"], 0) + 1
        incoming_count.setdefault(record["source_id"], incoming_count.get(record["source_id"], 0))

    lines: List[str] = []
    seen = set()
    fallback_lines: List[str] = []

    used_keys = set()
    for record in filtered_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        if incoming_count.get(record["source_id"], 0) != 0 or len(outgoing.get(record["source_id"], [])) != 1:
            continue
        chain = [record]
        next_id = record["target_id"]
        temp_used = {record["key"]}
        while len(outgoing.get(next_id, [])) == 1 and incoming_count.get(next_id, 0) == 1:
            next_record = outgoing[next_id][0]
            if next_record["key"] in temp_used:
                break
            chain.append(next_record)
            temp_used.add(next_record["key"])
            next_id = next_record["target_id"]
        if len(chain) < 2:
            continue
        segments = [
            f"{_chronology_fact_label(chain[0]['source_fact']).lower()} on {chain[0]['source_date']}"
        ]
        segments.extend(
            f"{_chronology_fact_label(item['target_fact']).lower()} on {item['target_date']}"
            for item in chain
        )
        line = f"The chronology shows {_join_chronology_segments(segments)} in sequence."
        last_target = chain[-1]["target_fact"]
        target_context = last_target.get("temporal_context") if isinstance(last_target.get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f" The later date is derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        used_keys.update(temp_used)
        lines.append(line)

    for record in filtered_records:
        if record["key"] in used_keys:
            continue
        source_label = _chronology_fact_label(record["source_fact"])
        target_label = _chronology_fact_label(record["target_fact"]).lower()
        line = f"The chronology shows {source_label.lower()} on {record['source_date']} before {target_label} on {record['target_date']}."
        target_context = record["target_fact"].get("temporal_context") if isinstance(record["target_fact"].get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f" The later date is derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        fallback_lines.append(line)

    if lines:
        return lines
    if fallback_lines:
        return fallback_lines[:1]
    return []


def _claim_temporal_gap_focus(claim_type: str, claim_name: str) -> Dict[str, set[str]]:
    combined = " ".join([str(claim_type or ""), str(claim_name or "")]).strip().lower()
    issue_families = {"timeline"}
    element_tags = {"timeline"}
    objectives = {"timeline", "exact_dates"}
    if any(token in combined for token in ("retaliat", "reprisal", "protected activity")):
        issue_families.update({"causation", "protected_activity", "adverse_action"})
        element_tags.update({"causation", "protected_activity", "adverse_action"})
        objectives.update({"causation_link", "causation_sequence", "anchor_adverse_action"})
    if any(token in combined for token in ("due process", "grievance", "hearing", "appeal", "review", "notice")):
        issue_families.update({"notice_chain", "hearing_process", "response_timeline"})
        element_tags.update({"notice", "hearing", "appeal", "response", "review"})
        objectives.update({"anchor_appeal_rights", "hearing_request_timing", "response_dates"})
    if any(token in combined for token in ("accommodation", "disabil", "discrimination", "termination", "denial")):
        issue_families.update({"adverse_action", "notice_chain", "response_timeline"})
        element_tags.update({"adverse_action", "response", "notice"})
        objectives.update({"response_dates", "anchor_adverse_action"})
    return {
        "issue_families": issue_families,
        "element_tags": element_tags,
        "objectives": objectives,
    }


def _build_claim_temporal_gap_hints(
    intake_case_file: Dict[str, Any],
    *,
    claim_type: str,
    claim_name: str,
    limit: int = 3,
) -> List[Dict[str, str]]:
    if not isinstance(intake_case_file, dict):
        return []

    focus = _claim_temporal_gap_focus(claim_type, claim_name)
    normalized_claim_type = str(claim_type or "").strip().lower()
    hints: List[Dict[str, str]] = []
    seen = set()

    def _append_hint(summary: str, *, label: str = "Chronology gap") -> None:
        cleaned_summary = _clean_text(summary)
        if not cleaned_summary:
            return
        key = (label.lower(), cleaned_summary.lower())
        if key in seen:
            return
        seen.add(key)
        hints.append(
            {
                "name": label,
                "citation": "",
                "suggested_action": cleaned_summary,
            }
        )

    for issue in _listify(intake_case_file.get("temporal_issue_registry")):
        if not isinstance(issue, dict):
            continue
        status = str(issue.get("status") or "open").strip().lower()
        if status not in {"open", "blocking", "warning"}:
            continue
        issue_claim_types = {str(item).strip().lower() for item in _listify(issue.get("claim_types")) if str(item).strip()}
        issue_element_tags = {str(item).strip().lower() for item in _listify(issue.get("element_tags")) if str(item).strip()}
        if issue_claim_types:
            if normalized_claim_type not in issue_claim_types:
                continue
        elif issue_element_tags and not (issue_element_tags & focus["element_tags"]):
            continue
        _append_hint(str(issue.get("summary") or ""))

    blocker_summary = intake_case_file.get("blocker_follow_up_summary") if isinstance(intake_case_file.get("blocker_follow_up_summary"), dict) else {}
    for blocker in _listify(blocker_summary.get("blocking_items")):
        if not isinstance(blocker, dict):
            continue
        issue_family = str(blocker.get("issue_family") or "").strip().lower()
        primary_objective = str(blocker.get("primary_objective") or "").strip().lower()
        blocker_objectives = {str(item).strip().lower() for item in _listify(blocker.get("blocker_objectives")) if str(item).strip()}
        matched_objectives = ({primary_objective} if primary_objective else set()) | blocker_objectives
        if issue_family:
            if issue_family not in focus["issue_families"] and not (matched_objectives & focus["objectives"]):
                continue
        elif not (matched_objectives & focus["objectives"]):
            continue
        _append_hint(str(blocker.get("reason") or ""))

    return hints[:limit]


def _normalize_claim_type_key(value: Any) -> str:
    return str(value or "").strip().lower().replace("_", " ")


def _extract_validation_claim_for_type(
    validation_summary: Any,
    claim_type: str,
) -> Dict[str, Any]:
    if not isinstance(validation_summary, dict):
        return {}
    claims = validation_summary.get("claims")
    if not isinstance(claims, dict):
        return {}
    direct = claims.get(claim_type)
    if isinstance(direct, dict):
        return direct
    normalized_claim_type = _normalize_claim_type_key(claim_type)
    for key, value in claims.items():
        if _normalize_claim_type_key(key) == normalized_claim_type and isinstance(value, dict):
            return value
    return {}


def _normalize_theorem_export_metadata_for_drafting(value: Any) -> Dict[str, Any]:
    metadata = value if isinstance(value, dict) else {}
    unresolved_issue_ids = [
        str(item).strip()
        for item in _listify(metadata.get("unresolved_temporal_issue_ids"))
        if str(item).strip()
    ]
    proof_objectives = [
        str(item).strip()
        for item in _listify(metadata.get("temporal_proof_objectives"))
        if str(item).strip()
    ]
    unresolved_count = int(
        metadata.get("unresolved_temporal_issue_count", len(unresolved_issue_ids)) or 0
    )
    chronology_task_count = int(metadata.get("chronology_task_count", 0) or 0)
    return {
        "unresolved_count": unresolved_count,
        "chronology_task_count": chronology_task_count,
        "unresolved_issue_ids": unresolved_issue_ids,
        "proof_objectives": proof_objectives,
        "has_signal": bool(
            unresolved_count
            or chronology_task_count
            or unresolved_issue_ids
            or proof_objectives
            or metadata.get("chronology_blocked")
        ),
    }


def _contains_temporal_sequence_marker(value: Any) -> bool:
    text = str(value or "")
    if not text:
        return False
    return bool(
        _contains_date_anchor(text)
        or re.search(
            r"\b(before|after|preced(?:e|ed|es|ing)|follow(?:ed|ing|s)?|sequence|timeline|chronolog(?:y|ical))\b",
            text,
            flags=re.IGNORECASE,
        )
    )


def _build_claim_reasoning_temporal_fallbacks(
    validation_claim: Any,
    *,
    claim_type: str,
    claim_name: str,
    limit: int = 2,
) -> Tuple[List[str], List[Dict[str, str]]]:
    claim_validation = validation_claim if isinstance(validation_claim, dict) else {}
    elements = claim_validation.get("elements") if isinstance(claim_validation.get("elements"), list) else []
    support_lines: List[str] = []
    gap_hints: List[Dict[str, str]] = []
    seen_support = set()
    seen_hints = set()

    for element in elements:
        if not isinstance(element, dict):
            continue
        reasoning = element.get("reasoning_diagnostics")
        if not isinstance(reasoning, dict):
            reasoning = {}
        hybrid_reasoning = reasoning.get("hybrid_reasoning")
        if not isinstance(hybrid_reasoning, dict):
            hybrid_reasoning = {}
        hybrid_result = hybrid_reasoning.get("result")
        if not isinstance(hybrid_result, dict):
            hybrid_result = {}
        proof_artifact = hybrid_result.get("proof_artifact")
        if not isinstance(proof_artifact, dict):
            proof_artifact = {}
        temporal_proof_bundle = reasoning.get("temporal_proof_bundle")
        if not isinstance(temporal_proof_bundle, dict):
            temporal_proof_bundle = {}
        theorem_exports = temporal_proof_bundle.get("theorem_exports")
        if not isinstance(theorem_exports, dict):
            theorem_exports = {}

        proof_metadata = _normalize_theorem_export_metadata_for_drafting(
            proof_artifact.get("theorem_export_metadata")
        )
        bundle_metadata = _normalize_theorem_export_metadata_for_drafting(
            theorem_exports.get("theorem_export_metadata")
        )
        active_metadata = proof_metadata if proof_metadata["has_signal"] else bundle_metadata

        proof_sentence = _clean_sentence(proof_artifact.get("sentence") or "")
        if (
            proof_sentence
            and active_metadata["has_signal"]
            and _contains_temporal_sequence_marker(proof_sentence)
        ):
            marker = proof_sentence.lower()
            if marker not in seen_support:
                seen_support.add(marker)
                support_lines.append(proof_sentence)

        unresolved_count = int(active_metadata.get("unresolved_count", 0) or 0)
        chronology_task_count = int(active_metadata.get("chronology_task_count", 0) or 0)
        if unresolved_count <= 0 and chronology_task_count <= 0:
            continue

        element_label = _clean_text(
            element.get("element_text") or element.get("element_id") or claim_name or claim_type
        ) or "This claim element"
        summary_parts: List[str] = []
        if unresolved_count > 0:
            summary_parts.append(f"{unresolved_count} unresolved chronology issue(s)")
        if chronology_task_count > 0:
            summary_parts.append(f"{chronology_task_count} chronology task(s)")
        objective_text = ""
        if active_metadata.get("proof_objectives"):
            objective_text = f" Focus on {str(active_metadata['proof_objectives'][0]).replace('_', ' ')}."
        summary = (
            f"{element_label} still carries {' and '.join(summary_parts)} in the proof handoff."
            f"{objective_text}"
        ).strip()
        marker = summary.lower()
        if marker in seen_hints:
            continue
        seen_hints.add(marker)
        gap_hints.append(
            {
                "name": "Chronology gap",
                "citation": "",
                "suggested_action": summary,
            }
        )

    return support_lines[:limit], gap_hints[:limit]


def _listify(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return list(value)
    if isinstance(value, tuple):
        return list(value)
    return [value]


def _roman_numeral(number: int) -> str:
    values: Sequence[Tuple[int, str]] = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    remaining = max(int(number or 0), 1)
    output: List[str] = []
    for integer, numeral in values:
        while remaining >= integer:
            output.append(numeral)
            remaining -= integer
    return "".join(output)


def _exhibit_label(index: int) -> str:
    ordinal = max(index, 1)
    letters: List[str] = []
    while ordinal > 0:
        ordinal -= 1
        letters.append(chr(ord("A") + (ordinal % 26)))
        ordinal //= 26
    return f"Exhibit {''.join(reversed(letters))}"


class ComplaintDocumentBuilder:
    """Build and export a court-style complaint draft from mediator state."""

    def __init__(self, mediator):
        self.mediator = mediator

    def build(
        self,
        *,
        court_name: Optional[str] = None,
        district: Optional[str] = None,
        county: Optional[str] = None,
        division: Optional[str] = None,
        court_header_override: Optional[str] = None,
        case_number: Optional[str] = None,
        lead_case_number: Optional[str] = None,
        related_case_number: Optional[str] = None,
        assigned_judge: Optional[str] = None,
        courtroom: Optional[str] = None,
        title_override: Optional[str] = None,
        plaintiff_names: Optional[List[str]] = None,
        defendant_names: Optional[List[str]] = None,
        requested_relief: Optional[List[str]] = None,
        jury_demand: Optional[bool] = None,
        jury_demand_text: Optional[str] = None,
        signer_name: Optional[str] = None,
        signer_title: Optional[str] = None,
        signer_firm: Optional[str] = None,
        signer_bar_number: Optional[str] = None,
        signer_contact: Optional[str] = None,
        additional_signers: Optional[List[Dict[str, str]]] = None,
        declarant_name: Optional[str] = None,
        service_method: Optional[str] = None,
        service_recipients: Optional[List[str]] = None,
        service_recipient_details: Optional[List[Dict[str, str]]] = None,
        signature_date: Optional[str] = None,
        verification_date: Optional[str] = None,
        service_date: Optional[str] = None,
        affidavit_title: Optional[str] = None,
        affidavit_intro: Optional[str] = None,
        affidavit_facts: Optional[List[str]] = None,
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]] = None,
        affidavit_include_complaint_exhibits: Optional[bool] = None,
        affidavit_venue_lines: Optional[List[str]] = None,
        affidavit_jurat: Optional[str] = None,
        affidavit_notary_block: Optional[List[str]] = None,
        user_id: Optional[str] = None,
        base_formal_complaint: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        kg = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "knowledge_graph") if phase_manager else None
        dg = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "dependency_graph") if phase_manager else None
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        legal_graph = phase_manager.get_phase_data(ComplaintPhase.FORMALIZATION, "legal_graph") if phase_manager else None
        matching_results = phase_manager.get_phase_data(ComplaintPhase.FORMALIZATION, "matching_results") if phase_manager else None

        resolved_user_id = user_id or getattr(self.mediator.state, "username", None) or getattr(self.mediator.state, "hashed_username", None) or "anonymous"
        base = dict(base_formal_complaint or {})
        evidence_records = self._safe_call("get_user_evidence", resolved_user_id, default=[])
        authority_records = self._safe_call("get_legal_authorities", resolved_user_id, default=[])
        support_links = self._safe_call("get_claim_support", resolved_user_id, default=[])
        intake_summary = self._collect_intake_summary()
        parties = self._build_parties(
            kg,
            base.get("parties"),
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
        )
        jurisdiction = str(base.get("jurisdiction") or self._infer_jurisdiction(legal_graph) or "federal")
        caption = self._build_caption(
            parties=parties,
            jurisdiction=jurisdiction,
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            court_header_override=court_header_override,
            case_number=case_number,
            lead_case_number=lead_case_number,
            related_case_number=related_case_number,
            assigned_judge=assigned_judge,
            courtroom=courtroom,
            title=title_override or str(base.get("title") or "").strip(),
        )
        factual_allegations = self._build_factual_allegations(
            base.get("factual_allegations"),
            kg,
            intake_summary,
        )
        exhibits = self._build_exhibits(evidence_records, support_links)
        claims = self._build_claims(
            matching_results=matching_results,
            dependency_graph=dg,
            legal_graph=legal_graph,
            authority_records=authority_records,
            exhibits=exhibits,
            factual_allegations=factual_allegations,
            user_id=resolved_user_id,
        )
        requested_relief = self._build_requested_relief(requested_relief or base.get("prayer_for_relief"))
        jury_demand_block = self._build_jury_demand(jury_demand=jury_demand, jury_demand_text=jury_demand_text)
        caption["jury_demand_notice"] = "JURY TRIAL DEMANDED" if jury_demand_block else None
        nature_of_action = self._build_nature_of_action(base.get("statement_of_claim"), claims)
        jurisdiction_statement = self._build_jurisdiction_statement(jurisdiction, authority_records)
        venue_statement = self._build_venue_statement(district, county, division, jurisdiction)
        signature_block = self._build_signature_block(
            parties,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            signature_date=signature_date,
        )
        verification = self._build_verification(
            parties,
            declarant_name=declarant_name,
            signer_name=signer_name,
            verification_date=verification_date,
        )
        certificate_of_service = self._build_certificate_of_service(
            parties,
            signer_name=signer_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            service_date=service_date,
            jurisdiction=jurisdiction,
        )
        affidavit = self._build_affidavit(
            parties=parties,
            caption=caption,
            factual_allegations=factual_allegations,
            claims=claims,
            exhibits=exhibits,
            verification=verification,
            signature_block=signature_block,
            affidavit_overrides=self._build_affidavit_overrides(
                affidavit_title=affidavit_title,
                affidavit_intro=affidavit_intro,
                affidavit_facts=affidavit_facts,
                affidavit_supporting_exhibits=affidavit_supporting_exhibits,
                affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
                affidavit_venue_lines=affidavit_venue_lines,
                affidavit_jurat=affidavit_jurat,
                affidavit_notary_block=affidavit_notary_block,
            ),
        )
        legal_standards = [
            {
                "claim_name": claim.get("claim_name", ""),
                "claim_type": claim.get("claim_type", ""),
                "standard": claim.get("legal_standard", ""),
                "citations": [
                    item.get("citation", "")
                    for item in claim.get("legal_standard_elements", [])
                    if item.get("citation")
                ],
            }
            for claim in claims
        ]
        anchored_chronology_summary = _build_anchored_chronology_summary(
            intake_case_file if isinstance(intake_case_file, dict) else {}
        )

        draft = {
            **base,
            "generated_at": _utc_now_isoformat(),
            "court_header": caption["court_header"],
            "caption": caption,
            "case_number": caption["case_number"],
            "title": caption["case_title"],
            "parties": parties,
            "nature_of_action": nature_of_action,
            "jurisdiction": jurisdiction,
            "jurisdiction_statement": jurisdiction_statement,
            "venue_statement": venue_statement,
            "statement_of_claim": str(base.get("statement_of_claim") or nature_of_action),
            "summary_of_facts": factual_allegations[: min(len(factual_allegations), 6)],
            "factual_allegations": factual_allegations,
            "factual_allegation_groups": _build_factual_allegation_groups(factual_allegations),
            "anchored_chronology_summary": anchored_chronology_summary,
            "legal_standards": legal_standards,
            "legal_claims": claims,
            "claims_for_relief": claims,
            "prayer_for_relief": requested_relief,
            "requested_relief": requested_relief,
            "jury_demand": jury_demand_block,
            "supporting_documents": exhibits,
            "exhibits": exhibits,
            "supporting_exhibits": exhibits,
            "intake_summary": intake_summary,
            "signature_block": signature_block,
            "verification": verification,
            "certificate_of_service": certificate_of_service,
            "affidavit": affidavit,
        }
        draft["draft_text"] = self.render_text(draft)
        return draft

    def export(self, draft: Dict[str, Any], output_path: str, *, format: Optional[str] = None) -> Dict[str, Any]:
        destination = Path(output_path)
        resolved_format = (format or destination.suffix.lstrip(".") or "txt").lower()
        if resolved_format == "txt":
            destination.write_text(draft.get("draft_text") or self.render_text(draft), encoding="utf-8")
        elif resolved_format == "docx":
            self._write_docx(draft, destination)
        elif resolved_format == "pdf":
            self._write_pdf(draft, destination)
        else:
            raise ValueError(f"Unsupported complaint export format: {resolved_format}")

        return {
            "path": str(destination),
            "format": resolved_format,
            "bytes_written": destination.stat().st_size,
        }

    def render_text(self, draft: Dict[str, Any]) -> str:
        caption = draft.get("caption", {}) if isinstance(draft.get("caption"), dict) else {}
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        intake_summary = [item for item in _listify(draft.get("intake_summary")) if isinstance(item, dict)]
        lines: List[str] = []

        def _append_blank() -> None:
            if lines and lines[-1] != "":
                lines.append("")

        def _format_id_list(values: Any, *, prefix: str = "") -> str:
            entries = [_clean_text(value) for value in _listify(values) if _clean_text(value)]
            return f"{prefix}{', '.join(entries)}" if entries else ""

        def _claim_rows(claim: Dict[str, Any]) -> List[Dict[str, Any]]:
            rows = [
                row
                for row in _listify(claim.get("supporting_fact_provenance"))
                if isinstance(row, dict) and _clean_text(row.get("text") or "")
            ]
            if rows:
                return rows
            return [
                row
                for row in _listify(claim.get("supporting_fact_entries"))
                if isinstance(row, dict) and _clean_text(row.get("text") or "")
            ]

        def _intake_clarifications() -> List[str]:
            focus_tokens = {
                "when",
                "date",
                "time",
                "timeline",
                "after",
                "before",
                "who",
                "actor",
                "title",
                "because",
                "caus",
                "grievance",
                "hearing",
                "appeal",
                "retaliat",
                "notice",
            }
            clarifications: List[str] = []
            seen = set()
            for item in intake_summary:
                question = _clean_text(item.get("question") or "")
                answer = _clean_text(item.get("answer") or "")
                if not answer:
                    continue
                lowered = f"{question} {answer}".lower()
                if not any(token in lowered for token in focus_tokens):
                    continue
                sentence = _clean_sentence(answer)
                if not sentence:
                    continue
                marker = sentence.lower()
                if marker in seen:
                    continue
                seen.add(marker)
                clarifications.append(sentence)
            return clarifications

        lines.append(str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"))
        if caption.get("division_line"):
            lines.append(str(caption["division_line"]))
        if caption.get("county_line"):
            lines.append(str(caption["county_line"]))
        lines.append("")

        case_number = draft.get("case_number") or caption.get("case_number") or "________________"
        lines.extend(
            build_case_caption_text_lines(
                parties,
                {"case_number": case_number, **caption},
                assigned_judge_label="Assigned to:",
            )
        )
        lines.append("")
        lines.append("COMPLAINT")
        if caption.get("jury_demand_notice"):
            lines.append(str(caption["jury_demand_notice"]))

        _append_blank()
        lines.append("NATURE OF THE ACTION")
        lines.append(_clean_sentence(draft.get("nature_of_action")))

        _append_blank()
        lines.append("PARTIES")
        plaintiffs_list = _listify(parties.get("plaintiffs"))
        defendants_list = _listify(parties.get("defendants"))
        for index, party in enumerate(plaintiffs_list, 1):
            lines.append(f"{index}. Plaintiff {party} is an aggrieved party bringing this action.")
        start_index = len(plaintiffs_list)
        for offset, party in enumerate(defendants_list, 1):
            lines.append(f"{start_index + offset}. Defendant {party} is alleged to be responsible for the acts described below.")

        _append_blank()
        lines.append("JURISDICTION AND VENUE")
        lines.append(_clean_sentence(draft.get("jurisdiction_statement")))
        lines.append(_clean_sentence(draft.get("venue_statement")))

        _append_blank()
        lines.append("FACTUAL ALLEGATIONS")
        allegation_counter = 0
        allegation_groups = draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else []
        if allegation_groups:
            for group in allegation_groups:
                if not isinstance(group, dict):
                    continue
                if group.get("title"):
                    lines.append(str(group["title"]).upper())
                for entry in _listify(group.get("paragraphs")):
                    if not isinstance(entry, dict):
                        continue
                    text = _clean_sentence(entry.get("text"))
                    if not text:
                        continue
                    allegation_counter += 1
                    number = entry.get("number") if entry.get("number") not in (None, "") else allegation_counter
                    lines.append(f"{number}. {text}")
        else:
            allegation_entries = [
                entry
                for entry in _listify(draft.get("factual_allegation_paragraphs"))
                if isinstance(entry, dict) and _clean_text(entry.get("text") or "")
            ]
            if allegation_entries:
                for index, entry in enumerate(allegation_entries, 1):
                    text = _clean_sentence(entry.get("text"))
                    if not text:
                        continue
                    lines.append(f"{index}. {text}")
                    trace_bits = []
                    fact_ids = _format_id_list(entry.get("fact_ids"), prefix="fact ids: ")
                    source_ids = _format_id_list(entry.get("source_artifact_ids"), prefix="artifact ids: ")
                    if fact_ids:
                        trace_bits.append(fact_ids)
                    if source_ids:
                        trace_bits.append(source_ids)
                    if trace_bits:
                        lines.append(f"   Support Trace: {'; '.join(trace_bits)}")
            else:
                for index, allegation in enumerate(_listify(draft.get("factual_allegations")), 1):
                    lines.append(f"{index}. {_clean_sentence(allegation)}")

        incorporated_clarifications = _intake_clarifications()
        if incorporated_clarifications:
            _append_blank()
            lines.append("INTAKE CLARIFICATIONS INCORPORATED")
            for index, sentence in enumerate(incorporated_clarifications, 1):
                lines.append(f"{index}. {sentence}")

        chronology_lines = [
            _clean_sentence(item)
            for item in _listify(draft.get("anchored_chronology_summary"))
            if _clean_text(item)
        ]
        if chronology_lines:
            _append_blank()
            lines.append("ANCHORED CHRONOLOGY")
            for index, chronology_line in enumerate(chronology_lines, 1):
                lines.append(f"{index}. {chronology_line}")

        claims = _listify(draft.get("legal_claims")) or _listify(draft.get("claims_for_relief"))
        if claims:
            _append_blank()
            lines.append("CLAIMS FOR RELIEF")
        for claim in claims:
            if not isinstance(claim, dict):
                continue
            _append_blank()
            lines.append(str(claim.get("title") or claim.get("count_title") or claim.get("claim_name") or "Claim"))
            if claim.get("description"):
                lines.append(_clean_sentence(claim.get("description")))
            if claim.get("legal_standard"):
                lines.append("Legal Standard:")
                lines.append(_clean_sentence(claim.get("legal_standard")))
            for item in _listify(claim.get("legal_standard_elements")):
                if not isinstance(item, dict):
                    continue
                element_text = _clean_sentence(item.get("element") or item.get("description"))
                citation = str(item.get("citation") or "").strip()
                if citation and element_text:
                    lines.append(f"- {element_text} ({citation})")
                elif element_text:
                    lines.append(f"- {element_text}")

            supporting_facts = [_clean_sentence(fact) for fact in _listify(claim.get("supporting_facts")) if _clean_text(fact)]
            if supporting_facts:
                lines.append("Supporting Facts:")
                for fact in supporting_facts:
                    lines.append(f"- {fact}")
                rows = _claim_rows(claim)
                if rows:
                    row_lookup = {}
                    for row in rows:
                        text_key = _clean_text(row.get("text") or "").lower()
                        if text_key and text_key not in row_lookup:
                            row_lookup[text_key] = row
                    for fact in supporting_facts:
                        row = row_lookup.get(_clean_text(fact).rstrip(".").lower())
                        if not isinstance(row, dict):
                            row = row_lookup.get(_clean_text(fact).lower())
                        if not isinstance(row, dict):
                            continue
                        trace_parts = []
                        fact_ids = _format_id_list(row.get("fact_ids"), prefix="fact ids: ")
                        trace_ids = _format_id_list(row.get("support_trace_ids"), prefix="trace ids: ")
                        artifact_ids = _format_id_list(row.get("source_artifact_ids"), prefix="artifact ids: ")
                        if fact_ids:
                            trace_parts.append(fact_ids)
                        if trace_ids:
                            trace_parts.append(trace_ids)
                        if artifact_ids:
                            trace_parts.append(artifact_ids)
                        if trace_parts:
                            lines.append(f"  Support Trace: {'; '.join(trace_parts)}")

            authorities = _listify(claim.get("supporting_authorities"))
            if authorities:
                lines.append("Supporting Authorities:")
                for authority in authorities:
                    if not isinstance(authority, dict):
                        continue
                    authority_line = authority.get("citation") or authority.get("title") or "Authority"
                    if authority.get("title") and authority.get("citation"):
                        authority_line = f"{authority['citation']} - {authority['title']}"
                    lines.append(f"- {_clean_text(authority_line)}")

            missing_requirements = [item for item in _listify(claim.get("missing_requirements")) if isinstance(item, dict)]
            if missing_requirements:
                lines.append("Unresolved Gaps:")
                for requirement in missing_requirements:
                    name = _clean_text(requirement.get("name") or "Gap")
                    citation = _clean_text(requirement.get("citation") or "")
                    action = _clean_text(requirement.get("suggested_action") or "")
                    if citation and action:
                        lines.append(f"- {name} ({citation}): {action}")
                    elif citation:
                        lines.append(f"- {name} ({citation})")
                    elif action:
                        lines.append(f"- {name}: {action}")
                    else:
                        lines.append(f"- {name}")

            supporting_exhibits = _listify(claim.get("supporting_exhibits"))
            if supporting_exhibits:
                exhibit_labels = ", ".join(exhibit.get("label", "") for exhibit in supporting_exhibits if isinstance(exhibit, dict) and exhibit.get("label"))
                if exhibit_labels:
                    lines.append(f"Exhibits Incorporated by Reference: {exhibit_labels}.")

        _append_blank()
        lines.append("PRAYER FOR RELIEF")
        for item in _listify(draft.get("requested_relief")):
            lines.append(f"- {_clean_sentence(item)}")

        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            _append_blank()
            lines.append(str(jury_demand.get("title") or "JURY DEMAND").upper())
            if jury_demand.get("text"):
                lines.append(_clean_sentence(jury_demand.get("text")))

        exhibits = [item for item in _listify(draft.get("exhibits")) if isinstance(item, dict)]
        if exhibits:
            _append_blank()
            lines.append("EXHIBITS")
            for exhibit in exhibits:
                reference = _clean_text(exhibit.get("reference") or exhibit.get("source_url") or exhibit.get("cid") or "")
                summary = _clean_text(exhibit.get("summary") or exhibit.get("description") or "")
                title = _clean_text(exhibit.get("title") or exhibit.get("description") or "Supporting exhibit")
                exhibit_line = f"{exhibit.get('label', 'Exhibit')} - {title}"
                context_parts: List[str] = []
                if reference:
                    context_parts.append(reference)
                modality = _clean_text(exhibit.get("type") or exhibit.get("source_kind") or "")
                if modality in {"policy_document", "file_evidence"}:
                    context_parts.append(f"modality: {modality}")
                if context_parts:
                    exhibit_line = f"{exhibit_line} ({'; '.join(context_parts)})"
                lines.append(exhibit_line)
                if summary:
                    lines.append(f"  {_clean_sentence(summary)}")
                exhibit_fact_ids = _format_id_list(exhibit.get("fact_ids") or exhibit.get("canonical_fact_ids"), prefix="fact ids: ")
                exhibit_trace_ids = _format_id_list(exhibit.get("support_trace_ids"), prefix="trace ids: ")
                if exhibit_fact_ids or exhibit_trace_ids:
                    trace_parts = [item for item in [exhibit_fact_ids, exhibit_trace_ids] if item]
                    lines.append(f"  Support Trace: {'; '.join(trace_parts)}")

        drafting_readiness = draft.get("drafting_readiness", {}) if isinstance(draft.get("drafting_readiness"), dict) else {}
        document_provenance_summary = (
            draft.get("document_provenance_summary", {})
            if isinstance(draft.get("document_provenance_summary"), dict)
            else {}
        )
        if drafting_readiness or document_provenance_summary:
            _append_blank()
            lines.append("DRAFTING READINESS")
            if drafting_readiness:
                lines.append(f"Status: {_clean_text(drafting_readiness.get('status') or 'unknown').upper() or 'UNKNOWN'}")
                warning_count = drafting_readiness.get("warning_count")
                if warning_count not in (None, ""):
                    lines.append(f"Warning count: {int(warning_count)}")
                for section in _listify(drafting_readiness.get("sections")):
                    if not isinstance(section, dict):
                        continue
                    section_title = _clean_text(section.get("title") or section.get("name") or "")
                    section_status = _clean_text(section.get("status") or "")
                    if section_title and section_status:
                        lines.append(f"- {section_title}: {section_status}")
            if document_provenance_summary:
                ratio = document_provenance_summary.get("fact_backed_ratio")
                if ratio is not None:
                    lines.append(f"Fact-backed ratio: {float(ratio):.2f}")
                summary_fact_backed = int(document_provenance_summary.get("summary_fact_backed_count") or 0)
                summary_fact_total = int(document_provenance_summary.get("summary_fact_count") or 0)
                claim_fact_backed = int(document_provenance_summary.get("claim_supporting_fact_backed_count") or 0)
                claim_fact_total = int(document_provenance_summary.get("claim_supporting_fact_count") or 0)
                lines.append(f"Summary facts grounded: {summary_fact_backed}/{summary_fact_total}")
                lines.append(f"Claim support grounded: {claim_fact_backed}/{claim_fact_total}")
                if document_provenance_summary.get("low_grounding_flag"):
                    lines.append("Grounding warning: canonical fact and artifact support should be increased before filing.")

        explicit_anchor_markers = " ".join(lines).lower()
        unresolved_anchor_questions: List[str] = []
        if "grievance_hearing" not in explicit_anchor_markers and "grievance hearing" not in explicit_anchor_markers:
            unresolved_anchor_questions.append("What grievance hearing request was made, when was it made, and who responded on behalf of HACC?")
        if "appeal_rights" not in explicit_anchor_markers and "appeal rights" not in explicit_anchor_markers:
            unresolved_anchor_questions.append("What appeal rights notice was provided, on what date, and by which HACC actor?")
        if unresolved_anchor_questions:
            _append_blank()
            lines.append("UNRESOLVED FACTUAL OR LEGAL GAPS BEFORE FORMALIZATION")
            for index, question in enumerate(unresolved_anchor_questions, 1):
                lines.append(f"{index}. {_clean_sentence(question)}")

        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            _append_blank()
            lines.append(str(affidavit.get("title") or "AFFIDAVIT IN SUPPORT OF COMPLAINT").upper())
            for venue_line in _listify(affidavit.get("venue_lines")):
                cleaned_venue_line = _clean_text(venue_line)
                if cleaned_venue_line:
                    lines.append(cleaned_venue_line)
            if affidavit.get("intro"):
                lines.append(_clean_sentence(affidavit.get("intro")))
            if affidavit.get("knowledge_graph_note"):
                lines.append(_clean_sentence(affidavit.get("knowledge_graph_note")))
            if affidavit.get("facts"):
                lines.append("Affiant states as follows:")
                for index, fact in enumerate(_listify(affidavit.get("facts")), 1):
                    cleaned_fact = _clean_sentence(fact)
                    if cleaned_fact:
                        lines.append(f"{index}. {cleaned_fact}")
            for exhibit in _listify(affidavit.get("supporting_exhibits")):
                if not isinstance(exhibit, dict):
                    continue
                exhibit_line = f"{exhibit.get('label', 'Exhibit')} - {_clean_text(exhibit.get('title') or 'Supporting exhibit')}"
                if exhibit.get("link"):
                    exhibit_line = f"{exhibit_line} ({exhibit['link']})"
                lines.append(exhibit_line)
            if affidavit.get("dated"):
                lines.append(str(affidavit["dated"]))
            if affidavit.get("signature_line"):
                lines.append(str(affidavit["signature_line"]))
            if affidavit.get("jurat"):
                lines.append(str(affidavit["jurat"]))
            for notary_line in _listify(affidavit.get("notary_block")):
                cleaned_notary_line = _clean_text(notary_line)
                if cleaned_notary_line:
                    lines.append(cleaned_notary_line)

        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            _append_blank()
            lines.append(str(verification.get("title") or "VERIFICATION").upper())
            if verification.get("text"):
                lines.append(_clean_sentence(verification.get("text")))
            if verification.get("dated"):
                lines.append(str(verification["dated"]))
            if verification.get("signature_line"):
                lines.append(str(verification["signature_line"]))

        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            _append_blank()
            lines.append(str(certificate_of_service.get("title") or "CERTIFICATE OF SERVICE").upper())
            if certificate_of_service.get("text"):
                lines.append(_clean_sentence(certificate_of_service.get("text")))
            for detail_line in _listify(certificate_of_service.get("detail_lines")):
                if detail_line:
                    lines.append(str(detail_line))
            if certificate_of_service.get("dated"):
                lines.append(str(certificate_of_service["dated"]))
            if certificate_of_service.get("signature_line"):
                lines.append(str(certificate_of_service["signature_line"]))

        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        _append_blank()
        lines.append("SIGNATURE BLOCK")
        lines.append("Respectfully submitted,")
        lines.extend(self._signature_block_lines(signature_block))
        return "\n".join(line for line in lines if line is not None)

    def _safe_call(self, method_name: str, *args, default=None, **kwargs):
        method = getattr(self.mediator, method_name, None)
        if not callable(method):
            return default
        try:
            return method(*args, **kwargs)
        except Exception:
            return default

    def _collect_intake_summary(self) -> List[Dict[str, str]]:
        summary: List[Dict[str, str]] = []
        seen = set()
        for inquiry in _listify(getattr(self.mediator.state, "inquiries", [])):
            if not isinstance(inquiry, dict):
                continue
            question = _clean_text(inquiry.get("question"))
            answer = _clean_text(inquiry.get("answer"))
            if not question or not answer:
                continue
            marker = (question, answer)
            if marker in seen:
                continue
            seen.add(marker)
            summary.append({"question": question, "answer": answer})

        answered_questions = getattr(self.mediator.state, "answered_questions", {})
        if isinstance(answered_questions, dict):
            for question, answer in answered_questions.items():
                if str(question) == "last_question":
                    continue
                question_text = _clean_text(question)
                answer_text = _clean_text(answer)
                if not question_text or not answer_text:
                    continue
                marker = (question_text, answer_text)
                if marker in seen:
                    continue
                seen.add(marker)
                summary.append({"question": question_text, "answer": answer_text})
        return summary

    def _build_parties(
        self,
        knowledge_graph,
        base_parties: Any,
        *,
        plaintiff_names: Optional[List[str]] = None,
        defendant_names: Optional[List[str]] = None,
    ) -> Dict[str, List[str]]:
        plaintiffs: List[str] = []
        defendants: List[str] = []
        others: List[str] = []

        plaintiffs.extend(_clean_text(name) for name in _listify(plaintiff_names))
        defendants.extend(_clean_text(name) for name in _listify(defendant_names))

        if isinstance(base_parties, dict):
            plaintiffs.extend(_clean_text(name) for name in _listify(base_parties.get("plaintiffs")))
            defendants.extend(_clean_text(name) for name in _listify(base_parties.get("defendants")))

        if knowledge_graph is not None:
            for entity in knowledge_graph.get_entities_by_type("person"):
                role = str(entity.attributes.get("role") or "").lower()
                name = _clean_text(entity.name)
                if not name:
                    continue
                if any(token in role for token in ("plaintiff", "complainant", "claimant", "petitioner")):
                    plaintiffs.append(name)
                elif any(token in role for token in ("defendant", "respondent", "manager", "supervisor", "owner", "employer")):
                    defendants.append(name)
                else:
                    others.append(name)

            organizations = [_clean_text(entity.name) for entity in knowledge_graph.get_entities_by_type("organization") if _clean_text(entity.name)]
            defendants.extend(organizations)

        plaintiffs = self._dedupe(plaintiffs) or ["Plaintiff"]
        defendants = self._dedupe(defendants) or ["Defendant"]
        return {
            "plaintiffs": plaintiffs,
            "defendants": defendants,
            "other_parties": self._dedupe(others),
        }

    def _build_caption(
        self,
        *,
        parties: Dict[str, List[str]],
        jurisdiction: str,
        court_name: Optional[str],
        district: Optional[str],
        county: Optional[str],
        division: Optional[str],
        court_header_override: Optional[str],
        case_number: Optional[str],
        lead_case_number: Optional[str],
        related_case_number: Optional[str],
        assigned_judge: Optional[str],
        courtroom: Optional[str],
        title: str,
    ) -> Dict[str, str]:
        district_text = _clean_text(district).upper()
        county_text = _clean_text(county).upper()
        division_text = _clean_text(division).upper()
        jurisdiction_text = str(jurisdiction).lower()
        if court_header_override:
            court_header = _clean_text(court_header_override).upper()
        elif court_name and county_text and jurisdiction_text not in {"federal", "us", "united states"}:
            county_suffix = county_text if county_text.startswith("COUNTY OF ") else f"COUNTY OF {county_text}"
            court_header = f"IN THE {_clean_text(court_name).upper()} FOR THE {county_suffix}"
        elif court_name:
            court_header = _clean_text(court_name).upper()
        elif district_text:
            court_header = f"IN THE UNITED STATES DISTRICT COURT FOR THE DISTRICT OF {district_text}"
        elif jurisdiction_text in {"federal", "us", "united states"}:
            court_header = "IN THE UNITED STATES DISTRICT COURT"
        else:
            court_header = "IN THE COURT OF COMPETENT JURISDICTION"
        resolved_title = title or f"{parties['plaintiffs'][0]} v. {parties['defendants'][0]}"
        return {
            "court_header": court_header,
            "division_line": division_text,
            "county_line": county_text if county_text and county_text not in court_header else "",
            "case_title": resolved_title,
            "case_number": _clean_text(case_number) or "________________",
            "lead_case_number": _clean_text(lead_case_number),
            "related_case_number": _clean_text(related_case_number),
            "assigned_judge": _clean_text(assigned_judge),
            "courtroom": _clean_text(courtroom),
        }

    def _build_factual_allegations(self, base_allegations: Any, knowledge_graph, intake_summary: List[Dict[str, str]]) -> List[str]:
        allegations = _expand_allegation_sources(base_allegations, limit=8)
        if knowledge_graph is not None:
            fact_entities = []
            for entity in knowledge_graph.entities.values():
                if getattr(entity, "type", "") == "fact":
                    fact_entities.extend(_expand_allegation_sources(entity.name, limit=2))
            allegations.extend(fact_entities)
        fallback = []
        for item in intake_summary:
            question = item.get("question", "")
            answer = item.get("answer", "")
            if answer:
                fallback.append(f"{question}: {answer}" if question else answer)
        allegations.extend(_expand_allegation_sources(fallback, limit=8))
        allegations.extend(_prioritized_intake_statements(intake_summary))
        deduped = self._dedupe(allegations)
        combined = _synthesize_narrative_allegations(deduped)
        pruned = _prune_subsumed_narrative_clauses(deduped)
        merged = self._dedupe(combined + pruned)
        if merged and not any(_contains_date_anchor(line) for line in merged):
            merged.append("On or about [date], HACC communicated the adverse action described in this complaint.")
        if merged and not any(_contains_actor_marker(line) for line in merged):
            merged.append("HACC decision-makers for intake, review, hearing, and adverse-action steps should be identified by name or title.")
        if merged and not any(_contains_causation_marker(line) for line in merged):
            merged.append("After Plaintiff engaged in protected activity, HACC took adverse action, and the available timeline supports a causal connection.")
        if merged and not any(_contains_hearing_timing_marker(line) for line in merged):
            merged.append(
                "Plaintiff requested an informal review or hearing on [date], and the complaint should state when that request was made in relation to each adverse-action step."
            )
        if merged and not any(_contains_response_date_marker(line) for line in merged):
            merged.append(
                "HACC response dates for notice, hearing/review requests, and final decision communications should be identified with exact dates."
            )
        if merged and not any(_contains_staff_identity_marker(line) for line in merged):
            merged.append(
                "For each key event, the complaint should identify the HACC staff member by name and title, or by the best-known title if the name is not yet confirmed."
            )
        if merged and not any("days after" in str(line).lower() or "weeks after" in str(line).lower() for line in merged):
            merged.append(
                "The complaint should describe the sequence between protected activity and adverse treatment using concrete timing, including whether action occurred days or weeks after protected activity."
            )
        return _prune_near_duplicate_allegations(merged)[:18] or ["Plaintiff will supplement the factual record with additional detail."]

    def _build_affidavit(
        self,
        *,
        parties: Dict[str, List[str]],
        caption: Dict[str, str],
        factual_allegations: List[str],
        claims: List[Dict[str, Any]],
        exhibits: List[Dict[str, Any]],
        verification: Dict[str, Any],
        signature_block: Dict[str, Any],
        affidavit_overrides: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        overrides = affidavit_overrides if isinstance(affidavit_overrides, dict) else {}
        declarant_name = self._derive_affidavit_declarant_name(parties, verification, signature_block)
        facts = list(overrides.get("facts") or self._collect_affidavit_facts(parties, factual_allegations, claims))
        supporting_exhibits = []
        for exhibit in _listify(exhibits):
            if not isinstance(exhibit, dict):
                continue
            supporting_exhibits.append(
                {
                    "label": str(exhibit.get("label") or "Exhibit").strip(),
                    "title": _clean_text(exhibit.get("title") or exhibit.get("description") or "Supporting exhibit"),
                    "link": _clean_text(exhibit.get("reference") or exhibit.get("source_url") or ""),
                }
            )
        return {
            "title": str(overrides.get("title") or f"Affidavit of {declarant_name.upper()} in Support of Complaint"),
            "declarant_name": declarant_name,
            "intro": str(
                overrides.get("intro")
                or (
                    f"I, {declarant_name}, declare under penalty of perjury that I am competent to testify to the matters stated below, "
                    "that these statements are based on my personal knowledge and the complaint intake knowledge graph assembled from the facts, records, and exhibits provided in support of this action, and that the following facts are true and correct."
                )
            ),
            "knowledge_graph_note": "This affidavit is generated from the complaint intake knowledge graph and supporting records rather than a turn-by-turn chat transcript.",
            "venue_lines": list(overrides.get("venue_lines") or self._build_affidavit_venue_lines(caption)),
            "facts": facts,
            "supporting_exhibits": list(
                overrides.get("supporting_exhibits")
                or ([] if overrides.get("include_complaint_exhibits") is False else supporting_exhibits)
            ),
            "dated": verification.get("dated") or signature_block.get("dated") or self._format_dated_line("Executed on", None),
            "signature_line": verification.get("signature_line") or signature_block.get("signature_line") or f"/s/ {declarant_name}",
            "jurat": str(overrides.get("jurat") or f"Subscribed and sworn to (or affirmed) before me on __________________ by {declarant_name}."),
            "notary_block": list(
                overrides.get("notary_block")
                or [
                    "__________________________________",
                    "Notary Public",
                    "My commission expires: __________________",
                ]
            ),
            "case_number": caption.get("case_number") or "________________",
        }

    def _build_affidavit_overrides(
        self,
        *,
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        normalized_facts: List[str] = []
        for value in _listify(affidavit_facts):
            cleaned = self._sanitize_affidavit_fact(value)
            if cleaned:
                normalized_facts.append(cleaned)
        normalized_supporting_exhibits: List[Dict[str, str]] = []
        for exhibit in _listify(affidavit_supporting_exhibits):
            if not isinstance(exhibit, dict):
                continue
            normalized = {
                "label": _clean_text(exhibit.get("label") or "Exhibit"),
                "title": _clean_text(exhibit.get("title") or exhibit.get("summary") or "Supporting exhibit"),
                "link": _clean_text(exhibit.get("link") or exhibit.get("reference") or ""),
                "summary": _clean_text(exhibit.get("summary") or ""),
            }
            if any(normalized.values()):
                normalized_supporting_exhibits.append(normalized)
        return {
            "title": _clean_text(affidavit_title) or None,
            "intro": _clean_text(affidavit_intro) or None,
            "facts": normalized_facts,
            "supporting_exhibits": normalized_supporting_exhibits,
            "include_complaint_exhibits": affidavit_include_complaint_exhibits,
            "venue_lines": [_clean_text(item) for item in _listify(affidavit_venue_lines) if _clean_text(item)],
            "jurat": _clean_text(affidavit_jurat) or None,
            "notary_block": [_clean_text(item) for item in _listify(affidavit_notary_block) if _clean_text(item)],
        }

    def _build_affidavit_venue_lines(self, caption: Dict[str, str]) -> List[str]:
        county = _clean_text(caption.get("county_line") or caption.get("county") or "")
        forum_type = _clean_text(caption.get("forum_type") or "").lower()
        lines: List[str] = []
        if county:
            lines.append(f"County: {county}")
        elif forum_type == "state":
            lines.append("County: __________________")
        elif forum_type == "federal":
            lines.append("State/District: __________________")
        return lines or ["Venue: __________________"]

    def _derive_affidavit_declarant_name(self, parties: Dict[str, List[str]], verification: Dict[str, Any], signature_block: Dict[str, Any]) -> str:
        signature_line = str(verification.get("signature_line") or "").strip()
        if signature_line.startswith("/s/ "):
            return signature_line[4:].strip() or str(signature_block.get("name") or "Plaintiff")
        return _clean_text(signature_block.get("name") or (parties.get("plaintiffs") or ["Plaintiff"])[0]) or "Plaintiff"

    def _collect_affidavit_facts(self, parties: Dict[str, List[str]], factual_allegations: List[str], claims: List[Dict[str, Any]]) -> List[str]:
        candidates: List[str] = []
        plaintiffs = [_clean_text(name) for name in _listify(parties.get("plaintiffs")) if _clean_text(name)]
        if plaintiffs:
            candidates.append(f"I am {plaintiffs[0]}, the plaintiff in this action.")
        candidates.extend(_clean_sentence(item) for item in _listify(factual_allegations) if _clean_text(item))

        normalized: List[str] = []
        seen = set()
        for item in candidates:
            fact = self._sanitize_affidavit_fact(item)
            if not fact:
                continue
            marker = fact.lower()
            if marker in seen:
                continue
            seen.add(marker)
            normalized.append(fact)
            if len(normalized) >= 12:
                break
        return normalized or ["Additional fact development is required before the affidavit can be finalized."]

    def _sanitize_affidavit_fact(self, value: Any) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            return ""
        text = re.sub(r"^As to [^,]+,\s*", "", text, flags=re.IGNORECASE)
        if ": " in text:
            prefix, suffix = text.split(": ", 1)
            prefix_lower = prefix.strip().lower()
            if (
                prefix.strip().endswith("?")
                or prefix_lower.startswith(("what ", "when ", "where ", "why ", "how ", "who ", "describe ", "explain "))
                or prefix_lower in {"what happened", "what relief do you want"}
            ):
                text = suffix.strip()
        lowered = text.lower()
        if lowered.startswith("plaintiff repeats and realleges"):
            return ""
        if not _is_factual_allegation_candidate(text) and not lowered.startswith("i am "):
            return ""
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        cleaned = _clean_sentence(text)
        return cleaned if len(cleaned) >= 12 else ""

    def _build_exhibits(self, evidence_records: Any, support_links: Any) -> List[Dict[str, Any]]:
        support_by_ref: Dict[str, Dict[str, Any]] = {}
        for link in _listify(support_links):
            if not isinstance(link, dict):
                continue
            support_ref = str(link.get("support_ref") or "").strip()
            if support_ref and support_ref not in support_by_ref:
                support_by_ref[support_ref] = link

        exhibits: List[Dict[str, Any]] = []
        for index, record in enumerate(_listify(evidence_records), 1):
            if not isinstance(record, dict):
                continue
            cid = str(record.get("cid") or "").strip()
            metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}
            parse_metadata = record.get("parse_metadata", {}) if isinstance(record.get("parse_metadata"), dict) else {}
            graph_summary = metadata.get("document_graph_summary", {}) if isinstance(metadata.get("document_graph_summary"), dict) else {}
            filename = _clean_text(metadata.get("filename") or parse_metadata.get("filename") or "")
            source_url = _clean_text(record.get("source_url") or metadata.get("source_url") or "")
            reference = source_url or (f"ipfs://{cid}" if cid else "")
            support_link = support_by_ref.get(cid) or support_by_ref.get(source_url) or {}
            summary = _clean_text(record.get("parsed_text_preview") or "")
            if not summary and record.get("id") not in (None, ""):
                get_evidence_facts = getattr(self.mediator, "get_evidence_facts", None)
                if callable(get_evidence_facts):
                    try:
                        fact_rows = get_evidence_facts(record.get("id")) or []
                    except Exception:
                        fact_rows = []
                    fact_lines = [
                        _clean_text(item.get("text") or "")
                        for item in _listify(fact_rows)
                        if isinstance(item, dict) and _clean_text(item.get("text") or "")
                    ]
                    if fact_lines:
                        summary = _clean_text("; ".join(fact_lines[:2]))
            if graph_summary and (graph_summary.get("entity_count") or graph_summary.get("relationship_count")):
                graph_text = _clean_text(
                    f"Graph extraction: {int(graph_summary.get('entity_count') or 0)} entities, {int(graph_summary.get('relationship_count') or 0)} relationships."
                )
                summary = _clean_text(" ".join(part for part in [summary, graph_text] if part))
            exhibits.append(
                {
                    "label": _exhibit_label(index),
                    "title": _clean_text(record.get("description") or filename or record.get("type") or f"Exhibit {index}"),
                    "description": _clean_text(record.get("description") or ""),
                    "claim_type": _clean_text(record.get("claim_type") or support_link.get("claim_type") or "").lower(),
                    "claim_element": _clean_text(record.get("claim_element") or support_link.get("claim_element_text") or ""),
                    "reference": reference,
                    "source_url": source_url,
                    "cid": cid,
                    "summary": summary,
                    "fact_count": int(record.get("fact_count") or 0),
                }
            )
        return exhibits

    def _build_claims(
        self,
        *,
        matching_results: Any,
        dependency_graph,
        legal_graph,
        authority_records: Any,
        exhibits: List[Dict[str, Any]],
        factual_allegations: List[str],
        user_id: str,
    ) -> List[Dict[str, Any]]:
        claim_entries = []
        raw_claims = []
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        if isinstance(matching_results, dict):
            raw_claims.extend(_listify(matching_results.get("claims")))
        if not raw_claims and dependency_graph is not None:
            for node in dependency_graph.get_nodes_by_type(NodeType.CLAIM):
                raw_claims.append(
                    {
                        "claim_name": node.name,
                        "claim_type": node.attributes.get("claim_type", node.name),
                        "satisfied_requirements": 0,
                        "legal_requirements": 0,
                        "requirements": [],
                        "missing_requirements": [],
                    }
                )

        for index, claim in enumerate(raw_claims, 1):
            if not isinstance(claim, dict):
                continue
            claim_name = _clean_text(claim.get("claim_name") or claim.get("title") or f"Claim {index}")
            claim_type = _clean_text(claim.get("claim_type") or claim_name).lower().replace(" ", "_")
            validation_summary = self._safe_call(
                "get_claim_support_validation",
                claim_type=claim_type,
                user_id=user_id,
                default={},
            )
            validation_claim = _extract_validation_claim_for_type(validation_summary, claim_type)
            reasoning_support_facts, reasoning_gap_hints = _build_claim_reasoning_temporal_fallbacks(
                validation_claim,
                claim_type=claim_type,
                claim_name=claim_name,
            )
            legal_requirements = legal_graph.get_requirements_for_claim_type(claim_type) if legal_graph is not None else []
            claim_authorities = self._filter_authorities(authority_records, claim_type)
            supporting_facts = self._build_supporting_facts(
                claim_type,
                claim_name,
                factual_allegations,
                user_id,
                reasoning_support_facts=reasoning_support_facts,
            )
            supporting_exhibits = [exhibit for exhibit in exhibits if not exhibit.get("claim_type") or exhibit.get("claim_type") == claim_type]
            legal_standard_elements = []
            for requirement in legal_requirements:
                legal_standard_elements.append(
                    {
                        "element": _clean_text(requirement.description or requirement.name),
                        "citation": _clean_text(requirement.citation),
                    }
                )
            if not legal_standard_elements:
                for requirement in _listify(claim.get("requirements")):
                    if not isinstance(requirement, dict):
                        continue
                    legal_standard_elements.append(
                        {
                            "element": _clean_text(requirement.get("requirement_description") or requirement.get("requirement_name") or ""),
                            "citation": _clean_text(requirement.get("citation") or ""),
                        }
                    )

            legal_standard = self._compose_legal_standard(claim_name, legal_standard_elements, claim_authorities)
            missing_requirements = [
                {
                    "name": _clean_text(item.get("requirement_name") or ""),
                    "citation": _clean_text(item.get("citation") or ""),
                    "suggested_action": _clean_text(item.get("suggested_action") or ""),
                }
                for item in _listify(claim.get("missing_requirements"))
                if isinstance(item, dict)
            ]
            missing_requirements.extend(
                _build_claim_temporal_gap_hints(
                    intake_case_file if isinstance(intake_case_file, dict) else {},
                    claim_type=claim_type,
                    claim_name=claim_name,
                )
            )
            missing_requirements.extend(reasoning_gap_hints)
            deduped_missing_requirements = []
            seen_missing_requirements = set()
            for requirement in missing_requirements:
                if not isinstance(requirement, dict):
                    continue
                marker = (
                    _clean_text(requirement.get("name") or "").lower(),
                    _clean_text(requirement.get("citation") or "").lower(),
                    _clean_text(requirement.get("suggested_action") or "").lower(),
                )
                if marker in seen_missing_requirements:
                    continue
                seen_missing_requirements.add(marker)
                deduped_missing_requirements.append(requirement)
            claim_entries.append(
                {
                    "count": index,
                    "title": f"COUNT {_roman_numeral(index)} - {claim_name.upper()}",
                    "claim_name": claim_name,
                    "claim_type": claim_type,
                    "description": _clean_sentence(claim.get("description") or f"Plaintiff realleges the foregoing paragraphs and asserts {claim_name}."),
                    "elements_satisfied": f"{int(claim.get('satisfied_requirements') or 0)}/{int(claim.get('legal_requirements') or 0)}",
                    "legal_standard": legal_standard,
                    "legal_standard_elements": legal_standard_elements,
                    "supporting_facts": supporting_facts,
                    "supporting_authorities": claim_authorities,
                    "supporting_exhibits": supporting_exhibits,
                    "missing_requirements": deduped_missing_requirements,
                }
            )
        return claim_entries

    def _build_supporting_facts(
        self,
        claim_type: str,
        claim_name: str,
        factual_allegations: List[str],
        user_id: str,
        *,
        reasoning_support_facts: Optional[List[str]] = None,
    ) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        chronology_support = _build_claim_chronology_support(
            intake_case_file if isinstance(intake_case_file, dict) else {},
            claim_type=claim_type,
            claim_name=claim_name,
        )
        reasoning_support = [
            _clean_sentence(item)
            for item in _listify(reasoning_support_facts)
            if _clean_text(item)
        ]
        facts = self._safe_call("get_claim_support_facts", claim_type, user_id, default=[])
        supporting_facts = []
        for fact in _listify(facts):
            if not isinstance(fact, dict):
                continue
            text = _clean_sentence(fact.get("text") or fact.get("fact_text") or "")
            if text:
                supporting_facts.append(text)
        if supporting_facts:
            return self._dedupe(chronology_support + reasoning_support + supporting_facts)
        lowered_claim_name = claim_name.lower()
        filtered = [
            allegation for allegation in factual_allegations
            if claim_type.replace("_", " ") in allegation.lower() or lowered_claim_name in allegation.lower()
        ]
        return self._dedupe(chronology_support + reasoning_support + (filtered or factual_allegations[:3]))

    def _filter_authorities(self, authority_records: Any, claim_type: str) -> List[Dict[str, Any]]:
        matched = []
        for authority in _listify(authority_records):
            if not isinstance(authority, dict):
                continue
            authority_claim_type = _clean_text(authority.get("claim_type") or "").lower()
            if authority_claim_type and authority_claim_type != claim_type:
                continue
            matched.append(
                {
                    "citation": _clean_text(authority.get("citation") or ""),
                    "title": _clean_text(authority.get("title") or authority.get("authority_type") or "Authority"),
                    "url": _clean_text(authority.get("url") or ""),
                    "relevance_score": authority.get("relevance_score"),
                }
            )
        return matched[:5]

    def _build_requested_relief(self, base_relief: Any) -> List[str]:
        relief_items = [_clean_sentence(item) for item in _listify(base_relief) if _clean_text(item)]
        if relief_items:
            return relief_items
        return [
            "Compensatory damages in an amount to be proven at trial",
            "Injunctive and declaratory relief sufficient to stop the challenged conduct",
            "Costs, fees, and any statutory fee-shifting relief authorized by law",
            "Such other and further relief as the Court deems just and proper",
        ]

    def _build_jury_demand(
        self,
        *,
        jury_demand: Optional[bool] = None,
        jury_demand_text: Optional[str] = None,
    ) -> Dict[str, str]:
        text = _clean_sentence(jury_demand_text)
        if text:
            return {
                "title": "Jury Demand",
                "text": text,
            }
        if jury_demand:
            return {
                "title": "Jury Demand",
                "text": "Plaintiff demands a trial by jury on all issues so triable.",
            }
        return {}

    def _build_nature_of_action(self, statement_of_claim: Any, claims: List[Dict[str, Any]]) -> str:
        statement = _clean_sentence(statement_of_claim)
        if statement:
            return statement
        claim_names = [claim.get("claim_name", "") for claim in claims if claim.get("claim_name")]
        if claim_names:
            if len(claim_names) == 1:
                return f"This civil action arises from {claim_names[0]} and seeks relief for the injuries caused by that conduct."
            joined = ", ".join(claim_names[:-1]) + f", and {claim_names[-1]}"
            return f"This civil action arises from {joined} and seeks relief for the injuries caused by that conduct."
        return "This civil action seeks relief for unlawful conduct described in the factual allegations below."

    def _build_jurisdiction_statement(self, jurisdiction: str, authority_records: Any) -> str:
        citations = []
        for authority in _listify(authority_records):
            if not isinstance(authority, dict):
                continue
            citation = _clean_text(authority.get("citation") or "")
            if citation:
                citations.append(citation)
        jurisdiction_lower = str(jurisdiction or "").lower()
        if jurisdiction_lower in {"federal", "us", "united states"}:
            if citations:
                return f"This Court has subject-matter jurisdiction under federal law, including {citations[0]}, because Plaintiff alleges violations arising under the Constitution, laws, or treaties of the United States."
            return "This Court has subject-matter jurisdiction under 28 U.S.C. § 1331 because Plaintiff alleges claims arising under federal law."
        return "This Court has subject-matter jurisdiction because the claims arise under the governing law identified in this pleading."

    def _build_venue_statement(
        self,
        district: Optional[str],
        county: Optional[str],
        division: Optional[str],
        jurisdiction: Optional[str],
    ) -> str:
        district_text = _clean_text(district)
        county_text = _clean_text(county)
        division_text = _clean_text(division)
        if county_text and str(jurisdiction).lower() not in {"federal", "us", "united states"}:
            return f"Venue is proper in this Court because a substantial part of the events or omissions giving rise to these claims occurred in {county_text}."
        if district_text and division_text:
            return f"Venue is proper in the {division_text} Division of the District of {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
        if district_text:
            return f"Venue is proper in the District of {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
        return "Venue is proper in this Court because a substantial part of the events or omissions giving rise to these claims occurred in this judicial district."

    def _build_signature_block(
        self,
        parties: Dict[str, List[str]],
        *,
        signer_name: Optional[str] = None,
        signer_title: Optional[str] = None,
        signer_firm: Optional[str] = None,
        signer_bar_number: Optional[str] = None,
        signer_contact: Optional[str] = None,
        additional_signers: Optional[List[Dict[str, str]]] = None,
        signature_date: Optional[str] = None,
    ) -> Dict[str, Any]:
        plaintiff_name = _clean_text(signer_name) or parties.get("plaintiffs", ["Plaintiff"])[0]
        return {
            "name": plaintiff_name,
            "signature_line": f"/s/ {plaintiff_name}",
            "title": _clean_text(signer_title) or "Plaintiff, Pro Se",
            "firm": _clean_text(signer_firm),
            "bar_number": _clean_text(signer_bar_number),
            "contact": _clean_text(signer_contact) or "Mailing address, telephone number, and email address to be completed before filing.",
            "additional_signers": self._normalize_additional_signers(additional_signers),
            "dated": self._format_dated_line("Dated", signature_date),
        }

    def _normalize_additional_signers(self, values: Any) -> List[Dict[str, str]]:
        normalized: List[Dict[str, str]] = []
        seen: set[tuple[str, str, str, str, str]] = set()
        for item in _listify(values):
            if not isinstance(item, dict):
                continue
            name = _clean_text(item.get("name") or item.get("signer_name"))
            title = _clean_text(item.get("title") or item.get("signer_title"))
            firm = _clean_text(item.get("firm") or item.get("signer_firm"))
            bar_number = _clean_text(item.get("bar_number") or item.get("signer_bar_number"))
            contact = _clean_text(item.get("contact") or item.get("signer_contact"))
            if not any((name, title, firm, bar_number, contact)):
                continue
            key = (name, title, firm, bar_number, contact)
            if key in seen:
                continue
            seen.add(key)
            normalized.append(
                {
                    "name": name or "Additional Counsel",
                    "signature_line": f"/s/ {name}" if name else "",
                    "title": title,
                    "firm": firm,
                    "bar_number": bar_number,
                    "contact": contact,
                }
            )
        return normalized

    def _signature_block_lines(self, signature_block: Dict[str, Any]) -> List[str]:
        lines: List[str] = []
        if signature_block.get("signature_line"):
            lines.append(str(signature_block["signature_line"]))
        lines.append(str(signature_block.get("name") or "Plaintiff"))
        if signature_block.get("title"):
            lines.append(str(signature_block["title"]))
        if signature_block.get("firm"):
            lines.append(str(signature_block["firm"]))
        if signature_block.get("bar_number"):
            lines.append(f"Bar No. {signature_block['bar_number']}")
        if signature_block.get("contact"):
            lines.append(str(signature_block["contact"]))
        for signer in _listify(signature_block.get("additional_signers")):
            if not isinstance(signer, dict):
                continue
            lines.append("")
            if signer.get("signature_line"):
                lines.append(str(signer["signature_line"]))
            lines.append(str(signer.get("name") or "Additional Counsel"))
            if signer.get("title"):
                lines.append(str(signer["title"]))
            if signer.get("firm"):
                lines.append(str(signer["firm"]))
            if signer.get("bar_number"):
                lines.append(f"Bar No. {signer['bar_number']}")
            if signer.get("contact"):
                lines.append(str(signer["contact"]))
        if signature_block.get("dated"):
            lines.append(str(signature_block["dated"]))
        return lines

    def _build_verification(
        self,
        parties: Dict[str, List[str]],
        *,
        declarant_name: Optional[str] = None,
        signer_name: Optional[str] = None,
        verification_date: Optional[str] = None,
    ) -> Dict[str, str]:
        plaintiff_name = _clean_text(declarant_name) or _clean_text(signer_name) or parties.get("plaintiffs", ["Plaintiff"])[0]
        return {
            "title": "Verification",
            "text": (
                f"I, {plaintiff_name}, declare under penalty of perjury that I have reviewed this Complaint "
                "and that the factual allegations stated in it are true and correct to the best of my knowledge, "
                "information, and belief."
            ),
            "dated": self._format_dated_line("Executed on", verification_date),
            "signature_line": f"/s/ {plaintiff_name}",
        }

    def _build_certificate_of_service(
        self,
        parties: Dict[str, List[str]],
        *,
        signer_name: Optional[str] = None,
        service_method: Optional[str] = None,
        service_recipients: Optional[List[str]] = None,
        service_recipient_details: Optional[List[Dict[str, str]]] = None,
        service_date: Optional[str] = None,
        jurisdiction: Optional[str] = None,
    ) -> Dict[str, Any]:
        plaintiff_name = _clean_text(signer_name) or parties.get("plaintiffs", ["Plaintiff"])[0]
        recipient_details = self._normalize_service_recipient_details(service_recipient_details)
        recipient_names = [detail["recipient"] for detail in recipient_details if detail.get("recipient")]
        recipients_list = self._dedupe(
            list(_clean_text(item) for item in _listify(service_recipients)) + recipient_names
        ) or parties.get("defendants", []) or ["all defendants"]
        recipients = ", ".join(recipients_list)
        method_text = _clean_text(service_method) or "a method authorized by the applicable rules of civil procedure"
        detail_lines = [self._format_service_recipient_detail(detail) for detail in recipient_details]
        if detail_lines:
            text = (
                ("I declare that a true and correct copy of this Complaint will be served promptly after filing on the following recipients:"
                if str(jurisdiction or "").strip().lower() == "state"
                else "I certify that a true and correct copy of this Complaint will be served promptly after filing on the following recipients:")
            )
        else:
            text = (
                (("I declare that a true and correct copy of this Complaint will be served on "
                if str(jurisdiction or "").strip().lower() == "state"
                else "I certify that a true and correct copy of this Complaint will be served on ")
                + f"{recipients} using {method_text} promptly after filing.")
            )
        return {
            "title": "Proof of Service" if str(jurisdiction or "").strip().lower() == "state" else "Certificate of Service",
            "text": text,
            "recipients": recipients_list,
            "recipient_details": recipient_details,
            "detail_lines": detail_lines,
            "dated": self._format_dated_line("Service date", service_date),
            "signature_line": f"/s/ {plaintiff_name}",
        }

    def _normalize_service_recipient_details(self, values: Any) -> List[Dict[str, str]]:
        details: List[Dict[str, str]] = []
        seen: set[tuple[str, str, str, str]] = set()
        for item in _listify(values):
            if not isinstance(item, dict):
                continue
            detail = {
                "recipient": _clean_text(item.get("recipient")),
                "method": _clean_text(item.get("method")),
                "address": _clean_text(item.get("address")),
                "notes": _clean_text(item.get("notes")),
            }
            if not any(detail.values()):
                continue
            key = (detail["recipient"], detail["method"], detail["address"], detail["notes"])
            if key in seen:
                continue
            seen.add(key)
            details.append(detail)
        return details

    def _format_service_recipient_detail(self, detail: Dict[str, str]) -> str:
        segments = [detail.get("recipient") or "Recipient"]
        if detail.get("method"):
            segments.append(f"Method: {detail['method']}")
        if detail.get("address"):
            segments.append(f"Address: {detail['address']}")
        if detail.get("notes"):
            segments.append(f"Notes: {detail['notes']}")
        return " | ".join(segment for segment in segments if segment)

    def _format_dated_line(self, label: str, value: Optional[str]) -> str:
        cleaned = _clean_text(value)
        return f"{label}: {cleaned}" if cleaned else f"{label}: __________________"

    def _compose_legal_standard(self, claim_name: str, legal_standard_elements: List[Dict[str, str]], authorities: List[Dict[str, Any]]) -> str:
        if legal_standard_elements:
            citations = [item.get("citation", "") for item in legal_standard_elements if item.get("citation")]
            if citations:
                return f"To state {claim_name}, Plaintiff must allege facts satisfying the governing elements recognized by {citations[0]}."
            return f"To state {claim_name}, Plaintiff must allege facts satisfying each essential element recognized by governing law."
        if authorities:
            authority = authorities[0]
            reference = authority.get("citation") or authority.get("title") or "governing authority"
            return f"The legal standard for {claim_name} is supplied by {reference}."
        return f"Plaintiff alleges facts sufficient to satisfy the legal standard for {claim_name}."

    def _infer_jurisdiction(self, legal_graph) -> str:
        if legal_graph is None:
            return "federal"
        jurisdictions = []
        for element in getattr(legal_graph, "elements", {}).values():
            jurisdiction = _clean_text(getattr(element, "jurisdiction", ""))
            if jurisdiction:
                jurisdictions.append(jurisdiction)
        if not jurisdictions:
            return "federal"
        first = jurisdictions[0].lower()
        if first in {"us", "united states"}:
            return "federal"
        return first

    def _dedupe(self, values: Iterable[str]) -> List[str]:
        seen = set()
        deduped = []
        for value in values:
            cleaned = _clean_text(value)
            if not cleaned:
                continue
            marker = cleaned.lower()
            if marker in seen:
                continue
            seen.add(marker)
            deduped.append(cleaned)
        return deduped

    def _write_docx(self, draft: Dict[str, Any], destination: Path) -> None:
        if not HAS_DOCX:
            raise RuntimeError("DOCX export requires python-docx to be installed")
        document = DocxDocument()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        for line in [draft.get("court_header"), draft.get("caption", {}).get("division_line")]:
            if not line:
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(line))
            run.bold = True

        caption = draft.get("caption", {}) if isinstance(draft.get("caption"), dict) else {}
        table = document.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(4.75)
        table.columns[1].width = Inches(2.0)
        resolved_case_number = caption.get("case_number") or draft.get("case_number") or "________________"
        table.cell(0, 0).text = "\n".join(build_caption_party_block_lines(draft.get("parties", {})))
        table.cell(0, 1).text = "\n".join(
            build_case_detail_lines({"case_number": resolved_case_number, **caption})
        )

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("COMPLAINT").bold = True
        if caption.get("jury_demand_notice"):
            jury_notice = document.add_paragraph()
            jury_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            jury_notice.add_run(str(caption["jury_demand_notice"])).bold = True

        self._docx_section(document, "Nature of the Action", [draft.get("nature_of_action")])
        party_lines = []
        for plaintiff in draft.get("parties", {}).get("plaintiffs", []):
            party_lines.append(f"Plaintiff {plaintiff} is an aggrieved party bringing this action.")
        for defendant in draft.get("parties", {}).get("defendants", []):
            party_lines.append(f"Defendant {defendant} is alleged to be responsible for the conduct described below.")
        self._docx_section(document, "Parties", party_lines)
        self._docx_section(document, "Jurisdiction and Venue", [draft.get("jurisdiction_statement"), draft.get("venue_statement")])
        allegation_groups = draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else []
        if allegation_groups:
            heading_paragraph = document.add_paragraph()
            heading_paragraph.add_run("Factual Allegations").bold = True
            for group in allegation_groups:
                if not isinstance(group, dict):
                    continue
                if group.get("title"):
                    group_heading = document.add_paragraph()
                    group_heading.add_run(str(group["title"])).italic = True
                for entry in _listify(group.get("paragraphs")):
                    if not isinstance(entry, dict):
                        continue
                    text = _clean_sentence(entry.get("text"))
                    if text:
                        document.add_paragraph(f"{entry.get('number')}. {text}")
        else:
            self._docx_section(document, "Factual Allegations", draft.get("factual_allegations", []), numbered=True)

        chronology_lines = [
            _clean_sentence(item)
            for item in _listify(draft.get("anchored_chronology_summary"))
            if _clean_text(item)
        ]
        if chronology_lines:
            self._docx_section(document, "Anchored Chronology", chronology_lines, numbered=True)

        claims_heading = document.add_paragraph()
        claims_heading.add_run("Claims for Relief").bold = True
        for claim in _listify(draft.get("legal_claims")):
            claim_heading = document.add_paragraph()
            claim_heading.add_run(str(claim.get("title") or "Claim")).bold = True
            if claim.get("description"):
                document.add_paragraph(_clean_sentence(claim.get("description")))
            if claim.get("legal_standard"):
                paragraph = document.add_paragraph()
                paragraph.add_run("Legal Standard: ").bold = True
                paragraph.add_run(_clean_sentence(claim.get("legal_standard")))
            for item in _listify(claim.get("legal_standard_elements")):
                paragraph = document.add_paragraph(style="List Bullet")
                text = _clean_sentence(item.get("element") or "")
                citation = _clean_text(item.get("citation") or "")
                paragraph.add_run(f"{text} ({citation})" if citation else text)
            for fact in _listify(claim.get("supporting_facts")):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(_clean_sentence(fact))
            exhibit_labels = ", ".join(exhibit.get("label", "") for exhibit in _listify(claim.get("supporting_exhibits")) if exhibit.get("label"))
            if exhibit_labels:
                document.add_paragraph(f"Supported by {exhibit_labels}.")

        self._docx_section(document, "Prayer for Relief", draft.get("requested_relief", []), bulleted=True)
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            self._docx_section(document, jury_demand.get("title") or "Jury Demand", [jury_demand.get("text")])
        exhibit_lines = []
        for exhibit in _listify(draft.get("exhibits")):
            line = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
            if exhibit.get("reference"):
                line = f"{line} ({exhibit['reference']})"
            exhibit_lines.append(line)
        self._docx_section(document, "Exhibits", exhibit_lines)
        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            affidavit_lines = list(_listify(affidavit.get("venue_lines"))) + [affidavit.get("intro"), affidavit.get("knowledge_graph_note")]
            affidavit_lines.extend(
                f"{index}. {_clean_sentence(fact)}"
                for index, fact in enumerate(_listify(affidavit.get("facts")), 1)
                if _clean_text(fact)
            )
            affidavit_lines.extend(
                f"{exhibit.get('label', 'Exhibit')} - {_clean_text(exhibit.get('title') or 'Supporting exhibit')}"
                + (f" ({exhibit['link']})" if exhibit.get("link") else "")
                for exhibit in _listify(affidavit.get("supporting_exhibits"))
                if isinstance(exhibit, dict)
            )
            affidavit_lines.extend([affidavit.get("dated"), affidavit.get("signature_line"), affidavit.get("jurat")])
            affidavit_lines.extend(_listify(affidavit.get("notary_block")))
            self._docx_section(document, affidavit.get("title") or "Affidavit in Support of Complaint", affidavit_lines)
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            self._docx_section(
                document,
                verification.get("title") or "Verification",
                [verification.get("text"), verification.get("dated"), verification.get("signature_line")],
            )
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            self._docx_section(
                document,
                certificate_of_service.get("title") or "Certificate of Service",
                [certificate_of_service.get("text")]
                + _listify(certificate_of_service.get("detail_lines"))
                + [certificate_of_service.get("dated"), certificate_of_service.get("signature_line")],
            )
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        self._docx_section(
            document,
            "Signature Block",
            ["Respectfully submitted,", *self._signature_block_lines(signature_block)],
        )
        document.save(str(destination))

    def _docx_section(self, document, heading: str, paragraphs: Sequence[Any], *, numbered: bool = False, bulleted: bool = False) -> None:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.add_run(heading).bold = True
        for index, value in enumerate(_listify(paragraphs), 1):
            text = _clean_sentence(value)
            if not text:
                continue
            if numbered:
                document.add_paragraph(f"{index}. {text}")
            elif bulleted:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(text)
            else:
                document.add_paragraph(text)

    def _write_pdf(self, draft: Dict[str, Any], destination: Path) -> None:
        if not HAS_REPORTLAB:
            raise RuntimeError("PDF export requires reportlab to be installed")
        caption = draft.get("caption", {}) if isinstance(draft.get("caption"), dict) else {}
        styles = getSampleStyleSheet()
        centered = ParagraphStyle("ComplaintCentered", parent=styles["Heading2"], alignment=TA_CENTER)
        section = ParagraphStyle("ComplaintSection", parent=styles["Heading3"], spaceBefore=10, spaceAfter=6)
        body = ParagraphStyle("ComplaintBody", parent=styles["BodyText"], leading=15)
        story = []
        story.append(Paragraph(str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"), centered))
        if caption.get("division_line"):
            story.append(Paragraph(str(caption["division_line"]), centered))
        story.append(Spacer(1, 0.2 * inch))

        story.append(
            build_case_caption_table(
                draft.get("parties", {}),
                {"case_number": draft.get("case_number") or caption.get("case_number"), **caption},
                body,
                left_width=4.5 * inch,
                right_width=2.0 * inch,
            )
        )
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("COMPLAINT", centered))
        if caption.get("jury_demand_notice"):
            story.append(Paragraph(str(caption["jury_demand_notice"]), centered))

        self._pdf_section(story, section, body, "Nature of the Action", [draft.get("nature_of_action")])
        party_lines = []
        for plaintiff in draft.get("parties", {}).get("plaintiffs", []):
            party_lines.append(f"Plaintiff {plaintiff} is an aggrieved party bringing this action.")
        for defendant in draft.get("parties", {}).get("defendants", []):
            party_lines.append(f"Defendant {defendant} is alleged to be responsible for the conduct described below.")
        self._pdf_section(story, section, body, "Parties", party_lines)
        self._pdf_section(story, section, body, "Jurisdiction and Venue", [draft.get("jurisdiction_statement"), draft.get("venue_statement")])
        allegation_groups = draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else []
        if allegation_groups:
            story.append(Paragraph("Factual Allegations", section))
            for group in allegation_groups:
                if not isinstance(group, dict):
                    continue
                if group.get("title"):
                    story.append(Paragraph(str(group["title"]), body))
                for entry in _listify(group.get("paragraphs")):
                    if not isinstance(entry, dict):
                        continue
                    text = _clean_sentence(entry.get("text"))
                    if text:
                        story.append(Paragraph(f"{entry.get('number')}. {text}", body))
        else:
            self._pdf_section(story, section, body, "Factual Allegations", draft.get("factual_allegations", []), numbered=True)

        chronology_lines = [
            _clean_sentence(item)
            for item in _listify(draft.get("anchored_chronology_summary"))
            if _clean_text(item)
        ]
        if chronology_lines:
            self._pdf_section(story, section, body, "Anchored Chronology", chronology_lines, numbered=True)

        story.append(Paragraph("Claims for Relief", section))
        for claim in _listify(draft.get("legal_claims")):
            story.append(Paragraph(str(claim.get("title") or "Claim"), section))
            if claim.get("description"):
                story.append(Paragraph(_clean_sentence(claim.get("description")), body))
            if claim.get("legal_standard"):
                story.append(Paragraph(f"<b>Legal Standard:</b> {_clean_sentence(claim.get('legal_standard'))}", body))
            for item in _listify(claim.get("legal_standard_elements")):
                text = _clean_sentence(item.get("element") or "")
                citation = _clean_text(item.get("citation") or "")
                story.append(Paragraph(f"• {text} ({citation})" if citation else f"• {text}", body))
            for fact in _listify(claim.get("supporting_facts")):
                story.append(Paragraph(f"• {_clean_sentence(fact)}", body))
            exhibit_labels = ", ".join(exhibit.get("label", "") for exhibit in _listify(claim.get("supporting_exhibits")) if exhibit.get("label"))
            if exhibit_labels:
                story.append(Paragraph(f"Supported by {exhibit_labels}.", body))

        self._pdf_section(story, section, body, "Prayer for Relief", draft.get("requested_relief", []), bulleted=True)
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            self._pdf_section(story, section, body, jury_demand.get("title") or "Jury Demand", [jury_demand.get("text")])
        exhibit_lines = []
        for exhibit in _listify(draft.get("exhibits")):
            line = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
            if exhibit.get("reference"):
                line = f"{line} ({exhibit['reference']})"
            exhibit_lines.append(line)
        self._pdf_section(story, section, body, "Exhibits", exhibit_lines)
        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            affidavit_lines = list(_listify(affidavit.get("venue_lines"))) + [affidavit.get("intro"), affidavit.get("knowledge_graph_note")]
            affidavit_lines.extend(
                f"{index}. {_clean_sentence(fact)}"
                for index, fact in enumerate(_listify(affidavit.get("facts")), 1)
                if _clean_text(fact)
            )
            affidavit_lines.extend(
                f"{exhibit.get('label', 'Exhibit')} - {_clean_text(exhibit.get('title') or 'Supporting exhibit')}"
                + (f" ({exhibit['link']})" if exhibit.get("link") else "")
                for exhibit in _listify(affidavit.get("supporting_exhibits"))
                if isinstance(exhibit, dict)
            )
            affidavit_lines.extend([affidavit.get("dated"), affidavit.get("signature_line"), affidavit.get("jurat")])
            affidavit_lines.extend(_listify(affidavit.get("notary_block")))
            self._pdf_section(story, section, body, affidavit.get("title") or "Affidavit in Support of Complaint", affidavit_lines)
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            self._pdf_section(
                story,
                section,
                body,
                verification.get("title") or "Verification",
                [verification.get("text"), verification.get("dated"), verification.get("signature_line")],
            )
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            self._pdf_section(
                story,
                section,
                body,
                certificate_of_service.get("title") or "Certificate of Service",
                [certificate_of_service.get("text")]
                + _listify(certificate_of_service.get("detail_lines"))
                + [certificate_of_service.get("dated"), certificate_of_service.get("signature_line")],
            )
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        self._pdf_section(
            story,
            section,
            body,
            "Signature Block",
            ["Respectfully submitted,", *self._signature_block_lines(signature_block)],
        )
        document = SimpleDocTemplate(str(destination), pagesize=letter)
        footer = make_page_footer_renderer("Complaint")
        document.build(story, onFirstPage=footer, onLaterPages=footer)

    def _pdf_section(self, story, section_style, body_style, heading: str, paragraphs: Sequence[Any], *, numbered: bool = False, bulleted: bool = False) -> None:
        story.append(Paragraph(heading, section_style))
        for index, value in enumerate(_listify(paragraphs), 1):
            text = _clean_sentence(value)
            if not text:
                continue
            if numbered:
                story.append(Paragraph(f"{index}. {text}", body_style))
            elif bulleted:
                story.append(Paragraph(f"• {text}", body_style))
            else:
                story.append(Paragraph(text, body_style))
